"""Build + store the downloadable compliance report (JSON / CSV / DOCX / PDF).

Sync (called via ``run_sync`` from the audit graph's report_build node). Embeds
matplotlib charts into a python-docx report and converts to PDF via LibreOffice
when available. Returns a dict of stored-artifact URIs keyed by format.
"""
from __future__ import annotations

import csv
import datetime as dt
import io
import json
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from app.agents.nodes.common import DOCX_MIME, PDF_MIME
from app.core.logging import get_logger
from app.services.compliance import charts, scoring
from app.storage import get_storage

log = get_logger(__name__)

_BRAND = RGBColor(0x05, 0x1D, 0x60)
_STATUS_COLOR = {
    "compliant": RGBColor(0x05, 0x96, 0x69),
    "partial": RGBColor(0xB4, 0x53, 0x09),
    "non_compliant": RGBColor(0xBE, 0x12, 0x3B),
    "not_applicable": RGBColor(0x64, 0x74, 0x8B),
}
_SEV_RANK = {"critical": 0, "major": 1, "minor": 2, "info": 3}


# ── JSON / CSV ───────────────────────────────────────────────────────────────
def _build_json(run_id: str, compliance: dict, findings: list[dict], doc_meta: dict) -> bytes:
    payload = {
        "run_id": run_id,
        "generated_at": dt.datetime.now(dt.timezone.utc).isoformat(),
        "document": doc_meta,
        "guideline": compliance.get("guideline"),
        "overall_score": compliance.get("overall_score"),
        "status_label": compliance.get("status_label"),
        "per_dimension": compliance.get("per_dimension"),
        "per_section": compliance.get("per_section"),
        "severity_counts": compliance.get("severity_counts"),
        "summary": compliance.get("summary"),
        "findings": findings,
    }
    return json.dumps(payload, indent=2, default=str).encode("utf-8")


def _build_csv(findings: list[dict]) -> bytes:
    cols = [
        "section", "dimension", "status", "severity", "confidence",
        "requirement_title", "doc_location", "evidence", "rationale",
        "suggested_fix", "citation_section", "citation_quote",
    ]
    buf = io.StringIO()
    w = csv.DictWriter(buf, fieldnames=cols, extrasaction="ignore")
    w.writeheader()
    for f in findings:
        cit = f.get("citation") or {}
        w.writerow(
            {
                "section": f.get("section") or f.get("section_no"),
                "dimension": f.get("dimension"),
                "status": f.get("status"),
                "severity": f.get("severity"),
                "confidence": f.get("confidence"),
                "requirement_title": f.get("requirement_title"),
                "doc_location": f.get("doc_location"),
                "evidence": f.get("evidence"),
                "rationale": f.get("rationale"),
                "suggested_fix": f.get("suggested_fix"),
                "citation_section": cit.get("guideline_section"),
                "citation_quote": cit.get("quote"),
            }
        )
    return buf.getvalue().encode("utf-8")


# ── DOCX ─────────────────────────────────────────────────────────────────────
def _add_picture(doc: Document, png: bytes, width_in: float) -> None:
    try:
        doc.add_picture(io.BytesIO(png), width=Inches(width_in))
    except Exception as e:  # noqa: BLE001
        log.warning("chart embed failed: %s", e)


def _build_docx(run_id: str, compliance: dict, findings: list[dict], doc_meta: dict) -> bytes:
    doc = Document()
    g = compliance.get("guideline") or {}
    score = float(compliance.get("overall_score") or 0.0)
    label = compliance.get("status_label", "n/a")

    h = doc.add_heading("Compliance Report", level=0)
    sub = doc.add_paragraph()
    run = sub.add_run(
        f"{g.get('code','Guideline')} — {g.get('title','')}"
        + (f" ({g.get('version')})" if g.get("version") else "")
    )
    run.bold = True
    meta_line = doc.add_paragraph()
    meta_line.add_run(
        f"Document: {doc_meta.get('title') or doc_meta.get('source_file') or 'Uploaded document'}    "
        f"Generated: {dt.datetime.now(dt.timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    ).font.size = Pt(9)

    # Overall verdict + donut
    verdict = doc.add_paragraph()
    vr = verdict.add_run(f"Overall: {round(score * 100)}%  —  {label.upper()}")
    vr.bold = True
    vr.font.size = Pt(14)
    vr.font.color.rgb = _BRAND
    try:
        _add_picture(doc, charts.overall_donut(score), 2.6)
    except Exception as e:  # noqa: BLE001
        log.warning("donut failed: %s", e)

    if compliance.get("summary"):
        doc.add_heading("Executive summary", level=1)
        doc.add_paragraph(compliance["summary"])

    # Dimension + severity charts
    doc.add_heading("Scores by dimension", level=1)
    try:
        _add_picture(doc, charts.dimension_bars(compliance.get("per_dimension") or {}), 5.2)
    except Exception as e:  # noqa: BLE001
        log.warning("dimension chart failed: %s", e)

    sev = compliance.get("severity_counts") or {}
    doc.add_heading("Open issues by severity", level=1)
    doc.add_paragraph(
        f"Critical: {sev.get('critical',0)}  •  Major: {sev.get('major',0)}  •  "
        f"Minor: {sev.get('minor',0)}  •  Info: {sev.get('info',0)}"
    )
    try:
        _add_picture(doc, charts.severity_bar(sev), 4.4)
    except Exception as e:  # noqa: BLE001
        log.warning("severity chart failed: %s", e)

    # Per-section table + chart
    per_section = compliance.get("per_section") or []
    if per_section:
        doc.add_heading("Compliance by section", level=1)
        try:
            _add_picture(doc, charts.section_bars(per_section), 5.6)
        except Exception as e:  # noqa: BLE001
            log.warning("section chart failed: %s", e)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Section", "Score", "Status", "Open issues"
        for s in per_section:
            cells = table.add_row().cells
            cells[0].text = f"{s.get('section','')}. {s.get('title','')[:40]}"
            cells[1].text = f"{round((s.get('score') or 0) * 100)}%" if s.get("score") is not None else "—"
            cells[2].text = str(s.get("status", ""))
            cells[3].text = str(s.get("findings_count", 0))

    # Findings grouped by top-level section
    doc.add_heading("Detailed findings", level=1)
    by_top: dict[str, list[dict]] = {}
    for f in findings:
        by_top.setdefault(scoring.top_section(f.get("section") or f.get("section_no")), []).append(f)

    def _sort_top(sec: str):
        return (0, int(sec)) if sec.isdigit() else (1, 0)

    for top in sorted(by_top, key=_sort_top):
        group = sorted(by_top[top], key=lambda x: _SEV_RANK.get(x.get("severity", "minor"), 9))
        doc.add_heading(f"Section {top}", level=2)
        for f in group:
            p = doc.add_paragraph(style="List Bullet")
            sr = p.add_run(f"[{(f.get('status') or '').upper()}] ")
            sr.bold = True
            sr.font.color.rgb = _STATUS_COLOR.get(f.get("status"), _BRAND)
            p.add_run(f"{f.get('requirement_title','')}  ")
            tag = p.add_run(f"({f.get('severity','')} · {f.get('dimension','')})")
            tag.font.size = Pt(8)
            tag.italic = True
            if f.get("evidence"):
                ev = doc.add_paragraph()
                ev.paragraph_format.left_indent = Inches(0.5)
                ev.add_run("Evidence: ").bold = True
                ev.add_run(str(f["evidence"])[:400]).italic = True
            cit = f.get("citation") or {}
            if cit.get("quote"):
                cp = doc.add_paragraph()
                cp.paragraph_format.left_indent = Inches(0.5)
                cp.add_run(f"Guideline [{cit.get('guideline_section','')}]: ").bold = True
                cp.add_run(str(cit.get("quote"))[:400])
            if f.get("suggested_fix"):
                fx = doc.add_paragraph()
                fx.paragraph_format.left_indent = Inches(0.5)
                r = fx.add_run("Fix: ")
                r.bold = True
                fx.add_run(str(f["suggested_fix"])[:400])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── orchestrator ─────────────────────────────────────────────────────────────
def build_and_store_report(
    project_id: Any,
    run_id: Any,
    compliance: dict,
    findings: list[dict],
    doc_meta: dict,
) -> dict[str, str]:
    """Produce JSON/CSV/DOCX/PDF, store each, return {fmt: uri}."""
    run_id = str(run_id)
    storage = get_storage()

    def _store(data: bytes, fmt: str, mime: str) -> str:
        key = storage.make_key(
            project_id=str(project_id or "_compliance"),
            kind="compliance_report",
            filename=f"compliance-report.{fmt}",
        )
        return storage.put(data, key=key, content_type=mime).uri

    out: dict[str, str] = {}
    try:
        out["json"] = _store(_build_json(run_id, compliance, findings, doc_meta), "json", "application/json")
    except Exception as e:  # noqa: BLE001
        log.warning("report json failed: %s", e)
    try:
        out["csv"] = _store(_build_csv(findings), "csv", "text/csv")
    except Exception as e:  # noqa: BLE001
        log.warning("report csv failed: %s", e)

    docx_bytes: bytes | None = None
    try:
        docx_bytes = _build_docx(run_id, compliance, findings, doc_meta)
        out["docx"] = _store(docx_bytes, "docx", DOCX_MIME)
    except Exception as e:  # noqa: BLE001
        log.warning("report docx failed: %s", e)

    if docx_bytes:
        try:
            from app.services.office.office_pipeline import available, export_pdf

            if available():
                out["pdf"] = _store(export_pdf(docx_bytes), "pdf", PDF_MIME)
        except Exception as e:  # noqa: BLE001
            log.warning("report pdf failed: %s", e)

    return out
