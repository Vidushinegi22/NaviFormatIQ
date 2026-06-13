"""
Word Document Extractor
=======================
Extracts content structure and styling from .docx files into JSON.

Produces two outputs:
  1. DocumentContent  — structured content (headings, paragraphs, images, tables, lists)
  2. DocumentStyling  — reusable style definitions (fonts, sizes, colors, spacing)

Uses python-docx for deep inspection of every formatting property.
"""

from __future__ import annotations

import base64
import hashlib
import io
import math
import os
import re
from datetime import datetime, timezone
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_UNDERLINE,
)
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

from app.schemas.document_model import (
    Alignment,
    BorderDef,
    BorderStyle,
    CellBorders,
    CellStyle,
    ContentElement,
    DocumentContent,
    DocumentMetadata,
    DocumentStyling,
    ElementType,
    ImagePosition,
    IndentStyle,
    ListType,
    PageMargins,
    PageStyle,
    ParagraphStyle,
    RunStyle,
    StyleMetadata,
    TableCell,
    TableRow,
    TableStyle,
    TextRun,
    UnderlineType,
    VerticalAlignment,
    Orientation,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_ALIGNMENT_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: Alignment.LEFT,
    WD_ALIGN_PARAGRAPH.CENTER: Alignment.CENTER,
    WD_ALIGN_PARAGRAPH.RIGHT: Alignment.RIGHT,
    WD_ALIGN_PARAGRAPH.JUSTIFY: Alignment.JUSTIFY,
}

_UNDERLINE_MAP: dict = {
    True: UnderlineType.SINGLE,
    False: UnderlineType.NONE,
    None: None,
    WD_UNDERLINE.SINGLE: UnderlineType.SINGLE,
    WD_UNDERLINE.DOUBLE: UnderlineType.DOUBLE,
    WD_UNDERLINE.DOTTED: UnderlineType.DOTTED,
    WD_UNDERLINE.DASH: UnderlineType.DASHED,
    WD_UNDERLINE.WAVY: UnderlineType.WAVY,
    WD_UNDERLINE.THICK: UnderlineType.THICK,
    WD_UNDERLINE.WORDS: UnderlineType.WORDS_ONLY,
}

_VERT_ALIGN_MAP = {
    WD_ALIGN_VERTICAL.TOP: VerticalAlignment.TOP,
    WD_ALIGN_VERTICAL.CENTER: VerticalAlignment.CENTER,
    WD_ALIGN_VERTICAL.BOTTOM: VerticalAlignment.BOTTOM,
}


def _emu_to_inches(emu: Optional[int]) -> Optional[float]:
    if emu is None:
        return None
    return round(emu / 914400, 4)


def _pt_value(val) -> Optional[float]:
    """Convert a python-docx Length (or plain number) to points.

    NB: python-docx Length subclasses (Emu, Pt, Inches, etc.) are *int*
    subclasses whose raw integer value is in EMU (1pt = 12700 EMU).
    We therefore have to check for the ``.pt`` property *before* the
    int/float branch — otherwise we'd return the raw EMU integer as if
    it were already in points and the caller would re-multiply it by
    another 12700 when wrapping with ``Pt(...)``.
    """
    if val is None:
        return None
    if isinstance(val, bool):
        return None
    # Length objects expose a ``.pt`` float — always prefer that.
    pt_attr = getattr(val, "pt", None)
    if pt_attr is not None:
        try:
            return round(float(pt_attr), 2)
        except (TypeError, ValueError):
            pass
    if isinstance(val, (int, float)):
        return round(float(val), 2)
    return None


def _rgb_to_hex(rgb: Optional[RGBColor]) -> Optional[str]:
    if rgb is None:
        return None
    return f"#{rgb}"


def _make_style_key(prefix: str, *parts: str) -> str:
    """Create a deterministic, filesystem-safe style key."""
    raw = f"{prefix}_{'_'.join(parts)}"
    # Sanitize
    raw = re.sub(r"[^a-zA-Z0-9_]", "_", raw)
    raw = re.sub(r"_+", "_", raw).strip("_").lower()
    if len(raw) > 80:
        raw = raw[:60] + "_" + hashlib.md5(raw.encode()).hexdigest()[:8]
    return raw


def _fingerprint_run_style(rs: RunStyle) -> str:
    """Return a hashable fingerprint for a RunStyle to enable deduplication."""
    parts = []
    for field_name in rs.model_fields:
        val = getattr(rs, field_name)
        if val is not None:
            parts.append(f"{field_name}={val}")
    return "|".join(parts) if parts else "__default__"


def _fingerprint_para_style(ps: ParagraphStyle) -> str:
    parts = []
    for field_name in ps.model_fields:
        val = getattr(ps, field_name)
        if val is not None:
            if isinstance(val, IndentStyle):
                parts.append(f"indent={val.model_dump_json()}")
            else:
                parts.append(f"{field_name}={val}")
    return "|".join(parts) if parts else "__default__"


def _count_words(elements: list) -> int:
    """Count words across body text + table cells for a page-count estimate."""
    total = 0
    for el in elements:
        for run in el.content or []:
            total += len((run.text or "").split())
        for row in el.rows or []:
            for cell in row.cells:
                for run in cell.content or []:
                    total += len((run.text or "").split())
    return total


# ---------------------------------------------------------------------------
# Page-count resolution
# ---------------------------------------------------------------------------
#
# A .docx body carries NO rendered page count — Word computes pagination at
# render time from fonts, spacing, images, tables and page breaks. The old code
# guessed ``word_count // 350``, which was wrong for almost every real document
# (e.g. a 2-3 page SOP reported as 1 page). We resolve the page count from the
# best authoritative source available, in order:
#
#   1. Microsoft Word COM on Windows — exact, using Word's own layout engine.
#   2. LibreOffice render → PDF page count — accurate for documents that carry
#      no app.xml page count (e.g. files generated by python-docx). Cached
#      upstream, so this one-time cost is acceptable; skipped if soffice is
#      unavailable.
#   3. ``docProps/app.xml`` ``<Pages>`` — a stored Office statistic, but it can
#      be missing or stale in generated / Word-for-web documents.
#   4. A layout-aware estimate from page geometry, wrapping, paragraph spacing,
#      images and tables.
#   5. A conservative word-density estimate — last resort only.

def _pages_via_word_com(raw: Optional[bytes]) -> Optional[int]:
    """Use Microsoft Word's COM interface to get the exact page count.

    Only works on Windows hosts with Word installed. Returns ``None`` elsewhere.
    """
    if not raw:
        return None
    import sys

    if sys.platform != "win32":
        return None

    tmp_path = None
    word = None
    doc = None
    try:
        import tempfile
        import win32com.client as win32  # type: ignore[import-untyped]

        fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        with open(tmp_path, "wb") as fh:
            fh.write(raw)

        word = win32.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(tmp_path, ReadOnly=True, AddToRecentFiles=False)
        doc.Repaginate()

        sel = word.Selection
        sel.EndKey(6)  # wdStory
        count = sel.Information(3)  # wdActiveEndPageNumber
        return int(count) if count and count > 0 else None
    except Exception:  # noqa: BLE001
        return None
    finally:
        try:
            if doc:
                doc.Close(False)
        except Exception:  # noqa: BLE001
            pass
        try:
            if word:
                word.Quit()
        except Exception:  # noqa: BLE001
            pass
        try:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)
        except Exception:  # noqa: BLE001
            pass

def _pages_from_app_xml(raw: Optional[bytes]) -> Optional[int]:
    """Return the rendered page count Word stamps into ``docProps/app.xml``.

    The count is trusted only when ``<Words>`` is non-zero, i.e. an Office app
    actually computed the document statistics. Tools like python-docx ship a
    template whose app.xml carries placeholder ``<Pages>1</Pages>`` /
    ``<Words>0</Words>`` regardless of the real content — trusting that would
    report multi-page generated documents as a single page."""
    if not raw:
        return None
    import zipfile

    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            if "docProps/app.xml" not in z.namelist():
                return None
            xml = z.read("docProps/app.xml").decode("utf-8", "ignore")
    except Exception:
        return None
    m = re.search(r"<Pages>\s*(\d+)\s*</Pages>", xml)
    if not m:
        return None
    pages = int(m.group(1))
    if pages <= 0:
        return None
    wm = re.search(r"<Words>\s*(\d+)\s*</Words>", xml)
    words = int(wm.group(1)) if wm else 0
    # Trust the count only if statistics were really computed (Words > 0) or it
    # already reports more than one page (no tool defaults to a >1 placeholder).
    if words > 0 or pages > 1:
        return pages
    return None


def _pages_from_rendered_breaks(raw: Optional[bytes]) -> Optional[int]:
    """Count Word's stored rendered-page markers when present."""
    if not raw:
        return None
    import zipfile

    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            if "word/document.xml" not in z.namelist():
                return None
            xml = z.read("word/document.xml").decode("utf-8", "ignore")
    except Exception:
        return None

    count = xml.count("lastRenderedPageBreak")
    return count + 1 if count > 0 else None


def _pages_via_libreoffice(raw: Optional[bytes]) -> Optional[int]:
    """Render to PDF with LibreOffice and count the pages (best-effort)."""
    if not raw:
        return None
    try:
        from app.services.office.office_pipeline import available, convert

        if not available():
            return None
        pdf = convert(raw, "pdf")
        import fitz  # PyMuPDF

        doc = fitz.open(stream=pdf, filetype="pdf")
        n = doc.page_count
        doc.close()
        return n if n and n > 0 else None
    except Exception:
        return None


def _length_to_pt(val) -> Optional[float]:
    """Convert a python-docx Length/raw EMU value to points."""
    if val is None or isinstance(val, bool):
        return None
    pt_attr = getattr(val, "pt", None)
    if pt_attr is not None:
        try:
            return float(pt_attr)
        except (TypeError, ValueError):
            return None
    if isinstance(val, (int, float)):
        # python-docx stores most lengths as EMU integers.
        return float(val) / 12700.0
    return None


def _section_content_box_pt(doc: Document) -> tuple[float, float]:
    section = doc.sections[0] if doc.sections else None
    if not section:
        return 432.0, 720.0
    page_w = _length_to_pt(section.page_width) or 612.0
    page_h = _length_to_pt(section.page_height) or 792.0
    left = _length_to_pt(section.left_margin) or 72.0
    right = _length_to_pt(section.right_margin) or 72.0
    top = _length_to_pt(section.top_margin) or 72.0
    bottom = _length_to_pt(section.bottom_margin) or 72.0
    return max(144.0, page_w - left - right), max(144.0, page_h - top - bottom)


def _effective_font_size_pt(p: Paragraph) -> float:
    sizes: list[float] = []
    for run in p.runs:
        size = _pt_value(run.font.size)
        if size:
            sizes.append(size)
    try:
        size = _pt_value(p.style.font.size) if p.style and p.style.font.size else None
        if size:
            sizes.append(size)
    except Exception:
        pass
    try:
        if p.style and "heading" in (p.style.name or "").lower():
            sizes.append(16.0)
    except Exception:
        pass
    return max(sizes) if sizes else 12.0


def _line_spacing_multiplier(p: Paragraph, font_size: float) -> float:
    for fmt in (p.paragraph_format, getattr(p.style, "paragraph_format", None)):
        if not fmt:
            continue
        spacing = fmt.line_spacing
        if isinstance(spacing, (int, float)) and not isinstance(spacing, bool):
            return max(float(spacing), 1.0)
        spacing_pt = _pt_value(spacing)
        if spacing_pt:
            return max(spacing_pt / max(font_size, 1.0), 1.0)
    return 1.15


def _paragraph_spacing_pt(p: Paragraph) -> tuple[float, float]:
    before = _pt_value(p.paragraph_format.space_before)
    after = _pt_value(p.paragraph_format.space_after)
    try:
        style_fmt = p.style.paragraph_format if p.style else None
    except Exception:
        style_fmt = None
    if style_fmt:
        before = before if before is not None else _pt_value(style_fmt.space_before)
        after = after if after is not None else _pt_value(style_fmt.space_after)
    return before or 0.0, after or 0.0


def _paragraph_image_height_pt(p: Paragraph) -> float:
    total = 0.0
    try:
        drawings = p._p.xpath(".//w:drawing")
    except Exception:
        return 0.0
    for drawing in drawings:
        try:
            extents = drawing.xpath(".//wp:extent")
        except Exception:
            continue
        for extent in extents:
            cy = extent.get("cy")
            if cy:
                try:
                    total += int(cy) / 12700.0
                except ValueError:
                    pass
    return total


def _paragraph_page_breaks(p: Paragraph) -> int:
    try:
        return len(p._p.xpath('.//w:br[@w:type="page"]'))
    except Exception:
        return 0


def _paragraph_break_before(p: Paragraph) -> bool:
    try:
        if p.paragraph_format.page_break_before:
            return True
        return bool(p.style and p.style.paragraph_format.page_break_before)
    except Exception:
        return False


def _estimate_paragraph_height_pt(p: Paragraph, usable_width_pt: float) -> float:
    font_size = _effective_font_size_pt(p)
    line_height = font_size * _line_spacing_multiplier(p, font_size)
    before, after = _paragraph_spacing_pt(p)
    image_height = _paragraph_image_height_pt(p)
    text = " ".join((p.text or "").split())

    if image_height:
        return before + after + image_height + (line_height * 0.5)
    if not text:
        return before + after + line_height

    # 0.52em average glyph width is intentionally conservative for reports:
    # it avoids the old dense-word fallback undercounting double-spaced docs.
    chars_per_line = max(8.0, usable_width_pt / max(font_size * 0.52, 1.0))
    text_lines = max(1, math.ceil(len(text) / chars_per_line))
    return before + after + (text_lines * line_height)


def _iter_table_paragraphs(table: DocxTable):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _estimate_pages_from_docx_layout(raw: Optional[bytes]) -> Optional[int]:
    """Estimate pages from layout when no renderer is available.

    This is still an estimate, but it is much better than word density for
    report-like documents with blank lines, double spacing, large headings,
    images, or tables.
    """
    if not raw:
        return None
    try:
        doc = Document(io.BytesIO(raw))
    except Exception:
        return None

    usable_width, usable_height = _section_content_box_pt(doc)
    used = 0.0
    pages = 1
    saw_content = False

    paragraphs: list[Paragraph] = list(doc.paragraphs)
    for table in doc.tables:
        paragraphs.extend(_iter_table_paragraphs(table))

    for paragraph in paragraphs:
        height = _estimate_paragraph_height_pt(paragraph, usable_width)
        if height <= 0:
            continue
        saw_content = True

        if _paragraph_break_before(paragraph) and used > 0:
            pages += 1
            used = 0.0

        while used > 0 and used + height > usable_height:
            pages += 1
            used = 0.0
        while height > usable_height:
            pages += 1
            height -= usable_height
        used += height

        for _ in range(_paragraph_page_breaks(paragraph)):
            pages += 1
            used = 0.0

    if not saw_content:
        return None
    return max(1, pages)


def _docx_page_count(raw: Optional[bytes], word_count: int) -> Optional[int]:
    """Best-available DOCX page count."""
    for resolver in (
        _pages_via_word_com,
        _pages_via_libreoffice,
    ):
        n = resolver(raw)
        if n:
            return n

    app_pages = _pages_from_app_xml(raw)
    layout_pages = _estimate_pages_from_docx_layout(raw)
    if app_pages and layout_pages:
        if layout_pages >= app_pages + 3 and layout_pages >= app_pages * 1.75:
            return layout_pages
        return app_pages
    if app_pages:
        return app_pages
    if layout_pages:
        return layout_pages

    n = _pages_from_rendered_breaks(raw)
    if n:
        return n

    return max(1, round(word_count / 380)) if word_count else None


# ---------------------------------------------------------------------------
# Extractor class
# ---------------------------------------------------------------------------

class WordExtractor:
    """Extract content and styling from a Word (.docx) document."""

    def __init__(self):
        self._run_style_cache: dict[str, str] = {}  # fingerprint -> style_ref
        self._para_style_cache: dict[str, str] = {}
        self._run_style_counter = 0
        self._para_style_counter = 0
        self._run_styles: dict[str, RunStyle] = {}
        self._para_styles: dict[str, ParagraphStyle] = {}
        self._table_styles: dict[str, TableStyle] = {}
        self._cell_styles: dict[str, CellStyle] = {}
        self._num_to_abstract: dict[str, str] = {}
        self._numbering_levels: dict[tuple[str, int], dict[str, str]] = {}


    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def extract(
        self,
        file_path: Optional[str] = None,
        file_stream: Optional[io.BytesIO] = None,
        filename: Optional[str] = None,
    ) -> tuple[DocumentContent, DocumentStyling]:
        """
        Extract content and styling from a .docx.

        Provide either ``file_path`` (path on disk) or ``file_stream`` (in-memory bytes).
        """
        raw_bytes: Optional[bytes] = None
        if file_path:
            with open(file_path, "rb") as _fh:
                raw_bytes = _fh.read()
            doc = Document(io.BytesIO(raw_bytes))
            source_name = os.path.basename(file_path)
        elif file_stream:
            file_stream.seek(0)
            raw_bytes = file_stream.read()
            doc = Document(io.BytesIO(raw_bytes))
            source_name = filename or "uploaded.docx"
        else:
            raise ValueError("Provide either file_path or file_stream")

        self._load_numbering_lookup(doc)

        # ---- Page style from first section ----
        page_style = self._extract_page_style(doc)

        # ---- Body elements ----
        elements = self._extract_body(doc)

        # ---- Headers / footers (from first section) ----
        header_footer_elements = self._extract_headers_footers(doc)
        elements = header_footer_elements + elements

        # ---- Metadata ----
        core = doc.core_properties
        word_count = _count_words(elements)
        metadata = DocumentMetadata(
            source_file=source_name,
            source_type="docx",
            extracted_at=datetime.now(timezone.utc).isoformat(),
            page_count=_docx_page_count(raw_bytes, word_count),
            author=core.author if core.author else None,
            title=core.title if core.title else None,
        )

        content = DocumentContent(metadata=metadata, elements=elements)

        styling = DocumentStyling(
            metadata=StyleMetadata(
                source_file=source_name,
                created_at=datetime.now(timezone.utc).isoformat(),
            ),
            page_style=page_style,
            paragraph_styles=self._para_styles,
            run_styles=self._run_styles,
            table_styles=self._table_styles,
            cell_styles=self._cell_styles,
        )

        return content, styling

    def _load_numbering_lookup(self, doc: Document) -> None:
        """Cache numbering.xml so list extraction can distinguish bullets from
        decimal lists even when every paragraph uses the generic List Paragraph
        style."""
        self._num_to_abstract = {}
        self._numbering_levels = {}
        try:
            numbering = doc.part.numbering_part.element
        except Exception:
            return
        if numbering is None:
            return

        for num in numbering.findall(qn("w:num")):
            num_id = num.get(qn("w:numId"))
            abs_ref = num.find(qn("w:abstractNumId"))
            abs_id = abs_ref.get(qn("w:val")) if abs_ref is not None else None
            if num_id and abs_id:
                self._num_to_abstract[num_id] = abs_id

        for abstract in numbering.findall(qn("w:abstractNum")):
            abs_id = abstract.get(qn("w:abstractNumId"))
            if not abs_id:
                continue
            for lvl in abstract.findall(qn("w:lvl")):
                try:
                    ilvl = int(lvl.get(qn("w:ilvl"), "0"))
                except ValueError:
                    ilvl = 0
                fmt_el = lvl.find(qn("w:numFmt"))
                txt_el = lvl.find(qn("w:lvlText"))
                self._numbering_levels[(abs_id, ilvl)] = {
                    "fmt": fmt_el.get(qn("w:val")) if fmt_el is not None else "",
                    "text": txt_el.get(qn("w:val")) if txt_el is not None else "",
                }

    # ------------------------------------------------------------------
    # Page style
    # ------------------------------------------------------------------

    def _extract_page_style(self, doc: Document) -> PageStyle:
        section = doc.sections[0] if doc.sections else None
        if section is None:
            return PageStyle()

        orientation = Orientation.PORTRAIT
        if section.orientation == WD_ORIENT.LANDSCAPE:
            orientation = Orientation.LANDSCAPE

        return PageStyle(
            width_inches=round(_emu_to_inches(section.page_width) or 8.5, 4),
            height_inches=round(_emu_to_inches(section.page_height) or 11.0, 4),
            orientation=orientation,
            margins=PageMargins(
                top_inches=round(_emu_to_inches(section.top_margin) or 1.0, 4),
                bottom_inches=round(_emu_to_inches(section.bottom_margin) or 1.0, 4),
                left_inches=round(_emu_to_inches(section.left_margin) or 1.25, 4),
                right_inches=round(_emu_to_inches(section.right_margin) or 1.25, 4),
            ),
        )

    # ------------------------------------------------------------------
    # Body iteration
    # ------------------------------------------------------------------

    def _extract_body(self, doc: Document) -> list[ContentElement]:
        """Walk the document body XML to extract paragraphs, tables, and images in order."""
        elements: list[ContentElement] = []
        body = doc.element.body

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                para = Paragraph(child, doc)
                elem = self._process_paragraph(para, doc)
                if elem is not None:
                    elements.append(elem)
            elif tag == "tbl":
                table = DocxTable(child, doc)
                elem = self._process_table(table)
                elements.append(elem)
            elif tag == "sectPr":
                pass  # section properties — already handled via doc.sections

        # Promote numbered pseudo-headings in FLAT documents (plain, non-bold
        # "1. Purpose" / "3.1 Scope" lines). Needs document-level adjacency
        # context, so it runs as a post-pass over the element list.
        self._promote_flat_headings(elements)

        # Normalize heading hierarchy so levels don't jump unexpectedly
        # (e.g. H1 → H4 → H2 — the H4 should become H2).
        self._normalize_heading_levels(elements)

        return elements

    @classmethod
    def _promote_flat_headings(cls, elements: list[ContentElement]) -> None:
        """Promote numbered section headings authored as plain paragraphs.

        The per-paragraph heuristic (:meth:`_heuristic_heading_level`) only
        promotes BOLD numbered lines — safe without context. Flat documents
        (everything in ``Normal``, no bold) would yield a single section, which
        starves the mapping agent. This pass adds the missing context:

          * multi-level numbers ("3.1 Title") are headings at their depth;
          * single-level numbers ("1. Title") are headings only when they are
            NOT adjacent to another single-level-numbered paragraph (such a
            run is a numbered LIST) and don't read like prose.

        Only applies to documents with fewer than 2 real headings — in a
        structured document a plain numbered line is a list item or clause.
        """
        existing = sum(1 for el in elements if el.type == ElementType.HEADING)
        if existing >= 2:
            return

        def para_text(el: ContentElement) -> str:
            return "".join((r.text or "") for r in (el.content or [])).strip()

        # idx -> section-number string for every numbered candidate. Besides
        # plain paragraphs this includes NUMBERED list items whose number is
        # LITERAL TEXT (the heuristic in _get_list_info tags flat "1. Purpose"
        # lines as list items — a real numPr list carries no literal number).
        nums: dict[int, str] = {}
        for i, el in enumerate(elements):
            if el.type == ElementType.PARAGRAPH:
                pass
            elif (
                el.type == ElementType.LIST_ITEM
                and el.list_type == ListType.NUMBERED
            ):
                pass
            else:
                continue
            text = para_text(el)
            if not text or len(text) > 120:
                continue
            if "|" in text or "\t" in text:
                continue  # a delimited table row, not a heading
            m = cls._RE_NUMBERED_HEADING.match(text)
            if not m:
                continue
            title = m.group(2).strip()
            if not title or len(title.split()) > 14:
                continue
            nums[i] = m.group(1)

        def promote(el: ContentElement, level: int) -> None:
            el.type = ElementType.HEADING
            el.level = level
            el.list_type = None
            el.list_level = None
            el.bullet_char = None
            el.number_format = None

        for i, num in nums.items():
            el = elements[i]
            text = para_text(el)
            if "." in num:  # multi-level ("3.1") — always a heading
                promote(el, min(num.count(".") + 1, 4))
                continue
            # Single-level: require an explicit "1." / "1)" separator so a
            # sentence starting with a bare number ("2021 was…") is left alone.
            if len(text) <= len(num) or text[len(num)] not in ".)":
                continue
            # An adjacent run of "1."/"2."/"3." lines is a numbered list,
            # not a stack of headings.
            if (i - 1 in nums and "." not in nums[i - 1]) or (
                i + 1 in nums and "." not in nums[i + 1]
            ):
                continue
            if text.rstrip().endswith((".", "!", "?", ";", ",")):
                continue  # reads like prose
            promote(el, 1)

    @staticmethod
    def _normalize_heading_levels(elements: list[ContentElement]) -> None:
        """Ensure heading levels form a valid hierarchy without skips.

        Walks the heading elements in order and re-maps levels so that no
        heading jumps more than one level deeper than the previous heading.
        """
        level_map: dict[int, int] = {}  # original_level -> normalized_level
        max_seen = 0
        for el in elements:
            if el.type != ElementType.HEADING or el.level is None:
                continue
            orig = el.level
            if orig not in level_map:
                # First time seeing this level — assign the next available slot
                # but don't exceed one more than the current maximum.
                level_map[orig] = min(orig, max_seen + 1)
            el.level = level_map[orig]
            max_seen = max(max_seen, el.level)

    # ------------------------------------------------------------------
    # Paragraph processing
    # ------------------------------------------------------------------

    def _process_paragraph(self, para: Paragraph, doc: Document) -> Optional[ContentElement]:
        """Process a single paragraph into a ContentElement."""

        # Check if it's a page break
        if self._is_page_break(para):
            return ContentElement(type=ElementType.PAGE_BREAK)

        # Determine element type
        heading_level = self._get_heading_level(para)
        list_info = self._get_list_info(para)
        is_heading = heading_level is not None

        # Extract inline content (runs + hyperlinks)
        runs = self._extract_runs(para, doc)

        # Skip empty paragraphs with no text
        if not runs and not is_heading:
            return None

        # Extract paragraph style
        para_style = self._extract_paragraph_style(para)
        para_style_ref = self._register_para_style(para, para_style)

        if is_heading:
            return ContentElement(
                type=ElementType.HEADING,
                level=heading_level,
                content=runs,
                style_ref=para_style_ref,
            )
        elif list_info:
            return ContentElement(
                type=ElementType.LIST_ITEM,
                content=runs,
                style_ref=para_style_ref,
                list_type=list_info["list_type"],
                list_level=list_info["level"],
                bullet_char=list_info.get("bullet_char"),
                number_format=list_info.get("number_format"),
            )
        else:
            return ContentElement(
                type=ElementType.PARAGRAPH,
                content=runs,
                style_ref=para_style_ref,
            )

    def _is_page_break(self, para: Paragraph) -> bool:
        """Check if the paragraph consists solely of a page break."""
        xml = para._element.xml
        if "w:br" in xml and 'w:type="page"' in xml:
            # Only count as page break if there's no real text
            if not para.text.strip():
                return True
        return False

    def _get_heading_level(self, para: Paragraph) -> Optional[int]:
        """Return heading level (1-9) or None.

        Checks in order:
          1. Built-in Word Heading styles (``Heading 1`` … ``Heading 9``).
          2. The ``Title`` style.
          3. XML ``outlineLvl`` property.
          4. **Heuristic detection** for pseudo-headings authored as ``Normal``
             with bold formatting + numbering patterns, ALL CAPS, or TitleCase.
        """
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip())
                return level
            except ValueError:
                pass

        # The "Title" style (e.g. add_heading(level=0)) is the document's top
        # heading — treat it as level 1 so it lands in the outline / slots.
        if style_name.strip().lower() == "title":
            return 1

        # Check outline level from XML directly (paragraph_format doesn't expose it)
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            outline_lvl = pPr.find(qn("w:outlineLvl"))
            if outline_lvl is not None:
                try:
                    level = int(outline_lvl.get(qn("w:val"), "9"))
                    if 0 <= level <= 8 and not self._outline_level_is_spurious(para):
                        return level + 1  # outline_level is 0-indexed
                except (ValueError, TypeError):
                    pass

        # ----- Heuristic detection for pseudo-headings -----
        return self._heuristic_heading_level(para)

    # ------------------------------------------------------------------
    # Heuristic heading helpers
    # ------------------------------------------------------------------

    # Matches multi-level numbering: "1.", "2.1.", "10.3.2" etc.
    _RE_NUMBERED_HEADING = re.compile(
        r"^(\d+(?:\.\d+)*)[\.\)]?\s+(.+)", re.DOTALL
    )

    @staticmethod
    def _outline_level_is_spurious(para: Paragraph) -> bool:
        """True when a direct ``outlineLvl`` contradicts what the paragraph IS.

        Authors (and some export tools) leave outline levels on paragraphs
        that are plainly not headings — e.g. bold ``Label: value`` field lines
        inside a List Paragraph. Trusting them turns metadata into template
        slots and fragments the section. Only paragraphs NOT in a real
        Heading/Title style are second-guessed."""
        style_name = (para.style.name if para.style else "").strip().lower()
        if style_name.startswith("heading") or style_name == "title":
            return False  # a real heading style — trust the outline level
        text = (para.text or "").strip()
        if not text:
            return True  # an empty "heading" is never real
        # "Label: value" field line — colon with substantive content after it.
        _head, sep, tail = text.partition(":")
        if sep and tail.strip():
            return True
        return False

    @staticmethod
    def _is_all_bold(para: Paragraph) -> bool:
        """True if every text-bearing run in the paragraph is bold."""
        runs = [r for r in para.runs if (r.text or "").strip()]
        if not runs:
            return False
        bold_chars = sum(len(r.text) for r in runs if r.bold)
        total = sum(len(r.text) for r in runs)
        return total > 0 and bold_chars / total >= 0.7

    def _heuristic_heading_level(self, para: Paragraph) -> Optional[int]:
        """Detect headings that aren't tagged with a Heading style.

        Returns a heading level (1-6) or ``None``.

        Covers:
        * **Numbered section headings** — "1. BACKGROUND", "2.1. Services"
        * **ALL-CAPS bold lines** — "STATEMENT OF WORK"
        * **TitleCase bold lines** — "Brief Description of Project"
        """
        text = (para.text or "").strip()
        if not text or len(text) > 120:
            return None

        # Quick reject: too many words → almost certainly body text
        words = text.split()
        if len(words) > 14:
            return None

        # Reject lines ending with sentence-terminal punctuation
        # (allow trailing colon for label-style headings like "Services:")
        if text.rstrip().endswith((".", ";", ",")) and not self._RE_NUMBERED_HEADING.match(text):
            return None

        is_bold = self._is_all_bold(para)

        # 1) Numbered section heading: "1.", "2.1.", "10.3.2 TITLE"
        m = self._RE_NUMBERED_HEADING.match(text)
        if m and is_bold:
            numbering = m.group(1)  # e.g. "2.1"
            depth = numbering.count(".") + 1  # "1" → 1, "2.1" → 2, "2.1.3" → 3
            # Cap at level 4 so very deep numbering doesn't produce H7, H8…
            return min(depth, 4)

        # The remaining heuristics require bold
        if not is_bold:
            return None

        # A "Label: value" line ("Table: stg_mass_email", "Schema : cdm") is a
        # field, not a heading — promoting it fragments the section and turns
        # metadata into template slots. A *trailing* colon ("Results:") is
        # still a label-style heading and stays eligible.
        head, sep, tail = text.partition(":")
        if sep and tail.strip():
            return None

        has_alpha = sum(1 for c in text if c.isalpha()) >= 2

        # 2) ALL CAPS bold heading: "STATEMENT OF WORK", "SERVICES AND DELIVERABLES"
        if text.upper() == text and has_alpha and len(text) <= 80:
            return 1

        # 3) TitleCase bold heading: "Brief Description of Project"
        #    Accept if most words start with uppercase (allowing small
        #    prepositions / articles to be lowercase).
        if len(words) <= 10 and has_alpha and len(text) <= 80:
            upper_words = sum(1 for w in words if w[0:1].isupper())
            if upper_words >= max(1, len(words) - 2):  # tolerate 1-2 lowercase words
                return 2

        return None

    def _get_list_info(self, para: Paragraph) -> Optional[dict]:
        """Detect if paragraph is a list item and extract list metadata.

        NOTE: This is called **after** ``_get_heading_level``. If the
        paragraph was already detected as a heading the caller never
        reaches here, so we don't need to re-check heading status.
        However, the *heuristic* branch below (no ``numPr``) must avoid
        classifying numbered section headings ("1. TITLE") as list items.
        """
        # Check XML for numbering properties, first on the paragraph and then
        # on its paragraph style. Many Word files use the generic "List
        # Paragraph" style for both bullets and numbered steps, so style-name
        # heuristics alone are not reliable.
        numPr = self._paragraph_num_pr(para)
        if numPr is None:
            # Heuristic: check if text starts with bullet characters
            text = para.text.strip()
            bullet_chars = ["•", "◦", "▪", "▸", "‣", "-", "–", "—", "●", "○"]
            for bc in bullet_chars:
                if text.startswith(bc):
                    return {
                        "list_type": ListType.BULLET,
                        "level": 0,
                        "bullet_char": bc,
                    }
            # Check for numbered list pattern — but skip if it looks like
            # a section heading (bold + short + uppercase/titlecase).
            num_match = re.match(r"^(\d+)[.)]\s", text)
            if num_match:
                # Guard: don't classify short bold lines as list items when
                # they look like numbered section headings.
                words = text.split()
                is_bold = self._is_all_bold(para)
                looks_like_heading = (
                    is_bold
                    and len(words) <= 14
                    and len(text) <= 120
                    and (text.isupper() or text.istitle()
                         or sum(1 for w in words if w[0:1].isupper()) >= max(1, len(words) - 2))
                )
                if not looks_like_heading:
                    return {
                        "list_type": ListType.NUMBERED,
                        "level": 0,
                        "number_format": "decimal",
                    }
            return None

        level = self._num_pr_level(numPr)
        num_id = self._num_pr_num_id(numPr)
        resolved = self._resolve_numbering(num_id, level)
        if resolved:
            return resolved

        # Determine bullet vs numbered from style name heuristic
        style_name = (para.style.name if para.style else "").lower()
        if "bullet" in style_name or "list bullet" in style_name:
            return {
                "list_type": ListType.BULLET,
                "level": level,
                "bullet_char": "•",
            }
        else:
            return {
                "list_type": ListType.NUMBERED,
                "level": level,
                "number_format": "decimal",
            }

    def _paragraph_num_pr(self, para: Paragraph):
        pPr = para._element.find(qn("w:pPr"))
        numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
        if numPr is not None:
            return numPr
        try:
            style_el = para.style.element if para.style is not None else None
        except Exception:
            style_el = None
        if style_el is None:
            return None
        style_pPr = style_el.find(qn("w:pPr"))
        return style_pPr.find(qn("w:numPr")) if style_pPr is not None else None

    @staticmethod
    def _num_pr_level(numPr) -> int:
        ilvl_elem = numPr.find(qn("w:ilvl"))
        if ilvl_elem is None:
            return 0
        try:
            return int(ilvl_elem.get(qn("w:val"), "0"))
        except ValueError:
            return 0

    @staticmethod
    def _num_pr_num_id(numPr) -> Optional[str]:
        num_id_el = numPr.find(qn("w:numId"))
        return num_id_el.get(qn("w:val")) if num_id_el is not None else None

    def _resolve_numbering(self, num_id: Optional[str], level: int) -> Optional[dict]:
        if not num_id:
            return None
        abs_id = self._num_to_abstract.get(str(num_id))
        if not abs_id:
            return None
        lvl = self._numbering_levels.get((abs_id, level)) or self._numbering_levels.get((abs_id, 0))
        if not lvl:
            return None
        fmt = (lvl.get("fmt") or "").lower()
        if fmt == "bullet":
            glyph = (lvl.get("text") or "•").strip() or "•"
            return {
                "list_type": ListType.BULLET,
                "level": level,
                "bullet_char": glyph[:2],
            }
        return {
            "list_type": ListType.NUMBERED,
            "level": level,
            "number_format": fmt or "decimal",
        }

    # ------------------------------------------------------------------
    # Run extraction
    # ------------------------------------------------------------------

    def _extract_runs(self, para: Paragraph, doc: Document) -> list[TextRun]:
        """Extract text runs and inline images from a paragraph."""
        runs: list[TextRun] = []

        for child in para._element:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "r":
                # Regular run
                run_obj = None
                for r in para.runs:
                    if r._element is child:
                        run_obj = r
                        break

                if run_obj is None:
                    # Fallback: extract text directly from XML
                    text_elems = child.findall(qn("w:t"))
                    text = "".join(t.text or "" for t in text_elems)
                    if text:
                        runs.append(TextRun(text=text))
                    continue

                # Check for inline image
                drawing = child.find(qn("w:drawing"))
                if drawing is not None:
                    img_elem = self._extract_inline_image(drawing, doc)
                    if img_elem is not None:
                        runs.append(img_elem)
                    continue

                text = run_obj.text
                if not text:
                    continue

                run_style = self._extract_run_style(run_obj)
                style_ref = self._register_run_style(run_style)
                runs.append(TextRun(text=text, style_ref=style_ref))

            elif tag == "hyperlink":
                # Hyperlink
                href = self._get_hyperlink_url(child, doc)
                for r_elem in child.findall(qn("w:r")):
                    text_elems = r_elem.findall(qn("w:t"))
                    text = "".join(t.text or "" for t in text_elems)
                    if text:
                        runs.append(TextRun(text=text, hyperlink_url=href))

        return runs

    def _get_hyperlink_url(self, hyperlink_elem, doc: Document) -> Optional[str]:
        """Extract the URL from a hyperlink element."""
        r_id = hyperlink_elem.get(qn("r:id"))
        if r_id:
            try:
                rel = doc.part.rels.get(r_id)
                if rel and rel.target_ref:
                    return str(rel.target_ref)
            except Exception:
                pass
        return None

    def _extract_inline_image(self, drawing, doc: Document) -> Optional[TextRun]:
        """Extract an image from a drawing element, return as a special marker run."""
        # This is a content-level image — we'll create a placeholder TextRun
        # and also a separate ContentElement for it (handled at paragraph level)
        # For now, we extract the image data.
        blip = drawing.find(".//" + qn("a:blip"))
        if blip is None:
            return None

        r_embed = blip.get(qn("r:embed"))
        if r_embed is None:
            return None

        try:
            image_part = doc.part.rels[r_embed].target_part
            image_bytes = image_part.blob
            img_b64 = base64.b64encode(image_bytes).decode("utf-8")

            # Get dimensions
            extent = drawing.find(".//" + qn("wp:extent"))
            width_inches = None
            height_inches = None
            if extent is not None:
                cx = extent.get("cx")
                cy = extent.get("cy")
                if cx:
                    width_inches = round(int(cx) / 914400, 4)
                if cy:
                    height_inches = round(int(cy) / 914400, 4)

            # Determine format
            content_type = image_part.content_type
            fmt = "png"
            if "jpeg" in content_type or "jpg" in content_type:
                fmt = "jpeg"
            elif "gif" in content_type:
                fmt = "gif"
            elif "bmp" in content_type:
                fmt = "bmp"
            elif "tiff" in content_type:
                fmt = "tiff"

            # We return this as a special TextRun with image marker
            return TextRun(
                text=f"[IMAGE:{fmt}:{width_inches}x{height_inches}]",
                style_ref="__image__",
                image_data_b64=img_b64,
            )
        except Exception:
            return None

    # ------------------------------------------------------------------
    # Style extraction
    # ------------------------------------------------------------------

    def _extract_run_style(self, run) -> RunStyle:
        """Extract all formatting from a python-docx Run object."""
        font = run.font

        # Underline
        underline = None
        if font.underline is not None:
            underline = _UNDERLINE_MAP.get(font.underline, UnderlineType.SINGLE)

        # Color
        color_hex = None
        if font.color and font.color.rgb:
            color_hex = _rgb_to_hex(font.color.rgb)

        return RunStyle(
            font_name=font.name,
            font_size_pt=_pt_value(font.size) if font.size else None,
            bold=font.bold,
            italic=font.italic,
            underline=underline,
            strikethrough=font.strike,
            double_strikethrough=font.double_strike,
            superscript=font.superscript,
            subscript=font.subscript,
            color_hex=color_hex,
            highlight_color=str(font.highlight_color) if font.highlight_color else None,
            all_caps=font.all_caps,
            small_caps=font.small_caps,
            character_spacing_pt=None,  # character spacing requires XML-level extraction
        )

    def _extract_paragraph_style(self, para: Paragraph) -> ParagraphStyle:
        """Extract all formatting from a python-docx Paragraph."""
        pf = para.paragraph_format

        # Alignment
        alignment = None
        try:
            if pf.alignment is not None:
                alignment = _ALIGNMENT_MAP.get(pf.alignment)
        except (ValueError, KeyError):
            pass  # unmapped alignment values (e.g. 'start') — leave as None

        # Indentation
        indent = IndentStyle(
            left_inches=_emu_to_inches(pf.left_indent) if pf.left_indent else None,
            right_inches=_emu_to_inches(pf.right_indent) if pf.right_indent else None,
            first_line_inches=_emu_to_inches(pf.first_line_indent) if pf.first_line_indent else None,
        )
        # Hanging indent (negative first_line_indent)
        if pf.first_line_indent and pf.first_line_indent < 0:
            indent.hanging_inches = abs(_emu_to_inches(pf.first_line_indent))
            indent.first_line_inches = None

        # Line spacing
        line_spacing = None
        line_spacing_rule = None
        if pf.line_spacing is not None:
            line_spacing = float(pf.line_spacing)
        if pf.line_spacing_rule is not None:
            rule_map = {
                WD_LINE_SPACING.SINGLE: "single",
                WD_LINE_SPACING.ONE_POINT_FIVE: "one_point_five",
                WD_LINE_SPACING.DOUBLE: "double",
                WD_LINE_SPACING.AT_LEAST: "at_least",
                WD_LINE_SPACING.EXACTLY: "exactly",
                WD_LINE_SPACING.MULTIPLE: "multiple",
            }
            line_spacing_rule = rule_map.get(pf.line_spacing_rule, "single")

        # Tab stops
        tab_stops = None
        if pf.tab_stops:
            tab_stops = [
                round(_emu_to_inches(ts.position) or 0, 4)
                for ts in pf.tab_stops
            ]

        return ParagraphStyle(
            alignment=alignment,
            space_before_pt=_pt_value(pf.space_before) if pf.space_before else None,
            space_after_pt=_pt_value(pf.space_after) if pf.space_after else None,
            line_spacing=line_spacing,
            line_spacing_rule=line_spacing_rule,
            indent=indent if any(
                v is not None
                for v in [indent.left_inches, indent.right_inches, indent.first_line_inches, indent.hanging_inches]
            ) else None,
            keep_with_next=pf.keep_with_next,
            keep_together=pf.keep_together,
            widow_control=pf.widow_control,
            outline_level=None,  # handled separately in heading detection
            tab_stops=tab_stops,
        )

    # ------------------------------------------------------------------
    # Table processing
    # ------------------------------------------------------------------

    def _process_table(self, table: DocxTable) -> ContentElement:
        """Process a Word table into a ContentElement."""
        rows: list[TableRow] = []

        # Extract table-level style
        table_style = self._extract_table_style(table)
        table_style_key = _make_style_key("table", str(len(self._table_styles)))
        self._table_styles[table_style_key] = table_style

        for row_idx, row in enumerate(table.rows):
            cells: list[TableCell] = []
            for cell in row.cells:
                cell_runs: list[TextRun] = []
                for para in cell.paragraphs:
                    para_runs = self._extract_runs(para, table._parent)
                    cell_runs.extend(para_runs)

                # Extract cell style
                cell_style = self._extract_cell_style(cell)
                cell_style_key = _make_style_key(
                    "cell", str(len(self._cell_styles))
                )
                self._cell_styles[cell_style_key] = cell_style

                cells.append(
                    TableCell(
                        content=cell_runs,
                        style_ref=cell_style_key,
                        inline_style=cell_style,
                    )
                )
            rows.append(TableRow(cells=cells, is_header=(row_idx == 0)))

        return ContentElement(
            type=ElementType.TABLE,
            rows=rows,
            table_style_ref=table_style_key,
            inline_table_style=table_style,
        )

    def _extract_table_style(self, table: DocxTable) -> TableStyle:
        """Extract table-level formatting."""
        alignment = None
        try:
            tbl_pr = table._tbl.find(qn("w:tblPr"))
            if tbl_pr is not None:
                jc = tbl_pr.find(qn("w:jc"))
                if jc is not None:
                    val = jc.get(qn("w:val"), "left")
                    align_map = {"left": Alignment.LEFT, "center": Alignment.CENTER, "right": Alignment.RIGHT}
                    alignment = align_map.get(val)
        except Exception:
            pass

        return TableStyle(alignment=alignment)

    def _extract_cell_style(self, cell) -> CellStyle:
        """Extract cell-level formatting."""
        width_inches = None
        shading_hex = None
        vert_align = None

        try:
            tc = cell._tc
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is not None:
                # Width
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is not None:
                    w_val = tc_w.get(qn("w:w"))
                    w_type = tc_w.get(qn("w:type"), "dxa")
                    if w_val and w_type == "dxa":
                        width_inches = round(int(w_val) / 1440, 4)

                # Shading
                shd = tc_pr.find(qn("w:shd"))
                if shd is not None:
                    fill = shd.get(qn("w:fill"))
                    if fill and fill != "auto":
                        shading_hex = f"#{fill}"

                # Vertical alignment
                v_align = tc_pr.find(qn("w:vAlign"))
                if v_align is not None:
                    va_val = v_align.get(qn("w:val"), "top")
                    va_map = {
                        "top": VerticalAlignment.TOP,
                        "center": VerticalAlignment.CENTER,
                        "bottom": VerticalAlignment.BOTTOM,
                    }
                    vert_align = va_map.get(va_val)
        except Exception:
            pass

        # Detect cell merge spans
        col_span = 1
        row_span = 1
        try:
            tc = cell._tc
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is not None:
                grid_span = tc_pr.find(qn("w:gridSpan"))
                if grid_span is not None:
                    col_span = int(grid_span.get(qn("w:val"), "1"))
                v_merge = tc_pr.find(qn("w:vMerge"))
                if v_merge is not None:
                    merge_val = v_merge.get(qn("w:val"), "continue")
                    if merge_val == "restart":
                        row_span = 1  # Start of vertical merge
        except Exception:
            pass

        return CellStyle(
            width_inches=width_inches,
            shading_color_hex=shading_hex,
            vertical_alignment=vert_align,
            col_span=col_span,
            row_span=row_span,
        )

    # ------------------------------------------------------------------
    # Headers & Footers
    # ------------------------------------------------------------------

    def _extract_headers_footers(self, doc: Document) -> list[ContentElement]:
        """Extract header and footer content from the first section."""
        elements: list[ContentElement] = []
        if not doc.sections:
            return elements

        section = doc.sections[0]

        # Header
        try:
            header = section.header
            if header and not header.is_linked_to_previous:
                for para in header.paragraphs:
                    if para.text.strip():
                        runs = self._extract_runs(para, doc)
                        para_style = self._extract_paragraph_style(para)
                        style_ref = self._register_para_style(para, para_style)
                        elements.append(
                            ContentElement(
                                type=ElementType.HEADER,
                                content=runs,
                                style_ref=style_ref,
                            )
                        )
        except Exception:
            pass

        # Footer
        try:
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                for para in footer.paragraphs:
                    if para.text.strip():
                        runs = self._extract_runs(para, doc)
                        para_style = self._extract_paragraph_style(para)
                        style_ref = self._register_para_style(para, para_style)
                        elements.append(
                            ContentElement(
                                type=ElementType.FOOTER,
                                content=runs,
                                style_ref=style_ref,
                            )
                        )
        except Exception:
            pass

        return elements

    # ------------------------------------------------------------------
    # Style registration (deduplication)
    # ------------------------------------------------------------------

    def _register_run_style(self, run_style: RunStyle) -> str:
        """Register a run style and return its reference key, deduplicating."""
        fp = _fingerprint_run_style(run_style)
        if fp in self._run_style_cache:
            return self._run_style_cache[fp]

        # Generate a meaningful key
        parts = []
        if run_style.font_name:
            parts.append(run_style.font_name.lower().replace(" ", ""))
        if run_style.font_size_pt:
            parts.append(f"{run_style.font_size_pt}pt")
        if run_style.bold:
            parts.append("bold")
        if run_style.italic:
            parts.append("italic")
        if not parts:
            parts.append(f"run_{self._run_style_counter}")
        self._run_style_counter += 1

        key = _make_style_key("run", *parts, str(self._run_style_counter))
        self._run_style_cache[fp] = key
        self._run_styles[key] = run_style
        return key

    def _register_para_style(self, para: Paragraph, para_style: ParagraphStyle) -> str:
        """Register a paragraph style and return its reference key."""
        fp = _fingerprint_para_style(para_style)
        if fp in self._para_style_cache:
            return self._para_style_cache[fp]

        # Use the Word style name as part of the key
        word_style_name = para.style.name if para.style else "default"
        self._para_style_counter += 1
        key = _make_style_key("para", word_style_name, str(self._para_style_counter))

        self._para_style_cache[fp] = key
        self._para_styles[key] = para_style
        return key


# ---------------------------------------------------------------------------
# Convenience function
# ---------------------------------------------------------------------------

def extract_word_document(
    file_path: Optional[str] = None,
    file_stream: Optional[io.BytesIO] = None,
    filename: Optional[str] = None,
) -> tuple[DocumentContent, DocumentStyling]:
    """
    Extract content and styling from a Word document.

    Returns:
        (DocumentContent, DocumentStyling) — the two JSON-serializable objects.
    """
    extractor = WordExtractor()
    return extractor.extract(file_path=file_path, file_stream=file_stream, filename=filename)


# ---------------------------------------------------------------------------
# Template fingerprinting (hackathon spec §2)
# ---------------------------------------------------------------------------

import re as _re
import zipfile as _zipfile
from xml.etree import ElementTree as _ET

from app.schemas.document_model import (
    DraftSection,
    DraftStructure,
    HeaderFooterSection,
    HeadingSlot,
    NumberingDef,
    TableSchema,
    TemplateFingerprint,
)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS = {"w": _W_NS}


def _slugify(text: str, fallback: str) -> str:
    if not text:
        return fallback
    slug = _re.sub(r"[^a-zA-Z0-9]+", "_", text.strip().lower()).strip("_")
    return slug or fallback


def _parse_numbering_xml(docx_bytes: bytes) -> list[NumberingDef]:
    """Pull list-numbering definitions from word/numbering.xml."""
    defs: list[NumberingDef] = []
    try:
        with _zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
            if "word/numbering.xml" not in z.namelist():
                return defs
            xml = z.read("word/numbering.xml")
    except Exception:
        return defs

    try:
        root = _ET.fromstring(xml)
    except _ET.ParseError:
        return defs

    abstracts: dict[str, dict] = {}
    for an in root.findall("w:abstractNum", _NS):
        an_id = an.get(f"{{{_W_NS}}}abstractNumId", "")
        levels: dict[str, dict] = {}
        for lvl in an.findall("w:lvl", _NS):
            ilvl = lvl.get(f"{{{_W_NS}}}ilvl", "0")
            fmt_el = lvl.find("w:numFmt", _NS)
            text_el = lvl.find("w:lvlText", _NS)
            levels[ilvl] = {
                "num_fmt": fmt_el.get(f"{{{_W_NS}}}val") if fmt_el is not None else None,
                "lvl_text": text_el.get(f"{{{_W_NS}}}val") if text_el is not None else None,
            }
        abstracts[an_id] = {"levels": levels}

    for num in root.findall("w:num", _NS):
        num_id = num.get(f"{{{_W_NS}}}numId", "")
        ref = num.find("w:abstractNumId", _NS)
        an_id = ref.get(f"{{{_W_NS}}}val") if ref is not None else None
        abstract = abstracts.get(an_id or "", {"levels": {}})
        defs.append(
            NumberingDef(
                num_id=num_id,
                abstract_num_id=an_id,
                levels=abstract["levels"],
            )
        )
    return defs


def _detect_toc_paragraph_index(doc) -> Optional[int]:
    """Find the paragraph index where the TOC field begins, if any."""
    for idx, para in enumerate(doc.paragraphs):
        for fld in para._p.iter(qn("w:fldChar")):
            # walk siblings looking for an instrText with "TOC"
            pass
        for instr in para._p.iter(qn("w:instrText")):
            text = instr.text or ""
            if text.strip().upper().startswith("TOC"):
                return idx
        # also check for SDT-based TOCs
        for sdt in para._p.iter(qn("w:sdt")):
            for instr in sdt.iter(qn("w:instrText")):
                if (instr.text or "").strip().upper().startswith("TOC"):
                    return idx
    return None


def _extract_jinja_marker(text: str) -> Optional[str]:
    """Return the first `{{ ... }}` marker in the paragraph text, if any."""
    m = _re.search(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}", text or "")
    return m.group(0) if m else None


def _derive_heading_slots(doc, extractor: "WordExtractor") -> list[HeadingSlot]:
    """Walk the doc's paragraphs and build HeadingSlot entries from headings."""
    slots: list[HeadingSlot] = []
    used_ids: set[str] = set()
    counter = 0
    next_marker: Optional[str] = None

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        marker = _extract_jinja_marker(text)
        level = extractor._get_heading_level(para)
        if level is not None and text:
            counter += 1
            base = _slugify(text, f"slot_{counter}")
            slot_id = base
            i = 1
            while slot_id in used_ids:
                i += 1
                slot_id = f"{base}_{i}"
            used_ids.add(slot_id)
            slots.append(
                HeadingSlot(
                    slot_id=slot_id,
                    level=level,
                    title=text,
                    required=True,
                    expected_keywords=[w.lower() for w in _re.findall(r"[A-Za-z]{4,}", text)],
                    placeholder_marker=next_marker or marker,
                )
            )
            next_marker = None
        elif marker and slots:
            # Jinja marker that follows a heading paragraph → attach to last slot
            slots[-1].placeholder_marker = marker
        elif marker:
            # Marker before any heading → stash for the next one
            next_marker = marker

    return _drop_preface_heading_slots(slots)


def _drop_preface_heading_slots(slots: list[HeadingSlot]) -> list[HeadingSlot]:
    """Remove title/subtitle masthead slots before the first numbered section.

    The heuristic heading detector intentionally recognizes bold title blocks,
    but Flow 1 should not treat those as editable body sections. If a document
    has numbered sections, the opening unnumbered title/subtitle area is
    masthead content and must be preserved in place. Jinja-marked templates are
    left untouched because the user explicitly marked those slots.
    """
    first_numbered = None
    for idx, slot in enumerate(slots):
        if re.match(r"^\s*\d+(?:\.\d+)*[\.\)]?\s+\S", slot.title or ""):
            first_numbered = idx
            break
    if not first_numbered:
        return slots
    if any(s.placeholder_marker for s in slots[:first_numbered]):
        return slots
    return slots[first_numbered:]


def _extract_table_schemas(doc) -> list[TableSchema]:
    schemas: list[TableSchema] = []
    for t in doc.tables:
        if not t.rows:
            continue
        header_row = [(c.text or "").strip() for c in t.rows[0].cells]
        schemas.append(
            TableSchema(
                header_row=header_row,
                expected_columns=len(header_row),
                is_dynamic=len(t.rows) <= 2,
            )
        )
    return schemas


def _extract_headers_footers_all(doc) -> list[HeaderFooterSection]:
    out: list[HeaderFooterSection] = []
    for i, section in enumerate(doc.sections):
        header_text = None
        footer_text = None
        try:
            if section.header is not None:
                header_text = "\n".join(
                    (p.text or "").strip() for p in section.header.paragraphs if (p.text or "").strip()
                ) or None
        except Exception:
            pass
        try:
            if section.footer is not None:
                footer_text = "\n".join(
                    (p.text or "").strip() for p in section.footer.paragraphs if (p.text or "").strip()
                ) or None
        except Exception:
            pass
        out.append(
            HeaderFooterSection(
                section_index=i,
                header_text=header_text,
                footer_text=footer_text,
            )
        )
    return out


def fingerprint_word_template(
    file_path: Optional[str] = None,
    file_stream: Optional[io.BytesIO] = None,
    filename: Optional[str] = None,
    include_template_bytes: bool = True,
) -> TemplateFingerprint:
    """
    Build a TemplateFingerprint from a .docx template.

    The fingerprint captures heading hierarchy, numbering.xml definitions,
    table schemas, headers/footers across all sections, TOC location, and a
    base64 copy of the original .docx so the docxtpl emitter can render
    directly into the same file (preserving every original style/asset).
    """
    if file_path:
        with open(file_path, "rb") as fh:
            raw = fh.read()
        if not filename:
            filename = os.path.basename(file_path)
    elif file_stream is not None:
        file_stream.seek(0)
        raw = file_stream.read()
    else:
        raise ValueError("Either file_path or file_stream must be provided")

    extractor = WordExtractor()
    content, styling = extractor.extract(
        file_stream=io.BytesIO(raw),
        filename=filename or "template.docx",
    )

    doc = Document(io.BytesIO(raw))

    heading_slots = _derive_heading_slots(doc, extractor)
    numbering = _parse_numbering_xml(raw)
    tables = _extract_table_schemas(doc)
    headers_footers = _extract_headers_footers_all(doc)
    toc_idx = _detect_toc_paragraph_index(doc)

    template_b64 = base64.b64encode(raw).decode("ascii") if include_template_bytes else None

    return TemplateFingerprint(
        metadata=content.metadata,
        page_style=styling.page_style,
        heading_hierarchy=heading_slots,
        numbering_defs=numbering,
        table_schemas=tables,
        headers_footers=headers_footers,
        toc_location=toc_idx,
        style_registry=styling,
        template_b64=template_b64,
        source_format="docx",
    )


def _to_roman(num: int) -> str:
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
    ]
    syb = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV",
        "I"
    ]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syb[i]
            num -= val[i]
        i += 1
    return roman_num


def _to_letter(num: int) -> str:
    res = ""
    while num > 0:
        num, rem = divmod(num - 1, 26)
        res = chr(65 + rem) + res
    return res


def _format_list_number(num: int, fmt: str) -> str:
    if fmt == "upperRoman":
        return _to_roman(num) + "."
    elif fmt == "lowerRoman":
        return _to_roman(num).lower() + "."
    elif fmt == "upperLetter":
        return _to_letter(num) + "."
    elif fmt == "lowerLetter":
        return _to_letter(num).lower() + "."
    else:
        return f"{num}."


# A literal list leader at the start of a list item's TEXT (bullet glyph or
# "1." / "(a)" / "iv)" style number) — present only when the list was detected
# from text patterns rather than real Word numbering.
_LITERAL_LIST_LEADER_RE = re.compile(
    r"^\s*(?:[•◦▪▸‣●○·\-–—*]|\(?\d{1,3}\s*[.)\]]|\(?[a-zA-Z]\s*[.)\]])\s+"
)


def structure_word_draft(
    file_path: Optional[str] = None,
    file_stream: Optional[io.BytesIO] = None,
    filename: Optional[str] = None,
) -> DraftStructure:
    """
    Convert a .docx draft into a section-tagged JSON tree.

    Sections are split at every heading. Each DraftSection accumulates the
    body text (and any tables) that fall between consecutive headings.
    """
    content, _ = extract_word_document(
        file_path=file_path, file_stream=file_stream, filename=filename
    )

    sections: list[DraftSection] = []
    current = DraftSection(index=0, heading=None, level=0, text="")
    list_counters: dict[int, int] = {}
    last_list_kind: Optional[ListType] = None

    def _flush():
        if current.heading or current.text.strip() or current.tables:
            sections.append(current)

    def _reset_list_state():
        nonlocal last_list_kind
        list_counters.clear()
        last_list_kind = None

    def _append_text(line: str):
        current.text = (current.text + "\n" + line).strip() if current.text else line

    def _list_marker(el: ContentElement) -> str:
        nonlocal last_list_kind
        level = max(0, min(el.list_level or 0, 8))
        indent = "  " * level
        if el.list_type == ListType.NUMBERED:
            if last_list_kind != ListType.NUMBERED:
                list_counters.clear()
            list_counters[level] = list_counters.get(level, 0) + 1
            for k in list(list_counters):
                if k > level:
                    del list_counters[k]
            last_list_kind = ListType.NUMBERED
            fmt = el.number_format or "decimal"
            marker = _format_list_number(list_counters[level], fmt)
            return f"{indent}{marker}"
        last_list_kind = ListType.BULLET
        return f"{indent}-"

    for el in content.elements:
        if el.type == ElementType.HEADING:
            _flush()
            _reset_list_state()
            heading_text = "".join((r.text or "") for r in (el.content or []))
            current = DraftSection(
                index=len(sections),
                heading=heading_text.strip(),
                level=el.level or 1,
                text="",
            )
            for r in (el.content or []):
                if getattr(r, "image_data_b64", None):
                    current.images_b64.append(r.image_data_b64)
        elif el.type == ElementType.PARAGRAPH:
            text = "".join((r.text or "") for r in (el.content or []))
            if text:
                _append_text(text)
            _reset_list_state()
            for r in (el.content or []):
                if getattr(r, "image_data_b64", None):
                    current.images_b64.append(r.image_data_b64)
        elif el.type == ElementType.LIST_ITEM:
            text = "".join((r.text or "") for r in (el.content or []))
            # Text-detected list items keep their literal leader ("1. ", "• ")
            # — strip it so the synthesized marker isn't doubled ("1. 1. …").
            text = _LITERAL_LIST_LEADER_RE.sub("", text, count=1).strip()
            if text:
                _append_text(f"{_list_marker(el)} {text}")
        elif el.type == ElementType.TABLE and el.rows:
            _reset_list_state()
            grid: list[list[str]] = []
            for row in el.rows:
                grid.append(
                    [
                        "".join((r.text or "") for r in (cell.content or []))
                        for cell in row.cells
                    ]
                )
            current.tables.append(grid)
            # Inline the table as markdown rows AT ITS POSITION in the text —
            # the rewrite pipeline protects these rows behind placeholders and
            # the emitter uses them to keep prose on the correct side of each
            # real table when a section is regenerated.
            md_rows = [
                "| " + " | ".join((c or "").replace("\n", " ").strip() for c in row) + " |"
                for row in grid
            ]
            _append_text("\n".join(md_rows))
        elif el.type == ElementType.IMAGE and el.data_base64:
            _reset_list_state()
            current.images_b64.append(el.data_base64)

    _flush()

    return DraftStructure(metadata=content.metadata, sections=sections)


# ---------------------------------------------------------------------------
# CLI entry point (for standalone testing)
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) < 2:
        print("Usage: python word_ext.py <path-to-docx>")
        sys.exit(1)

    path = sys.argv[1]
    content, styling = extract_word_document(file_path=path)

    # Write outputs
    base = os.path.splitext(path)[0]

    content_path = f"{base}_content.json"
    with open(content_path, "w", encoding="utf-8") as f:
        json.dump(content.model_dump(exclude_none=True), f, indent=2, ensure_ascii=False)
    print(f"Content JSON written to: {content_path}")

    styling_path = f"{base}_styling.json"
    with open(styling_path, "w", encoding="utf-8") as f:
        json.dump(styling.model_dump(exclude_none=True), f, indent=2, ensure_ascii=False)
    print(f"Styling JSON written to: {styling_path}")
