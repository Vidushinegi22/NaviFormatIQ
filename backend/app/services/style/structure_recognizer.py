"""
Document structure recognition — understand a flat document semantically.
=========================================================================

The Flow-2 style engine can only *apply* a look as well as it *understands* the
content document. A "raw" content doc is frequently flat: section headings are
plain ``Normal`` paragraphs ("1. Purpose"), list items are plain paragraphs,
tables are encoded as pipe-delimited text, and the masthead (title / subtitle /
document number) is just three more paragraphs. A pure style-layer swap leaves
all of that flat — headings never become headings, lists never become lists.

This module reads the content document and produces a :class:`StructurePlan`:
for every body paragraph it decides a *semantic role* —

    title · subtitle · metadata · heading(level) · list_item(kind, level) · body

— and it detects *table blocks* (runs of delimited-text paragraphs that should
become real Word tables). :mod:`structure_apply` then upgrades the live OOXML in
place, so the document keeps 100% of its content but gains real structure that
can carry the target's styling.

Recognition is **LLM-first with a strong deterministic fallback**. The heuristic
alone handles numbered headings, literal-leader lists, parallel-block lists and
pipe/tab-delimited tables; the LLM (when configured) refines ambiguous cases
(unnumbered headings, prose-vs-list, masthead) and the deterministic
numbering-based heading levels are always layered on top as ground truth.

The plan is **safe and idempotent**: paragraphs that are already real headings
or real lists are preserved, and only flat paragraphs are ever upgraded.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Optional

from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Plan data structures
# ---------------------------------------------------------------------------

# Semantic roles a paragraph can take.
TITLE = "title"
SUBTITLE = "subtitle"
METADATA = "metadata"
HEADING = "heading"
LIST_ITEM = "list_item"
BODY = "body"
TABLE_CELL = "table_cell"  # consumed by a TableBlock — not rendered as a paragraph

BULLET = "bullet"
NUMBERED = "numbered"


@dataclass
class ParaRole:
    """The decided role of a single body paragraph."""
    kind: str = BODY
    level: int = 1               # heading level (1-4) or list nesting level (0-based)
    list_kind: str = BULLET      # for list_item: "bullet" | "numbered"
    source: str = "heuristic"    # "heuristic" | "llm" | "numbering" | "existing"
    masthead: bool = False       # metadata that sits in the centred title block


@dataclass
class TableBlock:
    """A run of consecutive paragraphs that should become one Word table."""
    start_idx: int               # paragraph index of the first row
    end_idx: int                 # paragraph index of the last row (inclusive)
    rows: list[list[str]]
    has_header: bool = True


@dataclass
class StructurePlan:
    roles: dict[int, ParaRole] = field(default_factory=dict)   # para idx -> role
    tables: list[TableBlock] = field(default_factory=list)
    consumed_idxs: set[int] = field(default_factory=set)       # idxs inside a table block
    method: str = "heuristic"                                  # how the plan was built
    # (paragraphs sent, total non-empty) when the LLM refinement only saw a
    # prefix of a very long document — surfaced so partial coverage is visible.
    llm_truncated: Optional[tuple[int, int]] = None

    def role(self, idx: int) -> ParaRole:
        return self.roles.get(idx, ParaRole(kind=BODY))

    def summary(self) -> dict:
        counts = {TITLE: 0, SUBTITLE: 0, METADATA: 0, HEADING: 0, LIST_ITEM: 0, BODY: 0}
        heading_levels: dict[int, int] = {}
        list_kinds = {BULLET: 0, NUMBERED: 0}
        for idx, r in self.roles.items():
            if idx in self.consumed_idxs:
                continue
            counts[r.kind] = counts.get(r.kind, 0) + 1
            if r.kind == HEADING:
                heading_levels[r.level] = heading_levels.get(r.level, 0) + 1
            elif r.kind == LIST_ITEM:
                list_kinds[r.list_kind] = list_kinds.get(r.list_kind, 0) + 1
        out = {
            "headings": counts[HEADING],
            "heading_levels": dict(sorted(heading_levels.items())),
            "list_items": counts[LIST_ITEM],
            "list_bullet": list_kinds[BULLET],
            "list_numbered": list_kinds[NUMBERED],
            "tables": len(self.tables),
            "title_block": counts[TITLE] + counts[SUBTITLE] + counts[METADATA],
            "method": self.method,
        }
        if self.llm_truncated:
            sent, total = self.llm_truncated
            out["llm_coverage"] = (
                f"partial — LLM refinement saw the first {sent} of {total} "
                f"paragraphs; the rest used heuristics only"
            )
        return out


# ---------------------------------------------------------------------------
# Paragraph inspection
# ---------------------------------------------------------------------------

@dataclass
class _ParaInfo:
    idx: int
    text: str
    style_name: str
    is_empty: bool
    bold: bool
    all_caps: bool
    alignment: Optional[int]
    existing_heading_level: Optional[int]   # from a real Heading style
    existing_list: bool                     # has numPr or List* style


_HEADING_STYLE_RE = re.compile(r"heading\s*(\d)", re.I)


def _para_infos(doc) -> list[_ParaInfo]:
    infos: list[_ParaInfo] = []
    for idx, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        try:
            sname = p.style.name or ""
        except Exception:
            sname = ""
        # existing heading?
        ex_level = None
        m = _HEADING_STYLE_RE.search(sname)
        if m:
            ex_level = int(m.group(1))
        elif sname.lower() in ("title",):
            ex_level = 0
        # existing list? (numPr lives under pPr, so search descendants)
        has_numpr = p._p.find(".//" + qn("w:numPr")) is not None
        ex_list = has_numpr or sname.lower().startswith("list")
        # dominant bold
        runs = [r for r in p.runs if (r.text or "").strip()]
        bold_chars = sum(len(r.text) for r in runs if r.bold)
        total = sum(len(r.text) for r in runs) or 1
        bold = total > 0 and bold_chars / total >= 0.6
        infos.append(_ParaInfo(
            idx=idx,
            text=text,
            style_name=sname,
            is_empty=not text,
            bold=bold,
            all_caps=bool(text) and text.upper() == text and any(c.isalpha() for c in text),
            alignment=int(p.alignment) if p.alignment is not None else None,
            existing_heading_level=ex_level,
            existing_list=ex_list,
        ))
    return infos


# ---------------------------------------------------------------------------
# Heuristic recognisers
# ---------------------------------------------------------------------------

# A leading "1." / "3.1" / "3.1.2" section number.
_SECTION_NUM_RE = re.compile(r"^\s*(\d+(?:\.\d+)*)\.?\s+\S")
# A single-level numbered-list leader: "1." "2)" "(3)".
_NUM_LEADER_RE = re.compile(r"^\s*\(?\s*(\d{1,3})\s*[.)\]]\s+\S")
# An alpha / roman leader: "a)" "(iv)" "A.".
_ALPHA_LEADER_RE = re.compile(r"^\s*\(?\s*(?:[ivxlcdm]{1,6}|[a-z])\s*[.)\]]\s+\S", re.I)
# A bullet-glyph leader.
_BULLET_LEADER_RE = re.compile(r"^\s*[•●○◦▪■□◆◇‣·\-–—*]\s+\S")

# Masthead metadata signals.
_META_RE = re.compile(
    r"(document\s*(?:no|number)|version|effective\s*date|reference|revision|"
    r"\bdoc[-\s]?id\b|prepared\s*by|reviewed\s*by|approved\s*by|department\s*:)",
    re.I,
)

# Sentence-final punctuation that suggests flowing prose, not a list item.
_PROSE_TAIL = (".", "!", "?")

_MAX_HEADING_WORDS = 12
_MAX_HEADING_CHARS = 90
_MAX_LIST_CHARS = 320


def _section_level(text: str) -> Optional[int]:
    """Return the heading level implied by a leading section number, else None."""
    m = _SECTION_NUM_RE.match(text)
    if not m:
        return None
    parts = m.group(1).split(".")
    return min(len(parts), 4)


def _looks_like_heading(info: _ParaInfo) -> bool:
    """Unnumbered heading heuristic: short, title-ish, often bold/all-caps."""
    t = info.text
    if not t or len(t) > _MAX_HEADING_CHARS:
        return False
    words = t.split()
    if len(words) > _MAX_HEADING_WORDS:
        return False
    # A trailing period (other than a colon) reads like prose.
    if t.rstrip().endswith(_PROSE_TAIL):
        return False
    # A "Label: value" line ("Table: stg_mass_email") is a field, not a
    # heading; a *trailing* colon ("Results:") is still heading-eligible.
    _head, sep, tail = t.partition(":")
    if sep and tail.strip():
        return False
    if info.all_caps:
        return True
    if info.bold and (t.istitle() or sum(1 for w in words if w[:1].isupper()) >= max(1, len(words) - 1)):
        return True
    return False


# Common column delimiters, in priority order.
_DELIMITERS = ("|", "\t")


def _split_row(text: str, delim: str) -> list[str]:
    cells = [c.strip() for c in text.split(delim)]
    # tab-delimited rows sometimes have a leading/trailing empty cell — keep
    # interior structure but drop purely empty edges.
    while cells and cells[0] == "":
        cells.pop(0)
    while cells and cells[-1] == "":
        cells.pop()
    return cells


def _detect_tables(infos: list[_ParaInfo]) -> list[TableBlock]:
    """Find runs of >=2 consecutive non-empty paragraphs that share a delimiter
    and a consistent column count (>=2 columns)."""
    blocks: list[TableBlock] = []
    i = 0
    n = len(infos)
    while i < n:
        info = infos[i]
        if info.is_empty or info.existing_list or info.existing_heading_level is not None:
            i += 1
            continue
        # Pick the delimiter that yields the most columns for this line.
        best_delim = None
        best_cols = 1
        for d in _DELIMITERS:
            if d in info.text:
                ncols = len(_split_row(info.text, d))
                if ncols > best_cols:
                    best_cols, best_delim = ncols, d
        if best_delim is None or best_cols < 2:
            i += 1
            continue
        # Greedily extend the run while the column count stays consistent.
        rows = [_split_row(info.text, best_delim)]
        j = i + 1
        while j < n:
            nxt = infos[j]
            if nxt.is_empty or best_delim not in nxt.text:
                break
            cols = _split_row(nxt.text, best_delim)
            # Allow a +/-1 column wobble (ragged trailing cells) but require overlap.
            if abs(len(cols) - best_cols) > 1 or len(cols) < 2:
                break
            rows.append(cols)
            j += 1
        if len(rows) >= 2:
            width = max(len(r) for r in rows)
            norm = [r + [""] * (width - len(r)) for r in rows]
            blocks.append(TableBlock(start_idx=i, end_idx=j - 1, rows=norm, has_header=True))
            i = j
        else:
            i += 1
    return blocks


def _detect_masthead(infos: list[_ParaInfo]) -> dict[int, ParaRole]:
    """Classify the masthead — the title/subtitle/metadata that opens a document.

    The masthead is the first block of non-empty paragraphs, bounded by the
    first blank line OR the first numbered section heading (whichever comes
    first). The opening paragraph is the title; a following non-metadata line is
    the subtitle; document-number / version / date lines are metadata.
    """
    roles: dict[int, ParaRole] = {}
    end = len(infos)
    for info in infos:
        if info.is_empty or _section_level(info.text) is not None:
            end = info.idx
            break
    leading = [info for info in infos if info.idx < end and not info.is_empty]
    if not leading:
        return roles
    roles[leading[0].idx] = ParaRole(kind=TITLE, source="heuristic", masthead=True)
    seen_subtitle = False
    for info in leading[1:]:
        if _META_RE.search(info.text) or _looks_like_doc_meta(info.text):
            roles[info.idx] = ParaRole(kind=METADATA, source="heuristic", masthead=True)
        elif not seen_subtitle and len(info.text.split()) <= 18:
            roles[info.idx] = ParaRole(kind=SUBTITLE, source="heuristic", masthead=True)
            seen_subtitle = True
        else:
            roles[info.idx] = ParaRole(kind=METADATA, source="heuristic", masthead=True)
    return roles


def _detect_metadata_blocks(infos: list[_ParaInfo], plan: StructurePlan) -> None:
    """Mark runs of >=2 consecutive 'Label: value' metadata lines (e.g.
    Department / Prepared By / Reviewed By / Approved By) so they render as a
    clean metadata block instead of being mistaken for a bullet list."""
    n = len(infos)
    i = 0
    while i < n:
        info = infos[i]
        if (info.is_empty or info.idx in plan.roles or info.idx in plan.consumed_idxs
                or not _META_RE.search(info.text)):
            i += 1
            continue
        run = [info.idx]
        j = i + 1
        while j < n and not infos[j].is_empty and infos[j].idx not in plan.roles \
                and _META_RE.search(infos[j].text):
            run.append(infos[j].idx)
            j += 1
        if len(run) >= 2:
            for k in run:
                plan.roles[k] = ParaRole(kind=METADATA, source="heuristic")
        i = j if len(run) >= 2 else i + 1


def _looks_like_doc_meta(text: str) -> bool:
    # Lines with date-ish or "X: Y | A: B" patterns read as metadata.
    if text.count("|") >= 1 and ":" in text:
        return True
    if re.search(r"\b\d{1,2}[-/][A-Za-z]{3,9}[-/]\d{2,4}\b", text):
        return True
    return False


def _heuristic_plan(infos: list[_ParaInfo]) -> StructurePlan:
    plan = StructurePlan(method="heuristic")

    # 1) Tables first — they consume their paragraphs.
    plan.tables = _detect_tables(infos)
    for tb in plan.tables:
        for k in range(tb.start_idx, tb.end_idx + 1):
            plan.consumed_idxs.add(k)
            plan.roles[k] = ParaRole(kind=TABLE_CELL, source="heuristic")

    # 2) Masthead (title / subtitle / metadata) + standalone metadata blocks.
    for idx, role in _detect_masthead(infos).items():
        if idx not in plan.consumed_idxs:
            plan.roles[idx] = role
    _detect_metadata_blocks(infos, plan)

    # 3) Headings — numbered (deterministic) + existing styles + unnumbered.
    #    Single-level numbers that form an adjacent run are a numbered LIST, not
    #    a stack of headings; multi-level numbers (3.1) are always headings.
    numbered_singletons: dict[int, _ParaInfo] = {}
    for info in infos:
        if info.idx in plan.consumed_idxs or info.is_empty or info.idx in plan.roles:
            continue
        if info.existing_heading_level is not None and info.existing_heading_level >= 1:
            plan.roles[info.idx] = ParaRole(
                kind=HEADING, level=info.existing_heading_level, source="existing"
            )
            continue
        lvl = _section_level(info.text)
        if lvl is None:
            continue
        m = _SECTION_NUM_RE.match(info.text)
        is_multi = "." in (m.group(1) if m else "")
        if is_multi:
            plan.roles[info.idx] = ParaRole(kind=HEADING, level=lvl, source="numbering")
        else:
            numbered_singletons[info.idx] = info  # decide list-vs-heading below

    # Decide single-level numbered lines: adjacency (ignoring blanks) => list.
    singleton_idxs = sorted(numbered_singletons)
    consumed_as_list: set[int] = set()
    for idx in singleton_idxs:
        if idx in consumed_as_list:
            continue
        # Walk forward over adjacent single-level-numbered paragraphs.
        run = [idx]
        cur = idx
        while True:
            nxt = _next_nonempty_idx(infos, cur)
            if nxt is not None and nxt in numbered_singletons:
                run.append(nxt)
                cur = nxt
            else:
                break
        if len(run) >= 2:
            for k in run:
                consumed_as_list.add(k)
                plan.roles[k] = ParaRole(kind=LIST_ITEM, level=0, list_kind=NUMBERED, source="heuristic")
        else:
            info = numbered_singletons[idx]
            # Isolated single number → heading if it's short/title-ish.
            if len(info.text.split()) <= _MAX_HEADING_WORDS and len(info.text) <= _MAX_HEADING_CHARS:
                plan.roles[idx] = ParaRole(kind=HEADING, level=1, source="numbering")
            else:
                plan.roles[idx] = ParaRole(kind=LIST_ITEM, level=0, list_kind=NUMBERED, source="heuristic")

    # Unnumbered headings (bold / all-caps / title-ish), only if not already set.
    for info in infos:
        if info.idx in plan.roles or info.idx in plan.consumed_idxs or info.is_empty:
            continue
        if info.existing_list:
            continue
        if _looks_like_heading(info):
            lvl = 1 if info.all_caps and len(info.text.split()) <= 6 else 2
            plan.roles[info.idx] = ParaRole(kind=HEADING, level=lvl, source="heuristic")

    # 4) Lists — literal leaders + parallel blocks among the remaining body paras.
    _detect_lists(infos, plan)

    # 5) Everything else non-empty that is still unset → body.
    for info in infos:
        if info.is_empty or info.idx in plan.consumed_idxs:
            continue
        if info.idx not in plan.roles:
            if info.existing_list:
                plan.roles[info.idx] = ParaRole(
                    kind=LIST_ITEM, level=0,
                    list_kind=NUMBERED if "number" in info.style_name.lower() else BULLET,
                    source="existing",
                )
            else:
                plan.roles[info.idx] = ParaRole(kind=BODY, source="heuristic")

    return plan


def _next_nonempty_idx(infos: list[_ParaInfo], idx: int) -> Optional[int]:
    for j in range(idx + 1, len(infos)):
        if not infos[j].is_empty:
            return j
    return None


def _detect_lists(infos: list[_ParaInfo], plan: StructurePlan) -> None:
    """Mark list items via literal leaders and parallel-block detection.

    Literal leaders are unambiguous. For unmarked paragraphs, a *parallel block*
    is a maximal run of >=2 consecutive (blank-tolerant) body paragraphs that are
    short, single-clause, and not headings — typical of an SOP's bullet lists.
    """
    n = len(infos)

    # Pass A — literal leaders.
    for info in infos:
        if info.idx in plan.roles or info.idx in plan.consumed_idxs or info.is_empty:
            continue
        kind = _leader_kind(info.text)
        if kind is not None:
            plan.roles[info.idx] = ParaRole(kind=LIST_ITEM, level=0, list_kind=kind, source="heuristic")

    # Pass B — parallel blocks of plain paragraphs.
    i = 0
    while i < n:
        info = infos[i]
        if (info.is_empty or info.idx in plan.consumed_idxs
                or info.idx in plan.roles):
            i += 1
            continue
        # Collect a run of consecutive (blank-tolerant) unmarked body paragraphs.
        run: list[int] = []
        j = i
        while j < n:
            cur = infos[j]
            if cur.is_empty:
                # A single blank is tolerated only if the next non-empty is also a candidate.
                nxt = _next_nonempty_idx(infos, j)
                if nxt is not None and _is_list_candidate(infos[nxt], plan):
                    j = nxt
                    continue
                break
            if cur.idx in plan.roles or cur.idx in plan.consumed_idxs:
                break
            if not _is_list_candidate(cur, plan):
                break
            run.append(cur.idx)
            j += 1
        # A run of >=2 short parallel items → bullet list (only when they look
        # list-like: most items are short and don't read as multi-sentence prose).
        if len(run) >= 2 and _run_is_listy([infos[k] for k in run]):
            for k in run:
                plan.roles[k] = ParaRole(kind=LIST_ITEM, level=0, list_kind=BULLET, source="heuristic")
            i = j
        else:
            i = max(j, i + 1)


def _leader_kind(text: str) -> Optional[str]:
    if _BULLET_LEADER_RE.match(text):
        return BULLET
    if _NUM_LEADER_RE.match(text):
        return NUMBERED
    if _ALPHA_LEADER_RE.match(text):
        return NUMBERED
    return None


def _is_list_candidate(info: _ParaInfo, plan: StructurePlan) -> bool:
    if info.is_empty or info.idx in plan.consumed_idxs:
        return False
    if info.existing_heading_level is not None:
        return False
    if len(info.text) > _MAX_LIST_CHARS:
        return False
    if _section_level(info.text) is not None:
        return False  # a section number — handled as heading/numbered-list already
    # A line ending with a colon is a list LEAD-IN (e.g. "...as follows:"), not a
    # list item — leaving it as body starts the list cleanly on the next line.
    if info.text.rstrip().endswith(":"):
        return False
    return True


def _run_is_listy(items: list[_ParaInfo]) -> bool:
    """A run reads as a list when items are short and rarely multi-sentence."""
    if len(items) < 2:
        return False
    multi_sentence = 0
    for it in items:
        # crude sentence count: interior periods followed by a space + capital.
        interior = len(re.findall(r"\.\s+[A-Z]", it.text))
        if interior >= 2 or len(it.text) > 240:
            multi_sentence += 1
    return multi_sentence <= max(1, len(items) // 4)


# ---------------------------------------------------------------------------
# LLM refinement (optional, layered on top of the heuristic plan)
# ---------------------------------------------------------------------------

# Cap on the paragraphs sent for LLM refinement — beyond this the heuristic
# plan stands alone for the tail and the plan summary notes partial coverage.
_LLM_MAX_PARAS = 800

_LLM_SYSTEM = (
    "You are a meticulous document-structure analyst. You receive an ordered "
    "list of a document's non-empty paragraphs (index + text + current style). "
    "The document is 'flat' — section headings, list items and tables may be "
    "plain paragraphs. Classify the STRUCTURE so it can be re-formatted.\n\n"
    "Return ONLY JSON with this shape:\n"
    "{\n"
    '  "title_idx": <int|null>,            // the document title paragraph\n'
    '  "subtitle_idx": <int|null>,         // the subtitle, if any\n'
    '  "metadata_idxs": [<int>, ...],      // doc-number / version / date lines\n'
    '  "headings": [{"idx":<int>,"level":<1-4>}, ...],\n'
    '  "lists": [{"idx":<int>,"kind":"bullet"|"numbered","level":<0-3>}, ...],\n'
    '  "tables": [{"start_idx":<int>,"end_idx":<int>,"header":true}]\n'
    "}\n"
    "Rules: a numbered SECTION HEADING like '1. Purpose' or '3.1 Scope' is a "
    "heading (level = depth of its number). A back-to-back run of '1.'..'4.' "
    "sentences is a numbered LIST, not headings. Short parallel phrases under a "
    "heading are a bullet list. A run of pipe('|') or tab delimited lines is a "
    "table. Do not list plain body paragraphs. Indices MUST come from the input."
)


def _llm_refine(infos: list[_ParaInfo], plan: StructurePlan) -> StructurePlan:
    try:
        from app.llm.adapters import chat_json, llm_available
    except Exception:
        return plan
    if not llm_available():
        return plan

    non_empty = [info for info in infos if not info.is_empty]
    listing = [
        {"idx": info.idx, "style": info.style_name or "Normal", "text": info.text[:160]}
        for info in non_empty[:_LLM_MAX_PARAS]
    ]
    user = (
        "Paragraphs (idx, style, text):\n"
        + _compact_json(listing)
        + "\n\nReturn the structure JSON now."
    )
    raw = chat_json(_LLM_SYSTEM, user, temperature=0.0, max_tokens=3000)
    if not isinstance(raw, dict):
        return plan

    valid = {info.idx for info in infos if not info.is_empty}
    merged = StructurePlan(method="llm+heuristic")
    if len(non_empty) > _LLM_MAX_PARAS:
        merged.llm_truncated = (_LLM_MAX_PARAS, len(non_empty))
    # Start from heuristic tables/roles, then let the LLM override where it is
    # confident. Deterministic numbering levels are re-asserted at the end.
    merged.tables = list(plan.tables)
    merged.consumed_idxs = set(plan.consumed_idxs)
    merged.roles = dict(plan.roles)

    # Tables from the LLM (only if the heuristic didn't already find overlapping ones).
    for t in raw.get("tables") or []:
        try:
            s, e = int(t.get("start_idx")), int(t.get("end_idx"))
        except (TypeError, ValueError):
            continue
        if s not in valid or e not in valid or e < s:
            continue
        if any(not (e < tb.start_idx or s > tb.end_idx) for tb in merged.tables):
            continue  # overlaps an existing block
        rows = _rows_from_idxs(infos, s, e)
        if len(rows) >= 2 and max(len(r) for r in rows) >= 2:
            tb = TableBlock(start_idx=s, end_idx=e, rows=rows, has_header=bool(t.get("header", True)))
            merged.tables.append(tb)
            for k in range(s, e + 1):
                merged.consumed_idxs.add(k)
                merged.roles[k] = ParaRole(kind=TABLE_CELL, source="llm")

    # The centred masthead is the first block, bounded by the first blank line
    # or first heading; a later metadata block (Department / …) stays left.
    heading_idxs = [int(h["idx"]) for h in (raw.get("headings") or [])
                    if isinstance(h, dict) and isinstance(h.get("idx"), int)]
    first_heading = min(heading_idxs) if heading_idxs else len(infos)
    first_blank = next((info.idx for info in infos if info.is_empty), len(infos))
    masthead_end = min(first_blank, first_heading)

    # Title block.
    ti = raw.get("title_idx")
    if isinstance(ti, int) and ti in valid and ti not in merged.consumed_idxs:
        merged.roles[ti] = ParaRole(kind=TITLE, source="llm", masthead=True)
    si = raw.get("subtitle_idx")
    if isinstance(si, int) and si in valid and si not in merged.consumed_idxs:
        merged.roles[si] = ParaRole(kind=SUBTITLE, source="llm", masthead=True)
    for mi in raw.get("metadata_idxs") or []:
        if isinstance(mi, int) and mi in valid and mi not in merged.consumed_idxs:
            merged.roles[mi] = ParaRole(kind=METADATA, source="llm", masthead=mi < masthead_end)

    # Headings.
    for h in raw.get("headings") or []:
        try:
            idx, level = int(h.get("idx")), int(h.get("level", 2))
        except (TypeError, ValueError):
            continue
        if idx in valid and idx not in merged.consumed_idxs:
            merged.roles[idx] = ParaRole(kind=HEADING, level=max(1, min(level, 4)), source="llm")

    # Lists.
    for l in raw.get("lists") or []:
        try:
            idx = int(l.get("idx"))
        except (TypeError, ValueError):
            continue
        if idx not in valid or idx in merged.consumed_idxs:
            continue
        kind = NUMBERED if str(l.get("kind", "")).lower().startswith("num") else BULLET
        try:
            level = max(0, min(int(l.get("level", 0)), 3))
        except (TypeError, ValueError):
            level = 0
        merged.roles[idx] = ParaRole(kind=LIST_ITEM, level=level, list_kind=kind, source="llm")

    # Fill any non-empty unset paragraphs as body.
    for info in infos:
        if not info.is_empty and info.idx not in merged.roles:
            merged.roles[info.idx] = ParaRole(kind=BODY, source="llm")

    return merged


def _rows_from_idxs(infos: list[_ParaInfo], start: int, end: int) -> list[list[str]]:
    by_idx = {info.idx: info for info in infos}
    texts = [by_idx[k].text for k in range(start, end + 1) if k in by_idx and not by_idx[k].is_empty]
    delim = "|" if any("|" in t for t in texts) else ("\t" if any("\t" in t for t in texts) else "|")
    rows = [_split_row(t, delim) for t in texts]
    rows = [r for r in rows if len(r) >= 2]
    if not rows:
        return []
    width = max(len(r) for r in rows)
    return [r + [""] * (width - len(r)) for r in rows]


def _compact_json(obj) -> str:
    import json
    return json.dumps(obj, ensure_ascii=False, separators=(",", ":"))


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def recognize_structure(doc, *, use_llm: Optional[bool] = None) -> StructurePlan:
    """Analyse ``doc`` (a python-docx Document) and return a :class:`StructurePlan`.

    ``use_llm`` defaults to auto (uses the LLM when configured). The heuristic
    plan is always computed first as a robust foundation; the LLM only refines.
    Deterministic section-number heading levels are re-asserted as ground truth.
    """
    infos = _para_infos(doc)
    plan = _heuristic_plan(infos)

    if use_llm is None:
        try:
            from app.llm.adapters import llm_available
            use_llm = llm_available()
        except Exception:
            use_llm = False
    if use_llm:
        try:
            plan = _llm_refine(infos, plan)
        except Exception:
            pass  # never block on the LLM

    # Re-assert deterministic heading levels from explicit section numbering —
    # the most reliable signal, and immune to LLM mis-levelling.
    by_idx = {info.idx: info for info in infos}
    for idx, role in plan.roles.items():
        if role.kind != HEADING or idx in plan.consumed_idxs:
            continue
        info = by_idx.get(idx)
        if not info:
            continue
        lvl = _section_level(info.text)
        if lvl is not None and "." in (_SECTION_NUM_RE.match(info.text).group(1)):
            role.level = lvl
            role.source = "numbering"

    return plan
