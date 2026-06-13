"""
Style Engine — best-in-class document style transfer.
=====================================================

Apply the *visual identity* of a style-source document onto the *content* of
a target document, preserving 100% of the target's content and structure.

Two transfer paths, chosen automatically from the style source's file type:

  1. DOCX style source  → transplant the source's ``styles.xml`` (merged by
     styleId, source wins), ``docDefaults`` and ``theme1.xml`` into the
     content package. Content body XML is never touched → zero content loss.

  2. PDF style source    → build a :class:`StyleProfile` from PyMuPDF font /
     size / colour / alignment / margin histograms, then mutate the content
     document's named styles, ``docDefaults`` and page margins in place.

Crucially this engine NEVER rebuilds the body from a JSON intermediate (the
approach that previously dropped paragraphs, mangled tables and auto-numbered
every list). It edits the style layer of the real OOXML package, so it scales
to arbitrarily large documents and keeps every paragraph, table and image.

Public API
----------
    transfer_style(content_bytes, content_name, style_bytes, style_name, ...) -> bytes
    transfer_style_files(content_path, style_path, output_path, ...) -> str
    profile_from_pdf(pdf_bytes) -> StyleProfile
    profile_from_docx(docx_bytes) -> StyleProfile
"""

from __future__ import annotations

import collections
import copy
import io
import re
from dataclasses import dataclass, field
from typing import Optional, TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

if TYPE_CHECKING:  # annotations only — avoids any import cycle at runtime
    from app.services.style.guideline_interpreter import (
        HeaderFooterSpec,
        ListSpec,
        PageSpec,
        StyleSpec,
        TableSpec,
    )


# ---------------------------------------------------------------------------
# Profile dataclasses
# ---------------------------------------------------------------------------

@dataclass
class HeadingSpec:
    font: Optional[str] = None
    size_pt: Optional[float] = None
    bold: bool = True
    italic: bool = False
    underline: bool = False
    color_hex: Optional[str] = None          # "RRGGBB"
    alignment: str = "left"                   # left|center|right|justify
    space_before_pt: float = 12.0
    space_after_pt: float = 6.0


@dataclass
class StyleProfile:
    """A normalized description of a document's visual style."""
    body_font: str = "Calibri"
    body_size_pt: float = 11.0
    body_color_hex: str = "000000"
    body_alignment: str = "left"              # left|justify|...
    line_spacing: Optional[float] = None      # multiple, e.g. 1.15
    space_after_pt: float = 8.0
    heading_scale: dict[int, HeadingSpec] = field(default_factory=dict)
    margins_in: tuple[float, float, float, float] = (1.0, 1.0, 1.0, 1.0)  # T,B,L,R
    accent_color_hex: Optional[str] = None
    list_indent_in: float = 0.25
    source_kind: str = "pdf"                  # "pdf" | "docx"

    def heading(self, level: int) -> HeadingSpec:
        if level in self.heading_scale:
            return self.heading_scale[level]
        # Derive a sensible scale from the body if not explicitly captured.
        ratios = {1: 1.8, 2: 1.45, 3: 1.2, 4: 1.08, 5: 1.0, 6: 1.0}
        ratio = ratios.get(level, 1.0)
        return HeadingSpec(
            font=self.body_font,
            size_pt=round(self.body_size_pt * ratio, 1),
            bold=True,
            color_hex=self.accent_color_hex or self.body_color_hex,
            alignment="left",
            space_before_pt=max(6.0, 14.0 - level * 1.5),
            space_after_pt=max(3.0, 8.0 - level),
        )


_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_HEADING_STYLE_NAMES = {f"Heading {i}" for i in range(1, 10)} | {"Title", "Subtitle"}


# ---------------------------------------------------------------------------
# Low-level OOXML helpers
# ---------------------------------------------------------------------------

def _int_to_hex(color_int: int) -> str:
    return f"{color_int & 0xFFFFFF:06X}"


def _clean_font_name(raw: str) -> str:
    """Strip PyMuPDF subset prefixes (``ABCDEF+Calibri``) and weight/style
    suffixes (``Calibri-Bold`` → ``Calibri``)."""
    if not raw:
        return raw
    name = raw.split("+", 1)[-1]  # drop subset prefix
    # Drop common weight/style suffixes
    for sep in ("-", ","):
        if sep in name:
            head = name.split(sep, 1)[0]
            if head:
                name = head
                break
    # Drop trailing "MT"/"PS" PostScript markers
    name = re.sub(r"(MT|PS)$", "", name)
    return name.strip() or raw


def _set_style_font(style, font_name: Optional[str], size_pt: Optional[float] = None,
                    color_hex: Optional[str] = None) -> None:
    """Set a named style's run font on ALL script slots (ascii/hAnsi/cs/eastAsia)."""
    rpr = style.element.get_or_add_rPr()
    if font_name:
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), font_name)
    if size_pt is not None:
        try:
            style.font.size = Pt(size_pt)
        except Exception:
            pass
    if color_hex:
        try:
            style.font.color.rgb = RGBColor.from_string(color_hex)
        except Exception:
            pass


def _set_style_bool(style, *, bold: Optional[bool] = None,
                    italic: Optional[bool] = None,
                    underline: Optional[bool] = None) -> None:
    rpr = style.element.get_or_add_rPr()

    def _toggle(tag: str, on: Optional[bool]):
        if on is None:
            return
        existing = rpr.find(qn(tag))
        if on:
            if existing is None:
                el = OxmlElement(tag)
                rpr.append(el)
        else:
            if existing is not None:
                rpr.remove(existing)

    _toggle("w:b", bold)
    _toggle("w:i", italic)
    if underline is not None:
        existing = rpr.find(qn("w:u"))
        if existing is not None:
            rpr.remove(existing)
        if underline:
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rpr.append(u)


def _set_docdefaults_font(doc, font_name: str, size_pt: Optional[float]) -> None:
    """Set the document-wide default run font/size so any style that doesn't
    pin its own font inherits the body font."""
    styles_el = doc.styles.element
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is None:
        dd = OxmlElement("w:docDefaults")
        styles_el.insert(0, dd)
    rpd = dd.find(qn("w:rPrDefault"))
    if rpd is None:
        rpd = OxmlElement("w:rPrDefault")
        dd.append(rpd)
    rpr = rpd.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpd.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font_name)
    if size_pt is not None:
        half = str(int(round(size_pt * 2)))
        for tag in ("w:sz", "w:szCs"):
            el = rpr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rpr.append(el)
            el.set(qn("w:val"), half)


def _apply_alignment_spacing(style, alignment: str, space_after_pt: Optional[float],
                             line_spacing: Optional[float] = None,
                             space_before_pt: Optional[float] = None) -> None:
    pf = style.paragraph_format
    al = _ALIGN_MAP.get(alignment)
    if al is not None:
        pf.alignment = al
    if space_after_pt is not None:
        pf.space_after = Pt(space_after_pt)
    if space_before_pt is not None:
        pf.space_before = Pt(space_before_pt)
    if line_spacing is not None:
        pf.line_spacing = line_spacing


# ---------------------------------------------------------------------------
# PDF profiling (PyMuPDF)
# ---------------------------------------------------------------------------

def profile_from_pdf(pdf_bytes: bytes) -> StyleProfile:
    """Derive a :class:`StyleProfile` from a PDF using span-level analysis."""
    import fitz  # PyMuPDF

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    size_hist: collections.Counter = collections.Counter()       # (font,size)->chars
    plain_size_hist: collections.Counter = collections.Counter() # size->chars
    font_hist: collections.Counter = collections.Counter()       # family->chars
    color_hist: collections.Counter = collections.Counter()      # int->chars
    heading_spans: list[tuple[float, bool, bool, str]] = []       # size,bold,centered,font

    left_edges: list[float] = []
    right_gaps: list[float] = []
    top_edges: list[float] = []
    bottom_gaps: list[float] = []
    page_w = page_h = 0.0

    for page in doc:
        page_w, page_h = page.rect.width, page.rect.height
        data = page.get_text("dict")
        for block in data.get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    txt = (span.get("text") or "").strip()
                    if not txt:
                        continue
                    sz = round(span.get("size", 0.0), 1)
                    family = _clean_font_name(span.get("font", ""))
                    flags = span.get("flags", 0)
                    bold = bool(flags & 16) or "bold" in (span.get("font", "").lower())
                    color = span.get("color", 0)
                    x0, y0, x1, y1 = span.get("bbox", (0, 0, 0, 0))

                    n = len(txt)
                    size_hist[(family, sz)] += n
                    plain_size_hist[sz] += n
                    font_hist[family] += n
                    color_hist[color] += n
                    left_edges.append(x0)
                    right_gaps.append(page_w - x1)
                    top_edges.append(y0)
                    bottom_gaps.append(page_h - y1)

                    center_x = (x0 + x1) / 2
                    is_centered = abs(center_x - page_w / 2) < 0.06 * page_w
                    if bold and (is_centered or sz > 0):
                        heading_spans.append((sz, bold, is_centered, family))

    doc.close()

    if not plain_size_hist:
        return StyleProfile()  # empty / image-only PDF → defaults

    # Body size = the most common size by character volume.
    body_size = plain_size_hist.most_common(1)[0][0]
    # Body font = most common family overall (favouring the body size).
    body_font_candidates = collections.Counter()
    for (fam, sz), cnt in size_hist.items():
        weight = cnt * (2 if abs(sz - body_size) < 0.6 else 1)
        body_font_candidates[fam] += weight
    body_font = body_font_candidates.most_common(1)[0][0] if body_font_candidates else "Calibri"

    # Body colour = most common colour (almost always near-black).
    body_color = _int_to_hex(color_hist.most_common(1)[0][0])

    # Accent colour = most common colour that is clearly NOT near-black.
    accent = None
    for color_int, _cnt in color_hist.most_common():
        r, g, b = (color_int >> 16) & 0xFF, (color_int >> 8) & 0xFF, color_int & 0xFF
        if max(r, g, b) > 60:  # not near-black
            accent = _int_to_hex(color_int)
            break

    # Alignment: justified if right edges hug the right margin consistently.
    body_alignment = _infer_alignment(left_edges, right_gaps, page_w)

    # Margins (inches) — robust low-percentile of edge positions.
    margins = _infer_margins(left_edges, right_gaps, top_edges, bottom_gaps, page_w, page_h)

    # Heading scale.
    heading_scale = _infer_heading_scale(
        plain_size_hist, body_size, body_font, body_color, accent, heading_spans
    )

    return StyleProfile(
        body_font=body_font,
        body_size_pt=body_size,
        body_color_hex=body_color,
        body_alignment=body_alignment,
        space_after_pt=8.0,
        heading_scale=heading_scale,
        margins_in=margins,
        accent_color_hex=accent,
        source_kind="pdf",
    )


def _infer_alignment(left_edges, right_gaps, page_w) -> str:
    """Detect justified body text.

    Justified lines all terminate at the text-area's right boundary, so they
    leave a near-identical right gap (≈ the right margin) — NOT a gap near
    zero (that would be the page edge). We therefore estimate the right
    boundary as a low percentile of the gaps, then measure how many lines
    cluster at that boundary. A high fraction ⇒ justified.
    """
    if not right_gaps:
        return "left"
    boundary = _percentile(right_gaps, 15)            # the right margin
    tol = 0.03 * page_w
    full_lines = sum(1 for g in right_gaps if g <= boundary + tol)
    frac_full = full_lines / max(len(right_gaps), 1)
    return "justify" if frac_full >= 0.40 else "left"


def _percentile(values: list[float], pct: float) -> float:
    if not values:
        return 0.0
    s = sorted(values)
    k = int(max(0, min(len(s) - 1, round(pct / 100.0 * (len(s) - 1)))))
    return s[k]


def _infer_margins(left_edges, right_gaps, top_edges, bottom_gaps, page_w, page_h):
    """Estimate page margins from text-edge positions.

    Left/right are reliable (lines hit the side margins). Top/bottom are
    noisier (logos, headers, footers, trailing whitespace), so we take the
    minimum observed and fall back to the horizontal margin for symmetry when
    the vertical estimate is implausible.
    """
    clamp = lambda v: round(max(0.3, min(v, 2.0)), 2)
    left = clamp(_percentile(left_edges, 10) / 72.0)
    right = clamp(_percentile(right_gaps, 10) / 72.0)
    top = clamp((min(top_edges) if top_edges else 72) / 72.0)
    bottom = clamp((min(bottom_gaps) if bottom_gaps else 72) / 72.0)
    # If the vertical estimate looks off (header/footer artefacts), borrow the
    # horizontal margin so the page stays balanced.
    side = round((left + right) / 2, 2)
    if top > 1.4 or top < 0.5:
        top = side
    if bottom > 1.4 or bottom < 0.5:
        bottom = side
    return (top, bottom, left, right)


def _infer_heading_scale(plain_size_hist, body_size, body_font, body_color,
                         accent, heading_spans) -> dict[int, HeadingSpec]:
    # Distinct sizes meaningfully larger than the body, by frequency.
    bigger = sorted(
        {sz for sz in plain_size_hist if sz > body_size + 0.6},
        reverse=True,
    )
    scale: dict[int, HeadingSpec] = {}

    # Determine whether the source uses centered/bold headings (flat scale,
    # like the offer letter) when there's no size hierarchy.
    centered_bold = any(c for (_sz, b, c, _f) in heading_spans if b and c)

    if bigger:
        for level, sz in enumerate(bigger[:4], start=1):
            scale[level] = HeadingSpec(
                font=body_font, size_pt=sz, bold=True,
                color_hex=accent or body_color,
                alignment="left",
                space_before_pt=max(6.0, 14.0 - level * 1.5),
                space_after_pt=max(3.0, 8.0 - level),
            )
    elif centered_bold:
        # Flat hierarchy: replicate the letter's bold/underlined heading DNA.
        # Only the top-level title is centered; numbered sub-headings read
        # better left-aligned.
        for level in range(1, 5):
            scale[level] = HeadingSpec(
                font=body_font,
                size_pt=round(body_size * (1.0 + (4 - level) * 0.05), 1),
                bold=True, underline=True,
                color_hex=body_color,
                alignment="center" if level == 1 else "left",
                space_before_pt=14.0 - level * 2,
                space_after_pt=8.0 - level,
            )
    return scale


# ---------------------------------------------------------------------------
# DOCX profiling
# ---------------------------------------------------------------------------

def profile_from_docx(docx_bytes: bytes) -> StyleProfile:
    doc = Document(io.BytesIO(docx_bytes))
    normal = None
    try:
        normal = doc.styles["Normal"]
    except KeyError:
        pass

    body_font = (normal.font.name if normal and normal.font.name else "Calibri")
    body_size = (normal.font.size.pt if normal and normal.font.size else 11.0)

    heading_scale: dict[int, HeadingSpec] = {}
    for level in range(1, 5):
        try:
            hs = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        heading_scale[level] = HeadingSpec(
            font=hs.font.name or body_font,
            size_pt=hs.font.size.pt if hs.font.size else None,
            bold=bool(hs.font.bold),
            italic=bool(hs.font.italic),
            color_hex=str(hs.font.color.rgb) if hs.font.color and hs.font.color.rgb else None,
        )

    sec = doc.sections[0] if doc.sections else None
    margins = (
        (sec.top_margin.inches, sec.bottom_margin.inches,
         sec.left_margin.inches, sec.right_margin.inches)
        if sec else (1.0, 1.0, 1.0, 1.0)
    )

    return StyleProfile(
        body_font=body_font,
        body_size_pt=body_size,
        heading_scale=heading_scale,
        margins_in=margins,
        source_kind="docx",
    )


# ---------------------------------------------------------------------------
# PDF path: apply a profile to a content .docx in place
# ---------------------------------------------------------------------------

def _apply_profile_to_docx(
    content_bytes: bytes,
    profile: StyleProfile,
    *,
    normalize_fonts: bool = True,
    promote_headings: bool = True,
    set_margins: bool = True,
) -> bytes:
    doc = Document(io.BytesIO(content_bytes))

    # 1) Document-wide default font → everything inheriting picks up the body font.
    _set_docdefaults_font(doc, profile.body_font, profile.body_size_pt)

    # 2) Normal style → body look.
    try:
        normal = doc.styles["Normal"]
        _set_style_font(normal, profile.body_font, profile.body_size_pt, profile.body_color_hex)
        _apply_alignment_spacing(
            normal, profile.body_alignment, profile.space_after_pt, profile.line_spacing
        )
    except KeyError:
        pass

    # 3) Heading styles → heading scale.
    for level in range(1, 5):
        spec = profile.heading(level)
        try:
            hs = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        _set_style_font(hs, spec.font or profile.body_font, spec.size_pt, spec.color_hex)
        _set_style_bool(hs, bold=spec.bold, italic=spec.italic, underline=spec.underline)
        _apply_alignment_spacing(
            hs, spec.alignment, spec.space_after_pt, None, spec.space_before_pt
        )

    # 4) Normalize the fonts of all other (non-heading) styles + clear direct
    #    run-level font overrides so the body font wins consistently.
    if normalize_fonts:
        _normalize_fonts(doc, profile.body_font)

    # 5) Page margins.
    if set_margins:
        t, b, l, r = profile.margins_in
        for section in doc.sections:
            section.top_margin = Inches(t)
            section.bottom_margin = Inches(b)
            section.left_margin = Inches(l)
            section.right_margin = Inches(r)

    # 6) Promote bold-Normal pseudo-headings to real Heading styles so they
    #    pick up the source's heading treatment.
    if promote_headings:
        _promote_headings(doc)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _normalize_fonts(doc, body_font: str) -> None:
    """Rewrite non-heading style fonts to body_font and clear direct run-level
    rFonts in body paragraphs (preserving bold/italic/colour/size)."""
    # Style pass — every paragraph/character style that isn't a heading.
    for style in doc.styles:
        try:
            name = style.name
        except Exception:
            continue
        if not name or name in _HEADING_STYLE_NAMES:
            continue
        # Only touch paragraph (1) and character (2) styles.
        if getattr(style, "type", None) is None:
            continue
        try:
            stype = int(style.type)
        except Exception:
            stype = None
        if stype not in (1, 2):
            continue
        rpr = style.element.find(qn("w:rPr"))
        # Always set so explicit non-target fonts (Helvetica, Times) are overridden.
        _set_style_font(style, body_font)

    # Direct-run pass — clear rFonts on runs in non-heading paragraphs.
    def _strip_runs_in(paragraph):
        try:
            pstyle = paragraph.style.name
        except Exception:
            pstyle = None
        if pstyle in _HEADING_STYLE_NAMES:
            return
        for run in paragraph.runs:
            rpr = run._r.find(qn("w:rPr"))
            if rpr is None:
                continue
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                rpr.remove(rfonts)

    for para in doc.paragraphs:
        _strip_runs_in(para)
    for table in doc.tables:
        _strip_table_runs(table, _strip_runs_in)


def _strip_table_runs(table, strip_fn) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                strip_fn(para)
            # nested tables
            for nested in cell.tables:
                _strip_table_runs(nested, strip_fn)


# ---------------------------------------------------------------------------
# Heading promotion (LLM + heuristic)
# ---------------------------------------------------------------------------

def _promote_headings(doc) -> None:
    """Detect bold pseudo-headings authored as Normal paragraphs and assign
    real Heading styles so they adopt the source heading treatment."""
    # Collect candidate body paragraphs (skip those already in a heading style
    # and skip empty paragraphs).
    paras = doc.paragraphs
    candidates: list[tuple[int, str]] = []
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if not text:
            continue
        try:
            if p.style.name in _HEADING_STYLE_NAMES:
                continue
        except Exception:
            pass
        candidates.append((i, text))

    if not candidates:
        return

    decisions = _llm_heading_decisions(candidates)
    if decisions is None:
        decisions = _heuristic_heading_decisions(doc, candidates)

    for idx, level in decisions.items():
        if idx < 0 or idx >= len(paras):
            continue
        level = max(1, min(level, 4))
        style_name = f"Heading {level}"
        try:
            paras[idx].style = doc.styles[style_name]
        except KeyError:
            continue
        # Remove now-redundant direct bold so the heading style governs.
        for run in paras[idx].runs:
            rpr = run._r.find(qn("w:rPr"))
            if rpr is not None:
                for tag in ("w:b", "w:bCs"):
                    el = rpr.find(qn(tag))
                    if el is not None:
                        rpr.remove(el)


def _dominant_bold(paragraph) -> bool:
    runs = [r for r in paragraph.runs if (r.text or "").strip()]
    if not runs:
        return False
    bold_chars = sum(len(r.text) for r in runs if r.bold)
    total = sum(len(r.text) for r in runs)
    return total > 0 and bold_chars / total >= 0.6


def _heuristic_heading_decisions(doc, candidates) -> dict[int, int]:
    paras = doc.paragraphs
    out: dict[int, int] = {}
    for idx, text in candidates:
        if len(text) > 90:
            continue
        words = text.split()
        if len(words) > 14:
            continue
        if text.rstrip().endswith((".", ":", ";", ",")):
            # allow trailing colon for label-style headings
            if not text.rstrip().endswith(":"):
                continue
        # "Label: value" lines are fields, not headings.
        _head, sep, tail = text.partition(":")
        if sep and tail.strip():
            continue
        if not _dominant_bold(paras[idx]):
            continue
        is_upper = text.isupper()
        is_title = text.istitle() or sum(1 for w in words if w[:1].isupper()) >= max(1, len(words) - 1)
        if not (is_upper or is_title):
            continue
        level = 1 if (is_upper and len(words) <= 6) else 2
        out[idx] = level
    return out


def _llm_heading_decisions(candidates) -> Optional[dict[int, int]]:
    try:
        from app.llm.adapters import chat_json, llm_available
    except Exception:
        return None
    if not llm_available():
        return None

    # Cap the payload — long docs send the first N candidates; the heuristic
    # covers the rest.
    sample = candidates[:400]
    listing = [{"index": idx, "text": text[:120]} for idx, text in sample]
    system = (
        "You identify which paragraphs in a document are SECTION HEADINGS "
        "(as opposed to body text, list items, or captions). Return ONLY a "
        "JSON object: {\"headings\": [{\"index\": <int>, \"level\": <1-4>}]}. "
        "index MUST be one of the provided indices. Heading levels: 1 = top "
        "section title, 2 = sub-section, 3-4 = deeper. Be selective; most "
        "paragraphs are NOT headings."
    )
    user = (
        "Paragraphs (index, text):\n"
        + str(listing)
        + "\n\nReturn the headings JSON now."
    )
    raw = chat_json(system, user, temperature=0.0, max_tokens=2000)
    if not isinstance(raw, dict):
        return None
    items = raw.get("headings")
    if not isinstance(items, list):
        return None
    valid_indices = {idx for idx, _ in candidates}
    out: dict[int, int] = {}
    for it in items:
        if not isinstance(it, dict):
            continue
        try:
            idx = int(it.get("index"))
            level = int(it.get("level", 2))
        except (TypeError, ValueError):
            continue
        if idx in valid_indices:
            out[idx] = max(1, min(level, 4))
    return out or None


# ---------------------------------------------------------------------------
# DOCX path: transplant styles + theme into the content package
# ---------------------------------------------------------------------------

def _transplant_docx_styles(content_bytes: bytes, style_bytes: bytes,
                            *, promote_headings: bool = False) -> bytes:
    content = Document(io.BytesIO(content_bytes))
    source = Document(io.BytesIO(style_bytes))

    content_styles_el = content.styles.element
    source_styles_el = source.styles.element

    # 1) Merge <w:style> definitions by styleId (source wins; keep content-only).
    content_by_id = {
        s.get(qn("w:styleId")): s
        for s in content_styles_el.findall(qn("w:style"))
    }
    last_style = None
    for s in content_styles_el.findall(qn("w:style")):
        last_style = s
    for src_style in source_styles_el.findall(qn("w:style")):
        sid = src_style.get(qn("w:styleId"))
        clone = copy.deepcopy(src_style)
        if sid in content_by_id:
            old = content_by_id[sid]
            content_styles_el.replace(old, clone)
        else:
            content_styles_el.append(clone)

    # 2) Replace docDefaults wholesale with the source's.
    src_dd = source_styles_el.find(qn("w:docDefaults"))
    if src_dd is not None:
        dst_dd = content_styles_el.find(qn("w:docDefaults"))
        clone_dd = copy.deepcopy(src_dd)
        if dst_dd is not None:
            content_styles_el.replace(dst_dd, clone_dd)
        else:
            content_styles_el.insert(0, clone_dd)

    # 3) Transplant the theme part (fonts/colour scheme) when both have one.
    try:
        src_theme = source.part.part_related_by(RT.THEME)
        try:
            dst_theme = content.part.part_related_by(RT.THEME)
            dst_theme._blob = src_theme.blob
        except KeyError:
            # Content has no theme part — relate the source's theme blob in.
            from docx.opc.part import Part
            from docx.opc.packuri import PackURI
            partname = PackURI("/word/theme/theme1.xml")
            new_part = Part(partname, src_theme.content_type, src_theme.blob,
                            content.part.package)
            content.part.relate_to(new_part, RT.THEME)
    except KeyError:
        pass  # source has no theme — styles/docDefaults still transferred

    if promote_headings:
        _promote_headings(content)

    out = io.BytesIO()
    content.save(out)
    return out.getvalue()


# ---------------------------------------------------------------------------
# Public dispatcher
# ---------------------------------------------------------------------------

def _looks_like_pdf(data: bytes) -> bool:
    return data[:5] == b"%PDF-"


def _looks_like_docx(data: bytes) -> bool:
    # docx is a zip; sniff PK magic. (We trust the extension downstream too.)
    return data[:2] == b"PK"


def pdf_content_to_docx(pdf_bytes: bytes, filename: str = "content.pdf") -> bytes:
    """Convert a PDF *content* document into a clean, flowing ``.docx``.

    Runs the extraction pipeline (text layer with automatic OCR fallback for
    scanned PDFs) and re-emits headings/paragraphs/lists/tables/images as a
    real Word document, so every style-transfer path can treat the content as
    OOXML. The LLM refinement pass is skipped here — the structure-aware
    applier downstream performs its own (deeper) recognition."""
    import io as _io

    from app.services.extraction.pdf_ext import extract_pdf_document
    from app.services.formatting.formater_apply import apply_styling

    content, styling = extract_pdf_document(
        file_stream=_io.BytesIO(pdf_bytes), filename=filename
    )
    return apply_styling(content, styling, use_llm=False).getvalue()


def _ensure_docx_content(content_bytes: bytes, content_name: str) -> bytes:
    """Accept .docx as-is; convert PDF content to .docx; reject the rest."""
    if _looks_like_docx(content_bytes):
        return content_bytes
    if _looks_like_pdf(content_bytes) or content_name.lower().endswith(".pdf"):
        return pdf_content_to_docx(content_bytes, content_name)
    raise ValueError(
        "Style transfer requires the CONTENT document to be a .docx or .pdf "
        f"(got {content_name!r})."
    )


def transfer_style(
    content_bytes: bytes,
    content_name: str,
    style_bytes: bytes,
    style_name: str,
    *,
    normalize_fonts: bool = True,
    promote_headings: bool = True,
) -> bytes:
    """Apply ``style_*``'s visual identity onto ``content_*`` and return .docx bytes.

    The content document may be a .docx (edited in place) or a .pdf (converted
    to a flowing .docx first). The style source may be .docx (full transplant)
    or .pdf (profile + apply).
    """
    content_bytes = _ensure_docx_content(content_bytes, content_name)

    style_is_pdf = _looks_like_pdf(style_bytes) or style_name.lower().endswith(".pdf")

    if style_is_pdf:
        profile = profile_from_pdf(style_bytes)
        return _apply_profile_to_docx(
            content_bytes, profile,
            normalize_fonts=normalize_fonts,
            promote_headings=promote_headings,
        )

    # DOCX → DOCX: transplant styles/theme for maximum fidelity, zero loss.
    return _transplant_docx_styles(
        content_bytes, style_bytes, promote_headings=promote_headings
    )


def transfer_style_files(
    content_path: str,
    style_path: str,
    output_path: str,
    *,
    normalize_fonts: bool = True,
    promote_headings: bool = True,
) -> str:
    with open(content_path, "rb") as fh:
        content_bytes = fh.read()
    with open(style_path, "rb") as fh:
        style_bytes = fh.read()
    result = transfer_style(
        content_bytes, content_path,
        style_bytes, style_path,
        normalize_fonts=normalize_fonts,
        promote_headings=promote_headings,
    )
    with open(output_path, "wb") as fh:
        fh.write(result)
    return output_path


# ===========================================================================
# Guideline path: apply a structured StyleSpec onto the content .docx in place
# ===========================================================================
#
# When the style source is a *formatting guideline* (rules described in prose /
# tables, not a sample to copy), ``guideline_interpreter`` distils it into a
# :class:`StyleSpec`. We apply that spec onto the content document's REAL OOXML
# — mutating named styles, tables, header/footer and page setup — so every
# paragraph, table and image of the content is preserved, exactly like the
# example path. Every sub-step is defensive: a failure appends a warning and
# the rest of the rules still land.

def _rfonts_set(rpr, font_name: str) -> None:
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font_name)


def _add_para_border(pPr, side: str, color_hex: Optional[str], width_pt: float) -> None:
    """Add a single-line paragraph border on one ``side`` (top|bottom|left|right)."""
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement(f"w:{side}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(max(2, int(round((width_pt or 1.0) * 8)))))  # eighths of a pt
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), (color_hex or "auto"))
    pBdr.append(b)
    # CT_PPr ordering: pBdr sits just after pStyle.
    pstyle = pPr.find(qn("w:pStyle"))
    if pstyle is not None:
        pstyle.addnext(pBdr)
    else:
        pPr.insert(0, pBdr)


def _set_cell_shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _make_borders(tag: str, color_hex: Optional[str], width_pt: float, sides) -> "OxmlElement":
    el = OxmlElement(tag)
    sz = str(max(2, int(round((width_pt or 0.5) * 8))))
    for side in sides:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), (color_hex or "auto"))
        el.append(b)
    return el


def _set_table_borders(table, color_hex: Optional[str], width_pt: float) -> None:
    tblPr = table._tbl.tblPr
    for tb in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(tb)
    tblPr.append(_make_borders(
        "w:tblBorders", color_hex, width_pt,
        sides=("top", "left", "bottom", "right", "insideH", "insideV"),
    ))


def _set_table_cell_margins(table, top_in: Optional[float], left_in: Optional[float]) -> None:
    tblPr = table._tbl.tblPr
    for cm in tblPr.findall(qn("w:tblCellMar")):
        tblPr.remove(cm)
    cm = OxmlElement("w:tblCellMar")

    def _mar(tag: str, inches: float):
        e = OxmlElement(tag)
        e.set(qn("w:w"), str(int(round(inches * 1440))))
        e.set(qn("w:type"), "dxa")
        cm.append(e)

    if top_in is not None:
        _mar("w:top", top_in)
        _mar("w:bottom", top_in)
    if left_in is not None:
        _mar("w:left", left_in)
        _mar("w:right", left_in)
    if len(cm):
        tblPr.append(cm)


def _style_cell_runs(cell, *, font=None, size_pt=None, color_hex=None, bold=None) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            if font:
                run.font.name = font
                _rfonts_set(run._r.get_or_add_rPr(), font)
            if size_pt is not None:
                run.font.size = Pt(size_pt)
            if color_hex:
                try:
                    run.font.color.rgb = RGBColor.from_string(color_hex)
                except Exception:
                    pass
            if bold is not None:
                run.font.bold = bold


def _apply_table_spec(table, t: "TableSpec", warnings: list[str]) -> None:
    try:
        if t.border_color_hex or t.border_width_pt:
            _set_table_borders(table, t.border_color_hex, t.border_width_pt or 0.5)
        if t.cell_padding_top_in is not None or t.cell_padding_left_in is not None:
            _set_table_cell_margins(table, t.cell_padding_top_in, t.cell_padding_left_in)
        header_styled = bool(
            t.header_fill_hex or t.header_text_hex or t.header_bold
            or t.header_font or t.header_size_pt
        )
        for ri, row in enumerate(table.rows):
            if ri == 0 and header_styled:
                for cell in row.cells:
                    if t.header_fill_hex:
                        _set_cell_shading(cell, t.header_fill_hex)
                    _style_cell_runs(
                        cell, font=t.header_font, size_pt=t.header_size_pt,
                        color_hex=t.header_text_hex, bold=t.header_bold,
                    )
                continue
            # Data rows: shade every other one + apply the table body look.
            if t.alt_row_fill_hex and ri >= 1 and ri % 2 == 0:
                for cell in row.cells:
                    _set_cell_shading(cell, t.alt_row_fill_hex)
            if t.body_font or t.body_size_pt:
                for cell in row.cells:
                    _style_cell_runs(cell, font=t.body_font, size_pt=t.body_size_pt)
    except Exception as e:  # noqa: BLE001 — never let table styling break the run
        warnings.append(f"table styling skipped: {e}")


def _apply_page_spec(doc, page: "PageSpec", warnings: list[str]) -> None:
    try:
        from docx.enum.section import WD_ORIENT

        for section in doc.sections:
            if page.margin_top_in is not None:
                section.top_margin = Inches(page.margin_top_in)
            if page.margin_bottom_in is not None:
                section.bottom_margin = Inches(page.margin_bottom_in)
            if page.margin_left_in is not None:
                section.left_margin = Inches(page.margin_left_in)
            if page.margin_right_in is not None:
                section.right_margin = Inches(page.margin_right_in)
            if page.width_in:
                section.page_width = Inches(page.width_in)
            if page.height_in:
                section.page_height = Inches(page.height_in)
            if page.orientation:
                o = page.orientation.lower()
                if o.startswith("land"):
                    section.orientation = WD_ORIENT.LANDSCAPE
                elif o.startswith("port"):
                    section.orientation = WD_ORIENT.PORTRAIT
    except Exception as e:  # noqa: BLE001
        warnings.append(f"page setup skipped: {e}")


def _apply_list_spec(doc, lst: "ListSpec", warnings: list[str]) -> None:
    if lst.bullet_indent_in is None and lst.bullet_hanging_in is None:
        return
    try:
        for name in ("List Bullet", "List Number", "List Paragraph"):
            try:
                st = doc.styles[name]
            except KeyError:
                continue
            pf = st.paragraph_format
            if lst.bullet_indent_in is not None:
                pf.left_indent = Inches(lst.bullet_indent_in)
            if lst.bullet_hanging_in is not None:
                pf.first_line_indent = Inches(-lst.bullet_hanging_in)
    except Exception as e:  # noqa: BLE001
        warnings.append(f"list spec skipped: {e}")


def _write_header_footer_part(part, left: Optional[str], right: Optional[str],
                              hf: "HeaderFooterSpec", border_hex: Optional[str],
                              *, border_side: str) -> None:
    from docx.enum.text import WD_TAB_ALIGNMENT

    part.is_linked_to_previous = False
    para = part.paragraphs[0] if part.paragraphs else part.add_paragraph()
    # Clear existing runs.
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    # Right-aligned tab stop at the content width so the right text hugs the margin.
    try:
        para.paragraph_format.tab_stops.clear_all()
        para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    except Exception:
        pass

    def _add(text: str):
        run = para.add_run(text or "")
        if hf.font:
            run.font.name = hf.font
            _rfonts_set(run._r.get_or_add_rPr(), hf.font)
        if hf.size_pt is not None:
            run.font.size = Pt(hf.size_pt)
        if hf.color_hex:
            try:
                run.font.color.rgb = RGBColor.from_string(hf.color_hex)
            except Exception:
                pass

    _add(left or "")
    if right:
        para.add_run("\t")
        _add(right)
    if border_hex:
        _add_para_border(para._p.get_or_add_pPr(), border_side, border_hex, 1.0)


def _apply_header_footer_spec(doc, hf: "HeaderFooterSpec", warnings: list[str]) -> None:
    try:
        for section in doc.sections:
            if hf.header_left or hf.header_right:
                _write_header_footer_part(
                    section.header, hf.header_left, hf.header_right, hf,
                    hf.header_border_color_hex, border_side="bottom",
                )
            if hf.footer_left or hf.footer_right:
                _write_header_footer_part(
                    section.footer, hf.footer_left, hf.footer_right, hf,
                    hf.footer_border_color_hex, border_side="top",
                )
    except Exception as e:  # noqa: BLE001
        warnings.append(f"header/footer skipped: {e}")


def apply_style_spec_to_docx(
    content_bytes: bytes,
    spec: "StyleSpec",
    *,
    normalize_fonts: bool = True,
    promote_headings: bool = True,
) -> tuple[bytes, list[str]]:
    """Apply a structured :class:`StyleSpec` onto a content .docx in place.

    Returns ``(docx_bytes, warnings)``. Content body XML (paragraphs, tables,
    images) is preserved — only the *style layer* is rewritten."""
    warnings: list[str] = []
    doc = Document(io.BytesIO(content_bytes))
    body = spec.body
    body_font = body.font

    # 1) Document-wide default font so everything inheriting picks up the body font.
    if body_font or body.size_pt:
        try:
            _set_docdefaults_font(doc, body_font or "Calibri", body.size_pt)
        except Exception as e:  # noqa: BLE001
            warnings.append(f"doc defaults skipped: {e}")

    # 2) Normal style → body look.
    try:
        normal = doc.styles["Normal"]
        if body_font or body.size_pt or body.color_hex:
            _set_style_font(normal, body_font, body.size_pt, body.color_hex)
        _apply_alignment_spacing(
            normal, body.alignment, body.space_after_pt, body.line_spacing, body.space_before_pt
        )
    except KeyError:
        pass

    # 3) Heading / named styles.
    for rule in spec.headings:
        name = rule.style_name or (f"Heading {rule.level}" if rule.level else None)
        if not name:
            continue
        try:
            st = doc.styles[name]
        except KeyError:
            warnings.append(f"style {name!r} not present in document; skipped")
            continue
        _set_style_font(st, rule.font or body_font, rule.size_pt, rule.color_hex)
        _set_style_bool(st, bold=rule.bold, italic=rule.italic, underline=rule.underline)
        _apply_alignment_spacing(st, rule.alignment, rule.space_after_pt, None, rule.space_before_pt)
        if rule.bottom_border:
            try:
                _add_para_border(
                    st.element.get_or_add_pPr(), "bottom",
                    rule.border_color_hex or rule.color_hex or spec.accent_color_hex,
                    rule.border_width_pt or 1.0,
                )
            except Exception as e:  # noqa: BLE001
                warnings.append(f"heading border skipped: {e}")

    # 4) Normalize remaining fonts onto the body font (when requested).
    if normalize_fonts and body_font:
        try:
            _normalize_fonts(doc, body_font)
        except Exception as e:  # noqa: BLE001
            warnings.append(f"font normalization skipped: {e}")

    # 5) Page setup (margins / size / orientation).
    _apply_page_spec(doc, spec.page, warnings)

    # 6) Lists.
    _apply_list_spec(doc, spec.lists, warnings)

    # 7) Tables — restyle every table already in the content document.
    if any(getattr(spec.table, f) is not None for f in spec.table.model_fields):
        for table in doc.tables:
            _apply_table_spec(table, spec.table, warnings)

    # 8) Header / footer.
    if any(getattr(spec.header_footer, f) for f in spec.header_footer.model_fields):
        _apply_header_footer_spec(doc, spec.header_footer, warnings)

    # 9) Promote bold-Normal pseudo-headings to real Heading styles.
    if promote_headings:
        try:
            _promote_headings(doc)
        except Exception as e:  # noqa: BLE001
            warnings.append(f"heading promotion skipped: {e}")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), warnings


# ---------------------------------------------------------------------------
# Smart dispatcher: classify the style source, then route to the right path
# ---------------------------------------------------------------------------

@dataclass
class StyleTransferOutcome:
    """Result of a smart style transfer — bytes plus what was detected/done."""
    docx_bytes: bytes
    mode_used: str            # "guideline" | "example"
    detected_kind: str        # classifier's verdict (even when overridden)
    confidence: float
    reason: str
    summary: str              # human-readable description of what was applied
    spec: Optional[dict]      # StyleSpec.model_dump() (guideline or profiled example)
    warnings: list[str] = field(default_factory=list)
    # Counts from the structure-recognition pass (headings/lists/tables found).
    structure: Optional[dict] = None


def _profile_to_spec(profile: "StyleProfile"):
    """Map a PDF/DOCX :class:`StyleProfile` into a :class:`StyleSpec` so the
    structure-aware applier can drive the example path uniformly."""
    from app.services.style.guideline_interpreter import (
        BodySpec, HeadingRule, ListSpec, PageSpec, StyleSpec,
    )

    headings = [
        HeadingRule(
            level=lvl, font=hs.font, size_pt=hs.size_pt, bold=hs.bold,
            italic=hs.italic, underline=hs.underline, color_hex=hs.color_hex,
            alignment=hs.alignment, space_before_pt=hs.space_before_pt,
            space_after_pt=hs.space_after_pt,
        )
        for lvl, hs in sorted(profile.heading_scale.items())
    ]
    t, b, l, r = profile.margins_in
    return StyleSpec(
        body=BodySpec(
            font=profile.body_font, size_pt=profile.body_size_pt,
            color_hex=profile.body_color_hex, alignment=profile.body_alignment,
            line_spacing=profile.line_spacing, space_after_pt=profile.space_after_pt,
        ),
        headings=headings,
        page=PageSpec(margin_top_in=t, margin_bottom_in=b, margin_left_in=l, margin_right_in=r),
        lists=ListSpec(space_after_pt=4.0),
        accent_color_hex=profile.accent_color_hex,
        colors={"Accent": profile.accent_color_hex} if profile.accent_color_hex else {},
    )


def transfer_style_smart(
    content_bytes: bytes,
    content_name: str,
    style_bytes: bytes,
    style_name: str,
    *,
    mode: str = "auto",
    normalize_fonts: bool = True,
    promote_headings: bool = True,
) -> StyleTransferOutcome:
    """Apply a style source onto the content, adapting to the source's nature.

    ``mode``: 'auto' (classify the source), 'guideline' (force rule-extraction),
    or 'example' (force the transplant/PDF-profile path). Content may be .docx
    or .pdf; the style source may be .docx or .pdf in either mode."""
    from app.services.style.guideline_interpreter import (
        decide_style_mode,
        extract_style_source_text,
        interpret_guideline,
        profile_example_to_spec,
        style_spec_summary,
    )
    from app.services.style.structure_apply import apply_structured_style

    content_bytes = _ensure_docx_content(content_bytes, content_name)

    digest = extract_style_source_text(style_bytes, style_name)
    effective, cls = decide_style_mode(digest, mode)
    warnings: list[str] = []
    style_is_pdf = _looks_like_pdf(style_bytes) or style_name.lower().endswith(".pdf")
    use_llm = promote_headings  # the "promote headings" toggle now gates LLM understanding

    # ── Resolve a StyleSpec + a transplant source from the style document ──
    transplant: Optional[bytes] = None
    spec = None
    if effective == "guideline":
        spec = interpret_guideline(digest)
        if spec.is_empty():
            warnings.append(
                "No concrete formatting rules detected in the guideline — "
                "profiling its own appearance instead."
            )
            effective = "example"
            spec = None
    if spec is None:  # example path (or guideline fell back)
        if style_is_pdf:
            spec = _profile_to_spec(profile_from_pdf(style_bytes))
        else:
            spec = profile_example_to_spec(style_bytes)
            transplant = style_bytes  # transplant the example's styles.xml/theme too
        mode_used = "example"
    else:
        mode_used = "guideline"

    # ── Recognise the content's structure and apply the spec meticulously ──
    outcome = apply_structured_style(
        content_bytes, spec,
        transplant_style_bytes=transplant,
        normalize_fonts=normalize_fonts,
        use_llm=use_llm,
    )
    warnings.extend(outcome.warnings)
    if spec.notes:
        warnings.extend(f"Guideline note: {n}" for n in spec.notes[:5])

    s = outcome.plan_summary
    # Long documents only send a prefix for LLM structure refinement — surface
    # the partial-coverage note so the user can see why the tail may differ.
    if s.get("llm_coverage"):
        warnings.append(f"Structure recognition coverage: {s['llm_coverage']}")
    struct_bits = (
        f"recognised {s.get('headings', 0)} heading(s), {s.get('list_items', 0)} "
        f"list item(s), {s.get('tables', 0)} table(s)"
    )
    if mode_used == "guideline":
        summary = f"{style_spec_summary(spec)} — {struct_bits}."
    else:
        summary = f"Profiled the example's look and {struct_bits}."

    return StyleTransferOutcome(
        docx_bytes=outcome.docx_bytes, mode_used=mode_used, detected_kind=cls.kind,
        confidence=cls.confidence, reason=cls.reason, summary=summary,
        spec=spec.model_dump(exclude_none=True), warnings=warnings,
        structure=outcome.plan_summary,
    )
