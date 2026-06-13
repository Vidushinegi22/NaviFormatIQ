"""
Style-source interpreter — turn ANY style template into actionable rules.
=========================================================================

A "Style Template" uploaded for the style-update flow can be one of two very
different things:

  1. **Example / sample document** — a real document that already *looks* the
     way the user wants. The style lives in the document's own OOXML (or, for
     a PDF, in its rendered appearance). → handled by the existing
     ``style_engine.transfer_style`` (transplant / PDF-profile).

  2. **Formatting guideline / specification** — a document that *describes*
     the formatting rules in prose and tables (fonts, sizes, colours,
     margins, heading treatments, table shading, header/footer text, …). Its
     own appearance is irrelevant; the *content* is the spec. A naive
     transplant of such a file copies the guideline's plain look, not the
     rules it dictates — the exact opposite of what the user wants.

This module makes the system *adaptive*: it reads the style source, decides
which kind it is (LLM classifier with a deterministic heuristic fallback),
and — for guidelines — distils the prose/tables into a structured
:class:`StyleSpec` that :func:`style_engine.apply_style_spec_to_docx` can
apply onto the target document's real OOXML (preserving 100% of its content).

Public API
----------
    extract_style_source_text(data, name)      -> StyleSourceDigest
    classify_style_source(digest, ...)         -> Classification
    interpret_guideline(digest, ...)           -> StyleSpec
    decide_style_mode(digest, requested, ...)  -> (effective_kind, Classification)
    style_spec_summary(spec)                   -> str
"""
from __future__ import annotations

import hashlib
import io
import re
import threading
from dataclasses import dataclass, field
from typing import Any, Optional

from docx.oxml.ns import qn
from pydantic import BaseModel, Field


# ---------------------------------------------------------------------------
# StyleSpec — the structured, applier-ready description of a formatting guide
# ---------------------------------------------------------------------------

class PageSpec(BaseModel):
    width_in: Optional[float] = None
    height_in: Optional[float] = None
    orientation: Optional[str] = None  # "portrait" | "landscape"
    margin_top_in: Optional[float] = None
    margin_bottom_in: Optional[float] = None
    margin_left_in: Optional[float] = None
    margin_right_in: Optional[float] = None


class BodySpec(BaseModel):
    font: Optional[str] = None
    size_pt: Optional[float] = None
    color_hex: Optional[str] = None
    alignment: Optional[str] = None        # left|center|right|justify
    line_spacing: Optional[float] = None   # multiple, e.g. 1.15
    space_before_pt: Optional[float] = None
    space_after_pt: Optional[float] = None


class HeadingRule(BaseModel):
    """One heading-style rule. Target the style by numeric ``level`` (→
    "Heading {level}") or by explicit ``style_name`` ("Title", "Subtitle")."""
    level: Optional[int] = None
    style_name: Optional[str] = None
    font: Optional[str] = None
    size_pt: Optional[float] = None
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    color_hex: Optional[str] = None
    alignment: Optional[str] = None
    space_before_pt: Optional[float] = None
    space_after_pt: Optional[float] = None
    bottom_border: Optional[bool] = None
    border_color_hex: Optional[str] = None
    border_width_pt: Optional[float] = None


class TableSpec(BaseModel):
    header_fill_hex: Optional[str] = None
    header_text_hex: Optional[str] = None
    header_bold: Optional[bool] = None
    header_font: Optional[str] = None
    header_size_pt: Optional[float] = None
    alt_row_fill_hex: Optional[str] = None
    border_color_hex: Optional[str] = None
    border_width_pt: Optional[float] = None
    cell_padding_top_in: Optional[float] = None
    cell_padding_left_in: Optional[float] = None
    body_font: Optional[str] = None
    body_size_pt: Optional[float] = None


class ListSpec(BaseModel):
    bullet_char: Optional[str] = None
    bullet_indent_in: Optional[float] = None
    bullet_hanging_in: Optional[float] = None
    number_format: Optional[str] = None
    space_before_pt: Optional[float] = None
    space_after_pt: Optional[float] = None


class HeaderFooterSpec(BaseModel):
    header_left: Optional[str] = None
    header_right: Optional[str] = None
    footer_left: Optional[str] = None
    footer_right: Optional[str] = None
    font: Optional[str] = None
    size_pt: Optional[float] = None
    color_hex: Optional[str] = None
    header_border_color_hex: Optional[str] = None
    footer_border_color_hex: Optional[str] = None


class TitleBlockSpec(BaseModel):
    """Look of one masthead element (document title / subtitle / metadata line)."""
    font: Optional[str] = None
    size_pt: Optional[float] = None
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    color_hex: Optional[str] = None
    alignment: Optional[str] = None        # left|center|right|justify
    space_before_pt: Optional[float] = None
    space_after_pt: Optional[float] = None

    def has_values(self) -> bool:
        return any(getattr(self, f) is not None for f in self.model_fields)


class StyleSpec(BaseModel):
    """A normalized, applier-ready description of a formatting guideline.

    Every field is optional — the applier only touches what was specified, so
    a partial spec degrades gracefully (e.g. fonts-only guides still work)."""
    page: PageSpec = Field(default_factory=PageSpec)
    body: BodySpec = Field(default_factory=BodySpec)
    headings: list[HeadingRule] = Field(default_factory=list)
    table: TableSpec = Field(default_factory=TableSpec)
    lists: ListSpec = Field(default_factory=ListSpec)
    header_footer: HeaderFooterSpec = Field(default_factory=HeaderFooterSpec)
    # Masthead looks — applied as direct formatting to the recognised title /
    # subtitle / metadata paragraphs of the content document.
    title: TitleBlockSpec = Field(default_factory=TitleBlockSpec)
    subtitle: TitleBlockSpec = Field(default_factory=TitleBlockSpec)
    metadata: TitleBlockSpec = Field(default_factory=TitleBlockSpec)
    colors: dict[str, str] = Field(default_factory=dict)
    accent_color_hex: Optional[str] = None
    # Anything the spec mentioned that has no structured home (e.g. required
    # section lists) — surfaced to the user, not applied automatically.
    notes: list[str] = Field(default_factory=list)

    def is_empty(self) -> bool:
        """True if nothing meaningful was extracted (applier would be a no-op)."""
        return not (
            self.headings
            or self.body.font or self.body.size_pt
            or any(getattr(self.page, f) is not None for f in self.page.model_fields)
            or self.colors or self.accent_color_hex
            or self.title.has_values() or self.subtitle.has_values()
            or any(getattr(self.table, f) is not None for f in self.table.model_fields)
        )


@dataclass
class Classification:
    kind: str          # "guideline" | "example"
    confidence: float  # 0..1
    reason: str
    method: str        # "llm" | "heuristic"


@dataclass
class StyleSourceDigest:
    """Plain-text + tabular view of a style source, ready for the LLM/heuristics."""
    text: str
    tables: list[list[list[str]]] = field(default_factory=list)
    source_kind: str = "docx"   # "docx" | "pdf"
    n_paragraphs: int = 0
    # sha256 of the source bytes (set by extract_style_source_text) — lets the
    # classification/interpretation caches key off the same template content.
    source_sha256: str = ""

    def as_prompt(self, *, max_chars: int = 7000, max_table_rows: int = 60) -> str:
        """Render a compact, LLM-friendly view of the document."""
        parts: list[str] = []
        if self.text.strip():
            parts.append("PROSE:\n" + self.text[:max_chars])
        rows_left = max_table_rows
        for ti, tbl in enumerate(self.tables):
            if rows_left <= 0:
                break
            lines = [f"\nTABLE {ti + 1}:"]
            for row in tbl[:rows_left]:
                lines.append(" | ".join((c or "").strip() for c in row))
            rows_left -= len(tbl)
            parts.append("\n".join(lines))
        return "\n\n".join(parts).strip() or "(empty document)"


# ---------------------------------------------------------------------------
# Memoization — users iterate against the same template, so the expensive
# per-document work (parsing, classification, LLM interpretation) is cached by
# sha256 of the style-source bytes. Bounded to the last few templates seen;
# eviction is insertion-ordered (plain dict + pop-oldest).
# ---------------------------------------------------------------------------

_CACHE_MAX_ENTRIES = 8
_CACHE_LOCK = threading.Lock()
_DIGEST_CACHE: dict[tuple[str, bool], "StyleSourceDigest"] = {}
_CLASSIFY_CACHE: dict[tuple[str, str, bool], "Classification"] = {}
_INTERPRET_CACHE: dict[tuple[str, str, bool], "StyleSpec"] = {}


def _cache_get(cache: dict, key):
    with _CACHE_LOCK:
        return cache.get(key)


def _cache_put(cache: dict, key, value) -> None:
    with _CACHE_LOCK:
        cache[key] = value
        while len(cache) > _CACHE_MAX_ENTRIES:
            cache.pop(next(iter(cache)))


# ---------------------------------------------------------------------------
# Extraction — pull text + tables out of the style source (docx / pdf)
# ---------------------------------------------------------------------------

def _looks_like_pdf(data: bytes) -> bool:
    return data[:5] == b"%PDF-"


def extract_style_source_text(data: bytes, name: str = "") -> StyleSourceDigest:
    """Extract prose + tables from a style source for classification/interpretation.

    Memoized by content hash — re-running the flow with the same template (the
    common iterate-and-retry case) skips the docx/pdf re-parse entirely."""
    is_pdf = _looks_like_pdf(data) or name.lower().endswith(".pdf")
    sha = hashlib.sha256(data).hexdigest()
    cached = _cache_get(_DIGEST_CACHE, (sha, is_pdf))
    if cached is not None:
        return cached
    digest = _digest_pdf(data) if is_pdf else _digest_docx(data)
    digest.source_sha256 = sha
    _cache_put(_DIGEST_CACHE, (sha, is_pdf), digest)
    return digest


def _digest_docx(data: bytes) -> StyleSourceDigest:
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        try:
            sty = p.style.name or ""
        except Exception:
            sty = ""
        # Prefix structural styles so the LLM sees the document's own hierarchy.
        prefix = f"[{sty}] " if sty and sty not in ("Normal", "Body Text") else ""
        lines.append(prefix + t)

    tables: list[list[list[str]]] = []
    for tbl in doc.tables:
        rows: list[list[str]] = []
        for r in tbl.rows:
            rows.append([(c.text or "").strip().replace("\n", " ") for c in r.cells])
        if rows:
            tables.append(rows)

    return StyleSourceDigest(
        text="\n".join(lines),
        tables=tables,
        source_kind="docx",
        n_paragraphs=len(lines),
    )


def _digest_pdf(data: bytes) -> StyleSourceDigest:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=data, filetype="pdf")
    text_parts: list[str] = []
    tables: list[list[list[str]]] = []
    for page in doc:
        try:
            text_parts.append(page.get_text("text") or "")
        except Exception:
            pass
        # find_tables() exists in PyMuPDF >= 1.23; tolerate older builds.
        try:
            found = page.find_tables()
            for tb in getattr(found, "tables", []) or []:
                rows = tb.extract()
                cleaned = [
                    [("" if c is None else str(c)).strip().replace("\n", " ") for c in row]
                    for row in rows
                ]
                if cleaned:
                    tables.append(cleaned)
        except Exception:
            pass
    doc.close()
    text = "\n".join(tp for tp in text_parts if tp.strip())
    return StyleSourceDigest(
        text=text,
        tables=tables,
        source_kind="pdf",
        n_paragraphs=text.count("\n") + 1,
    )


# ---------------------------------------------------------------------------
# Lightweight value coercion (LLM/heuristic outputs are messy)
# ---------------------------------------------------------------------------

_HEX_RE = re.compile(r"#?([0-9A-Fa-f]{6})\b")
_PT_RE = re.compile(r"(\d+(?:\.\d+)?)\s*(?:pt|point)", re.I)
_FONT_FAMILIES = (
    "Arial", "Calibri", "Times New Roman", "Times", "Helvetica", "Georgia",
    "Verdana", "Garamond", "Cambria", "Tahoma", "Century", "Book Antiqua",
    "Courier New", "Trebuchet", "Segoe UI", "Roboto", "Lato",
)


def _coerce_hex(v: Any) -> Optional[str]:
    if v is None:
        return None
    m = _HEX_RE.search(str(v))
    return m.group(1).upper() if m else None


def _coerce_num(v: Any) -> Optional[float]:
    if v is None:
        return None
    if isinstance(v, (int, float)):
        return float(v)
    m = re.search(r"-?\d+(?:\.\d+)?", str(v))
    return float(m.group()) if m else None


def _coerce_bool(v: Any) -> Optional[bool]:
    if isinstance(v, bool):
        return v
    if v is None:
        return None
    s = str(v).strip().lower()
    if s in ("true", "yes", "1", "bold", "on"):
        return True
    if s in ("false", "no", "0", "off", "none", "regular", "normal"):
        return False
    return None


def _clean_bullet_char(v: Any) -> Optional[str]:
    """A bullet glyph clean enough to store/show. Strips control/whitespace,
    maps Symbol-font bullets to '•', rejects alphanumerics (mis-extractions)."""
    if not v:
        return None
    s = "".join(c for c in str(v) if ord(c) >= 0x20 and not c.isspace())
    if not s:
        return None
    if s in ("\xb7", "·"):
        return "•"
    if len(s) > 2 or any(c.isalnum() for c in s):
        return None
    return s


def _coerce_align(v: Any) -> Optional[str]:
    if not v:
        return None
    s = str(v).strip().lower()
    for a in ("justify", "justified", "center", "centre", "right", "left"):
        if a in s:
            return {"justified": "justify", "centre": "center"}.get(a, a)
    return None


# ---------------------------------------------------------------------------
# Classification — guideline vs example
# ---------------------------------------------------------------------------

_GUIDELINE_KEYWORDS = (
    "formatting", "format ", "template", "style guide", "styling", "typography",
    "specification", "must ", "shall ", "convention", "standard operating",
    "margin", "heading 1", "heading 2", "font", "point size", "pt)", "hex",
    "colour", "color palette", "bold", "italic", "indent", "spacing", "border",
    "alignment", "twips", "dxa", "header and footer", "required", "section",
)
_TABLE_HEADER_SIGNALS = (
    "element", "font", "size", "style", "color", "colour", "property",
    "specification", "space before", "space after", "hex", "usage", "applied to",
    "weight", "pt",
)


def _heuristic_classify(digest: StyleSourceDigest) -> Classification:
    """Deterministic guideline detector used when the LLM is unavailable.

    Guidelines are dense with formatting *signals* (hex codes, point sizes,
    font names, rule keywords) and frequently carry "Element | Font | Size"
    style tables. Ordinary content documents almost never mention "11 pt" or
    "#2E5496". We score those signals and threshold conservatively.
    """
    text = digest.text or ""
    table_blob = " ".join(
        " ".join(cell for cell in row) for tbl in digest.tables for row in tbl
    )
    haystack = (text + " " + table_blob)
    low = haystack.lower()

    n_hex = len(set(_HEX_RE.findall(haystack)))
    n_pt = len(_PT_RE.findall(haystack))
    n_fonts = sum(1 for f in _FONT_FAMILIES if f.lower() in low)
    n_kw = sum(1 for kw in _GUIDELINE_KEYWORDS if kw in low)

    # Tables whose header row reads like a spec table ("Element/Font/Size/…").
    spec_tables = 0
    for tbl in digest.tables:
        if not tbl:
            continue
        header = " ".join(c.lower() for c in tbl[0])
        hits = sum(1 for s in _TABLE_HEADER_SIGNALS if s in header)
        if hits >= 2:
            spec_tables += 1

    # Weighted score; tuned so a clear guideline lands well above threshold and
    # a normal document (few/no signals) stays near zero.
    score = (
        min(n_hex, 8) * 1.5
        + min(n_pt, 10) * 1.2
        + min(n_fonts, 4) * 1.0
        + min(n_kw, 12) * 0.6
        + spec_tables * 3.0
    )

    is_guideline = (
        (n_hex >= 2 and n_pt >= 2)
        or spec_tables >= 1
        or score >= 9.0
    )
    confidence = max(0.5, min(0.97, 0.5 + score / 30.0)) if is_guideline else \
        max(0.5, min(0.95, 0.6 + (6 - score) / 20.0))
    reason = (
        f"{n_hex} hex colour(s), {n_pt} point-size mention(s), {n_fonts} font "
        f"name(s), {spec_tables} spec-style table(s)."
    )
    return Classification(
        kind="guideline" if is_guideline else "example",
        confidence=round(confidence, 2),
        reason=reason,
        method="heuristic",
    )


def classify_style_source(
    digest: StyleSourceDigest, *, use_llm: bool = True
) -> Classification:
    """Decide whether the style source is a *guideline* or an *example*.

    Uses the LLM when available (richer judgement), always with a deterministic
    heuristic fallback so the flow never depends on LLM availability. Results
    are memoized per template content hash — a transient LLM failure is never
    cached, so the heuristic fallback doesn't get pinned across runs."""
    cache_key = (
        (digest.source_sha256, digest.source_kind, use_llm)
        if digest.source_sha256 else None
    )
    if cache_key is not None:
        cached = _cache_get(_CLASSIFY_CACHE, cache_key)
        if cached is not None:
            return cached

    cls: Optional[Classification] = None
    cacheable = True
    if use_llm:
        try:
            from app.llm.adapters import chat_json, llm_available

            if llm_available():
                system = (
                    "You classify a document uploaded as a 'style template'. "
                    "Decide if it is (A) a FORMATTING GUIDELINE / STYLE "
                    "SPECIFICATION that DESCRIBES formatting rules in prose or "
                    "tables (fonts, sizes, colours, margins, heading/table "
                    "treatments) — its own look is irrelevant; or (B) an "
                    "EXAMPLE document whose OWN visual appearance is the style "
                    "to copy. Return ONLY JSON: "
                    '{"kind":"guideline"|"example","confidence":0..1,"reason":"..."}.'
                )
                user = (
                    "Document to classify:\n\n"
                    + digest.as_prompt()
                    + "\n\nReturn the classification JSON now."
                )
                raw = chat_json(system, user, temperature=0.0, max_tokens=300)
                if isinstance(raw, dict):
                    kind = str(raw.get("kind", "")).strip().lower()
                    if kind in ("guideline", "example"):
                        conf = _coerce_num(raw.get("confidence"))
                        cls = Classification(
                            kind=kind,
                            confidence=round(min(max(conf if conf is not None else 0.7, 0.0), 1.0), 2),
                            reason=str(raw.get("reason") or "")[:300] or "LLM classification.",
                            method="llm",
                        )
                if cls is None:
                    cacheable = False  # LLM responded but unusably — retry next run
        except Exception:
            cacheable = False  # transient LLM failure — don't pin the fallback
    if cls is None:
        cls = _heuristic_classify(digest)
    if cache_key is not None and cacheable:
        _cache_put(_CLASSIFY_CACHE, cache_key, cls)
    return cls


def decide_style_mode(
    digest: StyleSourceDigest, requested_mode: str = "auto", *, use_llm: bool = True
) -> tuple[str, Classification]:
    """Resolve the effective interpretation mode.

    ``requested_mode`` may be 'auto' (classify), 'guideline', or 'example'.
    A forced mode is honoured but classification still runs (when 'auto') so
    callers can show the user what was detected."""
    mode = (requested_mode or "auto").strip().lower()
    if mode in ("guideline", "example"):
        return mode, Classification(
            kind=mode, confidence=1.0, reason="User-selected mode.", method="forced"
        )
    cls = classify_style_source(digest, use_llm=use_llm)
    return cls.kind, cls


# ---------------------------------------------------------------------------
# Interpretation — guideline prose/tables -> StyleSpec
# ---------------------------------------------------------------------------

_SPEC_SCHEMA_HINT = """Return ONLY a JSON object with this shape (omit unknown fields or use null):
{
  "page": {"width_in":8.5,"height_in":11,"orientation":"portrait",
           "margin_top_in":1,"margin_bottom_in":1,"margin_left_in":1,"margin_right_in":1},
  "body": {"font":"Arial","size_pt":11,"color_hex":"000000","alignment":"left",
           "line_spacing":1.0,"space_before_pt":0,"space_after_pt":6},
  "headings": [
    {"level":1,"font":"Arial","size_pt":14,"bold":true,"color_hex":"2E5496",
     "alignment":"left","space_before_pt":14,"space_after_pt":6,
     "bottom_border":true,"border_color_hex":"2E5496","border_width_pt":2},
    {"style_name":"Title","font":"Arial","size_pt":16,"bold":true,"color_hex":"2E5496"}
  ],
  "table": {"header_fill_hex":"2E5496","header_text_hex":"FFFFFF","header_bold":true,
            "header_font":"Arial","header_size_pt":10,"alt_row_fill_hex":"F2F6FC",
            "border_color_hex":"CCCCCC","border_width_pt":0.5,
            "cell_padding_top_in":0.06,"cell_padding_left_in":0.08,
            "body_font":"Arial","body_size_pt":10},
  "title_block": {
     "title":   {"font":"Arial","size_pt":16,"bold":true,"color_hex":"2E5496","alignment":"center"},
     "subtitle":{"font":"Arial","size_pt":14,"bold":true,"color_hex":"000000","alignment":"center"},
     "metadata":{"font":"Arial","size_pt":10,"color_hex":"666666","alignment":"center"}},
  "lists": {"bullet_char":"\\u2022","bullet_indent_in":0.5,"bullet_hanging_in":0.25,
            "number_format":"decimal","space_before_pt":0,"space_after_pt":4},
  "header_footer": {"header_left":"...","header_right":"...","footer_left":"...",
                    "footer_right":"Page X of Y","font":"Arial","size_pt":9,
                    "color_hex":"666666","header_border_color_hex":"2E5496",
                    "footer_border_color_hex":"CCCCCC"},
  "colors": {"Primary Blue":"2E5496","Light Blue Row":"F2F6FC"},
  "accent_color_hex": "2E5496",
  "notes": ["free-text rules that don't fit above, e.g. required sections or contradictions"]
}
Convert any measurements to the units shown: sizes in points (pt), margins/indents
in inches (1 inch = 72 pt = 1440 twips = 1440 DXA; 20 twips = 1 pt). Strip the '#'
from hex colours. Named colours with an obvious hex equivalent count as explicit
values — translate them (e.g. "navy" -> 000080, "maroon" -> 800000, "teal" ->
008080) rather than dropping the rule. Capture EVERY explicit rule you can find."""


def interpret_guideline(
    digest: StyleSourceDigest, *, use_llm: bool = True
) -> StyleSpec:
    """Distil a formatting guideline into a structured :class:`StyleSpec`.

    LLM-first (it reads the prose/tables and fills the schema); falls back to a
    regex heuristic that lifts fonts, sizes, a colour palette and margins so
    the flow still produces *something* useful without an LLM. Interpretation
    is memoized per template content hash (callers receive their own copy);
    transient LLM failures are never cached."""
    cache_key = (
        (digest.source_sha256, digest.source_kind, use_llm)
        if digest.source_sha256 else None
    )
    if cache_key is not None:
        cached = _cache_get(_INTERPRET_CACHE, cache_key)
        if cached is not None:
            return cached.model_copy(deep=True)

    spec: Optional[StyleSpec] = None
    cacheable = True
    if use_llm:
        try:
            from app.llm.adapters import chat_json, llm_available

            if llm_available():
                system = (
                    "You are a document-formatting analyst. The user gives you a "
                    "FORMATTING GUIDELINE that describes how documents in a series "
                    "must look. Extract every concrete formatting rule into the "
                    "JSON schema below. Do not invent values that aren't stated.\n"
                    "Pay special attention to header/footer rules: what TEXT should "
                    "appear in headers/footers (running title, document code, "
                    "confidentiality line, page numbering)? Put it in "
                    "'header_footer'. If you find rules that CONTRADICT each other "
                    "(e.g. two different body sizes), apply the more specific one "
                    "and flag the conflict in 'notes' prefixed 'Contradiction:'.\n\n"
                    + _SPEC_SCHEMA_HINT
                )
                user = (
                    "Formatting guideline document:\n\n"
                    + digest.as_prompt()
                    + "\n\nReturn the StyleSpec JSON now."
                )
                raw = chat_json(system, user, temperature=0.0, max_tokens=2600)
                if isinstance(raw, dict):
                    candidate = _spec_from_dict(raw)
                    if not candidate.is_empty():
                        spec = _backfill_spec(candidate, digest)
                if spec is None:
                    cacheable = False  # LLM responded but unusably — retry next run
        except Exception:
            cacheable = False  # transient LLM failure — don't pin the fallback
    if spec is None:
        spec = _heuristic_interpret(digest)
    if cache_key is not None and cacheable:
        _cache_put(_INTERPRET_CACHE, cache_key, spec.model_copy(deep=True))
    return spec


def _spec_from_dict(raw: dict[str, Any]) -> StyleSpec:
    """Tolerantly coerce a loosely-typed LLM dict into a StyleSpec."""
    def num(d: dict, *keys):
        for k in keys:
            if k in d:
                return _coerce_num(d.get(k))
        return None

    page_in = raw.get("page") or {}
    page = PageSpec(
        width_in=num(page_in, "width_in", "width"),
        height_in=num(page_in, "height_in", "height"),
        orientation=(str(page_in.get("orientation")).lower()
                     if page_in.get("orientation") else None),
        margin_top_in=num(page_in, "margin_top_in", "margin_top", "top"),
        margin_bottom_in=num(page_in, "margin_bottom_in", "margin_bottom", "bottom"),
        margin_left_in=num(page_in, "margin_left_in", "margin_left", "left"),
        margin_right_in=num(page_in, "margin_right_in", "margin_right", "right"),
    )

    body_in = raw.get("body") or {}
    body = BodySpec(
        font=body_in.get("font") or None,
        size_pt=num(body_in, "size_pt", "size"),
        color_hex=_coerce_hex(body_in.get("color_hex") or body_in.get("color")),
        alignment=_coerce_align(body_in.get("alignment")),
        line_spacing=num(body_in, "line_spacing"),
        space_before_pt=num(body_in, "space_before_pt", "space_before"),
        space_after_pt=num(body_in, "space_after_pt", "space_after"),
    )

    headings: list[HeadingRule] = []
    for h in (raw.get("headings") or []):
        if not isinstance(h, dict):
            continue
        headings.append(HeadingRule(
            level=int(_coerce_num(h.get("level"))) if h.get("level") is not None else None,
            style_name=h.get("style_name") or h.get("name") or None,
            font=h.get("font") or None,
            size_pt=num(h, "size_pt", "size"),
            bold=_coerce_bool(h.get("bold")),
            italic=_coerce_bool(h.get("italic")),
            underline=_coerce_bool(h.get("underline")),
            color_hex=_coerce_hex(h.get("color_hex") or h.get("color")),
            alignment=_coerce_align(h.get("alignment")),
            space_before_pt=num(h, "space_before_pt", "space_before"),
            space_after_pt=num(h, "space_after_pt", "space_after"),
            bottom_border=_coerce_bool(h.get("bottom_border")),
            border_color_hex=_coerce_hex(h.get("border_color_hex")),
            border_width_pt=num(h, "border_width_pt"),
        ))

    t = raw.get("table") or {}
    table = TableSpec(
        header_fill_hex=_coerce_hex(t.get("header_fill_hex") or t.get("header_fill")),
        header_text_hex=_coerce_hex(t.get("header_text_hex") or t.get("header_text")),
        header_bold=_coerce_bool(t.get("header_bold")),
        header_font=t.get("header_font") or None,
        header_size_pt=num(t, "header_size_pt"),
        alt_row_fill_hex=_coerce_hex(t.get("alt_row_fill_hex") or t.get("alt_row_fill")),
        border_color_hex=_coerce_hex(t.get("border_color_hex")),
        border_width_pt=num(t, "border_width_pt"),
        cell_padding_top_in=num(t, "cell_padding_top_in"),
        cell_padding_left_in=num(t, "cell_padding_left_in"),
        body_font=t.get("body_font") or None,
        body_size_pt=num(t, "body_size_pt"),
    )

    l = raw.get("lists") or {}
    lists = ListSpec(
        bullet_char=_clean_bullet_char(l.get("bullet_char")),
        bullet_indent_in=num(l, "bullet_indent_in"),
        bullet_hanging_in=num(l, "bullet_hanging_in"),
        number_format=l.get("number_format") or None,
        space_before_pt=num(l, "space_before_pt", "space_before"),
        space_after_pt=num(l, "space_after_pt", "space_after"),
    )

    def _title_block(d: dict) -> TitleBlockSpec:
        return TitleBlockSpec(
            font=d.get("font") or None,
            size_pt=num(d, "size_pt", "size"),
            bold=_coerce_bool(d.get("bold")),
            italic=_coerce_bool(d.get("italic")),
            color_hex=_coerce_hex(d.get("color_hex") or d.get("color")),
            alignment=_coerce_align(d.get("alignment")),
            space_before_pt=num(d, "space_before_pt", "space_before"),
            space_after_pt=num(d, "space_after_pt", "space_after"),
        )

    tb_in = raw.get("title_block") or {}
    title = _title_block(tb_in.get("title") or {})
    subtitle = _title_block(tb_in.get("subtitle") or {})
    metadata = _title_block(tb_in.get("metadata") or {})

    hf = raw.get("header_footer") or {}
    header_footer = HeaderFooterSpec(
        header_left=hf.get("header_left") or None,
        header_right=hf.get("header_right") or None,
        footer_left=hf.get("footer_left") or None,
        footer_right=hf.get("footer_right") or None,
        font=hf.get("font") or None,
        size_pt=num(hf, "size_pt"),
        color_hex=_coerce_hex(hf.get("color_hex")),
        header_border_color_hex=_coerce_hex(hf.get("header_border_color_hex")),
        footer_border_color_hex=_coerce_hex(hf.get("footer_border_color_hex")),
    )

    colors: dict[str, str] = {}
    for k, v in (raw.get("colors") or {}).items():
        hx = _coerce_hex(v)
        if hx:
            colors[str(k)] = hx

    notes = [str(n) for n in (raw.get("notes") or []) if str(n).strip()]

    return StyleSpec(
        page=page, body=body, headings=headings, table=table, lists=lists,
        header_footer=header_footer, colors=colors,
        title=title, subtitle=subtitle, metadata=metadata,
        accent_color_hex=_coerce_hex(raw.get("accent_color_hex")),
        notes=notes,
    )


def _backfill_spec(spec: StyleSpec, digest: StyleSourceDigest) -> StyleSpec:
    """Fill a couple of high-value gaps the LLM may have left, from heuristics."""
    if not spec.accent_color_hex:
        # Pick the most prominent non-near-black/near-white palette colour.
        for hx in spec.colors.values():
            r, g, b = int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16)
            if 40 < max(r, g, b) and min(r, g, b) < 220:
                spec.accent_color_hex = hx
                break

    # Title block backfill from any heading rule named "Title"/"Subtitle" — many
    # guidelines describe the masthead as those named styles in the typography
    # table rather than in a dedicated title_block section.
    by_name = {(h.style_name or "").strip().lower(): h for h in spec.headings if h.style_name}
    if not spec.title.has_values() and "title" in by_name:
        h = by_name["title"]
        spec.title = TitleBlockSpec(
            font=h.font, size_pt=h.size_pt, bold=h.bold if h.bold is not None else True,
            color_hex=h.color_hex or spec.accent_color_hex, alignment=h.alignment or "center",
        )
    if not spec.subtitle.has_values() and ("subtitle" in by_name or "document subtitle" in by_name):
        h = by_name.get("subtitle") or by_name.get("document subtitle")
        spec.subtitle = TitleBlockSpec(
            font=h.font, size_pt=h.size_pt, bold=h.bold if h.bold is not None else True,
            color_hex=h.color_hex, alignment=h.alignment or "center",
        )
    # Drop the masthead pseudo-headings from the heading list so the applier
    # doesn't try to bind them to "Heading N" / "Title" styles in the content.
    spec.headings = [
        h for h in spec.headings
        if (h.style_name or "").strip().lower() not in ("title", "subtitle", "document subtitle", "document title")
    ]
    # Metadata colour from the palette ("Metadata / Footer #666666") if unset.
    if not spec.metadata.color_hex:
        for name, hx in spec.colors.items():
            if any(k in name.lower() for k in ("metadata", "footer", "muted", "gray", "grey")):
                spec.metadata.color_hex = hx
                spec.metadata.alignment = spec.metadata.alignment or "center"
                break
    return spec


def _heuristic_interpret(digest: StyleSourceDigest) -> StyleSpec:
    """Best-effort, LLM-free extraction. Lifts a palette, body font/size and a
    rough heading scale so a guideline still yields a usable spec offline."""
    haystack = digest.text + "\n" + "\n".join(
        " | ".join(row) for tbl in digest.tables for row in tbl
    )
    low = haystack.lower()

    # Colour palette: every distinct hex, keyed by a nearby label when possible.
    colors: dict[str, str] = {}
    for hx in dict.fromkeys(h.upper() for h in _HEX_RE.findall(haystack)):
        colors[hx] = hx
    accent = None
    for hx in colors:
        r, g, b = int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16)
        if max(r, g, b) > 60 and not (r > 230 and g > 230 and b > 230):
            accent = hx
            break

    # Body font: the most-mentioned known family.
    body_font = None
    best = 0
    for fam in _FONT_FAMILIES:
        c = low.count(fam.lower())
        if c > best:
            best, body_font = c, fam

    # Body size: the most-mentioned size in the 10–12pt band (body text is the
    # most frequent size); fall back to the smallest plausible body size.
    all_sizes = [float(s) for s in _PT_RE.findall(haystack)]
    sizes = sorted(set(all_sizes))
    band = [s for s in all_sizes if 10 <= s <= 12.5]
    if band:
        body_size = max(set(band), key=band.count)
    else:
        body_size = next((s for s in sizes if 9 <= s <= 12), None)

    # Heading scale: distinct sizes clearly larger than the body, descending.
    headings: list[HeadingRule] = []
    if sizes:
        body_ref = body_size or (sizes[0] if sizes else 11)
        bigger = sorted({s for s in sizes if s > (body_ref + 0.5)}, reverse=True)
        for lvl, sz in enumerate(bigger[:4], start=1):
            headings.append(HeadingRule(
                level=lvl, font=body_font, size_pt=sz, bold=True,
                color_hex=accent,
            ))

    # Margins: only an inch value that sits next to the word "margin" (avoids
    # grabbing unrelated inch mentions such as a table's content width).
    margin = None
    m = (re.search(r"margins?[^.\n]{0,40}?(\d+(?:\.\d+)?)\s*(?:inch|in\b|\")", low)
         or re.search(r"(\d+(?:\.\d+)?)\s*(?:inch|in\b|\")[^.\n]{0,20}?margin", low))
    if m:
        val = float(m.group(1))
        if 0.3 <= val <= 2.0:
            margin = val

    notes = []
    if not (body_font or headings or colors):
        notes.append("Could not confidently extract rules without an LLM; applied minimal defaults.")

    return StyleSpec(
        body=BodySpec(font=body_font, size_pt=body_size),
        headings=headings,
        page=PageSpec(
            margin_top_in=margin, margin_bottom_in=margin,
            margin_left_in=margin, margin_right_in=margin,
        ),
        colors=colors,
        accent_color_hex=accent,
        notes=notes,
    )


# ---------------------------------------------------------------------------
# Example profiling — read a formatted example .docx into a StyleSpec
# ---------------------------------------------------------------------------
#
# The example path transplants the example's styles.xml/theme for maximum
# fidelity, but the structural upgrade (masthead, tables, header/footer) still
# needs concrete values. This reads them straight off the example's OOXML so the
# *same* structure-aware applier drives both the example and guideline paths.

def _rgb_or_none(color) -> Optional[str]:
    try:
        if color is not None and color.type is not None and color.rgb is not None:
            return str(color.rgb)
    except Exception:
        return None
    return None


def _align_name(alignment) -> Optional[str]:
    if alignment is None:
        return None
    return {0: "left", 1: "center", 2: "right", 3: "justify"}.get(int(alignment))


def _pt(v) -> Optional[float]:
    try:
        return round(v.pt, 1) if v is not None else None
    except Exception:
        return None


def _read_docdefaults(doc) -> tuple[Optional[str], Optional[float]]:
    import re as _re

    try:
        dd = doc.styles.element.find(qn("w:docDefaults"))
        rpr = dd.find(qn("w:rPrDefault")).find(qn("w:rPr")) if dd is not None else None
    except Exception:
        rpr = None
    if rpr is None:
        return None, None
    font = None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is not None:
        font = rfonts.get(qn("w:ascii")) or rfonts.get(qn("w:hAnsi"))
    size = None
    sz = rpr.find(qn("w:sz"))
    if sz is not None and sz.get(qn("w:val")):
        try:
            size = float(sz.get(qn("w:val"))) / 2.0
        except (TypeError, ValueError):
            size = None
    return font, size


def _first_run(cell):
    for para in cell.paragraphs:
        for r in para.runs:
            if (r.text or "").strip():
                return r
    return None


def _cell_fill(cell) -> Optional[str]:
    shd = cell._tc.find(".//" + qn("w:shd"))
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    return fill.upper() if fill and fill.lower() != "auto" else None


def _has_bottom_border(style) -> Optional[bool]:
    try:
        pPr = style.element.find(qn("w:pPr"))
        pBdr = pPr.find(qn("w:pBdr")) if pPr is not None else None
        return pBdr is not None and pBdr.find(qn("w:bottom")) is not None
    except Exception:
        return None


def _table_border(tbl) -> tuple[Optional[str], Optional[float]]:
    try:
        borders = tbl._tbl.tblPr.find(qn("w:tblBorders"))
        if borders is None:
            return None, None
        edge = borders.find(qn("w:top"))
        if edge is None:
            edge = borders.find(qn("w:insideH"))
        if edge is None:
            return None, None
        color = edge.get(qn("w:color"))
        sz = edge.get(qn("w:sz"))
        width = float(sz) / 8.0 if sz else None
        return (color.upper() if color and color.lower() != "auto" else None), width
    except Exception:
        return None, None


def profile_example_to_spec(docx_bytes: bytes) -> StyleSpec:
    """Profile a formatted example .docx into a :class:`StyleSpec`.

    Unknown values are left ``None`` so the applier preserves whatever the
    transplanted styles already provide (None never clobbers a real value)."""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    dd_font, dd_size = _read_docdefaults(doc)

    try:
        normal = doc.styles["Normal"]
    except KeyError:
        normal = None
    body = BodySpec(
        font=(normal.font.name if normal and normal.font.name else dd_font),
        size_pt=(_pt(normal.font.size) if normal and normal.font.size else dd_size),
        color_hex=_rgb_or_none(normal.font.color) if normal else None,
        alignment=_align_name(normal.paragraph_format.alignment) if normal else None,
        space_after_pt=_pt(normal.paragraph_format.space_after) if normal else None,
    )

    headings: list[HeadingRule] = []
    accent = None
    for lvl in range(1, 5):
        try:
            st = doc.styles[f"Heading {lvl}"]
        except KeyError:
            continue
        pf = st.paragraph_format
        color = _rgb_or_none(st.font.color)
        if lvl == 1 and color:
            accent = color
        headings.append(HeadingRule(
            level=lvl,
            font=st.font.name or None,
            size_pt=_pt(st.font.size),
            bold=bool(st.font.bold) if st.font.bold is not None else None,
            color_hex=color,
            space_before_pt=_pt(pf.space_before),
            space_after_pt=_pt(pf.space_after),
            bottom_border=_has_bottom_border(st),
        ))

    # Masthead — the leading non-empty paragraphs before the first heading.
    title = subtitle = metadata = None
    seq = []
    for p in doc.paragraphs:
        sname = p.style.name if p.style else ""
        if sname.startswith("Heading"):
            break
        if not (p.text or "").strip():
            continue
        runs = [r for r in p.runs if (r.text or "").strip()]
        if runs:
            seq.append((p, runs[0]))
        if len(seq) >= 3:
            break

    def _look(pair) -> TitleBlockSpec:
        p, r = pair
        return TitleBlockSpec(
            font=r.font.name or None,
            size_pt=_pt(r.font.size),
            bold=bool(r.font.bold) if r.font.bold is not None else None,
            color_hex=_rgb_or_none(r.font.color),
            alignment=_align_name(p.paragraph_format.alignment) or "center",
        )

    title = _look(seq[0]) if len(seq) >= 1 else TitleBlockSpec()
    subtitle = _look(seq[1]) if len(seq) >= 2 else TitleBlockSpec()
    metadata = _look(seq[2]) if len(seq) >= 3 else TitleBlockSpec()
    if not accent and title.color_hex:
        accent = title.color_hex

    # Tables — header fill + text, an alternating row fill, borders, body font.
    table = TableSpec()
    if doc.tables:
        tbl = doc.tables[0]
        if tbl.rows:
            hdr = tbl.rows[0]
            table.header_fill_hex = _cell_fill(hdr.cells[0])
            hr = _first_run(hdr.cells[0])
            if hr is not None:
                table.header_text_hex = _rgb_or_none(hr.font.color)
                table.header_bold = bool(hr.font.bold) if hr.font.bold is not None else None
                table.header_font = hr.font.name or None
                table.header_size_pt = _pt(hr.font.size)
        for row in tbl.rows[1:]:
            f = _cell_fill(row.cells[0])
            if f and f != table.header_fill_hex:
                table.alt_row_fill_hex = f
                break
        if len(tbl.rows) > 1:
            br = _first_run(tbl.rows[1].cells[0])
            if br is not None:
                table.body_font = br.font.name or None
                table.body_size_pt = _pt(br.font.size)
        table.border_color_hex, table.border_width_pt = _table_border(tbl)

    # Header/footer styling (text is derived from the content, not copied).
    hf = HeaderFooterSpec()
    try:
        sec = doc.sections[0]
        hr = next((r for p in sec.header.paragraphs for r in p.runs if (r.text or "").strip()), None)
        fr = next((r for p in sec.footer.paragraphs for r in p.runs if (r.text or "").strip()), None)
        ref = hr or fr
        if ref is not None:
            hf.font = ref.font.name or None
            hf.size_pt = _pt(ref.font.size)
            hf.color_hex = _rgb_or_none(ref.font.color)
        ftxt = " ".join(p.text for p in sec.footer.paragraphs)
        if "confidential" in ftxt.lower():
            hf.footer_left = "CONFIDENTIAL — For Internal Use Only"
    except Exception:
        pass

    page = PageSpec()
    try:
        sec = doc.sections[0]
        page = PageSpec(
            width_in=round(sec.page_width.inches, 2),
            height_in=round(sec.page_height.inches, 2),
            orientation="portrait" if int(sec.orientation) == 0 else "landscape",
            margin_top_in=round(sec.top_margin.inches, 2),
            margin_bottom_in=round(sec.bottom_margin.inches, 2),
            margin_left_in=round(sec.left_margin.inches, 2),
            margin_right_in=round(sec.right_margin.inches, 2),
        )
    except Exception:
        pass

    colors = {"Accent": accent} if accent else {}
    return StyleSpec(
        page=page, body=body, headings=headings, table=table,
        lists=ListSpec(space_after_pt=4.0),
        header_footer=hf, title=title, subtitle=subtitle, metadata=metadata,
        accent_color_hex=accent, colors=colors,
    )


# ---------------------------------------------------------------------------
# Human-readable summary (for warnings + UI preview)
# ---------------------------------------------------------------------------

def style_spec_summary(spec: StyleSpec) -> str:
    bits: list[str] = []
    if spec.body.font or spec.body.size_pt:
        b = spec.body
        size = f"{b.size_pt:g}pt" if b.size_pt else ""
        bits.append(f"body {(_join_sp(b.font, size)) or '—'}")
    if spec.headings:
        bits.append(f"{len(spec.headings)} heading rule(s)")
    if spec.accent_color_hex:
        bits.append(f"accent #{spec.accent_color_hex}")
    elif spec.colors:
        bits.append(f"{len(spec.colors)} palette colour(s)")
    if any(getattr(spec.table, f) is not None for f in spec.table.model_fields):
        bits.append("table styling")
    margins = [spec.page.margin_top_in, spec.page.margin_left_in]
    if any(m is not None for m in margins):
        mv = next(m for m in margins if m is not None)
        bits.append(f"{mv:g}\" margins")
    if any(getattr(spec.header_footer, f) for f in spec.header_footer.model_fields):
        bits.append("header/footer")
    return ", ".join(bits) if bits else "no concrete rules detected"


def _join_sp(*parts: Optional[str]) -> str:
    return " ".join(p for p in parts if p)
