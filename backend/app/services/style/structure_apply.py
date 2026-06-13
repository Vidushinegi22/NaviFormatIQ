"""
Structure-aware style application — meticulous, in-place document formatting.
=============================================================================

Given a content ``.docx`` and a :class:`StyleSpec`, this module:

  1. **understands** the document (via :mod:`structure_recognizer`) — which
     paragraphs are the title / subtitle / metadata, which are headings (and at
     what level), which are list items (bullet vs numbered), and which runs of
     delimited text are really tables;
  2. **upgrades the live OOXML in place** so the document gains real structure
     while keeping 100% of its content — promoting headings to ``Heading N``,
     turning flat paragraphs into real bullet/numbered lists (with proper
     ``numPr`` numbering that restarts per list), rebuilding pipe/tab-delimited
     text into real Word tables, and centring + colouring the masthead;
  3. **applies the visual spec meticulously** to every element — body font,
     heading fonts/colours/borders, table header fill + alternating rows +
     borders, list indents + spacing, page margins, and a derived header/footer.

This is the engine that lets a *flat* raw document end up looking exactly like a
hand-formatted target — the gap the old style-layer-only transfer could not close.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import TYPE_CHECKING, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.services.formatting.text_safe import xml_safe
from app.services.style import style_engine as SE
from app.services.style.structure_recognizer import (
    BODY,
    BULLET,
    HEADING,
    LIST_ITEM,
    METADATA,
    NUMBERED,
    SUBTITLE,
    TITLE,
    StructurePlan,
    recognize_structure,
)

if TYPE_CHECKING:
    from app.services.style.guideline_interpreter import StyleSpec, TitleBlockSpec


# ---------------------------------------------------------------------------
# Numbering manager — mints well-formed bullet/decimal lists with restarts
# ---------------------------------------------------------------------------

_BULLET_GLYPHS = ["•", "◦", "▪", "•", "◦", "▪", "•", "◦", "▪"]


def _safe_glyph(ch: Optional[str]) -> Optional[str]:
    """A bullet glyph safe to embed in XML (no control chars). Maps the common
    Symbol-font code point to a real bullet; rejects empties/control chars."""
    if not ch:
        return None
    ch = "".join(c for c in ch if ord(c) >= 0x20 and not c.isspace())
    if not ch:
        return None
    if ch in ("\xb7", "·", ""):  # Symbol-font bullet variants
        return "•"
    # A bullet is a single non-alphanumeric symbol — reject words/digits
    # (e.g. a mis-extracted "7" from a Symbol-font code point).
    if len(ch) > 2 or any(c.isalnum() for c in ch):
        return None
    return ch


class _NumberingManager:
    """Creates dedicated bullet/decimal numbering definitions in the document's
    ``numbering.xml`` and attaches ``numPr`` to paragraphs. Each contiguous list
    group gets its own ``numId`` so numbered lists restart at 1 every section."""

    def __init__(self, doc, spec: "Optional[StyleSpec]" = None):
        self.doc = doc
        self.spec = spec
        self.numbering = self._numbering_el()
        self._abs: dict[str, str] = {}             # kind -> abstractNumId
        self._group_num: dict[tuple[str, int], int] = {}

    @property
    def _indent_in(self) -> float:
        return (self.spec.lists.bullet_indent_in if self.spec else None) or 0.5

    @property
    def _hang_in(self) -> float:
        return (self.spec.lists.bullet_hanging_in if self.spec else None) or 0.25

    @property
    def _bullet_char(self) -> Optional[str]:
        return self.spec.lists.bullet_char if self.spec else None

    def _numbering_el(self):
        try:
            return self.doc.part.numbering_part.element
        except Exception:
            return None

    def available(self) -> bool:
        return self.numbering is not None

    def _next_abstract_id(self) -> int:
        ids = [int(a.get(qn("w:abstractNumId")) or 0)
               for a in self.numbering.findall(qn("w:abstractNum"))]
        return (max(ids) if ids else 0) + 1

    def _next_num_id(self) -> int:
        ids = [int(n.get(qn("w:numId")) or 0) for n in self.numbering.findall(qn("w:num"))]
        return (max(ids) if ids else 0) + 1

    def _make_abstract(self, kind: str) -> str:
        nid = str(self._next_abstract_id())
        an = OxmlElement("w:abstractNum")
        an.set(qn("w:abstractNumId"), nid)
        mlt = OxmlElement("w:multiLevelType")
        mlt.set(qn("w:val"), "hybridMultilevel")
        an.append(mlt)
        base = self._indent_in
        hang = self._hang_in
        for ilvl in range(9):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(ilvl))
            start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
            fmt = OxmlElement("w:numFmt")
            fmt.set(qn("w:val"), "bullet" if kind == BULLET else "decimal")
            lvl.append(fmt)
            txt = OxmlElement("w:lvlText")
            if kind == BULLET:
                glyph = (_safe_glyph(self._bullet_char) if ilvl == 0 else None) or _BULLET_GLYPHS[ilvl]
                txt.set(qn("w:val"), glyph)
            else:
                txt.set(qn("w:val"), f"%{ilvl + 1}.")
            lvl.append(txt)
            jc = OxmlElement("w:lvlJc"); jc.set(qn("w:val"), "left"); lvl.append(jc)
            ppr = OxmlElement("w:pPr")
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(int(round((base + ilvl * 0.25) * 1440))))
            ind.set(qn("w:hanging"), str(int(round(hang * 1440))))
            ppr.append(ind)
            lvl.append(ppr)
            an.append(lvl)
        # Schema: all <w:abstractNum> must precede every <w:num>.
        first_num = self.numbering.find(qn("w:num"))
        if first_num is not None:
            first_num.addprevious(an)
        else:
            self.numbering.append(an)
        return nid

    def _ensure_abstract(self, kind: str) -> str:
        if kind not in self._abs:
            self._abs[kind] = self._make_abstract(kind)
        return self._abs[kind]

    def _make_num(self, abstract_id: str) -> int:
        nid = self._next_num_id()
        num = OxmlElement("w:num"); num.set(qn("w:numId"), str(nid))
        ref = OxmlElement("w:abstractNumId"); ref.set(qn("w:val"), abstract_id); num.append(ref)
        for ilvl in range(9):
            ov = OxmlElement("w:lvlOverride"); ov.set(qn("w:ilvl"), str(ilvl))
            so = OxmlElement("w:startOverride"); so.set(qn("w:val"), "1"); ov.append(so)
            num.append(ov)
        self.numbering.append(num)
        return nid

    def attach(self, p, kind: str, level: int, group_id: int) -> bool:
        if self.numbering is None:
            return False
        abs_id = self._ensure_abstract(kind)
        key = (kind, group_id)
        if key not in self._group_num:
            self._group_num[key] = self._make_num(abs_id)
        num_id = self._group_num[key]
        pPr = p._p.get_or_add_pPr()
        for ex in pPr.findall(qn("w:numPr")):
            pPr.remove(ex)
        numPr = OxmlElement("w:numPr")
        il = OxmlElement("w:ilvl"); il.set(qn("w:val"), str(max(0, min(level, 8)))); numPr.append(il)
        ni = OxmlElement("w:numId"); ni.set(qn("w:val"), str(num_id)); numPr.append(ni)
        # numPr belongs right after pStyle in CT_PPr ordering.
        pstyle = pPr.find(qn("w:pStyle"))
        if pstyle is not None:
            pstyle.addnext(numPr)
        else:
            pPr.insert(0, numPr)
        return True


# ---------------------------------------------------------------------------
# Small helpers
# ---------------------------------------------------------------------------

_LEADER_RE = re.compile(
    r"^\s*(?:"
    r"[•●○◦▪■□◆◇‣·\-–—*]"
    r"|\(?\s*(?:\d{1,3}|[ivxlcdmIVXLCDM]+|[a-zA-Z])\s*[.)\]]"
    r")\s+"
)


def _strip_leader_first_run(p) -> None:
    for r in p.runs:
        if (r.text or "").strip():
            r.text = _LEADER_RE.sub("", r.text, count=1)
            return


def _set_run_look(run, *, font=None, size_pt=None, bold=None, italic=None, color_hex=None) -> None:
    if font:
        run.font.name = font
        SE._rfonts_set(run._r.get_or_add_rPr(), font)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color_hex:
        try:
            run.font.color.rgb = RGBColor.from_string(color_hex.lstrip("#"))
        except Exception:
            pass


def _clear_direct_run_props(p, *, tags=("w:b", "w:bCs", "w:color", "w:sz", "w:szCs", "w:rFonts")) -> None:
    """Strip direct run formatting so the paragraph's *style* governs the look."""
    for r in p.runs:
        rpr = r._r.find(qn("w:rPr"))
        if rpr is None:
            continue
        for tag in tags:
            el = rpr.find(qn(tag))
            if el is not None:
                rpr.remove(el)


def _align_const(name: Optional[str]):
    return SE._ALIGN_MAP.get((name or "").lower())


# ---------------------------------------------------------------------------
# Masthead
# ---------------------------------------------------------------------------

def _apply_masthead(p, look: "TitleBlockSpec", *, default_align: str,
                    default_bold: Optional[bool] = None, align_override: Optional[str] = None) -> None:
    pf = p.paragraph_format
    al = _align_const(align_override or look.alignment or default_align)
    if al is not None:
        pf.alignment = al
    if look.space_before_pt is not None:
        pf.space_before = Pt(look.space_before_pt)
    if look.space_after_pt is not None:
        pf.space_after = Pt(look.space_after_pt)
    bold = look.bold if look.bold is not None else default_bold
    runs = [r for r in p.runs if (r.text or "")]
    for r in runs:
        _set_run_look(
            r, font=look.font, size_pt=look.size_pt, bold=bold,
            italic=look.italic, color_hex=look.color_hex,
        )


# ---------------------------------------------------------------------------
# Headings
# ---------------------------------------------------------------------------

def _assign_heading(doc, p, level: int) -> None:
    level = max(1, min(level, 4))
    for name in (f"Heading {level}", f"Heading {min(level, 3)}", "Heading 1"):
        try:
            p.style = doc.styles[name]
            break
        except KeyError:
            continue
    # Let the heading STYLE fully govern the look (colour/size/bold/font).
    _clear_direct_run_props(p)


# ---------------------------------------------------------------------------
# Lists
# ---------------------------------------------------------------------------

def _list_groups(plan: StructurePlan, ordered_idxs: list[int]) -> dict[int, int]:
    """Assign a group id to each list-item index. A new group starts when the
    previous element wasn't a list item, the kind changed, or the level dropped."""
    groups: dict[int, int] = {}
    gid = 0
    prev_was_list = False
    prev_kind = None
    prev_level = None
    for idx in ordered_idxs:
        role = plan.role(idx)
        if role.kind != LIST_ITEM:
            prev_was_list = False
            prev_kind = None
            prev_level = None
            continue
        start_new = (
            not prev_was_list
            or role.list_kind != prev_kind
            or (prev_level is not None and role.level < prev_level)
        )
        if start_new:
            gid += 1
        groups[idx] = gid
        prev_was_list = True
        prev_kind = role.list_kind
        prev_level = role.level
    return groups


def _apply_list_item(doc, p, role, group_id: int, numbering: _NumberingManager, spec: "StyleSpec") -> bool:
    """Turn a flat paragraph into a real list item. Returns False if it ended up
    empty (caller should drop it to avoid an empty bullet)."""
    _strip_leader_first_run(p)
    if not (p.text or "").strip():
        return False
    try:
        p.style = doc.styles["List Paragraph"]
    except KeyError:
        pass
    level = max(0, min(role.level, 8))
    ok = numbering.attach(p, role.list_kind, level, group_id)
    pf = p.paragraph_format
    base = spec.lists.bullet_indent_in or 0.5
    hang = spec.lists.bullet_hanging_in or 0.25
    pf.left_indent = Inches(base + level * 0.25)
    pf.first_line_indent = Inches(-hang)
    if spec.lists.space_after_pt is not None:
        pf.space_after = Pt(spec.lists.space_after_pt)
    if spec.lists.space_before_pt is not None:
        pf.space_before = Pt(spec.lists.space_before_pt)
    if not ok:
        # No numbering part — fall back to a literal glyph so the bullet still shows.
        glyph = (_safe_glyph(spec.lists.bullet_char) or "•") if role.list_kind == BULLET else "•"
        lead = p.add_run(glyph + " ")
        lead._r.getparent().insert(0, lead._r)  # move leader to the front
    return True


# ---------------------------------------------------------------------------
# Tables (rebuild delimited text into a real Word table)
# ---------------------------------------------------------------------------

def _build_table(doc, anchor_p_el, rows: list[list[str]], has_header: bool,
                 spec: "StyleSpec", warnings: list[str]):
    n_rows = len(rows)
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)  # appended at body end…
    try:
        table.style = doc.styles["Table Grid"]
    except KeyError:
        pass
    table.autofit = True
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            cell = table.cell(ri, ci)
            cell.text = xml_safe(row[ci]) if ci < len(row) else ""
    anchor_p_el.addprevious(table._tbl)  # …then moved to the right place.
    _style_table(table, has_header, spec, warnings)
    # A blank paragraph after the table keeps adjacent tables from merging and
    # gives breathing room.
    spacer = OxmlElement("w:p")
    table._tbl.addnext(spacer)
    return table


def _style_table(table, has_header: bool, spec: "StyleSpec", warnings: list[str]) -> None:
    t = spec.table
    accent = t.header_fill_hex or spec.accent_color_hex
    try:
        if t.border_color_hex or t.border_width_pt:
            SE._set_table_borders(table, t.border_color_hex, t.border_width_pt or 0.5)
        if t.cell_padding_top_in is not None or t.cell_padding_left_in is not None:
            SE._set_table_cell_margins(table, t.cell_padding_top_in, t.cell_padding_left_in)
        for ri, row in enumerate(table.rows):
            is_header = has_header and ri == 0
            for cell in row.cells:
                # vertical-center every cell
                try:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                except Exception:
                    pass
                if is_header:
                    if accent:
                        SE._set_cell_shading(cell, accent.lstrip("#"))
                    _style_cell(cell, font=t.header_font or t.body_font,
                                size_pt=t.header_size_pt or t.body_size_pt,
                                color_hex=t.header_text_hex or "FFFFFF",
                                bold=t.header_bold if t.header_bold is not None else True)
                else:
                    if t.alt_row_fill_hex and ri % 2 == 0:
                        SE._set_cell_shading(cell, t.alt_row_fill_hex.lstrip("#"))
                    _style_cell(cell, font=t.body_font, size_pt=t.body_size_pt)
    except Exception as e:  # noqa: BLE001 — table styling must never break the run
        warnings.append(f"table styling skipped: {e}")


def _style_cell(cell, *, font=None, size_pt=None, color_hex=None, bold=None) -> None:
    for para in cell.paragraphs:
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        for run in para.runs:
            _set_run_look(run, font=font, size_pt=size_pt, color_hex=color_hex, bold=bold)


# ---------------------------------------------------------------------------
# Metadata blocks → tables
# ---------------------------------------------------------------------------

_META_KV_RE = re.compile(r"^\s*([^:|]{2,48}?)\s*:\s*(.+?)\s*$")


def _metadata_table_groups(plan: StructurePlan, paras) -> list[list[int]]:
    """Runs of >=2 consecutive non-masthead ``Label: value`` metadata lines.

    An approval block (Department / Prepared By / Reviewed By / …) reads far
    better as a styled table than as loose paragraphs. Lines must be adjacent
    (at most one BLANK paragraph between them) and every line must parse as
    ``Label: value`` — otherwise the block is left alone."""
    eligible: list[int] = []
    for idx in sorted(plan.roles):
        role = plan.roles[idx]
        if (
            role.kind == METADATA
            and not role.masthead
            and idx not in plan.consumed_idxs
            and idx < len(paras)
            and _META_KV_RE.match((paras[idx].text or "").strip())
        ):
            eligible.append(idx)

    groups: list[list[int]] = []
    cur: list[int] = []
    prev: Optional[int] = None
    for idx in eligible:
        adjacent = prev is not None and (
            idx - prev == 1
            or (idx - prev == 2 and not (paras[prev + 1].text or "").strip())
        )
        if prev is not None and not adjacent:
            if len(cur) >= 2:
                groups.append(cur)
            cur = []
        cur.append(idx)
        prev = idx
    if len(cur) >= 2:
        groups.append(cur)
    # Cap at 6 entries — anything longer reads as a form, not an approval block.
    return [g for g in groups if 2 <= len(g) <= 6]


def _build_metadata_tables(doc, plan: StructurePlan, paras, spec: "StyleSpec",
                           warnings: list[str]) -> None:
    for group in sorted(_metadata_table_groups(plan, paras),
                        key=lambda g: g[0], reverse=True):
        pairs = []
        for idx in group:
            m = _META_KV_RE.match((paras[idx].text or "").strip())
            if m:
                pairs.append((m.group(1).strip(), m.group(2).strip()))
        if len(pairs) < 2:
            continue
        # Up to 5 entries fit across the page: labels as the header row,
        # values beneath. Longer blocks fall back to label|value rows.
        if len(pairs) <= 5:
            rows = [[p[0] for p in pairs], [p[1] for p in pairs]]
        else:
            rows = [["Field", "Value"]] + [[l, v] for l, v in pairs]
        try:
            _build_table(doc, paras[group[0]]._p, rows, True, spec, warnings)
        except Exception as e:  # noqa: BLE001
            warnings.append(f"metadata table skipped at idx {group[0]}: {e}")
            continue
        for idx in group:
            try:
                paras[idx]._p.getparent().remove(paras[idx]._p)
            except Exception:
                pass
            plan.consumed_idxs.add(idx)


# ---------------------------------------------------------------------------
# Header / footer — text DERIVED from the content, styled per the spec
# ---------------------------------------------------------------------------

_RULE_WORDS = ("pt", "arial", "calibri", "bold", "italic", "regular", "twips", "dxa", "#")
_DOC_CODE_RE = re.compile(r"\b([A-Z][A-Z0-9]*-[A-Z0-9]+(?:-[A-Z0-9]+)+)\b")


def _looks_like_rule_text(s: Optional[str]) -> bool:
    if not s:
        return True
    low = s.lower()
    return any(w in low for w in _RULE_WORDS) or bool(re.search(r"[0-9A-Fa-f]{6}\b", s) and "#" in s)


def _title_case(s: str) -> str:
    return s.title() if s.isupper() else s


def _add_field(para, instr: str) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f" {instr} ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    para._p.append(fld)


def _content_width_in(doc) -> float:
    """The usable text width (page minus margins) of the first section."""
    try:
        sec = doc.sections[0]
        width = (sec.page_width - sec.left_margin - sec.right_margin).inches
        return width if 2.0 <= width <= 20.0 else 6.5
    except Exception:
        return 6.5


def _write_hf(part, left: str, right: str, hf, *, page_field: bool,
              border_hex: Optional[str], border_side: str,
              tab_width_in: float = 6.5) -> None:
    from docx.enum.text import WD_TAB_ALIGNMENT

    part.is_linked_to_previous = False
    para = part.paragraphs[0] if part.paragraphs else part.add_paragraph()
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    try:
        para.paragraph_format.tab_stops.clear_all()
        para.paragraph_format.tab_stops.add_tab_stop(
            Inches(tab_width_in), WD_TAB_ALIGNMENT.RIGHT
        )
    except Exception:
        pass

    def styled_run(text: str):
        run = para.add_run(xml_safe(text or ""))
        _set_run_look(run, font=hf.font or "Arial", size_pt=hf.size_pt or 9,
                      color_hex=hf.color_hex or "666666")
        return run

    if left:
        styled_run(left)
    if right or page_field:
        styled_run("\t")
        if right:
            styled_run(right)
        if page_field:
            styled_run("Page ")
            _add_field(para, "PAGE")
            styled_run(" of ")
            _add_field(para, "NUMPAGES")
    if border_hex:
        SE._add_para_border(para._p.get_or_add_pPr(), border_side, border_hex, 1.0)


def _apply_header_footer(doc, spec: "StyleSpec", title_text: str, doc_code: str,
                         warnings: list[str]) -> None:
    hf = spec.header_footer
    # Footer confidentiality line: prefer clean literal text from the spec, else
    # a canonical phrase if the guideline mentioned confidentiality.
    footer_left = hf.footer_left if not _looks_like_rule_text(hf.footer_left) else None
    if not footer_left:
        blob = " ".join([hf.footer_left or "", hf.header_left or "", " ".join(spec.notes)]).lower()
        if "confidential" in blob:
            footer_left = "CONFIDENTIAL — For Internal Use Only"
    header_left = _title_case(title_text) if title_text else (
        hf.header_left if not _looks_like_rule_text(hf.header_left) else "")
    width = _content_width_in(doc)
    try:
        for section in doc.sections:
            if header_left or doc_code:
                _write_hf(section.header, header_left, doc_code, hf, page_field=False,
                          border_hex=hf.header_border_color_hex or spec.accent_color_hex,
                          border_side="bottom", tab_width_in=width)
            _write_hf(section.footer, footer_left or "", "", hf, page_field=True,
                      border_hex=hf.footer_border_color_hex, border_side="top",
                      tab_width_in=width)
    except Exception as e:  # noqa: BLE001
        warnings.append(f"header/footer skipped: {e}")


# ---------------------------------------------------------------------------
# Named-style configuration (mirrors guideline application's style layer)
# ---------------------------------------------------------------------------

def _configure_styles(doc, spec: "StyleSpec", warnings: list[str], normalize_fonts: bool) -> None:
    body = spec.body
    body_font = body.font
    if body_font or body.size_pt:
        try:
            SE._set_docdefaults_font(doc, body_font or "Calibri", body.size_pt)
        except Exception as e:  # noqa: BLE001
            warnings.append(f"doc defaults skipped: {e}")
    try:
        normal = doc.styles["Normal"]
        if body_font or body.size_pt or body.color_hex:
            SE._set_style_font(normal, body_font, body.size_pt, body.color_hex)
        SE._apply_alignment_spacing(normal, body.alignment, body.space_after_pt,
                                    body.line_spacing, body.space_before_pt)
    except KeyError:
        pass
    for rule in spec.headings:
        name = rule.style_name or (f"Heading {rule.level}" if rule.level else None)
        if not name:
            continue
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        SE._set_style_font(st, rule.font or body_font, rule.size_pt, rule.color_hex)
        SE._set_style_bool(st, bold=rule.bold, italic=rule.italic, underline=rule.underline)
        SE._apply_alignment_spacing(st, rule.alignment, rule.space_after_pt, None, rule.space_before_pt)
        if rule.bottom_border:
            try:
                SE._add_para_border(
                    st.element.get_or_add_pPr(), "bottom",
                    rule.border_color_hex or rule.color_hex or spec.accent_color_hex,
                    rule.border_width_pt or 1.0,
                )
            except Exception as e:  # noqa: BLE001
                warnings.append(f"heading border skipped: {e}")
    if normalize_fonts and body_font:
        try:
            SE._normalize_fonts(doc, body_font)
        except Exception as e:  # noqa: BLE001
            warnings.append(f"font normalization skipped: {e}")
    SE._apply_page_spec(doc, spec.page, warnings)


# ---------------------------------------------------------------------------
# Blank-line tidy
# ---------------------------------------------------------------------------

def _collapse_blanks(doc) -> None:
    """Collapse runs of 2+ empty paragraphs to one and drop trailing blanks, so
    the document never shows stacked blank lines / weird gaps."""
    prev_blank = False
    for p in list(doc.paragraphs):
        is_blank = not (p.text or "").strip() and p._p.find(".//" + qn("w:numPr")) is None
        if is_blank and prev_blank:
            p._p.getparent().remove(p._p)
            continue
        prev_blank = is_blank


# ---------------------------------------------------------------------------
# Orchestrator
# ---------------------------------------------------------------------------

@dataclass
class StructureOutcome:
    docx_bytes: bytes
    warnings: list[str]
    plan_summary: dict


def apply_structured_style(
    content_bytes: bytes,
    spec: "StyleSpec",
    *,
    transplant_style_bytes: Optional[bytes] = None,
    normalize_fonts: bool = True,
    use_llm: Optional[bool] = None,
) -> StructureOutcome:
    """Recognise the content's structure and apply ``spec`` to it in place.

    ``transplant_style_bytes`` (the example .docx) is merged at the style layer
    first so named styles (Heading N, List Paragraph, theme) carry the example's
    exact look before the structural upgrade lands.
    """
    warnings: list[str] = []

    if transplant_style_bytes is not None:
        try:
            content_bytes = SE._transplant_docx_styles(
                content_bytes, transplant_style_bytes, promote_headings=False
            )
        except Exception as e:  # noqa: BLE001
            warnings.append(f"style transplant skipped: {e}")

    doc = Document(io.BytesIO(content_bytes))
    paras = list(doc.paragraphs)

    # 1) Understand the document.
    plan = recognize_structure(doc, use_llm=use_llm)

    # 2) Capture masthead text BEFORE mutating (for the derived header/footer).
    title_text = ""
    doc_code = ""
    for idx, role in plan.roles.items():
        if idx >= len(paras):
            continue
        if role.kind == TITLE and not title_text:
            title_text = (paras[idx].text or "").strip()
        if role.kind == METADATA and not doc_code:
            m = _DOC_CODE_RE.search(paras[idx].text or "")
            if m:
                doc_code = m.group(1)

    # 3) Configure the named-style layer + page from the spec.
    _configure_styles(doc, spec, warnings, normalize_fonts)

    # 4) Masthead direct formatting.
    title_look = spec.title
    sub_look = spec.subtitle
    meta_look = spec.metadata
    for idx, role in plan.roles.items():
        if idx >= len(paras) or idx in plan.consumed_idxs:
            continue
        if role.kind == TITLE:
            _apply_masthead(paras[idx], title_look, default_align="center", default_bold=True)
        elif role.kind == SUBTITLE:
            _apply_masthead(paras[idx], sub_look, default_align="center", default_bold=True)
        elif role.kind == METADATA:
            # The doc-number line in the masthead is centred; a standalone
            # metadata block (Department / Prepared By / …) reads better left —
            # force it left even when the spec's metadata look says centre.
            _apply_masthead(paras[idx], meta_look, default_align="center",
                            align_override=None if role.masthead else "left")

    # 5) Headings.
    for idx, role in plan.roles.items():
        if idx < len(paras) and idx not in plan.consumed_idxs and role.kind == HEADING:
            _assign_heading(doc, paras[idx], role.level)

    # 6) Lists (with per-group restart numbering).
    numbering = _NumberingManager(doc, spec)
    ordered = sorted(i for i in plan.roles if i not in plan.consumed_idxs)
    groups = _list_groups(plan, ordered)
    if groups and not numbering.available():
        # _apply_list_item silently falls back to literal glyphs — say so once.
        warnings.append(
            "numbering part unavailable — list items rendered with literal "
            "bullet glyphs instead of real Word numbering"
        )
    drop: list = []
    for idx in ordered:
        role = plan.role(idx)
        if role.kind != LIST_ITEM or idx >= len(paras):
            continue
        if not _apply_list_item(doc, paras[idx], role, groups.get(idx, 0), numbering, spec):
            drop.append(paras[idx]._p)
    for p_el in drop:
        try:
            p_el.getparent().remove(p_el)
        except Exception:
            pass

    # 7) Tables — build last (mutates the paragraph set), using captured refs.
    for tb in sorted(plan.tables, key=lambda b: b.start_idx, reverse=True):
        if tb.start_idx >= len(paras):
            continue
        try:
            _build_table(doc, paras[tb.start_idx]._p, tb.rows, tb.has_header, spec, warnings)
        except Exception as e:  # noqa: BLE001
            warnings.append(f"table build skipped at idx {tb.start_idx}: {e}")
            continue
        for k in range(tb.start_idx, tb.end_idx + 1):
            if k < len(paras):
                try:
                    paras[k]._p.getparent().remove(paras[k]._p)
                except Exception:
                    pass

    # 7b) Approval/metadata blocks ("Department: QA" / "Prepared By: …")
    #     become styled tables — matching how hand-formatted SOPs render them.
    _build_metadata_tables(doc, plan, paras, spec, warnings)

    # 8) Header / footer (derived text, spec styling).
    _apply_header_footer(doc, spec, title_text, doc_code, warnings)

    # 9) Tidy blank lines.
    _collapse_blanks(doc)

    out = io.BytesIO()
    doc.save(out)
    return StructureOutcome(docx_bytes=out.getvalue(), warnings=warnings, plan_summary=plan.summary())
