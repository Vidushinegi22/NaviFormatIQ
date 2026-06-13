"""
Style-preserving Word emission via docxtpl.

Two render modes:

  1. **Jinja-marker templates** — the template's body contains ``{{ slot_id }}``
     placeholders (and optionally ``{%tr for row in table_xxx %}`` row loops).
     We render with ``docxtpl.DocxTemplate``; the template owns every style.

  2. **Heading-driven templates** — no markers. We open the original .docx
     with python-docx and, for each heading whose ``slot_id`` matches a key
     in ``rewritten``, replace the body paragraphs that follow the heading
     (up to the next heading) with the rewritten text. We never touch the
     heading paragraph itself, so its run-level formatting survives.

The output is always ``bytes`` (a complete .docx file).
"""

from __future__ import annotations

import base64
import io
import re
from copy import deepcopy
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from app.schemas.document_model import HeadingSlot, TemplateFingerprint


# Markdown-ish list markers the rewriter emits so structure survives regeneration.
_BULLET_MD = re.compile(r"^(\s*)[-*•◦▪‣]\s+(.*\S)\s*$")
_NUM_MD = re.compile(r"^(\s*)\(?\s*(?:\d{1,3}|[ivxlcdmIVXLCDM]+|[a-zA-Z])\s*[.)\]]\s+(.*\S)\s*$")


def _classify_md_line(raw: str) -> tuple[str, int, str]:
    """Classify one rewritten line as ('bullet'|'numbered'|'body', level, text)."""
    indent = len(raw) - len(raw.lstrip(" "))
    level = min(indent // 2, 3)
    m = _BULLET_MD.match(raw)
    if m:
        return "bullet", level, m.group(2).strip()
    m = _NUM_MD.match(raw)
    if m:
        return "numbered", level, m.group(2).strip()
    return "body", 0, raw.strip()


def _is_table_row(line: str) -> bool:
    """A pipe-delimited markdown row, e.g. ``| Version | Date |``."""
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _body_lines(text: str) -> list[str]:
    """Non-empty lines of body text, EXCLUDING markdown table rows.

    Table content is shown in the review/preview as ``| a | b |`` rows, but the
    real table is carried over verbatim by the renderer — so we drop those rows
    here to avoid duplicating the table as stray pipe-text paragraphs."""
    return [ln for ln in (text or "").split("\n") if ln.strip() and not _is_table_row(ln)]


def _remove_numpr(p: Paragraph) -> None:
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None:
        for np in pPr.findall(qn("w:numPr")):
            pPr.remove(np)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def render_template(
    fingerprint: TemplateFingerprint,
    rewritten: dict[str, str],
    tables: Optional[dict[str, list[list[str]]]] = None,
    edits: Optional[dict] = None,
    field_updates: Optional[dict] = None,
) -> bytes:
    """Render the template using ``rewritten`` (slot_id → body text).

    ``edits`` carries reviewer structural changes (all optional):
      - ``titles``       : {slot_id -> new heading text} (rename a heading)
      - ``removed``      : [slot_id, …]                   (drop heading + body)
      - ``new_sections`` : [{after_slot_id, title, level, text}, …]  (insert)
    """
    from app.services.formatting.text_safe import normalize_list_markers, strip_todo_placeholders, xml_safe

    # Scrub XML-illegal control characters (Word soft breaks, stray LLM bytes),
    # normalise odd bullet markers, and strip any [TODO] placeholders that the
    # RAG filler left behind BEFORE the text touches the document — otherwise a
    # single bad byte crashes the whole render or flattens a list, and leftover
    # TODOs leak into the final export.
    rewritten = {k: xml_safe(normalize_list_markers(strip_todo_placeholders(v))) for k, v in (rewritten or {}).items()}
    if tables:
        tables = {
            k: [[xml_safe(c) for c in row] for row in grid]
            for k, grid in tables.items()
        }
    edits = _sanitize_edits(edits)
    if not fingerprint.template_b64:
        # PDF-sourced template or callers stripped the bytes — fall back to
        # a brand-new doc that simulates the structure as best we can.
        return _render_synthesized(fingerprint, rewritten, tables or {}, edits)

    raw = base64.b64decode(fingerprint.template_b64)
    # Guard: the docxtpl / heading-walk renderers open these bytes as a .docx
    # (a ZIP package). If a non-OOXML payload slipped through (e.g. a PDF's
    # bytes), synthesize from the fingerprint rather than crash with BadZipFile.
    if raw[:4] != b"PK\x03\x04":
        return _render_synthesized(fingerprint, rewritten, tables or {}, edits)
    if _has_jinja_markers(fingerprint):
        return _render_via_docxtpl(raw, fingerprint, rewritten, tables or {})
    return _render_via_heading_walk(raw, fingerprint, rewritten, tables or {}, edits, field_updates)


def _sanitize_edits(edits: Optional[dict]) -> dict:
    """Normalise + XML-scrub the reviewer edit structure."""
    from app.services.formatting.text_safe import xml_safe

    base = {"titles": {}, "removed": [], "new_sections": []}
    if not edits:
        return base
    base["titles"] = {
        str(k): xml_safe(v) for k, v in (edits.get("titles") or {}).items() if v
    }
    base["removed"] = [str(s) for s in (edits.get("removed") or [])]
    for ns in edits.get("new_sections") or []:
        try:
            level = int(ns.get("level") or 1)
        except (TypeError, ValueError):
            level = 1
        base["new_sections"].append({
            "after_slot_id": ns.get("after_slot_id"),
            "title": xml_safe(ns.get("title") or "New Section"),
            "level": max(1, min(level, 9)),
            "text": xml_safe(ns.get("text") or ""),
        })
    return base


# ---------------------------------------------------------------------------
# Mode detection
# ---------------------------------------------------------------------------

def _has_jinja_markers(fingerprint: TemplateFingerprint) -> bool:
    return any(s.placeholder_marker for s in fingerprint.heading_hierarchy)


# ---------------------------------------------------------------------------
# Mode 1 — docxtpl rendering
# ---------------------------------------------------------------------------

def _render_via_docxtpl(
    template_bytes: bytes,
    fingerprint: TemplateFingerprint,
    rewritten: dict[str, str],
    tables: dict[str, list[list[str]]],
) -> bytes:
    from docxtpl import DocxTemplate

    tpl = DocxTemplate(io.BytesIO(template_bytes))
    context = {s.slot_id: rewritten.get(s.slot_id, "") for s in fingerprint.heading_hierarchy}
    # Tables expected as ``table_<slot_id>`` in the Jinja context
    for slot_id, grid in tables.items():
        context[f"table_{slot_id}"] = grid
    tpl.render(context)

    out = io.BytesIO()
    tpl.save(out)
    return out.getvalue()


# ---------------------------------------------------------------------------
# Mode 2 — heading-driven body replacement
# ---------------------------------------------------------------------------

def _render_via_heading_walk(
    template_bytes: bytes,
    fingerprint: TemplateFingerprint,
    rewritten: dict[str, str],
    tables: dict[str, list[list[str]]],
    edits: Optional[dict] = None,
    field_updates: Optional[dict] = None,
) -> bytes:
    from app.services.style.structure_apply import _NumberingManager

    edits = edits or {"titles": {}, "removed": [], "new_sections": []}
    titles: dict[str, str] = edits.get("titles") or {}
    removed: set[str] = set(edits.get("removed") or [])

    doc = Document(io.BytesIO(template_bytes))
    slots = fingerprint.heading_hierarchy
    slot_by_title = {s.title.strip(): s for s in slots}
    numbering = _NumberingManager(doc)

    body_paragraphs = list(doc.paragraphs)
    heading_indices: list[tuple[int, HeadingSlot]] = []
    for idx, para in enumerate(body_paragraphs):
        text = (para.text or "").strip()
        if text in slot_by_title:
            heading_indices.append((idx, slot_by_title[text]))

    # Stable heading-element refs + the kept order, for inserting new sections.
    slot_heading_el: dict[str, object] = {}
    kept_order: list[str] = []

    for n, (idx, slot) in enumerate(heading_indices):
        next_idx = (
            heading_indices[n + 1][0]
            if n + 1 < len(heading_indices)
            else len(body_paragraphs)
        )
        heading_para = body_paragraphs[idx]
        body = body_paragraphs[idx + 1 : next_idx]
        slot_heading_el[slot.slot_id] = heading_para._p

        # Reviewer removed this whole section → delete the heading AND its body.
        if slot.slot_id in removed:
            for p in [heading_para, *body]:
                try:
                    p._p.getparent().remove(p._p)
                except Exception:
                    pass
            continue

        kept_order.append(slot.slot_id)

        # Reviewer renamed the heading → rewrite its text, preserving the style.
        new_title = titles.get(slot.slot_id)
        if new_title and new_title.strip():
            _set_heading_text(heading_para, new_title.strip())

        new_text = rewritten.get(slot.slot_id, "")
        if not new_text:
            continue
        end_el = (
            body_paragraphs[next_idx]._p if next_idx < len(body_paragraphs) else None
        )
        _replace_body_region(
            heading_para, body, end_el, new_text, doc, numbering, group_base=n + 1
        )

    # Insert reviewer-added sections after the section they were placed under.
    for j, ns in enumerate(edits.get("new_sections") or []):
        after = ns.get("after_slot_id")
        anchor_el = None
        if after in kept_order:
            pos = kept_order.index(after)
            nxt = kept_order[pos + 1] if pos + 1 < len(kept_order) else None
            anchor_el = slot_heading_el.get(nxt) if nxt else None
        _insert_section(
            doc, anchor_el, ns["title"], ns["level"], ns["text"], numbering,
            group_base=900 + j,
        )

    # Auto field-updates (version bump, date, revision-history row).
    _apply_field_updates(doc, field_updates)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _set_paragraph_text(p: Paragraph, text: str) -> None:
    """Replace a paragraph's text in place, preserving its first run's style."""
    runs = list(p.runs)
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


# A version number immediately preceded by a version/revision keyword. Used to
# bump the masthead version robustly, regardless of the exact spacing/punctuation
# between the keyword and the number in the live document.
_VER_TOKEN_RE = re.compile(
    r"(?i)\b((?:version|revision|ver|rev)\.?\s*[:.\-]?\s*)(\d+(?:\.\d+){0,2})\b"
)


def _apply_one_replacement(text: str, r: dict) -> str:
    """Apply a single field replacement to ``text``.

    Version/date replacements carry structured ``kind``/``old``/``new`` so we can
    match them tolerantly: extracted masthead text (where the find-strings were
    derived) often differs from the live document in whitespace/punctuation, so a
    literal ``find in text`` misses. Falls back to the literal find/replace for
    any legacy replacement that lacks the structured fields.
    """
    kind, old, new = r.get("kind"), r.get("old"), r.get("new")
    if kind == "version" and old and new:
        return _VER_TOKEN_RE.sub(
            lambda m: (m.group(1) + new) if m.group(2) == old else m.group(0), text
        )
    if kind == "date" and old and new:
        if old in text:
            return text.replace(old, new)
        # Tolerant fallback: the day/month/year tokens are specific enough to be
        # safe, but the separators may differ (01-Jun-2025 vs 01 Jun 2025).
        toks = [t for t in re.split(r"[-/.\s]+", old) if t]
        if len(toks) >= 2:
            pat = re.compile(r"[-/.\s]+".join(re.escape(t) for t in toks))
            return pat.sub(lambda _m: new, text)
        return text
    find, rep = r.get("find"), r.get("replace")
    if find and rep is not None and find in text:
        return text.replace(find, rep)
    return text


def _apply_field_updates(doc, field_updates: Optional[dict]) -> None:
    """Apply version/date replacements and append a revision-history row.

    Scans the body, all tables, and every header/footer (mastheads with the
    Version/Effective-Date line frequently live in the running header), so the
    auto-updates land wherever the fields actually are.
    """
    if not field_updates:
        return
    from app.services.formatting.text_safe import xml_safe

    reps = field_updates.get("replacements") or []

    def _apply_to(paragraphs) -> None:
        for p in paragraphs:
            txt = p.text or ""
            new = txt
            for r in reps:
                new = _apply_one_replacement(new, r)
            if new != txt:
                _set_paragraph_text(p, xml_safe(new))

    def _scan_tables(tables) -> None:
        for t in tables:
            if _is_revision_history_table(t):
                continue
            for row in t.rows:
                for cell in row.cells:
                    _apply_to(cell.paragraphs)

    if reps:
        _apply_to(doc.paragraphs)
        _scan_tables(doc.tables)
        for section in doc.sections:
            for hf in (
                section.header,
                section.first_page_header,
                section.even_page_header,
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ):
                if hf is None:
                    continue
                try:
                    _apply_to(hf.paragraphs)
                    _scan_tables(hf.tables)
                except Exception:  # noqa: BLE001 — header/footer access is best-effort
                    continue

    rev = field_updates.get("revision")
    if rev and rev.get("row"):
        _append_revision_row(doc, rev.get("row"))


def _is_revision_history_table(table) -> bool:
    if not getattr(table, "rows", None):
        return False
    try:
        header = " ".join((c.text or "").strip().lower() for c in table.rows[0].cells)
    except Exception:
        return False
    return (
        "version" in header
        and "date" in header
        and any(k in header for k in ("change", "summary", "description", "author"))
    )


def _append_revision_row(doc, row: list[str]) -> None:
    """Append ``row`` to the document's revision-history table (matched by its
    header), styling the new cells to match the table body."""
    from app.services.formatting.text_safe import xml_safe

    target = None
    for t in doc.tables:
        if not t.rows:
            continue
        if _is_revision_history_table(t):
            target = t
            break
    if target is None:
        return
    try:
        ref_row = _revision_row_to_clone(target)
        new_tr = deepcopy(ref_row._tr)
        target.rows[-1]._tr.addnext(new_tr)
        new_cells = target.rows[-1].cells
        for i in range(min(len(new_cells), len(row))):
            _set_cell_text_preserving_style(new_cells[i], xml_safe(str(row[i])))
    except Exception:
        pass


def _cell_shading_fill(cell) -> Optional[str]:
    tc_pr = cell._tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def _revision_row_to_clone(table):
    """Pick a data row whose formatting should be reused for the appended row.

    If the last two body rows alternate shading, clone the second-last body row
    so the new row continues the visible striping. Otherwise clone the last row.
    """
    if len(table.rows) >= 3:
        last_fill = _cell_shading_fill(table.rows[-1].cells[0])
        prev_fill = _cell_shading_fill(table.rows[-2].cells[0])
        if last_fill != prev_fill:
            return table.rows[-2]
    return table.rows[-1]


def _set_cell_text_preserving_style(cell, text: str) -> None:
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        p = cell.add_paragraph()
    else:
        p = paragraphs[0]
        for extra in paragraphs[1:]:
            try:
                extra._p.getparent().remove(extra._p)
            except Exception:
                pass
    _set_paragraph_text(p, text)


def _replace_body_region(
    heading_para: Paragraph,
    body_paras: list[Paragraph],
    end_el,
    new_text: str,
    doc,
    numbering,
    *,
    group_base: int,
) -> None:
    """Replace a heading's body while keeping every REAL table in place.

    The naive replacement collapses all of a section's paragraphs to the top
    of the region; any paragraph that used to sit between two tables vanishes,
    the tables become adjacent siblings, and the LibreOffice field-refresh
    pass then merges them into one mangled grid. Here the region's paragraphs
    are segmented BY the tables between them, the rewritten text is segmented
    by its markdown table-row runs (the rewriter carries tables through in
    place as ``| a | b |`` lines), and segment k of the text lands in segment
    k of the paragraphs — so prose stays on the correct side of every table
    and adjacent tables always keep a separator."""
    region: list[tuple[str, object]] = []
    el = heading_para._p.getnext()
    while el is not None and el is not end_el:
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag in ("p", "tbl"):
            region.append((tag, el))
        el = el.getnext()

    body_by_el = {id(p._p): p for p in body_paras}
    para_segments: list[list[Paragraph]] = [[]]
    table_els: list[object] = []
    for tag, rel in region:
        if tag == "tbl":
            para_segments.append([])
            table_els.append(rel)
        elif id(rel) in body_by_el:
            para_segments[-1].append(body_by_el[id(rel)])

    if not table_els:
        _replace_body_paragraphs(body_paras, new_text, doc, numbering, group_base=group_base)
        return

    # Split the rewritten text at its table-row runs.
    text_segments: list[list[str]] = [[]]
    in_table = False
    for ln in (new_text or "").split("\n"):
        if _is_table_row(ln):
            if not in_table:
                text_segments.append([])
                in_table = True
            continue
        in_table = False
        text_segments[-1].append(ln)

    if len(text_segments) != len(para_segments):
        # The model merged/dropped/added a table marker — keep it safe: all
        # prose goes into the first segment; the others become empty spacers
        # (so adjacent tables still never touch).
        merged = [ln for seg in text_segments for ln in seg]
        text_segments = [merged] + [[] for _ in table_els]

    for k, (paras, lines) in enumerate(zip(para_segments, text_segments)):
        seg_text = "\n".join(lines).strip()
        if not paras:
            if not seg_text:
                continue
            # Nothing to anchor on (table directly after heading/table) —
            # mint an empty paragraph in the right spot.
            anchor_el = OxmlElement("w:p")
            if k == 0:
                heading_para._p.addnext(anchor_el)
            else:
                table_els[k - 1].addnext(anchor_el)
            paras = [Paragraph(anchor_el, heading_para._parent)]
        _replace_body_paragraphs(
            paras, seg_text, doc, numbering, group_base=group_base * 100 + k
        )


def _replace_body_paragraphs(
    paragraphs: list[Paragraph],
    new_text: str,
    doc,
    numbering,
    *,
    group_base: int,
) -> None:
    """Replace the body paragraphs between two headings with ``new_text``,
    reconstructing real bullet/numbered lists from markdown-style markers.
    """
    if not paragraphs:
        return

    items = [_classify_md_line(ln) for ln in _body_lines(new_text)]

    original_protos = []
    for p in paragraphs:
        original_protos.append({
            "p_xml": deepcopy(p._p),
            "text": (p.text or "").strip(),
            "runs": [{"text": getattr(r, "text", ""), "r_xml": r._r} for r in p.runs],
        })

    anchor = paragraphs[0]

    # Drop every template paragraph in the range except the anchor (used to place new ones).
    for p in paragraphs[1:]:
        try:
            p._p.getparent().remove(p._p)
        except Exception:
            pass

    if not items:
        # Clear anchor so it's an empty spacer
        for r in list(anchor.runs):
            r._r.getparent().remove(r._r)
        _remove_numpr(anchor)
        return

    group_ids = _group_ids(items, group_base)

    def _normalize(t):
        return re.sub(r'[^a-zA-Z0-9]', '', t.lower())

    cur = anchor
    for i, (kind, level, text) in enumerate(items):
        norm_t = _normalize(text)
        best_proto_info = None
        
        if norm_t:
            for p_info in original_protos:
                if _normalize(p_info["text"]) == norm_t:
                    best_proto_info = p_info
                    break
                    
        if not best_proto_info and norm_t:
            for p_info in original_protos:
                p_norm = _normalize(p_info["text"])
                if p_norm and (norm_t.startswith(p_norm) or p_norm.startswith(norm_t)):
                    best_proto_info = p_info
                    break

        if not best_proto_info:
            idx = min(i, len(original_protos) - 1)
            best_proto_info = original_protos[idx]

        new_p_xml = deepcopy(best_proto_info["p_xml"])
        cur._p.addnext(new_p_xml)
        p = Paragraph(new_p_xml, anchor._parent)
        
        for r in list(p.runs):
            r._r.getparent().remove(r._r)
        _remove_numpr(p)
        
        _write_body_line(p, kind, level, text, doc, numbering, group_ids[i], best_proto_info["runs"])
        cur = p

    try:
        anchor._p.getparent().remove(anchor._p)
    except Exception:
        pass


def _write_body_line(p, kind, level, text, doc, numbering, group_id, best_runs) -> None:
    if kind in ("bullet", "numbered"):
        if not getattr(p.style, 'name', '').startswith('List'):
            try:
                p.style = doc.styles["List Paragraph"]
            except KeyError:
                pass
        if numbering is not None and numbering.available():
            numbering.attach(p, "bullet" if kind == "bullet" else "numbered", level, group_id or 0)
        pf = p.paragraph_format
        pf.left_indent = Inches(0.5 + level * 0.25)
        pf.first_line_indent = Inches(-0.25)

    if text:
        # Fast path — the line is the same as (or a light edit of) its source
        # paragraph: write it as one run carrying the source's dominant run
        # formatting, so a bold label line stays bold even when the model
        # nudged its punctuation ("Schema : cdm" → "Schema: cdm").
        proto_full = "".join(r["text"] for r in best_runs)
        norm = lambda t: re.sub(r"[^a-zA-Z0-9]", "", t.lower())  # noqa: E731
        if best_runs and norm(proto_full) == norm(text):
            dominant = max(
                (r for r in best_runs if r["text"].strip()),
                key=lambda r: len(r["text"]),
                default=None,
            )
            rPr = (
                dominant["r_xml"].find(qn("w:rPr")) if dominant is not None else None
            )
            run = p.add_run(text)
            if rPr is not None:
                run._r.insert(0, deepcopy(rPr))
            return

        styled_chunks = []
        for r_info in best_runs:
            r_text = r_info["text"]
            if len(r_text.strip()) < 2:
                continue
            rPr = r_info["r_xml"].find(qn("w:rPr"))
            if rPr is not None:
                styled_chunks.append((r_text, rPr))
                
        styled_chunks.sort(key=lambda x: len(x[0]), reverse=True)
        
        segments = [(text, None)]
        for s_text, rPr in styled_chunks:
            pattern_parts = []
            if s_text and s_text[0].isalnum():
                pattern_parts.append(r'(?<![a-zA-Z0-9])')
            pattern_parts.append(re.escape(s_text))
            if s_text and s_text[-1].isalnum():
                pattern_parts.append(r'(?![a-zA-Z0-9])')
            
            pattern = "".join(pattern_parts)
            regex = re.compile(f'({pattern})')
            
            new_segments = []
            for seg_text, seg_rPr in segments:
                if seg_rPr is not None:
                    new_segments.append((seg_text, seg_rPr))
                    continue
                
                parts = regex.split(seg_text)
                for idx, part in enumerate(parts):
                    if not part:
                        continue
                    if idx % 2 == 1:
                        new_segments.append((part, rPr))
                    else:
                        new_segments.append((part, None))
            segments = new_segments
            
        for seg_text, rPr in segments:
            run = p.add_run(seg_text)
            if rPr is not None:
                run._r.append(deepcopy(rPr))


def _group_ids(items, group_base: int) -> list[Optional[int]]:
    """Assign list-numbering group ids: each contiguous numbered run gets its
    own group (so it restarts at 1); bullets share one (their number is unused)."""
    out: list[Optional[int]] = []
    counter = 0
    prev = None
    for kind, _lvl, _txt in items:
        if kind == "numbered":
            if prev != "numbered":
                counter += 1
            out.append(group_base * 1000 + counter)
        elif kind == "bullet":
            out.append(group_base * 1000)
        else:
            out.append(None)
        prev = kind
    return out


def _set_heading_text(heading_para: Paragraph, title: str) -> None:
    """Rewrite a heading's text in place, preserving its run/paragraph style."""
    runs = list(heading_para.runs)
    if runs:
        runs[0].text = title
        for r in runs[1:]:
            r.text = ""
    else:
        heading_para.add_run(title)


def _heading_style_name(doc, level: int) -> Optional[str]:
    """The best available 'Heading N' style name for a new section heading."""
    level = max(1, min(level, 9))
    for name in (f"Heading {level}", "Heading 1", "Heading 2"):
        try:
            doc.styles[name]
            return name
        except KeyError:
            continue
    return None


def _insert_section(doc, anchor_el, title: str, level: int, text: str, numbering, *, group_base: int) -> None:
    """Insert a brand-new section (heading + list-aware body). If ``anchor_el``
    is given the section is placed immediately before it; otherwise appended."""
    style_name = _heading_style_name(doc, level)
    created = []

    heading = doc.add_paragraph(title or "New Section")
    if style_name:
        try:
            heading.style = doc.styles[style_name]
        except KeyError:
            pass
    created.append(heading._p)

    items = [_classify_md_line(ln) for ln in _body_lines(text)]
    group_ids = _group_ids(items, group_base)
    for i, (kind, lvl, body) in enumerate(items):
        bp = doc.add_paragraph()
        _write_body_line(bp, kind, lvl, body, doc, numbering, group_ids[i], None)
        created.append(bp._p)

    # doc.add_paragraph appends at the end (before sectPr) in order. When an
    # anchor is supplied, relocate the new paragraphs to sit just before it.
    if anchor_el is not None:
        for el in created:
            anchor_el.addprevious(el)


# ---------------------------------------------------------------------------
# Fallback for PDF-sourced templates (no .docx bytes available)
# ---------------------------------------------------------------------------

def _render_synthesized(
    fingerprint: TemplateFingerprint,
    rewritten: dict[str, str],
    tables: dict[str, list[list[str]]],
    edits: Optional[dict] = None,
) -> bytes:
    """Last-resort rendering: build a new .docx from scratch.

    Style fidelity is necessarily limited — the source template was a PDF or
    its bytes were stripped — so we reuse the existing apply_styling path,
    which produces a reasonable Word document. Reviewer edits (renames,
    removals, new sections) are honoured here too.
    """
    from app.services.formatting.formater_apply import apply_styling
    from app.schemas.document_model import (
        ContentElement,
        DocumentContent,
        DocumentMetadata,
        ElementType,
        ListType,
        TextRun,
    )

    edits = edits or {"titles": {}, "removed": [], "new_sections": []}
    titles = edits.get("titles") or {}
    removed = set(edits.get("removed") or [])

    def _body_elements(text: str) -> list[ContentElement]:
        out: list[ContentElement] = []
        for raw in _body_lines(text):
            kind, level, body = _classify_md_line(raw)
            if kind in ("bullet", "numbered"):
                out.append(ContentElement(
                    type=ElementType.LIST_ITEM,
                    content=[TextRun(text=body)],
                    list_type=ListType.BULLET if kind == "bullet" else ListType.NUMBERED,
                    list_level=level,
                ))
            else:
                out.append(ContentElement(
                    type=ElementType.PARAGRAPH, content=[TextRun(text=body)]
                ))
        return out

    # New sections, grouped by the slot they should follow.
    new_after: dict[Optional[str], list[dict]] = {}
    for ns in edits.get("new_sections") or []:
        new_after.setdefault(ns.get("after_slot_id"), []).append(ns)

    def _new_section_elements(ns: dict) -> list[ContentElement]:
        els = [ContentElement(
            type=ElementType.HEADING, level=int(ns.get("level") or 1),
            content=[TextRun(text=ns.get("title") or "New Section")],
        )]
        els.extend(_body_elements(ns.get("text") or ""))
        return els

    elements: list[ContentElement] = []
    for slot in fingerprint.heading_hierarchy:
        if slot.slot_id in removed:
            continue
        elements.append(ContentElement(
            type=ElementType.HEADING,
            level=slot.level,
            content=[TextRun(text=titles.get(slot.slot_id) or slot.title)],
        ))
        elements.extend(_body_elements(rewritten.get(slot.slot_id, "")))
        for ns in new_after.get(slot.slot_id, []):
            elements.extend(_new_section_elements(ns))
    # Any new sections with no (or an unknown) anchor go at the end.
    placed = {id(ns) for lst in new_after.values() for ns in lst
              if ns.get("after_slot_id") in {s.slot_id for s in fingerprint.heading_hierarchy}}
    for ns in edits.get("new_sections") or []:
        if id(ns) not in placed:
            elements.extend(_new_section_elements(ns))

    content = DocumentContent(
        metadata=DocumentMetadata(
            source_file=fingerprint.metadata.source_file or "synthesized.docx"
        ),
        elements=elements,
    )
    stream = apply_styling(content, fingerprint.style_registry)
    stream.seek(0)
    return stream.read()
