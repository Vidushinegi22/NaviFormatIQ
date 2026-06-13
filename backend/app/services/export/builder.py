"""Catalog + producers for run exports.

Every function here is SYNC and pure of DB/HTTP. The route layer resolves the
``Run`` + its ``Artifact`` rows, computes a friendly ``base_name``, and passes
plain dicts in; it persists any generated PDF and streams the bytes back.

Deliverables (see ``list_exports``):
  - document_docx / document_pdf  — the finished document
  - formatting_json               — reusable styling/formatting spec
  - formatting_report             — human-readable .docx of the formatting
  - content_json                  — extracted structured content
  - change_report                 — section-by-section diff (regenerate/style)
  - compliance_report             — findings + coverage + score (compliance)
"""
from __future__ import annotations

import io
import json
import re
from dataclasses import asdict, dataclass
from typing import Any, Optional

from app.agents.nodes.common import DOCX_MIME, PDF_MIME, load_bytes

JSON_MIME = "application/json"

# Category keys — the frontend groups + orders sections by this list.
CAT_DOCUMENT = "document"
CAT_FORMATTING = "formatting"
CAT_DATA = "data"
CAT_REPORT = "report"
CATEGORY_ORDER = [CAT_DOCUMENT, CAT_FORMATTING, CAT_DATA, CAT_REPORT]


class ExportError(RuntimeError):
    """A deliverable could not be produced (e.g. no rendered document)."""


@dataclass
class ExportSpec:
    id: str
    category: str
    label: str
    description: str
    fmt: str                       # docx | pdf | json
    filename: str
    available: bool = True
    reason: Optional[str] = None   # why unavailable (shown to the user)
    artifact_id: Optional[str] = None
    size_bytes: Optional[int] = None

    def to_dict(self) -> dict:
        d = asdict(self)
        d["format"] = d.pop("fmt")
        return d


# ── filenames ────────────────────────────────────────────────────────────────

_FILENAMES = {
    "document_docx": "{base}.docx",
    "document_pdf": "{base}.pdf",
    "formatting_json": "{base}-formatting.json",
    "formatting_report": "{base}-formatting-report.docx",
    "content_json": "{base}-content.json",
    "change_report": "{base}-changes.docx",
    "compliance_report": "{base}-compliance.docx",
    # New audit engine — pre-built report artifacts (streamed as-is).
    "compliance_report_pdf": "{base}-compliance.pdf",
    "compliance_report_docx": "{base}-compliance.docx",
    "compliance_report_json": "{base}-compliance.json",
    "compliance_report_csv": "{base}-compliance.csv",
}

# Pre-built compliance report artifacts → (export id, category, label, description, format).
_COMPLIANCE_ARTIFACTS = [
    ("compliance_report_pdf", "compliance_report_pdf", CAT_REPORT, "Compliance report (PDF)",
     "The full audit: scores, charts, and every finding with guideline citations.", "pdf"),
    ("compliance_report_docx", "compliance_report_docx", CAT_REPORT, "Compliance report (Word)",
     "The full audit as an editable Word document.", "docx"),
    ("compliance_report_json", "compliance_report_json", CAT_DATA, "Compliance findings (JSON)",
     "Machine-readable scores + every finding for downstream tooling.", "json"),
    ("compliance_report_csv", "compliance_report_csv", CAT_DATA, "Compliance findings (CSV)",
     "All findings as a spreadsheet (status, severity, evidence, fix).", "csv"),
]

_SAFE_RE = re.compile(r"[^A-Za-z0-9._ -]+")


def _safe_base(base_name: str | None) -> str:
    base = (base_name or "document").strip()
    base = re.sub(r"\.(docx|pdf|doc|txt)$", "", base, flags=re.IGNORECASE)
    base = _SAFE_RE.sub("", base).strip(" .-")
    return base or "document"


def export_filename(export_id: str, base_name: str | None) -> str:
    base = _safe_base(base_name)
    return _FILENAMES.get(export_id, f"{base}-{export_id}").format(base=base)


# ── catalog ──────────────────────────────────────────────────────────────────

def _find_artifact(artifacts: list[dict], kind: str) -> Optional[dict]:
    for a in artifacts:
        if a.get("kind") == kind:
            return a
    return None


def _has_rendered_docx(state: dict, artifacts: list[dict]) -> bool:
    return bool(state.get("rendered_docx_uri")) or _find_artifact(artifacts, "rendered_docx") is not None


def _has_compliance(state: dict) -> bool:
    return bool(
        state.get("flags")
        or state.get("coverage")
        or state.get("compliance_score") is not None
    )


def list_exports(
    state: dict,
    artifacts: list[dict],
    *,
    base_name: str,
    soffice_available: bool,
) -> list[ExportSpec]:
    """Enumerate the deliverables this finished run can produce."""
    base = _safe_base(base_name)
    specs: list[ExportSpec] = []

    if _has_rendered_docx(state, artifacts):
        docx_art = _find_artifact(artifacts, "rendered_docx")
        specs.append(
            ExportSpec(
                id="document_docx",
                category=CAT_DOCUMENT,
                label="Word document",
                description="The finished document, editable in Microsoft Word.",
                fmt="docx",
                filename=export_filename("document_docx", base),
                artifact_id=str(docx_art["id"]) if docx_art else None,
                size_bytes=(docx_art or {}).get("size_bytes") or None,
            )
        )

        pdf_art = _find_artifact(artifacts, "rendered_pdf")
        pdf_ready = pdf_art is not None or bool(state.get("rendered_pdf_uri"))
        specs.append(
            ExportSpec(
                id="document_pdf",
                category=CAT_DOCUMENT,
                label="PDF document",
                description="A print-ready PDF copy of the finished document.",
                fmt="pdf",
                filename=export_filename("document_pdf", base),
                available=pdf_ready or soffice_available,
                reason=None
                if (pdf_ready or soffice_available)
                else "PDF conversion needs LibreOffice, which isn't available on the server.",
                artifact_id=str(pdf_art["id"]) if pdf_art else None,
                size_bytes=(pdf_art or {}).get("size_bytes") or None,
            )
        )

        specs.append(
            ExportSpec(
                id="formatting_json",
                category=CAT_FORMATTING,
                label="Styling & formatting spec (JSON)",
                description="Page setup, fonts, headings, colours, lists and tables as a reusable JSON spec.",
                fmt="json",
                filename=export_filename("formatting_json", base),
            )
        )
        specs.append(
            ExportSpec(
                id="formatting_report",
                category=CAT_FORMATTING,
                label="Formatting report (Word)",
                description="A human-readable summary of the document's styling and formatting.",
                fmt="docx",
                filename=export_filename("formatting_report", base),
            )
        )

        specs.append(
            ExportSpec(
                id="content_json",
                category=CAT_DATA,
                label="Extracted content (JSON)",
                description="The document's structured text, headings, lists and tables for downstream tools.",
                fmt="json",
                filename=export_filename("content_json", base),
            )
        )

    if state.get("diff"):
        specs.append(
            ExportSpec(
                id="change_report",
                category=CAT_REPORT,
                label="Change report (Word)",
                description="Section-by-section summary of what changed in this version.",
                fmt="docx",
                filename=export_filename("change_report", base),
            )
        )

    # New audit engine: surface the pre-built report artifacts (PDF/Word/JSON/CSV).
    has_audit_report = False
    for export_id, kind, category, label, description, fmt in _COMPLIANCE_ARTIFACTS:
        art = _find_artifact(artifacts, kind)
        if not art:
            continue
        has_audit_report = True
        specs.append(
            ExportSpec(
                id=export_id,
                category=category,
                label=label,
                description=description,
                fmt=fmt,
                filename=export_filename(export_id, base),
                artifact_id=str(art["id"]),
                size_bytes=art.get("size_bytes") or None,
            )
        )

    # Legacy compliance docx (only when the new artifacts aren't present).
    if not has_audit_report and _has_compliance(state):
        specs.append(
            ExportSpec(
                id="compliance_report",
                category=CAT_REPORT,
                label="Compliance report (Word)",
                description="Findings, coverage and the overall compliance score.",
                fmt="docx",
                filename=export_filename("compliance_report", base),
            )
        )

    return specs


# ── producers ─────────────────────────────────────────────────────────────────

def build_export(
    export_id: str,
    state: dict,
    artifacts: list[dict],
    *,
    base_name: str,
) -> tuple[bytes, str, str]:
    """Produce one deliverable. Returns ``(data, mime, filename)``.

    Assumes ``export_id`` is valid (the route validates it against the catalog).
    Raises ``ExportError`` for missing inputs; ``LibreOfficeUnavailable`` if PDF
    conversion is requested without LibreOffice.
    """
    filename = export_filename(export_id, base_name)

    if export_id == "document_docx":
        return _rendered_docx_bytes(state, artifacts), DOCX_MIME, filename

    if export_id == "document_pdf":
        return _document_pdf(state, artifacts), PDF_MIME, filename

    if export_id == "formatting_json":
        return _formatting_json(state, artifacts), JSON_MIME, filename

    if export_id == "formatting_report":
        return _formatting_report(state, artifacts), DOCX_MIME, filename

    if export_id == "content_json":
        return _content_json(state, artifacts), JSON_MIME, filename

    if export_id == "change_report":
        return _change_report(state), DOCX_MIME, filename

    if export_id == "compliance_report":
        return _compliance_report(state), DOCX_MIME, filename

    # Pre-built audit-report artifacts: stream the stored bytes as-is.
    _ARTIFACT_EXPORTS = {
        "compliance_report_pdf": ("compliance_report_pdf", PDF_MIME),
        "compliance_report_docx": ("compliance_report_docx", DOCX_MIME),
        "compliance_report_json": ("compliance_report_json", JSON_MIME),
        "compliance_report_csv": ("compliance_report_csv", "text/csv"),
    }
    if export_id in _ARTIFACT_EXPORTS:
        kind, mime = _ARTIFACT_EXPORTS[export_id]
        art = _find_artifact(artifacts, kind)
        if not art or not art.get("uri"):
            raise ExportError("This report artifact is not available.")
        return load_bytes(art["uri"]), mime, filename

    raise ExportError(f"unknown export {export_id!r}")


def _rendered_docx_bytes(state: dict, artifacts: list[dict]) -> bytes:
    uri = state.get("rendered_docx_uri")
    if not uri:
        art = _find_artifact(artifacts, "rendered_docx")
        uri = (art or {}).get("uri")
    if not uri:
        raise ExportError("This run did not produce a Word document.")
    return load_bytes(uri)


def _document_pdf(state: dict, artifacts: list[dict]) -> bytes:
    # Reuse a previously rendered PDF when one exists.
    uri = state.get("rendered_pdf_uri")
    if not uri:
        art = _find_artifact(artifacts, "rendered_pdf")
        uri = (art or {}).get("uri")
    if uri:
        return load_bytes(uri)
    # Otherwise convert the docx on demand (LibreOffice). May raise
    # LibreOfficeUnavailable, which the route maps to a 503.
    from app.services.office.office_pipeline import convert

    docx_bytes = _rendered_docx_bytes(state, artifacts)
    return convert(docx_bytes, "pdf")


def _extract_rendered(state: dict, artifacts: list[dict]):
    """Extract (content, styling) from the rendered .docx."""
    from app.services.extraction.word_ext import extract_word_document

    docx_bytes = _rendered_docx_bytes(state, artifacts)
    return extract_word_document(file_stream=io.BytesIO(docx_bytes), filename="output.docx")


def _formatting_json(state: dict, artifacts: list[dict]) -> bytes:
    _, styling = _extract_rendered(state, artifacts)
    data: dict[str, Any] = styling.model_dump(exclude_none=True)
    si = state.get("style_interpretation")
    if si:
        data["applied_style"] = {
            "detected_kind": si.get("detected_kind"),
            "mode_used": si.get("mode_used"),
            "confidence": si.get("confidence"),
            "summary": si.get("summary"),
            "spec": si.get("spec"),
            "structure": si.get("structure"),
        }
    payload = json.dumps(data, indent=2, ensure_ascii=False, default=str)
    return payload.encode("utf-8")


def _content_json(state: dict, artifacts: list[dict]) -> bytes:
    content, _ = _extract_rendered(state, artifacts)
    data: dict[str, Any] = content.model_dump(exclude_none=True)
    # Drop heavy inline image bytes — keep the structure lean and shareable.
    for el in data.get("elements", []):
        if isinstance(el, dict) and el.get("data_base64"):
            el["data_base64"] = "<omitted>"
    payload = json.dumps(data, indent=2, ensure_ascii=False, default=str)
    return payload.encode("utf-8")


# ── .docx report helpers ───────────────────────────────────────────────────────

def _docx_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _new_doc(title: str, subtitle: str | None = None):
    from docx import Document as _Docx

    doc = _Docx()
    doc.add_heading(title, level=0)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        if p.runs:
            p.runs[0].italic = True
    return doc


def _kv_table(doc, rows: list[tuple[str, str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=0, cols=2)
    try:
        table.style = "Light Grid Accent 1"
    except Exception:  # noqa: BLE001 — style may not exist in the default template
        pass
    for key, val in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = "" if val is None else str(val)


def _grid_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except Exception:  # noqa: BLE001
        pass
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i in range(len(headers)):
            cells[i].text = "" if i >= len(row) or row[i] is None else str(row[i])


def _truncate(text: str | None, limit: int = 1500) -> str:
    text = (text or "").strip()
    return text if len(text) <= limit else text[:limit] + " …"


def _fmt_bool(v: Any) -> str:
    return "yes" if v else "" if v is None else "no"


# ── formatting report ──────────────────────────────────────────────────────────

def _formatting_report(state: dict, artifacts: list[dict]) -> bytes:
    _, styling = _extract_rendered(state, artifacts)
    sd = styling.model_dump(exclude_none=True)

    doc = _new_doc("Formatting report", "Styling & formatting details of the finished document")

    # Page setup
    page = sd.get("page_style") or {}
    margins = page.get("margins") or {}
    doc.add_heading("Page setup", level=1)
    _kv_table(
        doc,
        [
            ("Size", f"{page.get('width_inches', '?')}\" × {page.get('height_inches', '?')}\""),
            ("Orientation", page.get("orientation", "portrait")),
            (
                "Margins",
                "top {t}\", bottom {b}\", left {l}\", right {r}\"".format(
                    t=margins.get("top_inches", "?"),
                    b=margins.get("bottom_inches", "?"),
                    l=margins.get("left_inches", "?"),
                    r=margins.get("right_inches", "?"),
                ),
            ),
        ],
    )

    # Text (run) styles
    run_styles = sd.get("run_styles") or {}
    doc.add_heading("Text styles", level=1)
    doc.add_paragraph(f"{len(run_styles)} named character style(s).")
    if run_styles:
        rows = []
        for name, rs in list(run_styles.items())[:25]:
            rows.append(
                [
                    name,
                    rs.get("font_name", ""),
                    f"{rs.get('font_size_pt')}pt" if rs.get("font_size_pt") else "",
                    _fmt_bool(rs.get("bold")),
                    _fmt_bool(rs.get("italic")),
                    f"#{rs.get('color_hex')}" if rs.get("color_hex") else "",
                ]
            )
        _grid_table(doc, ["Style", "Font", "Size", "Bold", "Italic", "Colour"], rows)

    # Paragraph styles
    para_styles = sd.get("paragraph_styles") or {}
    doc.add_heading("Paragraph styles", level=1)
    doc.add_paragraph(f"{len(para_styles)} named paragraph style(s).")
    if para_styles:
        rows = []
        for name, ps in list(para_styles.items())[:25]:
            rows.append(
                [
                    name,
                    ps.get("alignment", ""),
                    f"{ps.get('line_spacing')}" if ps.get("line_spacing") else "",
                    f"{ps.get('space_after_pt')}pt" if ps.get("space_after_pt") else "",
                ]
            )
        _grid_table(doc, ["Style", "Alignment", "Line spacing", "Space after"], rows)

    # Applied style (Flow 2 — style transfer)
    si = state.get("style_interpretation")
    if si:
        doc.add_heading("Applied style", level=1)
        rows = [
            ("Source read as", si.get("mode_used") or si.get("detected_kind") or "—"),
        ]
        if si.get("confidence") is not None:
            rows.append(("Confidence", f"{float(si['confidence']) * 100:.0f}%"))
        if si.get("summary"):
            rows.append(("Summary", si["summary"]))
        _kv_table(doc, rows)

        spec = si.get("spec") or {}
        _spec_section(doc, spec)

        structure = si.get("structure") or {}
        if structure:
            doc.add_heading("Recognised structure", level=2)
            _kv_table(
                doc,
                [
                    ("Headings", structure.get("headings", 0)),
                    ("List items", structure.get("list_items", 0)),
                    ("Tables", structure.get("tables", 0)),
                ],
            )

    return _docx_bytes(doc)


def _spec_section(doc, spec: dict) -> None:
    """Render an applied StyleSpec (from a formatting guideline) into the report."""
    if not spec:
        return
    body = spec.get("body") or {}
    if body.get("font") or body.get("size_pt"):
        doc.add_heading("Body text", level=2)
        _kv_table(
            doc,
            [
                ("Font", body.get("font", "—")),
                ("Size", f"{body.get('size_pt')}pt" if body.get("size_pt") else "—"),
                ("Line spacing", body.get("line_spacing", "—")),
            ],
        )

    headings = spec.get("headings") or []
    if headings:
        doc.add_heading("Heading rules", level=2)
        rows = []
        for h in headings:
            rows.append(
                [
                    f"H{h.get('level', '?')}",
                    h.get("font", ""),
                    f"{h.get('size_pt')}pt" if h.get("size_pt") else "",
                    _fmt_bool(h.get("bold")),
                    f"#{h.get('color_hex')}" if h.get("color_hex") else "",
                ]
            )
        _grid_table(doc, ["Level", "Font", "Size", "Bold", "Colour"], rows)

    colors = spec.get("colors") or {}
    accent = spec.get("accent_color_hex")
    if colors or accent:
        doc.add_heading("Palette", level=2)
        rows = []
        if accent:
            rows.append(("accent", f"#{accent}"))
        for name, hexv in colors.items():
            rows.append((name, f"#{hexv}"))
        _kv_table(doc, rows)


# ── change report ──────────────────────────────────────────────────────────────

def _change_report(state: dict) -> bytes:
    diff = state.get("diff") or []
    doc = _new_doc("Change report", f"{len(diff)} section(s) changed in this version")
    if not diff:
        doc.add_paragraph("No section-level changes were recorded for this run.")
        return _docx_bytes(doc)

    for i, d in enumerate(diff, 1):
        title = d.get("title") or d.get("slot_id") or f"Section {i}"
        doc.add_heading(f"{i}. {title}", level=1)
        sources = d.get("sources") or []
        if sources:
            p = doc.add_paragraph()
            r = p.add_run("Sources: " + ", ".join(str(s) for s in sources))
            r.italic = True
        doc.add_heading("Before", level=3)
        doc.add_paragraph(_truncate(d.get("original")) or "—")
        doc.add_heading("After", level=3)
        doc.add_paragraph(_truncate(d.get("proposed")) or "—")
    return _docx_bytes(doc)


# ── compliance report ────────────────────────────────────────────────────────

def _compliance_report(state: dict) -> bytes:
    score = state.get("compliance_score")
    flags = state.get("flags") or []
    coverage = state.get("coverage") or {}
    score_str = f"{float(score) * 100:.0f}%" if score is not None else "n/a"

    doc = _new_doc("Compliance report", f"Score: {score_str}")

    rows = [("Domain", state.get("domain_id") or "—"), ("Score", score_str)]
    _kv_table(doc, rows)

    if coverage:
        doc.add_heading("Coverage", level=1)
        filled = coverage.get("filled") or []
        missing = coverage.get("missing") or []
        _kv_table(
            doc,
            [
                ("Required sections", coverage.get("required_total", 0)),
                ("Filled", len(filled)),
                ("Missing", len(missing)),
            ],
        )
        if missing:
            doc.add_heading("Missing sections", level=2)
            for m in missing:
                doc.add_paragraph(str(m), style="List Bullet")

    doc.add_heading(f"Findings ({len(flags)})", level=1)
    if not flags:
        doc.add_paragraph("No issues found.")
    else:
        rows = [
            [str(i), f.get("kind", ""), f.get("slot_id", ""), f.get("note", "")]
            for i, f in enumerate(flags, 1)
        ]
        _grid_table(doc, ["#", "Kind", "Section", "Note"], rows)
    return _docx_bytes(doc)
