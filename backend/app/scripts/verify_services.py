"""Verify the ported document services + the Qdrant RAG path.

    python -m app.scripts.verify_services
"""
from __future__ import annotations

import io
import sys
import zipfile


def _docx(items: list[tuple[str, str]]) -> bytes:
    from docx import Document

    d = Document()
    for h, p in items:
        d.add_heading(h, 1)
        d.add_paragraph(p)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def check(name, fn) -> bool:
    try:
        print(f"  PASS  {name}: {fn()}")
        return True
    except Exception as e:  # noqa: BLE001
        print(f"  FAIL  {name}: {e!r}")
        return False


SAMPLE = _docx(
    [
        ("Executive Summary", "This report summarizes Q3 results for compound BP-217."),
        ("Methods", "A randomized, double-blind, placebo-controlled design was used."),
        ("Results", "The primary endpoint was met with statistical significance."),
        ("Conclusion", "BP-217 demonstrated efficacy and an acceptable safety profile."),
    ]
)
STYLE_DONOR = _docx([("Title", "Corporate Template"), ("Heading", "Branded look")])


def main() -> int:
    results: list[bool] = []

    def extract():
        from app.services.extraction.word_ext import extract_word_document

        c, st = extract_word_document(file_stream=io.BytesIO(SAMPLE), filename="s.docx")
        assert len(c.elements) >= 4, "too few elements"
        return f"{len(c.elements)} elements, {len(st.run_styles)} run / {len(st.paragraph_styles)} para styles"

    results.append(check("extract_word", extract))

    def apply():
        from app.services.extraction.word_ext import extract_word_document
        from app.services.formatting.formater_apply import apply_styling

        c, st = extract_word_document(file_stream=io.BytesIO(SAMPLE), filename="s.docx")
        out = apply_styling(c, st).getvalue()
        assert zipfile.is_zipfile(io.BytesIO(out)), "invalid docx"
        return f"{len(out)} bytes, valid docx"

    results.append(check("apply_styling", apply))

    def style():
        from app.services.style.style_engine import transfer_style

        out = transfer_style(
            SAMPLE, "content.docx", STYLE_DONOR, "style.docx",
            normalize_fonts=True, promote_headings=True,
        )
        assert zipfile.is_zipfile(io.BytesIO(out)), "invalid docx"
        return f"{len(out)} bytes, valid docx"

    results.append(check("transfer_style (docx->docx)", style))

    def fp():
        from app.services.orchestration.pipeline_steps import fingerprint_template, structure_draft

        f = fingerprint_template(SAMPLE, "s.docx")
        d = structure_draft(SAMPLE, "s.docx")
        assert f.heading_hierarchy, "no slots"
        assert d.sections, "no draft sections"
        return f"{len(f.heading_hierarchy)} template slots, {len(d.sections)} draft sections"

    results.append(check("fingerprint + structure", fp))

    def mapping():
        from app.services.mapping.section_mapper import map_sections
        from app.services.orchestration.pipeline_steps import fingerprint_template, structure_draft

        f = fingerprint_template(SAMPLE, "s.docx")
        d = structure_draft(SAMPLE, "s.docx")
        m = map_sections(f, d)
        assert m.mappings, "no mappings"
        actions = sorted({x.action.value for x in m.mappings})
        return f"{len(m.mappings)} mappings; actions={actions}"

    results.append(check("section_mapper (LLM/TF-IDF)", mapping))

    def rag():
        from app.rag.indexer import index_domain_profile
        from app.rag.retriever import load_domain_profile, retrieve

        n = index_domain_profile("pharma")
        prof = load_domain_profile("pharma")
        hits = retrieve("clinical safety and adverse events", prof, top_k=3)
        assert hits, "no RAG hits"
        return f"indexed {n} chunks; retrieved {len(hits)} via Qdrant (top {hits[0].score:.3f}, {hits[0].doc_id})"

    results.append(check("RAG: pharma index + Qdrant retrieve", rag))

    print(f"\n{sum(results)}/{len(results)} services passing")
    return 0 if all(results) else 1


if __name__ == "__main__":
    sys.exit(main())
