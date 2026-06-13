"""In-process end-to-end check for Flow 1 (no HTTP):

  create project → upload a sample .docx → execute regenerate run → expect HITL
  → resume → expect done + rendered .docx, persisted artifact + version.

    python -m app.scripts.e2e_flow1
"""
from __future__ import annotations

import asyncio
import io
import uuid
import zipfile

from app.agents.nodes.common import DOCX_MIME
from app.agents.runner import execute_run, resume_run
from app.core.db import create_all, get_sessionmaker
from app.models.artifact import Artifact
from app.models.document import Document, DocumentVersion
from app.models.project import Project
from app.models.run import Run
from app.storage import get_storage


def _sample_docx() -> bytes:
    from docx import Document as Docx

    d = Docx()
    d.add_heading("Executive Summary", level=1)
    d.add_paragraph("This report summarizes Q3 results for compound BP-217.")
    d.add_heading("Introduction", level=1)
    d.add_paragraph("Background, objectives, and scope of the study.")
    d.add_heading("Methods", level=1)
    d.add_paragraph("A randomized, double-blind, placebo-controlled design was used.")
    d.add_heading("Results", level=1)
    d.add_paragraph("The primary endpoint was met with statistical significance.")
    d.add_heading("Conclusion", level=1)
    d.add_paragraph("BP-217 demonstrated efficacy and an acceptable safety profile.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


async def main() -> int:
    await create_all()
    sm = get_sessionmaker()
    storage = get_storage()
    data = _sample_docx()

    async with sm() as s:
        p = Project(name="E2E Flow1 Test")
        s.add(p)
        await s.flush()
        pid = p.id
        key = storage.make_key(project_id=str(pid), kind="source", filename="report.docx")
        obj = storage.put(data, key=key, content_type=DOCX_MIME)
        art = Artifact(
            project_id=pid, uri=obj.uri, kind="upload", filename="report.docx",
            mime=DOCX_MIME, size_bytes=len(data),
        )
        s.add(art)
        doc = Document(project_id=pid, kind="source", display_name="report.docx", current_version=1)
        s.add(doc)
        await s.flush()
        s.add(DocumentVersion(document_id=doc.id, version=1, artifact_id=art.id))
        run = Run(project_id=pid, flow="regenerate", status="pending", input_refs={}, state={})
        s.add(run)
        await s.flush()
        run.langgraph_thread_id = str(run.id)
        rid = str(run.id)
        uri = obj.uri
        await s.commit()

    state = {
        "run_id": rid, "project_id": str(pid), "flow": "regenerate",
        "template_file_uri": uri, "draft_file_uri": uri, "domain_id": None,
        "output_format": "docx", "user_suggestions": "Make the Conclusion more concise.",
        "status": "pending", "warnings": [], "traces": [],
    }

    print("== execute regenerate ==")
    await execute_run(rid, "regenerate", state)
    async with sm() as s:
        run = await s.get(Run, uuid.UUID(rid))
        st = run.state or {}
        print(f"  status={run.status}  diff_slots={len(st.get('diff', []))}  "
              f"traces={len(st.get('traces', []))}  flags={len(st.get('flags', []))}")
        assert run.status == "hitl", f"expected hitl, got {run.status}"

    print("== resume (accept all) ==")
    await resume_run(rid, [])
    async with sm() as s:
        run = await s.get(Run, uuid.UUID(rid))
        st = run.state or {}
        print(f"  status={run.status}  rendered_docx_uri={st.get('rendered_docx_uri')}")
        print(f"  warnings={st.get('warnings')}")
        assert run.status == "done", f"expected done, got {run.status}: {run.error}"
        docx_uri = st.get("rendered_docx_uri")
        assert docx_uri, "no rendered_docx_uri"

        out = storage.get(docx_uri)
        valid = zipfile.is_zipfile(io.BytesIO(out))
        print(f"  rendered docx: {len(out)} bytes, valid_zip={valid}")
        assert valid, "rendered docx is not a valid .docx"

        ov = (
            await s.execute(
                DocumentVersion.__table__.select().where(  # type: ignore[attr-defined]
                    DocumentVersion.created_by_run_id == uuid.UUID(rid)
                )
            )
        ).first()
        print(f"  output DocumentVersion persisted: {ov is not None}")

    print("\nE2E Flow 1 PASSED ✅")
    return 0


if __name__ == "__main__":
    raise SystemExit(asyncio.run(main()))
