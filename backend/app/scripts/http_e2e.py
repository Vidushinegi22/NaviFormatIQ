"""Full HTTP integration test against a running server.

Start the server first:
    python -m uvicorn app.main:app --port 8099
Then:
    python -m app.scripts.http_e2e
"""
from __future__ import annotations

import io
import json
import sys
import time
import zipfile

import httpx

BASE = "http://127.0.0.1:8099"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _docx(items: list[tuple[str, str]]) -> bytes:
    from docx import Document

    d = Document()
    for h, p in items:
        d.add_heading(h, 1)
        d.add_paragraph(p)
    b = io.BytesIO()
    d.save(b)
    return b.getvalue()


def _poll(c: httpx.Client, rid: str, want: set[str], timeout: float = 180) -> dict:
    t0 = time.time()
    while time.time() - t0 < timeout:
        r = c.get(f"/api/v1/flows/{rid}").json()
        if r["status"] in want or r["status"] == "error":
            return r
        time.sleep(1.5)
    return c.get(f"/api/v1/flows/{rid}").json()


def main() -> int:
    c = httpx.Client(base_url=BASE, timeout=90)
    results: list[bool] = []

    def step(name, fn):
        try:
            print(f"  PASS  {name}: {fn()}")
            results.append(True)
        except Exception as e:  # noqa: BLE001
            print(f"  FAIL  {name}: {e!r}")
            results.append(False)

    step("healthz", lambda: c.get("/healthz").json()["checks"])

    pid = c.post("/api/v1/projects", json={"name": "HTTP e2e"}).json()["id"]
    content = _docx([
        ("Executive Summary", "BP-217 Q3 outcome summary."),
        ("Methods", "Randomized controlled trial."),
        ("Results", "Primary endpoint met; mild adverse events."),
        ("Conclusion", "Efficacy with acceptable safety."),
    ])
    donor = _docx([("Title", "Corporate Template"), ("Section", "Branded styling.")])

    up = c.post(
        f"/api/v1/projects/{pid}/uploads",
        files={"file": ("report.docx", content, DOCX_MIME)},
        data={"kind": "source"},
    ).json()
    doc_id = up["document_id"]
    step("upload → storage", lambda: up["uri"][:46] + "...")

    donor_doc = c.post(
        f"/api/v1/projects/{pid}/uploads",
        files={"file": ("donor.docx", donor, DOCX_MIME)},
        data={"kind": "style_donor"},
    ).json()["document_id"]

    # ── SSE: start a run and read its event stream ──
    def sse():
        rid = c.post(
            f"/api/v1/projects/{pid}/flows/compliance",
            json={"draft_document_id": doc_id, "domain_id": "pharma", "mode": "check"},
        ).json()["run_id"]
        seen: list[str] = []
        with httpx.Client(base_url=BASE, timeout=120) as sc, sc.stream(
            "GET", f"/api/v1/flows/{rid}/stream"
        ) as resp:
            for line in resp.iter_lines():
                if line.startswith("data: "):
                    ev = json.loads(line[6:])
                    seen.append(f"{ev.get('agent')}:{ev.get('status')}")
                    if ev.get("status") in ("done", "error", "hitl"):
                        break
        assert seen, "no SSE events"
        return f"{len(seen)} events e.g. {seen[:6]}"

    step("SSE stream", sse)

    # ── Flow 1: regenerate (start → hitl → resume → export → download) ──
    def flow1():
        rid = c.post(
            f"/api/v1/projects/{pid}/flows/regenerate",
            json={"draft_document_id": doc_id, "user_suggestions": "Make the Conclusion more concise.", "output_format": "docx"},
        ).json()["run_id"]
        r = _poll(c, rid, {"hitl"})
        assert r["status"] == "hitl", r["status"]
        assert r["diff"], "no diff at HITL"
        c.post(f"/api/v1/flows/{rid}/resume", json={"decisions": []})
        r = _poll(c, rid, {"done"})
        assert r["status"] == "done", r.get("error") or r["status"]
        assert r["rendered_docx_uri"], "no rendered output"
        aid = r["artifacts"][0]["id"]
        dl = c.get(f"/api/v1/artifacts/{aid}/download")
        assert zipfile.is_zipfile(io.BytesIO(dl.content)), "download not a valid docx"
        return f"hitl→done, {len(r['diff'])} diffs, {len(r['flags'])} flags, dl={len(dl.content)}B"

    step("Flow 1 regenerate", flow1)

    # ── Flow 2: style transfer ──
    def flow2():
        rid = c.post(
            f"/api/v1/projects/{pid}/flows/style",
            json={"content_document_id": doc_id, "style_document_id": donor_doc},
        ).json()["run_id"]
        r = _poll(c, rid, {"done"})
        assert r["status"] == "done", r.get("error") or r["status"]
        assert r["rendered_docx_uri"], "no rendered output"
        return "done, styled .docx produced"

    step("Flow 2 style transfer", flow2)

    # ── Flow 3: compliance check ──
    def flow3():
        rid = c.post(
            f"/api/v1/projects/{pid}/flows/compliance",
            json={"draft_document_id": doc_id, "domain_id": "pharma", "mode": "check"},
        ).json()["run_id"]
        r = _poll(c, rid, {"done"})
        assert r["status"] == "done", r.get("error") or r["status"]
        assert r["coverage"] is not None, "no coverage"
        return f"done, {len(r['flags'])} flags, coverage req={r['coverage'].get('required_total')}"

    step("Flow 3 compliance (check)", flow3)

    step("GET /domains", lambda: [d["slug"] for d in c.get("/api/v1/domains").json()])
    step("POST /domains/pharma/index", lambda: c.post("/api/v1/domains/pharma/index").json())

    # ── doc-chat: subject document ──
    def chat_doc():
        sid = c.post("/api/v1/chat/sessions", json={"project_id": pid, "subject_document_id": doc_id}).json()["id"]
        r = c.post(
            f"/api/v1/chat/sessions/{sid}/messages",
            json={"message": "List this document's headings and describe its margins."},
        ).json()
        assert r["answer"], "empty answer"
        return f"tools={[s['tool'] for s in r['steps']]} answer={r['answer'][:55]!r}"

    step("Doc-chat (subject doc)", chat_doc)

    # ── doc-chat: pharma RAG ──
    def chat_rag():
        sid = c.post("/api/v1/chat/sessions", json={"project_id": pid}).json()["id"]
        r = c.post(
            f"/api/v1/chat/sessions/{sid}/messages",
            json={"message": "Use the pharma domain to retrieve guidance about risk factors, then summarize it."},
        ).json()
        assert r["answer"], "empty answer"
        return f"tools={[s['tool'] for s in r['steps']]} answer={r['answer'][:55]!r}"

    step("Doc-chat (pharma RAG)", chat_rag)

    print(f"\n{sum(results)}/{len(results)} HTTP checks passing")
    return 0 if all(results) else 1


if __name__ == "__main__":
    sys.exit(main())
