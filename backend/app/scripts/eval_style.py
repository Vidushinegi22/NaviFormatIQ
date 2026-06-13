"""
Style-transfer eval harness.
============================

Runs the Flow-2 style engine on the three SOP-QA-0042 fixtures and scores the
output against the hand-formatted "gold" document. Lets us measure, concretely,
how well a RAW (flat) document is upgraded to match a target look — whether the
style source is the FORMATTED example or the TEMPLATE guideline.

Usage (from repo root):
    backend/.venv/bin/python -m app.scripts.eval_style            # uses default fixtures
    backend/.venv/bin/python -m app.scripts.eval_style /path/raw.docx /path/style.docx

It prints a per-source scorecard: number of real headings, real list items,
real tables, centered title-block paragraphs, header/footer presence, and
accent-colour usage — plus the gold document's numbers for reference.
"""
from __future__ import annotations

import io
import os
import sys

from docx import Document
from docx.oxml.ns import qn

NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

DEFAULT_DIR = "/Users/granthgaurav/Desktop/Navispark"
RAW = os.path.join(DEFAULT_DIR, "SOP-QA-0042_Raw-Content.docx")
FORMATTED = os.path.join(DEFAULT_DIR, "SOP-QA-0042_Formatted.docx")
TEMPLATE = os.path.join(DEFAULT_DIR, "SOP-QA-0042_Formatting-Template.docx")


def _has_numpr(p) -> bool:
    return p._p.find(".//" + qn("w:numPr")) is not None


def inspect(doc_bytes: bytes) -> dict:
    """Structural fingerprint of a rendered document."""
    doc = Document(io.BytesIO(doc_bytes))
    headings = {}
    list_items = 0
    centered_title = 0
    accent_runs = 0
    empty_paras = 0
    consecutive_empty = 0
    max_consecutive_empty = 0

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        sname = ""
        try:
            sname = p.style.name or ""
        except Exception:
            pass
        if sname.startswith("Heading"):
            headings[sname] = headings.get(sname, 0) + 1
        if _has_numpr(p) or sname.startswith("List"):
            if text:
                list_items += 1
        if not text and not _has_numpr(p):
            empty_paras += 1
            consecutive_empty += 1
            max_consecutive_empty = max(max_consecutive_empty, consecutive_empty)
        else:
            consecutive_empty = 0
        # centered title block (first ~4 paras, centered, has size/color)
        if p.alignment is not None and int(p.alignment) == 1 and text:
            centered_title += 1
        for r in p.runs:
            try:
                if r.font.color is not None and r.font.color.type and str(r.font.color.rgb) not in ("000000", "auto", "None"):
                    accent_runs += 1
                    break
            except Exception:
                pass

    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "headings": dict(sorted(headings.items())),
        "n_headings": sum(headings.values()),
        "list_items": list_items,
        "centered_paras": centered_title,
        "accent_color_paras": accent_runs,
        "empty_paras": empty_paras,
        "max_consecutive_empty": max_consecutive_empty,
        "header": _hf_text(doc, "header"),
        "footer": _hf_text(doc, "footer"),
    }


def _hf_text(doc, which: str) -> str:
    try:
        sec = doc.sections[0]
        part = sec.header if which == "header" else sec.footer
        return " | ".join(p.text for p in part.paragraphs if (p.text or "").strip())[:80]
    except Exception:
        return ""


def run_one(label: str, raw_path: str, style_path: str, mode: str = "auto") -> None:
    from app.services.style.style_engine import transfer_style_smart

    with open(raw_path, "rb") as fh:
        content_bytes = fh.read()
    with open(style_path, "rb") as fh:
        style_bytes = fh.read()

    outcome = transfer_style_smart(
        content_bytes,
        os.path.basename(raw_path),
        style_bytes,
        os.path.basename(style_path),
        mode=mode,
    )
    out_path = os.path.join("/tmp", f"eval_out_{label}.docx")
    with open(out_path, "wb") as fh:
        fh.write(outcome.docx_bytes)

    fp = inspect(outcome.docx_bytes)
    print("\n" + "=" * 78)
    print(f"SOURCE = {label}  (mode_used={outcome.mode_used}, detected={outcome.detected_kind} "
          f"@ {outcome.confidence:.0%})")
    print(f"  -> {out_path}")
    print("-" * 78)
    print(f"  summary       : {outcome.summary}")
    print(f"  paragraphs    : {fp['paragraphs']}   tables: {fp['tables']}")
    print(f"  headings      : {fp['n_headings']}  {fp['headings']}")
    print(f"  list items    : {fp['list_items']}")
    print(f"  centered paras: {fp['centered_paras']}   accent-color paras: {fp['accent_color_paras']}")
    print(f"  empty paras   : {fp['empty_paras']}  (max consecutive: {fp['max_consecutive_empty']})")
    print(f"  header        : {fp['header']!r}")
    print(f"  footer        : {fp['footer']!r}")
    if outcome.warnings:
        print(f"  warnings ({len(outcome.warnings)}):")
        for w in outcome.warnings[:12]:
            print(f"      - {w}")


def main() -> None:
    args = sys.argv[1:]
    if len(args) >= 2:
        run_one("custom", args[0], args[1], args[2] if len(args) > 2 else "auto")
        return

    print("\n########## GOLD (hand-formatted target) ##########")
    with open(FORMATTED, "rb") as fh:
        gold = inspect(fh.read())
    print(f"  paragraphs    : {gold['paragraphs']}   tables: {gold['tables']}")
    print(f"  headings      : {gold['n_headings']}  {gold['headings']}")
    print(f"  list items    : {gold['list_items']}")
    print(f"  centered paras: {gold['centered_paras']}   accent-color paras: {gold['accent_color_paras']}")
    print(f"  empty paras   : {gold['empty_paras']}  (max consecutive: {gold['max_consecutive_empty']})")
    print(f"  header        : {gold['header']!r}")
    print(f"  footer        : {gold['footer']!r}")

    run_one("example", RAW, FORMATTED)
    run_one("guideline", RAW, TEMPLATE)


if __name__ == "__main__":
    main()
