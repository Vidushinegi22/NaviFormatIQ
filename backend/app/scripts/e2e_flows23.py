"""In-process e2e for Flow 2 (style transfer) and Flow 3 (compliance apply/check).

    python -m app.scripts.e2e_flows23
"""
from __future__ import annotations

import asyncio
import io
import uuid
import zipfile

from app.agents.nodes.common import DOCX_MIME
from app.agents.runner import execute_run, resume_run
from app.core.db import create_all, get_sessionmaker
from app.models.artifact import Artifact
from app.models.project import Project
from app.models.run import Run
from app.storage import get_storage


def _docx(items: list[tuple[str, str]]) -> bytes:
    from docx import Document

    d = Document()
    for h, p in items:
        d.add_heading(h, 1)
        d.add_paragraph(p)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


async def _upload(s, storage, pid, data: bytes, name: str, kind: str) -> str:
    key = storage.make_key(project_id=str(pid), kind=kind, filename=name)
    obj = storage.put(data, key=key, content_type=DOCX_MIME)
    s.add(Artifact(project_id=pid, uri=obj.uri, kind="upload", filename=name, mime=DOCX_MIME, size_bytes=len(data)))
    await s.flush()
    return obj.uri


async def _mk_run(s, pid, flow, mode=None, domain_id=None) -> str:
    r = Run(project_id=pid, flow=flow, mode=mode, domain_id=domain_id, status="pending", input_refs={}, state={})
    s.add(r)
    await s.flush()
    r.langgraph_thread_id = str(r.id)
    return str(r.id)


async def _status(sm, rid):
    async with sm() as s:
        r = await s.get(Run, uuid.UUID(rid))
        return r.status, (r.state or {}), r.error


def _valid_docx_in_r2(storage, uri) -> bool:
    return bool(uri) and zipfile.is_zipfile(io.BytesIO(storage.get(uri)))


async def main() -> int:
    await create_all()
    sm = get_sessionmaker()
    storage = get_storage()

    content = _docx([
        ("Executive Summary", "Summary of compound BP-217 Q3 outcomes."),
        ("Methods", "Randomized, double-blind, placebo-controlled."),
        ("Results", "Primary endpoint met; adverse events were mild."),
        ("Conclusion", "Efficacy demonstrated with acceptable safety."),
    ])
    donor = _docx([("Title", "Corporate Template"), ("Section", "Branded styling.")])

    async with sm() as s:
        p = Project(name="E2E Flow2/3 Test")
        s.add(p)
        await s.flush()
        pid = p.id
        content_uri = await _upload(s, storage, pid, content, "content.docx", "source")
        style_uri = await _upload(s, storage, pid, donor, "donor.docx", "style_donor")
        r2id = await _mk_run(s, pid, "style")
        r3c = await _mk_run(s, pid, "compliance", mode="check", domain_id="pharma")
        r3a = await _mk_run(s, pid, "compliance", mode="apply", domain_id="pharma")
        await s.commit()

    ok = True

    # ── Flow 2: style transfer ──
    print("== Flow 2 (style transfer) ==")
    await execute_run(r2id, "style", {
        "run_id": r2id, "project_id": str(pid), "flow": "style",
        "template_file_uri": style_uri, "draft_file_uri": content_uri,
        "normalize_fonts": True, "promote_headings": True, "output_format": "docx",
        "status": "pending", "warnings": [], "traces": [],
    })
    st, state, err = await _status(sm, r2id)
    docx_uri = state.get("rendered_docx_uri")
    valid = _valid_docx_in_r2(storage, docx_uri)
    print(f"  status={st} rendered={docx_uri} valid_docx={valid} err={err}")
    ok = ok and st == "done" and valid

    # ── Flow 3: compliance CHECK (no rewrite/emit) ──
    print("== Flow 3 (compliance, check mode) ==")
    await execute_run(r3c, "compliance", {
        "run_id": r3c, "project_id": str(pid), "flow": "compliance", "mode": "check",
        "template_file_uri": content_uri, "draft_file_uri": content_uri,
        "domain_id": "pharma", "output_format": "docx",
        "status": "pending", "warnings": [], "traces": [],
    })
    st, state, err = await _status(sm, r3c)
    cov = state.get("coverage") or {}
    print(f"  status={st} flags={len(state.get('flags', []))} coverage(req={cov.get('required_total')},"
          f"missing={len(cov.get('missing', []))}) rendered={state.get('rendered_docx_uri')} err={err}")
    ok = ok and st == "done" and state.get("coverage") is not None and not state.get("rendered_docx_uri")

    # ── Flow 3: compliance APPLY (HITL → resume → emit) ──
    print("== Flow 3 (compliance, apply mode) ==")
    await execute_run(r3a, "compliance", {
        "run_id": r3a, "project_id": str(pid), "flow": "compliance", "mode": "apply",
        "template_file_uri": content_uri, "draft_file_uri": content_uri,
        "domain_id": "pharma", "output_format": "docx",
        "status": "pending", "warnings": [], "traces": [],
    })
    st, state, err = await _status(sm, r3a)
    print(f"  after execute: status={st} diff_slots={len(state.get('diff', []))} err={err}")
    ok = ok and st == "hitl"
    await resume_run(r3a, [])
    st, state, err = await _status(sm, r3a)
    docx_uri = state.get("rendered_docx_uri")
    valid = _valid_docx_in_r2(storage, docx_uri)
    print(f"  after resume: status={st} rendered={docx_uri} valid_docx={valid} err={err}")
    ok = ok and st == "done" and valid

    print("\n" + ("Flows 2 & 3 PASSED ✅" if ok else "Flows 2 & 3 FAILED ❌"))
    return 0 if ok else 1


if __name__ == "__main__":
    raise SystemExit(asyncio.run(main()))
