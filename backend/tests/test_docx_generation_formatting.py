from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree

from app.services.extraction.word_ext import (
    _count_words,
    extract_word_document,
    fingerprint_word_template,
    structure_word_draft,
)
from app.services.formatting.template_emitter import _apply_field_updates, render_template
from app.services.mapping import section_mapper


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _save(path: Path, doc: Document) -> Path:
    doc.save(path)
    return path


def _number_format(docx_path: Path, num_id: str) -> str:
    with ZipFile(docx_path) as zf:
        numbering = etree.fromstring(zf.read("word/numbering.xml"))
    abs_id = numbering.xpath(
        f'//w:num[@w:numId="{num_id}"]/w:abstractNumId/@w:val',
        namespaces=NS,
    )[0]
    return numbering.xpath(
        f'//w:abstractNum[@w:abstractNumId="{abs_id}"]/w:lvl[@w:ilvl="0"]/w:numFmt/@w:val',
        namespaces=NS,
    )[0]


def _direct_num_id(paragraph) -> str | None:
    num_id = paragraph._p.find(".//" + qn("w:numId"))
    return num_id.get(qn("w:val")) if num_id is not None else None


def _direct_indent(paragraph) -> tuple[str | None, str | None]:
    ind = paragraph._p.find(".//" + qn("w:ind"))
    if ind is None:
        return None, None
    return ind.get(qn("w:left")), ind.get(qn("w:hanging"))


def test_numbered_sop_preface_is_not_an_editable_slot(tmp_path: Path) -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("STANDARD OPERATING PROCEDURE").bold = True
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Handling and Disposal of Expired Products").bold = True
    doc.add_paragraph("Document No: SOP-QA-0042  |  Version: 2.1")
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph("Original purpose.")
    doc.add_heading("2. Procedure", level=1)
    doc.add_paragraph("Original procedure.")

    path = _save(tmp_path / "sop.docx", doc)
    fp = fingerprint_word_template(file_path=str(path), filename=path.name)

    assert [slot.title for slot in fp.heading_hierarchy] == ["1. Purpose", "2. Procedure"]


def test_draft_structure_preserves_bullet_and_numbered_markers(tmp_path: Path) -> None:
    doc = Document()
    doc.add_heading("1. Procedure", level=1)
    doc.add_paragraph("Follow these controls:")
    doc.add_paragraph("Confirm quarantine label", style="List Bullet")
    doc.add_paragraph("Record transfer in WIMS", style="List Bullet")
    doc.add_paragraph("Complete Section A", style="List Number")
    doc.add_paragraph("QA countersigns Section B", style="List Number")

    path = _save(tmp_path / "lists.docx", doc)
    draft = structure_word_draft(file_path=str(path), filename=path.name)
    section = next(s for s in draft.sections if s.heading == "1. Procedure")

    assert "- Confirm quarantine label" in section.text
    assert "- Record transfer in WIMS" in section.text
    assert "1. Complete Section A" in section.text
    assert "2. QA countersigns Section B" in section.text


def test_mapping_prefers_exact_numbered_headings_over_body_keywords(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("STANDARD OPERATING PROCEDURE").add_run().bold = True
    doc.add_paragraph("Document No: SOP-QA-0042  |  Version: 2.1")
    doc.add_heading("3.2 Warehouse Personnel", level=2)
    doc.add_paragraph("Perform weekly checks.")
    doc.add_heading("4.1 Identification of Expired Products", level=2)
    doc.add_paragraph("Warehouse personnel shall perform scheduled expiry checks as follows.")

    path = _save(tmp_path / "mapping.docx", doc)
    fp = fingerprint_word_template(file_path=str(path), filename=path.name)
    draft = structure_word_draft(file_path=str(path), filename=path.name)

    old_llm_available = section_mapper.llm_available
    section_mapper.llm_available = lambda: False
    try:
        mapping = section_mapper.map_sections(fp, draft)
    finally:
        section_mapper.llm_available = old_llm_available

    by_slot = {m.slot_id: m.draft_section_idx for m in mapping.mappings}
    sections = {s.index: s.heading for s in draft.sections}

    assert sections[by_slot["3_2_warehouse_personnel"]] == "3.2 Warehouse Personnel"
    assert sections[by_slot["4_1_identification_of_expired_products"]] == (
        "4.1 Identification of Expired Products"
    )


def test_sparse_docx_page_count_uses_layout_fallback(tmp_path: Path) -> None:
    doc = Document()
    for i in range(10):
        doc.add_heading(f"Section {i + 1}", level=1)
        for _ in range(2):
            p = doc.add_paragraph(
                ("This sparse report paragraph uses double spacing and paragraph gaps. ") * 7
            )
            p.paragraph_format.line_spacing = 2
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        for _ in range(2):
            blank = doc.add_paragraph("")
            blank.paragraph_format.line_spacing = 2
            blank.paragraph_format.space_after = Pt(12)

    path = _save(tmp_path / "sparse-report.docx", doc)
    content, _ = extract_word_document(file_path=str(path), filename=path.name)
    word_density_pages = max(1, round(_count_words(content.elements) / 380))

    assert content.metadata.page_count is not None
    assert content.metadata.page_count >= 8
    assert content.metadata.page_count > word_density_pages


def test_rendered_lists_use_real_word_numbering_and_standard_indent(tmp_path: Path) -> None:
    doc = Document()
    doc.add_heading("1. Procedure", level=1)
    doc.add_paragraph("Placeholder.")
    template = _save(tmp_path / "template.docx", doc)
    fp = fingerprint_word_template(file_path=str(template), filename=template.name)

    output = tmp_path / "rendered.docx"
    output.write_bytes(
        render_template(
            fp,
            {"1_procedure": "1. Complete Section A\n2. QA countersigns Section B\n- File EPDF-001"},
        )
    )
    rendered = Document(output)
    items = [p for p in rendered.paragraphs if _direct_num_id(p)]

    assert [p.text for p in items] == [
        "Complete Section A",
        "QA countersigns Section B",
        "File EPDF-001",
    ]
    assert _direct_num_id(items[0]) == _direct_num_id(items[1])
    assert _number_format(output, _direct_num_id(items[0])) == "decimal"
    assert _number_format(output, _direct_num_id(items[2])) == "bullet"
    assert _direct_indent(items[0]) == ("720", "360")
    assert _direct_indent(items[2]) == ("720", "360")


def test_field_updates_do_not_rewrite_existing_revision_history_rows(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Document No: SOP-QA-0042  |  Version: 2.1  |  Effective Date: 01-Jun-2025")
    table = doc.add_table(rows=2, cols=4)
    for cell, text in zip(table.rows[0].cells, ["Version", "Date", "Summary of Changes", "Author"]):
        cell.text = text
    for cell, text in zip(
        table.rows[1].cells,
        ["2.1", "01-Jun-2025", "Updated regulatory references", "R. Sharma"],
    ):
        cell.text = text

    _apply_field_updates(
        doc,
        {
            "replacements": [
                {"kind": "version", "old": "2.1", "new": "2.2"},
                {"kind": "date", "old": "01-Jun-2025", "new": "31-May-2026"},
            ],
            "revision": {
                "row": ["2.2", "31-May-2026", "Introduce chain-of-custody controls", "R. Sharma"]
            },
        },
    )

    assert doc.paragraphs[0].text.endswith("Version: 2.2  |  Effective Date: 31-May-2026")
    assert [cell.text for cell in table.rows[1].cells] == [
        "2.1",
        "01-Jun-2025",
        "Updated regulatory references",
        "R. Sharma",
    ]
    assert [cell.text for cell in table.rows[2].cells] == [
        "2.2",
        "31-May-2026",
        "Introduce chain-of-custody controls",
        "R. Sharma",
    ]
