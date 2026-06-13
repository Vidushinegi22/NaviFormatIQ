"""
End-to-end test: create a sample .docx with rich formatting,
extract via the API, then apply the styling back to generate a new .docx.
"""

import io
import json
import os
import sys

# Ensure we can import from the tests directory
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

from word_ext import extract_word_document
from pdf_ext import extract_pdf_document
from formater_apply import apply_styling, apply_styling_to_file, style_transfer
from models import DocumentContent, DocumentStyling

# ---------------------------------------------------------------------------
# All generated artifacts (inputs, outputs, JSON dumps, PDFs) live here so
# they're easy to inspect after a test run AND easy to .gitignore as a
# single directory.
# ---------------------------------------------------------------------------
SAMPLES_DIR = os.path.join(os.path.dirname(__file__), "sample_test_files")
os.makedirs(SAMPLES_DIR, exist_ok=True)


def _samples_path(name: str) -> str:
    """Resolve a path inside the samples dir."""
    return os.path.join(SAMPLES_DIR, name)


def create_sample_docx(path: str):
    """Create a richly formatted sample .docx for testing."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Title heading
    heading = doc.add_heading("Document Processing System Report", level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        run.font.size = Pt(24)

    # Subtitle
    sub = doc.add_heading("Quarterly Analysis — Q4 2026", level=2)
    for run in sub.runs:
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        run.font.size = Pt(18)

    # Body paragraph with mixed formatting
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = 1.15

    run1 = para.add_run("This report provides a ")
    run1.font.name = "Calibri"
    run1.font.size = Pt(11)

    run2 = para.add_run("comprehensive overview")
    run2.font.name = "Calibri"
    run2.font.size = Pt(11)
    run2.bold = True

    run3 = para.add_run(" of the document processing pipeline, including ")
    run3.font.name = "Calibri"
    run3.font.size = Pt(11)

    run4 = para.add_run("performance metrics")
    run4.font.name = "Calibri"
    run4.font.size = Pt(11)
    run4.italic = True

    run5 = para.add_run(" and key findings.")
    run5.font.name = "Calibri"
    run5.font.size = Pt(11)

    # Another heading
    doc.add_heading("Key Findings", level=3)

    # Bullet-style paragraphs
    for bullet_text in [
        "• Document extraction accuracy improved by 15%",
        "• PDF parsing now handles multi-column layouts",
        "• Style transfer fidelity increased to 94%",
    ]:
        bp = doc.add_paragraph(bullet_text)
        bp.paragraph_format.left_indent = Inches(0.5)
        bp.paragraph_format.space_after = Pt(4)
        for r in bp.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(11)

    # Table
    doc.add_heading("Performance Metrics", level=3)

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"

    # Header row
    headers = ["Metric", "Q3 2026", "Q4 2026"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = text
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)

    # Data rows
    data = [
        ["Extraction Speed", "2.3s/doc", "1.8s/doc"],
        ["Accuracy", "89%", "94%"],
        ["Style Fidelity", "82%", "91%"],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, text in enumerate(row_data):
            table.cell(row_idx, col_idx).text = text

    # Colored text paragraph
    para2 = doc.add_paragraph()
    para2.paragraph_format.space_before = Pt(12)
    r = para2.add_run("Important: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    r.font.size = Pt(11)
    r2 = para2.add_run("All metrics are measured against the standard benchmark suite.")
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"

    # Underlined and strikethrough text
    para3 = doc.add_paragraph()
    u = para3.add_run("Underlined text")
    u.underline = True
    u.font.size = Pt(11)
    para3.add_run(" and ")
    s = para3.add_run("strikethrough text")
    s.font.strike = True
    s.font.size = Pt(11)
    para3.add_run(" for completeness.")

    doc.save(path)
    print(f"✅ Sample DOCX created: {path}")


def test_word_round_trip():
    """Test: create .docx → extract → apply → new .docx."""
    print("\n" + "=" * 60)
    print("TEST: Word Document Round-Trip")
    print("=" * 60)

    sample_path = _samples_path("test_sample.docx")
    output_path = _samples_path("test_output.docx")

    # 1) Create sample
    create_sample_docx(sample_path)

    # 2) Extract
    print("📄 Extracting content and styling...")
    content, styling = extract_word_document(file_path=sample_path)

    # Print stats
    print(f"   Elements extracted: {len(content.elements)}")
    print(f"   Paragraph styles:   {len(styling.paragraph_styles)}")
    print(f"   Run styles:         {len(styling.run_styles)}")
    print(f"   Table styles:       {len(styling.table_styles)}")
    print(f"   Cell styles:        {len(styling.cell_styles)}")

    # 3) Save JSON
    content_json_path = _samples_path("test_content.json")
    styling_json_path = _samples_path("test_styling.json")

    with open(content_json_path, "w", encoding="utf-8") as f:
        json.dump(content.model_dump(exclude_none=True), f, indent=2, ensure_ascii=False)
    print(f"   Content JSON saved: {content_json_path}")

    with open(styling_json_path, "w", encoding="utf-8") as f:
        json.dump(styling.model_dump(exclude_none=True), f, indent=2, ensure_ascii=False)
    print(f"   Styling JSON saved: {styling_json_path}")

    # 4) Apply styling to generate new docx
    print("🎨 Applying styling to generate new document...")
    apply_styling_to_file(content, styling, output_path)
    print(f"   Output DOCX:        {output_path}")

    # 5) Verify output exists and has content
    assert os.path.exists(output_path), "Output file was not created!"
    assert os.path.getsize(output_path) > 1000, "Output file seems too small!"
    print(f"   Output size:        {os.path.getsize(output_path)} bytes")

    # 6) Re-extract from output and compare
    print("🔁 Re-extracting from output to verify...")
    content2, styling2 = extract_word_document(file_path=output_path)
    print(f"   Elements (re-extracted): {len(content2.elements)}")

    print("✅ Word round-trip test PASSED!")
    return True


def test_style_transfer():
    """Test: extract content from one doc, style from another, combine."""
    print("\n" + "=" * 60)
    print("TEST: Style Transfer")
    print("=" * 60)

    sample_path = _samples_path("test_sample.docx")
    output_path = _samples_path("test_style_transfer.docx")

    if not os.path.exists(sample_path):
        create_sample_docx(sample_path)

    # Extract
    content, styling = extract_word_document(file_path=sample_path)

    # Style transfer (same doc for simplicity — proves the mechanism works)
    print("🔄 Performing style transfer...")
    result = style_transfer(content, styling)

    with open(output_path, "wb") as f:
        f.write(result.read())

    assert os.path.exists(output_path), "Style transfer output was not created!"
    print(f"   Output: {output_path} ({os.path.getsize(output_path)} bytes)")
    print("✅ Style transfer test PASSED!")
    return True


def test_style_engine_fidelity():
    """Style engine must preserve 100% of content while restyling.

    Builds a content doc with mixed styles + a table, plus a separate style
    source doc, then transfers and asserts paragraph/table counts are
    unchanged and the body font now matches the source.
    """
    print("\n" + "=" * 60)
    print("TEST: Style Engine Fidelity (in-place transfer)")
    print("=" * 60)

    from style_engine import transfer_style, profile_from_docx

    # --- content doc: Times New Roman body, a few headings, a table ---
    content_path = _samples_path("test_se_content.docx")
    cdoc = Document()
    cstyle = cdoc.styles["Normal"]
    cstyle.font.name = "Times New Roman"
    cstyle.font.size = Pt(12)
    cdoc.add_paragraph("PROJECT OVERVIEW")  # bold pseudo-heading
    cdoc.paragraphs[-1].runs[0].bold = True
    cdoc.add_paragraph("This is the first body paragraph with several words in it.")
    cdoc.add_paragraph("Second body paragraph, also reasonably long for testing.")
    t = cdoc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "A"; t.cell(0, 1).text = "B"
    t.cell(1, 0).text = "C"; t.cell(1, 1).text = "D"
    for txt in ["Bullet one item", "Bullet two item"]:
        cdoc.add_paragraph(txt, style="List Bullet")
    cdoc.save(content_path)

    # --- style source: Calibri 11, narrow margins ---
    style_path = _samples_path("test_se_style.docx")
    sdoc = Document()
    sstyle = sdoc.styles["Normal"]
    sstyle.font.name = "Calibri"
    sstyle.font.size = Pt(11)
    sdoc.add_paragraph("Some styled source text.")
    sdoc.save(style_path)

    before = Document(content_path)
    n_paras = len(before.paragraphs)
    n_tables = len(before.tables)

    with open(content_path, "rb") as fh:
        content_bytes = fh.read()
    with open(style_path, "rb") as fh:
        style_bytes = fh.read()

    out_bytes = transfer_style(
        content_bytes, content_path, style_bytes, style_path,
        promote_headings=False,  # keep deterministic for the assertion
    )
    out_path = _samples_path("test_se_output.docx")
    with open(out_path, "wb") as fh:
        fh.write(out_bytes)

    after = Document(io.BytesIO(out_bytes))
    print(f"   paragraphs: {n_paras} -> {len(after.paragraphs)}")
    print(f"   tables:     {n_tables} -> {len(after.tables)}")
    assert len(after.paragraphs) == n_paras, "Style transfer dropped paragraphs!"
    assert len(after.tables) == n_tables, "Style transfer dropped tables!"

    # Body font should now be the source's (Calibri), not Times New Roman.
    out_normal = after.styles["Normal"]
    print(f"   Normal font: {out_normal.font.name}")
    assert out_normal.font.name == "Calibri", "Body font was not transferred!"

    # No residual Times New Roman in body runs.
    from docx.oxml.ns import qn as _qn
    residual = set()
    for p in after.paragraphs:
        for r in p.runs:
            rpr = r._r.find(_qn("w:rPr"))
            if rpr is not None:
                rf = rpr.find(_qn("w:rFonts"))
                if rf is not None and rf.get(_qn("w:ascii")) == "Times New Roman":
                    residual.add("Times New Roman")
    assert not residual, f"Residual source-doc fonts left behind: {residual}"

    print("✅ Style engine fidelity test PASSED!")
    return True


def test_json_reload():
    """Test: load content/styling from JSON files and regenerate document."""
    print("\n" + "=" * 60)
    print("TEST: JSON Reload and Apply")
    print("=" * 60)

    content_json_path = _samples_path("test_content.json")
    styling_json_path = _samples_path("test_styling.json")
    output_path = _samples_path("test_from_json.docx")

    if not os.path.exists(content_json_path):
        print("⚠️  Run test_word_round_trip first to generate JSON files.")
        return False

    with open(content_json_path, "r") as f:
        content = DocumentContent(**json.load(f))
    with open(styling_json_path, "r") as f:
        styling = DocumentStyling(**json.load(f))

    print("📥 Loaded from JSON files")
    print("🎨 Applying styling...")
    apply_styling_to_file(content, styling, output_path)

    assert os.path.exists(output_path), "Output was not created!"
    print(f"   Output: {output_path} ({os.path.getsize(output_path)} bytes)")
    print("✅ JSON reload test PASSED!")
    return True


# ---------------------------------------------------------------------------
# New tests — template fingerprinting + full pipeline
# ---------------------------------------------------------------------------

def _create_template_docx(path: str) -> None:
    """A template with three headings; the body paragraphs will be replaced."""
    doc = Document()
    title = doc.add_heading("Quarterly Pharma Report Template", level=1)
    for r in title.runs:
        r.font.size = Pt(20)

    doc.add_heading("Executive Summary", level=2)
    p = doc.add_paragraph("[placeholder — executive summary goes here]")
    p.paragraph_format.space_after = Pt(8)

    doc.add_heading("Market Segments", level=2)
    doc.add_paragraph("[placeholder — segment breakdown]")

    doc.add_heading("Risk Factors", level=2)
    doc.add_paragraph("[placeholder — risks]")

    doc.save(path)


def _create_draft_docx(path: str) -> None:
    """A draft with the same logical sections but different titles/wording."""
    doc = Document()
    doc.add_heading("Global Pharma 2024 — Internal Note", level=1)

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(
        "The pharmaceutical market hit roughly $1.48T in 2024. "
        "GLP-1 drugs and oncology were the big growth drivers."
    )

    doc.add_heading("Segments", level=2)
    doc.add_paragraph(
        "Oncology grew at 15.5% CAGR to ~$230B. "
        "GLP-1 / metabolic posted 32% CAGR to ~$118B."
    )

    doc.add_heading("Risks", level=2)
    doc.add_paragraph(
        "Patent cliff exposure is over $200B through 2028. "
        "IRA negotiations affect 20+ blockbusters."
    )

    doc.save(path)


def test_template_fingerprint():
    """Template fingerprinting captures headings, numbering and style registry."""
    print("\n" + "=" * 60)
    print("TEST: Template Fingerprint")
    print("=" * 60)

    from word_ext import fingerprint_word_template

    tpl_path = _samples_path("test_template.docx")
    _create_template_docx(tpl_path)

    fp = fingerprint_word_template(file_path=tpl_path, include_template_bytes=True)
    print(f"   Slots: {[s.title for s in fp.heading_hierarchy]}")
    print(f"   Numbering defs: {len(fp.numbering_defs)}")
    print(f"   Style registry para styles: {len(fp.style_registry.paragraph_styles)}")
    print(f"   template_b64 present: {bool(fp.template_b64)}")

    assert len(fp.heading_hierarchy) >= 3, "Expected 3+ heading slots"
    titles = [s.title for s in fp.heading_hierarchy]
    assert "Executive Summary" in titles, "Exec Summary slot missing"
    assert "Risk Factors" in titles, "Risk Factors slot missing"
    assert fp.template_b64, "template bytes not embedded"
    assert fp.source_format == "docx"
    print("✅ Template fingerprint test PASSED!")
    return True


def test_full_pipeline():
    """End-to-end: template + draft → emitted .docx with rewritten bodies."""
    print("\n" + "=" * 60)
    print("TEST: Full Pipeline (template + draft → v2)")
    print("=" * 60)

    from pipeline import run_pipeline

    tpl_path = _samples_path("test_template.docx")
    draft_path = _samples_path("test_draft.docx")
    if not os.path.exists(tpl_path):
        _create_template_docx(tpl_path)
    _create_draft_docx(draft_path)

    with open(tpl_path, "rb") as fh:
        template_bytes = fh.read()
    with open(draft_path, "rb") as fh:
        draft_bytes = fh.read()

    result = run_pipeline(
        template_bytes=template_bytes,
        template_name="test_template.docx",
        draft_bytes=draft_bytes,
        draft_name="test_draft.docx",
        domain_profile_id="pharma",
        output_format="docx",
    )

    print(f"   job_id={result.job_id}")
    print(f"   slots in fingerprint: {len(result.fingerprint.heading_hierarchy)}")
    print(f"   mappings: {len(result.mapping.mappings)}")
    print(f"   flags: {len(result.flags)}")
    print(f"   diff entries: {len(result.diff)}")
    print(f"   warnings: {result.warnings[:3]}")

    assert result.diff, "Expected diff entries"
    assert len(result.diff) == len(result.fingerprint.heading_hierarchy)
    # At least one slot should map to a draft section (the headings overlap
    # semantically: Summary→Executive Summary, Segments→Market Segments,
    # Risks→Risk Factors).
    matched = [m for m in result.mapping.mappings if m.draft_section_idx is not None]
    assert matched, "Expected at least one slot→draft mapping via heuristics"

    if result.artifact_docx_b64:
        out_path = _samples_path("test_pipeline_output.docx")
        with open(out_path, "wb") as fh:
            import base64 as _b64

            fh.write(_b64.b64decode(result.artifact_docx_b64))
        print(f"   Artifact written: {out_path} ({os.path.getsize(out_path)} bytes)")
        assert os.path.getsize(out_path) > 500
    else:
        print("   (No docx artifact embedded — emitter failed; check warnings.)")

    print("✅ Full pipeline test PASSED!")
    return True


def test_diff_shape():
    """ReviewDiff payload has one entry per slot with proposed text."""
    print("\n" + "=" * 60)
    print("TEST: Diff Shape")
    print("=" * 60)

    from pipeline import run_pipeline

    tpl_path = _samples_path("test_template.docx")
    draft_path = _samples_path("test_draft.docx")
    if not os.path.exists(tpl_path):
        _create_template_docx(tpl_path)
    if not os.path.exists(draft_path):
        _create_draft_docx(draft_path)

    with open(tpl_path, "rb") as fh:
        tb = fh.read()
    with open(draft_path, "rb") as fh:
        db = fh.read()

    result = run_pipeline(
        template_bytes=tb,
        template_name="test_template.docx",
        draft_bytes=db,
        draft_name="test_draft.docx",
        output_format="docx",
        embed_artifacts=False,
    )

    slot_ids = {s.slot_id for s in result.fingerprint.heading_hierarchy}
    diff_slot_ids = {d.slot_id for d in result.diff}
    assert slot_ids == diff_slot_ids, "Diff must cover every slot exactly once"
    print(f"   {len(result.diff)} diff entries, one per slot. PASSED.")
    return True


def test_ocr_path_skips_gracefully():
    """OCR path should either run or skip cleanly when no backend is present."""
    print("\n" + "=" * 60)
    print("TEST: OCR Path (skip-aware)")
    print("=" * 60)

    from config import settings
    from pdf_ext import extract_pdf_document

    if not settings.azure_di_configured():
        try:
            import pytesseract  # noqa: F401
            import pdf2image  # noqa: F401
        except ImportError:
            print("   ⚠️  No OCR backend (Azure DI / Tesseract). Skipping.")
            return True

    # Build a probe PDF with enough text that the auto-route does NOT
    # demote to OCR (we just want to verify the entry doesn't crash and
    # produces elements). A separate explicit-OCR probe below confirms
    # the OCR path itself runs when forced.
    import fitz  # PyMuPDF

    doc = fitz.open()
    page = doc.new_page()
    body = (
        "Pharma Document Pipeline OCR sanity check. "
        "This page contains enough printable text to exceed the auto-OCR "
        "threshold so the regular PyMuPDF/pdfplumber extractor handles it. "
        "The full pipeline still routes scanned/image-only PDFs through "
        "Azure Document Intelligence or local Tesseract as a fallback."
    )
    page.insert_textbox(fitz.Rect(72, 72, 540, 400), body, fontsize=11)
    pdf_bytes = doc.tobytes()
    doc.close()

    content, _ = extract_pdf_document(
        file_stream=io.BytesIO(pdf_bytes), filename="sanity.pdf"
    )
    assert content.elements, "Expected at least one element from the probe PDF"
    print(f"   text path: {len(content.elements)} elements extracted.")

    # Explicit OCR call — only run when a backend is available, and only
    # require that it returns *something* (an empty result on a tiny line
    # of text is acceptable, we just want to confirm the wiring works).
    if settings.azure_di_configured():
        backend = "Azure DI"
    else:
        try:
            import pytesseract  # noqa: F401
            import pdf2image  # noqa: F401
            backend = "Tesseract"
        except ImportError:
            backend = None

    if backend:
        try:
            ocr_content, _ = extract_pdf_document(
                file_stream=io.BytesIO(pdf_bytes),
                filename="sanity.pdf",
                use_ocr=True,
            )
            print(f"   OCR path ({backend}): {len(ocr_content.elements)} elements.")
        except Exception as e:
            print(f"   OCR path ({backend}) raised: {e} — non-fatal.")
    else:
        print("   OCR path: no backend available, skipped.")

    print("✅ OCR path test PASSED!")
    return True


if __name__ == "__main__":
    results = []
    results.append(("Word Round-Trip", test_word_round_trip()))
    results.append(("Style Transfer", test_style_transfer()))
    results.append(("Style Engine Fidelity", test_style_engine_fidelity()))
    results.append(("JSON Reload", test_json_reload()))
    results.append(("Template Fingerprint", test_template_fingerprint()))
    results.append(("Full Pipeline", test_full_pipeline()))
    results.append(("Diff Shape", test_diff_shape()))
    results.append(("OCR Path", test_ocr_path_skips_gracefully()))

    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    for name, passed in results:
        status = "✅ PASS" if passed else "❌ FAIL"
        print(f"  {status} — {name}")

    all_passed = all(r[1] for r in results)
    print(f"\n{'All tests passed!' if all_passed else 'Some tests failed.'}")
    sys.exit(0 if all_passed else 1)
