"""
Word Document Extractor
=======================
Extracts content structure and styling from .docx files into JSON.

Produces two outputs:
  1. DocumentContent  — structured content (headings, paragraphs, images, tables, lists)
  2. DocumentStyling  — reusable style definitions (fonts, sizes, colors, spacing)

Uses python-docx for deep inspection of every formatting property.
"""

from __future__ import annotations

import base64
import hashlib
import io
import os
import re
from datetime import datetime, timezone
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_UNDERLINE,
)
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

from models import (
    Alignment,
    BorderDef,
    BorderStyle,
    CellBorders,
    CellStyle,
    ContentElement,
    DocumentContent,
    DocumentMetadata,
    DocumentStyling,
    ElementType,
    ImagePosition,
    IndentStyle,
    ListType,
    PageMargins,
    PageStyle,
    ParagraphStyle,
    RunStyle,
    StyleMetadata,
    TableCell,
    TableRow,
    TableStyle,
    TextRun,
    UnderlineType,
    VerticalAlignment,
    Orientation,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_ALIGNMENT_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: Alignment.LEFT,
    WD_ALIGN_PARAGRAPH.CENTER: Alignment.CENTER,
    WD_ALIGN_PARAGRAPH.RIGHT: Alignment.RIGHT,
    WD_ALIGN_PARAGRAPH.JUSTIFY: Alignment.JUSTIFY,
}

_UNDERLINE_MAP: dict = {
    True: UnderlineType.SINGLE,
    False: UnderlineType.NONE,
    None: None,
    WD_UNDERLINE.SINGLE: UnderlineType.SINGLE,
    WD_UNDERLINE.DOUBLE: UnderlineType.DOUBLE,
    WD_UNDERLINE.DOTTED: UnderlineType.DOTTED,
    WD_UNDERLINE.DASH: UnderlineType.DASHED,
    WD_UNDERLINE.WAVY: UnderlineType.WAVY,
    WD_UNDERLINE.THICK: UnderlineType.THICK,
    WD_UNDERLINE.WORDS: UnderlineType.WORDS_ONLY,
}

_VERT_ALIGN_MAP = {
    WD_ALIGN_VERTICAL.TOP: VerticalAlignment.TOP,
    WD_ALIGN_VERTICAL.CENTER: VerticalAlignment.CENTER,
    WD_ALIGN_VERTICAL.BOTTOM: VerticalAlignment.BOTTOM,
}


def _emu_to_inches(emu: Optional[int]) -> Optional[float]:
    if emu is None:
        return None
    return round(emu / 914400, 4)


def _pt_value(val) -> Optional[float]:
    """Convert a python-docx Length (or plain number) to points.

    NB: python-docx Length subclasses (Emu, Pt, Inches, etc.) are *int*
    subclasses whose raw integer value is in EMU (1pt = 12700 EMU).
    We therefore have to check for the ``.pt`` property *before* the
    int/float branch — otherwise we'd return the raw EMU integer as if
    it were already in points and the caller would re-multiply it by
    another 12700 when wrapping with ``Pt(...)``.
    """
    if val is None:
        return None
    if isinstance(val, bool):
        return None
    # Length objects expose a ``.pt`` float — always prefer that.
    pt_attr = getattr(val, "pt", None)
    if pt_attr is not None:
        try:
            return round(float(pt_attr), 2)
        except (TypeError, ValueError):
            pass
    if isinstance(val, (int, float)):
        return round(float(val), 2)
    return None


def _rgb_to_hex(rgb: Optional[RGBColor]) -> Optional[str]:
    if rgb is None:
        return None
    return f"#{rgb}"


def _make_style_key(prefix: str, *parts: str) -> str:
    """Create a deterministic, filesystem-safe style key."""
    raw = f"{prefix}_{'_'.join(parts)}"
    # Sanitize
    raw = re.sub(r"[^a-zA-Z0-9_]", "_", raw)
    raw = re.sub(r"_+", "_", raw).strip("_").lower()
    if len(raw) > 80:
        raw = raw[:60] + "_" + hashlib.md5(raw.encode()).hexdigest()[:8]
    return raw


def _fingerprint_run_style(rs: RunStyle) -> str:
    """Return a hashable fingerprint for a RunStyle to enable deduplication."""
    parts = []
    for field_name in rs.model_fields:
        val = getattr(rs, field_name)
        if val is not None:
            parts.append(f"{field_name}={val}")
    return "|".join(parts) if parts else "__default__"


def _fingerprint_para_style(ps: ParagraphStyle) -> str:
    parts = []
    for field_name in ps.model_fields:
        val = getattr(ps, field_name)
        if val is not None:
            if isinstance(val, IndentStyle):
                parts.append(f"indent={val.model_dump_json()}")
            else:
                parts.append(f"{field_name}={val}")
    return "|".join(parts) if parts else "__default__"


def _count_words(elements: list) -> int:
    """Count words across body text + table cells for a page-count estimate."""
    total = 0
    for el in elements:
        for run in el.content or []:
            total += len((run.text or "").split())
        for row in el.rows or []:
            for cell in row.cells:
                for run in cell.content or []:
                    total += len((run.text or "").split())
    return total


# ---------------------------------------------------------------------------
# Extractor class
# ---------------------------------------------------------------------------

class WordExtractor:
    """Extract content and styling from a Word (.docx) document."""

    def __init__(self):
        self._run_style_cache: dict[str, str] = {}  # fingerprint -> style_ref
        self._para_style_cache: dict[str, str] = {}
        self._run_style_counter = 0
        self._para_style_counter = 0
        self._run_styles: dict[str, RunStyle] = {}
        self._para_styles: dict[str, ParagraphStyle] = {}
        self._table_styles: dict[str, TableStyle] = {}
        self._cell_styles: dict[str, CellStyle] = {}

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def extract(
        self,
        file_path: Optional[str] = None,
        file_stream: Optional[io.BytesIO] = None,
        filename: Optional[str] = None,
    ) -> tuple[DocumentContent, DocumentStyling]:
        """
        Extract content and styling from a .docx.

        Provide either ``file_path`` (path on disk) or ``file_stream`` (in-memory bytes).
        """
        if file_path:
            doc = Document(file_path)
            source_name = os.path.basename(file_path)
        elif file_stream:
            doc = Document(file_stream)
            source_name = filename or "uploaded.docx"
        else:
            raise ValueError("Provide either file_path or file_stream")

        # ---- Page style from first section ----
        page_style = self._extract_page_style(doc)

        # ---- Body elements ----
        elements = self._extract_body(doc)

        # ---- Headers / footers (from first section) ----
        header_footer_elements = self._extract_headers_footers(doc)
        elements = header_footer_elements + elements

        # ---- Metadata ----
        core = doc.core_properties
        word_count = _count_words(elements)
        metadata = DocumentMetadata(
            source_file=source_name,
            source_type="docx",
            extracted_at=datetime.now(timezone.utc).isoformat(),
            # .docx carries no rendered page count; estimate from body length
            # (~350 words/page) so downstream consumers have a sane figure.
            page_count=max(1, word_count // 350) if word_count else None,
            author=core.author if core.author else None,
            title=core.title if core.title else None,
        )

        content = DocumentContent(metadata=metadata, elements=elements)

        styling = DocumentStyling(
            metadata=StyleMetadata(
                source_file=source_name,
                created_at=datetime.now(timezone.utc).isoformat(),
            ),
            page_style=page_style,
            paragraph_styles=self._para_styles,
            run_styles=self._run_styles,
            table_styles=self._table_styles,
            cell_styles=self._cell_styles,
        )

        return content, styling

    # ------------------------------------------------------------------
    # Page style
    # ------------------------------------------------------------------

    def _extract_page_style(self, doc: Document) -> PageStyle:
        section = doc.sections[0] if doc.sections else None
        if section is None:
            return PageStyle()

        orientation = Orientation.PORTRAIT
        if section.orientation == WD_ORIENT.LANDSCAPE:
            orientation = Orientation.LANDSCAPE

        return PageStyle(
            width_inches=round(_emu_to_inches(section.page_width) or 8.5, 4),
            height_inches=round(_emu_to_inches(section.page_height) or 11.0, 4),
            orientation=orientation,
            margins=PageMargins(
                top_inches=round(_emu_to_inches(section.top_margin) or 1.0, 4),
                bottom_inches=round(_emu_to_inches(section.bottom_margin) or 1.0, 4),
                left_inches=round(_emu_to_inches(section.left_margin) or 1.25, 4),
                right_inches=round(_emu_to_inches(section.right_margin) or 1.25, 4),
            ),
        )

    # ------------------------------------------------------------------
    # Body iteration
    # ------------------------------------------------------------------

    def _extract_body(self, doc: Document) -> list[ContentElement]:
        """Walk the document body XML to extract paragraphs, tables, and images in order."""
        elements: list[ContentElement] = []
        body = doc.element.body

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                para = Paragraph(child, doc)
                elem = self._process_paragraph(para, doc)
                if elem is not None:
                    elements.append(elem)
            elif tag == "tbl":
                table = DocxTable(child, doc)
                elem = self._process_table(table)
                elements.append(elem)
            elif tag == "sectPr":
                pass  # section properties — already handled via doc.sections

        return elements

    # ------------------------------------------------------------------
    # Paragraph processing
    # ------------------------------------------------------------------

    def _process_paragraph(self, para: Paragraph, doc: Document) -> Optional[ContentElement]:
        """Process a single paragraph into a ContentElement."""

        # Check if it's a page break
        if self._is_page_break(para):
            return ContentElement(type=ElementType.PAGE_BREAK)

        # Determine element type
        heading_level = self._get_heading_level(para)
        list_info = self._get_list_info(para)
        is_heading = heading_level is not None

        # Extract inline content (runs + hyperlinks)
        runs = self._extract_runs(para, doc)

        # Skip empty paragraphs with no text
        if not runs and not is_heading:
            return None

        # Extract paragraph style
        para_style = self._extract_paragraph_style(para)
        para_style_ref = self._register_para_style(para, para_style)

        if is_heading:
            return ContentElement(
                type=ElementType.HEADING,
                level=heading_level,
                content=runs,
                style_ref=para_style_ref,
            )
        elif list_info:
            return ContentElement(
                type=ElementType.LIST_ITEM,
                content=runs,
                style_ref=para_style_ref,
                list_type=list_info["list_type"],
                list_level=list_info["level"],
                bullet_char=list_info.get("bullet_char"),
                number_format=list_info.get("number_format"),
            )
        else:
            return ContentElement(
                type=ElementType.PARAGRAPH,
                content=runs,
                style_ref=para_style_ref,
            )

    def _is_page_break(self, para: Paragraph) -> bool:
        """Check if the paragraph consists solely of a page break."""
        xml = para._element.xml
        if "w:br" in xml and 'w:type="page"' in xml:
            # Only count as page break if there's no real text
            if not para.text.strip():
                return True
        return False

    def _get_heading_level(self, para: Paragraph) -> Optional[int]:
        """Return heading level (1-9) or None."""
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip())
                return level
            except ValueError:
                pass

        # The "Title" style (e.g. add_heading(level=0)) is the document's top
        # heading — treat it as level 1 so it lands in the outline / slots.
        if style_name.strip().lower() == "title":
            return 1

        # Check outline level from XML directly (paragraph_format doesn't expose it)
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            outline_lvl = pPr.find(qn("w:outlineLvl"))
            if outline_lvl is not None:
                try:
                    level = int(outline_lvl.get(qn("w:val"), "9"))
                    if 0 <= level <= 8:
                        return level + 1  # outline_level is 0-indexed
                except (ValueError, TypeError):
                    pass

        return None

    def _get_list_info(self, para: Paragraph) -> Optional[dict]:
        """Detect if paragraph is a list item and extract list metadata."""
        # Check XML for numbering properties
        pPr = para._element.find(qn("w:pPr"))
        numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
        if numPr is None:
            # Heuristic: check if text starts with bullet characters
            text = para.text.strip()
            bullet_chars = ["•", "◦", "▪", "▸", "‣", "-", "–", "—", "●", "○"]
            for bc in bullet_chars:
                if text.startswith(bc):
                    return {
                        "list_type": ListType.BULLET,
                        "level": 0,
                        "bullet_char": bc,
                    }
            # Check for numbered list pattern
            num_match = re.match(r"^(\d+)[.)]\s", text)
            if num_match:
                return {
                    "list_type": ListType.NUMBERED,
                    "level": 0,
                    "number_format": "decimal",
                }
            return None

        # Extract level
        ilvl_elem = numPr.find(qn("w:ilvl"))
        level = 0
        if ilvl_elem is not None:
            level = int(ilvl_elem.get(qn("w:val"), "0"))

        # Determine bullet vs numbered from style name heuristic
        style_name = (para.style.name if para.style else "").lower()
        if "bullet" in style_name or "list bullet" in style_name:
            return {
                "list_type": ListType.BULLET,
                "level": level,
                "bullet_char": "•",
            }
        else:
            return {
                "list_type": ListType.NUMBERED,
                "level": level,
                "number_format": "decimal",
            }

    # ------------------------------------------------------------------
    # Run extraction
    # ------------------------------------------------------------------

    def _extract_runs(self, para: Paragraph, doc: Document) -> list[TextRun]:
        """Extract text runs and inline images from a paragraph."""
        runs: list[TextRun] = []

        for child in para._element:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "r":
                # Regular run
                run_obj = None
                for r in para.runs:
                    if r._element is child:
                        run_obj = r
                        break

                if run_obj is None:
                    # Fallback: extract text directly from XML
                    text_elems = child.findall(qn("w:t"))
                    text = "".join(t.text or "" for t in text_elems)
                    if text:
                        runs.append(TextRun(text=text))
                    continue

                # Check for inline image
                drawing = child.find(qn("w:drawing"))
                if drawing is not None:
                    img_elem = self._extract_inline_image(drawing, doc)
                    if img_elem is not None:
                        runs.append(img_elem)
                    continue

                text = run_obj.text
                if not text:
                    continue

                run_style = self._extract_run_style(run_obj)
                style_ref = self._register_run_style(run_style)
                runs.append(TextRun(text=text, style_ref=style_ref))

            elif tag == "hyperlink":
                # Hyperlink
                href = self._get_hyperlink_url(child, doc)
                for r_elem in child.findall(qn("w:r")):
                    text_elems = r_elem.findall(qn("w:t"))
                    text = "".join(t.text or "" for t in text_elems)
                    if text:
                        runs.append(TextRun(text=text, hyperlink_url=href))

        return runs

    def _get_hyperlink_url(self, hyperlink_elem, doc: Document) -> Optional[str]:
        """Extract the URL from a hyperlink element."""
        r_id = hyperlink_elem.get(qn("r:id"))
        if r_id:
            try:
                rel = doc.part.rels.get(r_id)
                if rel and rel.target_ref:
                    return str(rel.target_ref)
            except Exception:
                pass
        return None

    def _extract_inline_image(self, drawing, doc: Document) -> Optional[TextRun]:
        """Extract an image from a drawing element, return as a special marker run."""
        # This is a content-level image — we'll create a placeholder TextRun
        # and also a separate ContentElement for it (handled at paragraph level)
        # For now, we extract the image data.
        blip = drawing.find(".//" + qn("a:blip"))
        if blip is None:
            return None

        r_embed = blip.get(qn("r:embed"))
        if r_embed is None:
            return None

        try:
            image_part = doc.part.rels[r_embed].target_part
            image_bytes = image_part.blob
            img_b64 = base64.b64encode(image_bytes).decode("utf-8")

            # Get dimensions
            extent = drawing.find(".//" + qn("wp:extent"))
            width_inches = None
            height_inches = None
            if extent is not None:
                cx = extent.get("cx")
                cy = extent.get("cy")
                if cx:
                    width_inches = round(int(cx) / 914400, 4)
                if cy:
                    height_inches = round(int(cy) / 914400, 4)

            # Determine format
            content_type = image_part.content_type
            fmt = "png"
            if "jpeg" in content_type or "jpg" in content_type:
                fmt = "jpeg"
            elif "gif" in content_type:
                fmt = "gif"
            elif "bmp" in content_type:
                fmt = "bmp"
            elif "tiff" in content_type:
                fmt = "tiff"

            # We return this as a special TextRun with image marker
            return TextRun(
                text=f"[IMAGE:{fmt}:{width_inches}x{height_inches}]",
                style_ref="__image__",
                image_data_b64=img_b64,
            )
        except Exception:
            return None

    # ------------------------------------------------------------------
    # Style extraction
    # ------------------------------------------------------------------

    def _extract_run_style(self, run) -> RunStyle:
        """Extract all formatting from a python-docx Run object."""
        font = run.font

        # Underline
        underline = None
        if font.underline is not None:
            underline = _UNDERLINE_MAP.get(font.underline, UnderlineType.SINGLE)

        # Color
        color_hex = None
        if font.color and font.color.rgb:
            color_hex = _rgb_to_hex(font.color.rgb)

        return RunStyle(
            font_name=font.name,
            font_size_pt=_pt_value(font.size) if font.size else None,
            bold=font.bold,
            italic=font.italic,
            underline=underline,
            strikethrough=font.strike,
            double_strikethrough=font.double_strike,
            superscript=font.superscript,
            subscript=font.subscript,
            color_hex=color_hex,
            highlight_color=str(font.highlight_color) if font.highlight_color else None,
            all_caps=font.all_caps,
            small_caps=font.small_caps,
            character_spacing_pt=None,  # character spacing requires XML-level extraction
        )

    def _extract_paragraph_style(self, para: Paragraph) -> ParagraphStyle:
        """Extract all formatting from a python-docx Paragraph."""
        pf = para.paragraph_format

        # Alignment
        alignment = None
        if pf.alignment is not None:
            alignment = _ALIGNMENT_MAP.get(pf.alignment)

        # Indentation
        indent = IndentStyle(
            left_inches=_emu_to_inches(pf.left_indent) if pf.left_indent else None,
            right_inches=_emu_to_inches(pf.right_indent) if pf.right_indent else None,
            first_line_inches=_emu_to_inches(pf.first_line_indent) if pf.first_line_indent else None,
        )
        # Hanging indent (negative first_line_indent)
        if pf.first_line_indent and pf.first_line_indent < 0:
            indent.hanging_inches = abs(_emu_to_inches(pf.first_line_indent))
            indent.first_line_inches = None

        # Line spacing
        line_spacing = None
        line_spacing_rule = None
        if pf.line_spacing is not None:
            line_spacing = float(pf.line_spacing)
        if pf.line_spacing_rule is not None:
            rule_map = {
                WD_LINE_SPACING.SINGLE: "single",
                WD_LINE_SPACING.ONE_POINT_FIVE: "one_point_five",
                WD_LINE_SPACING.DOUBLE: "double",
                WD_LINE_SPACING.AT_LEAST: "at_least",
                WD_LINE_SPACING.EXACTLY: "exactly",
                WD_LINE_SPACING.MULTIPLE: "multiple",
            }
            line_spacing_rule = rule_map.get(pf.line_spacing_rule, "single")

        # Tab stops
        tab_stops = None
        if pf.tab_stops:
            tab_stops = [
                round(_emu_to_inches(ts.position) or 0, 4)
                for ts in pf.tab_stops
            ]

        return ParagraphStyle(
            alignment=alignment,
            space_before_pt=_pt_value(pf.space_before) if pf.space_before else None,
            space_after_pt=_pt_value(pf.space_after) if pf.space_after else None,
            line_spacing=line_spacing,
            line_spacing_rule=line_spacing_rule,
            indent=indent if any(
                v is not None
                for v in [indent.left_inches, indent.right_inches, indent.first_line_inches, indent.hanging_inches]
            ) else None,
            keep_with_next=pf.keep_with_next,
            keep_together=pf.keep_together,
            widow_control=pf.widow_control,
            outline_level=None,  # handled separately in heading detection
            tab_stops=tab_stops,
        )

    # ------------------------------------------------------------------
    # Table processing
    # ------------------------------------------------------------------

    def _process_table(self, table: DocxTable) -> ContentElement:
        """Process a Word table into a ContentElement."""
        rows: list[TableRow] = []

        # Extract table-level style
        table_style = self._extract_table_style(table)
        table_style_key = _make_style_key("table", str(len(self._table_styles)))
        self._table_styles[table_style_key] = table_style

        for row_idx, row in enumerate(table.rows):
            cells: list[TableCell] = []
            for cell in row.cells:
                cell_runs: list[TextRun] = []
                for para in cell.paragraphs:
                    para_runs = self._extract_runs(para, table._parent)
                    cell_runs.extend(para_runs)

                # Extract cell style
                cell_style = self._extract_cell_style(cell)
                cell_style_key = _make_style_key(
                    "cell", str(len(self._cell_styles))
                )
                self._cell_styles[cell_style_key] = cell_style

                cells.append(
                    TableCell(
                        content=cell_runs,
                        style_ref=cell_style_key,
                        inline_style=cell_style,
                    )
                )
            rows.append(TableRow(cells=cells, is_header=(row_idx == 0)))

        return ContentElement(
            type=ElementType.TABLE,
            rows=rows,
            table_style_ref=table_style_key,
            inline_table_style=table_style,
        )

    def _extract_table_style(self, table: DocxTable) -> TableStyle:
        """Extract table-level formatting."""
        alignment = None
        try:
            tbl_pr = table._tbl.find(qn("w:tblPr"))
            if tbl_pr is not None:
                jc = tbl_pr.find(qn("w:jc"))
                if jc is not None:
                    val = jc.get(qn("w:val"), "left")
                    align_map = {"left": Alignment.LEFT, "center": Alignment.CENTER, "right": Alignment.RIGHT}
                    alignment = align_map.get(val)
        except Exception:
            pass

        return TableStyle(alignment=alignment)

    def _extract_cell_style(self, cell) -> CellStyle:
        """Extract cell-level formatting."""
        width_inches = None
        shading_hex = None
        vert_align = None

        try:
            tc = cell._tc
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is not None:
                # Width
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is not None:
                    w_val = tc_w.get(qn("w:w"))
                    w_type = tc_w.get(qn("w:type"), "dxa")
                    if w_val and w_type == "dxa":
                        width_inches = round(int(w_val) / 1440, 4)

                # Shading
                shd = tc_pr.find(qn("w:shd"))
                if shd is not None:
                    fill = shd.get(qn("w:fill"))
                    if fill and fill != "auto":
                        shading_hex = f"#{fill}"

                # Vertical alignment
                v_align = tc_pr.find(qn("w:vAlign"))
                if v_align is not None:
                    va_val = v_align.get(qn("w:val"), "top")
                    va_map = {
                        "top": VerticalAlignment.TOP,
                        "center": VerticalAlignment.CENTER,
                        "bottom": VerticalAlignment.BOTTOM,
                    }
                    vert_align = va_map.get(va_val)
        except Exception:
            pass

        # Detect cell merge spans
        col_span = 1
        row_span = 1
        try:
            tc = cell._tc
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is not None:
                grid_span = tc_pr.find(qn("w:gridSpan"))
                if grid_span is not None:
                    col_span = int(grid_span.get(qn("w:val"), "1"))
                v_merge = tc_pr.find(qn("w:vMerge"))
                if v_merge is not None:
                    merge_val = v_merge.get(qn("w:val"), "continue")
                    if merge_val == "restart":
                        row_span = 1  # Start of vertical merge
        except Exception:
            pass

        return CellStyle(
            width_inches=width_inches,
            shading_color_hex=shading_hex,
            vertical_alignment=vert_align,
            col_span=col_span,
            row_span=row_span,
        )

    # ------------------------------------------------------------------
    # Headers & Footers
    # ------------------------------------------------------------------

    def _extract_headers_footers(self, doc: Document) -> list[ContentElement]:
        """Extract header and footer content from the first section."""
        elements: list[ContentElement] = []
        if not doc.sections:
            return elements

        section = doc.sections[0]

        # Header
        try:
            header = section.header
            if header and not header.is_linked_to_previous:
                for para in header.paragraphs:
                    if para.text.strip():
                        runs = self._extract_runs(para, doc)
                        para_style = self._extract_paragraph_style(para)
                        style_ref = self._register_para_style(para, para_style)
                        elements.append(
                            ContentElement(
                                type=ElementType.HEADER,
                                content=runs,
                                style_ref=style_ref,
                            )
                        )
        except Exception:
            pass

        # Footer
        try:
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                for para in footer.paragraphs:
                    if para.text.strip():
                        runs = self._extract_runs(para, doc)
                        para_style = self._extract_paragraph_style(para)
                        style_ref = self._register_para_style(para, para_style)
                        elements.append(
                            ContentElement(
                                type=ElementType.FOOTER,
                                content=runs,
                                style_ref=style_ref,
                            )
                        )
        except Exception:
            pass

        return elements

    # ------------------------------------------------------------------
    # Style registration (deduplication)
    # ------------------------------------------------------------------

    def _register_run_style(self, run_style: RunStyle) -> str:
        """Register a run style and return its reference key, deduplicating."""
        fp = _fingerprint_run_style(run_style)
        if fp in self._run_style_cache:
            return self._run_style_cache[fp]

        # Generate a meaningful key
        parts = []
        if run_style.font_name:
            parts.append(run_style.font_name.lower().replace(" ", ""))
        if run_style.font_size_pt:
            parts.append(f"{run_style.font_size_pt}pt")
        if run_style.bold:
            parts.append("bold")
        if run_style.italic:
            parts.append("italic")
        if not parts:
            parts.append(f"run_{self._run_style_counter}")
        self._run_style_counter += 1

        key = _make_style_key("run", *parts, str(self._run_style_counter))
        self._run_style_cache[fp] = key
        self._run_styles[key] = run_style
        return key

    def _register_para_style(self, para: Paragraph, para_style: ParagraphStyle) -> str:
        """Register a paragraph style and return its reference key."""
        fp = _fingerprint_para_style(para_style)
        if fp in self._para_style_cache:
            return self._para_style_cache[fp]

        # Use the Word style name as part of the key
        word_style_name = para.style.name if para.style else "default"
        self._para_style_counter += 1
        key = _make_style_key("para", word_style_name, str(self._para_style_counter))

        self._para_style_cache[fp] = key
        self._para_styles[key] = para_style
        return key


# ---------------------------------------------------------------------------
# Convenience function
# ---------------------------------------------------------------------------

def extract_word_document(
    file_path: Optional[str] = None,
    file_stream: Optional[io.BytesIO] = None,
    filename: Optional[str] = None,
) -> tuple[DocumentContent, DocumentStyling]:
    """
    Extract content and styling from a Word document.

    Returns:
        (DocumentContent, DocumentStyling) — the two JSON-serializable objects.
    """
    extractor = WordExtractor()
    return extractor.extract(file_path=file_path, file_stream=file_stream, filename=filename)


# ---------------------------------------------------------------------------
# Template fingerprinting (hackathon spec §2)
# ---------------------------------------------------------------------------

import re as _re
import zipfile as _zipfile
from xml.etree import ElementTree as _ET

from models import (
    DraftSection,
    DraftStructure,
    HeaderFooterSection,
    HeadingSlot,
    NumberingDef,
    TableSchema,
    TemplateFingerprint,
)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS = {"w": _W_NS}


def _slugify(text: str, fallback: str) -> str:
    if not text:
        return fallback
    slug = _re.sub(r"[^a-zA-Z0-9]+", "_", text.strip().lower()).strip("_")
    return slug or fallback


def _parse_numbering_xml(docx_bytes: bytes) -> list[NumberingDef]:
    """Pull list-numbering definitions from word/numbering.xml."""
    defs: list[NumberingDef] = []
    try:
        with _zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
            if "word/numbering.xml" not in z.namelist():
                return defs
            xml = z.read("word/numbering.xml")
    except Exception:
        return defs

    try:
        root = _ET.fromstring(xml)
    except _ET.ParseError:
        return defs

    abstracts: dict[str, dict] = {}
    for an in root.findall("w:abstractNum", _NS):
        an_id = an.get(f"{{{_W_NS}}}abstractNumId", "")
        levels: dict[str, dict] = {}
        for lvl in an.findall("w:lvl", _NS):
            ilvl = lvl.get(f"{{{_W_NS}}}ilvl", "0")
            fmt_el = lvl.find("w:numFmt", _NS)
            text_el = lvl.find("w:lvlText", _NS)
            levels[ilvl] = {
                "num_fmt": fmt_el.get(f"{{{_W_NS}}}val") if fmt_el is not None else None,
                "lvl_text": text_el.get(f"{{{_W_NS}}}val") if text_el is not None else None,
            }
        abstracts[an_id] = {"levels": levels}

    for num in root.findall("w:num", _NS):
        num_id = num.get(f"{{{_W_NS}}}numId", "")
        ref = num.find("w:abstractNumId", _NS)
        an_id = ref.get(f"{{{_W_NS}}}val") if ref is not None else None
        abstract = abstracts.get(an_id or "", {"levels": {}})
        defs.append(
            NumberingDef(
                num_id=num_id,
                abstract_num_id=an_id,
                levels=abstract["levels"],
            )
        )
    return defs


def _detect_toc_paragraph_index(doc) -> Optional[int]:
    """Find the paragraph index where the TOC field begins, if any."""
    for idx, para in enumerate(doc.paragraphs):
        for fld in para._p.iter(qn("w:fldChar")):
            # walk siblings looking for an instrText with "TOC"
            pass
        for instr in para._p.iter(qn("w:instrText")):
            text = instr.text or ""
            if text.strip().upper().startswith("TOC"):
                return idx
        # also check for SDT-based TOCs
        for sdt in para._p.iter(qn("w:sdt")):
            for instr in sdt.iter(qn("w:instrText")):
                if (instr.text or "").strip().upper().startswith("TOC"):
                    return idx
    return None


def _extract_jinja_marker(text: str) -> Optional[str]:
    """Return the first `{{ ... }}` marker in the paragraph text, if any."""
    m = _re.search(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}", text or "")
    return m.group(0) if m else None


def _derive_heading_slots(doc, extractor: "WordExtractor") -> list[HeadingSlot]:
    """Walk the doc's paragraphs and build HeadingSlot entries from headings."""
    slots: list[HeadingSlot] = []
    used_ids: set[str] = set()
    counter = 0
    next_marker: Optional[str] = None

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        marker = _extract_jinja_marker(text)
        level = extractor._get_heading_level(para)
        if level is not None and text:
            counter += 1
            base = _slugify(text, f"slot_{counter}")
            slot_id = base
            i = 1
            while slot_id in used_ids:
                i += 1
                slot_id = f"{base}_{i}"
            used_ids.add(slot_id)
            slots.append(
                HeadingSlot(
                    slot_id=slot_id,
                    level=level,
                    title=text,
                    required=True,
                    expected_keywords=[w.lower() for w in _re.findall(r"[A-Za-z]{4,}", text)],
                    placeholder_marker=next_marker or marker,
                )
            )
            next_marker = None
        elif marker and slots:
            # Jinja marker that follows a heading paragraph → attach to last slot
            slots[-1].placeholder_marker = marker
        elif marker:
            # Marker before any heading → stash for the next one
            next_marker = marker

    return slots


def _extract_table_schemas(doc) -> list[TableSchema]:
    schemas: list[TableSchema] = []
    for t in doc.tables:
        if not t.rows:
            continue
        header_row = [(c.text or "").strip() for c in t.rows[0].cells]
        schemas.append(
            TableSchema(
                header_row=header_row,
                expected_columns=len(header_row),
                is_dynamic=len(t.rows) <= 2,
            )
        )
    return schemas


def _extract_headers_footers_all(doc) -> list[HeaderFooterSection]:
    out: list[HeaderFooterSection] = []
    for i, section in enumerate(doc.sections):
        header_text = None
        footer_text = None
        try:
            if section.header is not None:
                header_text = "\n".join(
                    (p.text or "").strip() for p in section.header.paragraphs if (p.text or "").strip()
                ) or None
        except Exception:
            pass
        try:
            if section.footer is not None:
                footer_text = "\n".join(
                    (p.text or "").strip() for p in section.footer.paragraphs if (p.text or "").strip()
                ) or None
        except Exception:
            pass
        out.append(
            HeaderFooterSection(
                section_index=i,
                header_text=header_text,
                footer_text=footer_text,
            )
        )
    return out


def fingerprint_word_template(
    file_path: Optional[str] = None,
    file_stream: Optional[io.BytesIO] = None,
    filename: Optional[str] = None,
    include_template_bytes: bool = True,
) -> TemplateFingerprint:
    """
    Build a TemplateFingerprint from a .docx template.

    The fingerprint captures heading hierarchy, numbering.xml definitions,
    table schemas, headers/footers across all sections, TOC location, and a
    base64 copy of the original .docx so the docxtpl emitter can render
    directly into the same file (preserving every original style/asset).
    """
    if file_path:
        with open(file_path, "rb") as fh:
            raw = fh.read()
        if not filename:
            filename = os.path.basename(file_path)
    elif file_stream is not None:
        file_stream.seek(0)
        raw = file_stream.read()
    else:
        raise ValueError("Either file_path or file_stream must be provided")

    extractor = WordExtractor()
    content, styling = extractor.extract(
        file_stream=io.BytesIO(raw),
        filename=filename or "template.docx",
    )

    doc = Document(io.BytesIO(raw))

    heading_slots = _derive_heading_slots(doc, extractor)
    numbering = _parse_numbering_xml(raw)
    tables = _extract_table_schemas(doc)
    headers_footers = _extract_headers_footers_all(doc)
    toc_idx = _detect_toc_paragraph_index(doc)

    template_b64 = base64.b64encode(raw).decode("ascii") if include_template_bytes else None

    return TemplateFingerprint(
        metadata=content.metadata,
        page_style=styling.page_style,
        heading_hierarchy=heading_slots,
        numbering_defs=numbering,
        table_schemas=tables,
        headers_footers=headers_footers,
        toc_location=toc_idx,
        style_registry=styling,
        template_b64=template_b64,
        source_format="docx",
    )


def structure_word_draft(
    file_path: Optional[str] = None,
    file_stream: Optional[io.BytesIO] = None,
    filename: Optional[str] = None,
) -> DraftStructure:
    """
    Convert a .docx draft into a section-tagged JSON tree.

    Sections are split at every heading. Each DraftSection accumulates the
    body text (and any tables) that fall between consecutive headings.
    """
    content, _ = extract_word_document(
        file_path=file_path, file_stream=file_stream, filename=filename
    )

    sections: list[DraftSection] = []
    current = DraftSection(index=0, heading=None, level=0, text="")

    def _flush():
        if current.heading or current.text.strip() or current.tables:
            sections.append(current)

    for el in content.elements:
        if el.type == ElementType.HEADING:
            _flush()
            heading_text = " ".join((r.text or "") for r in (el.content or []))
            current = DraftSection(
                index=len(sections),
                heading=heading_text.strip(),
                level=el.level or 1,
                text="",
            )
            for r in (el.content or []):
                if getattr(r, "image_data_b64", None):
                    current.images_b64.append(r.image_data_b64)
        elif el.type == ElementType.PARAGRAPH:
            text = " ".join((r.text or "") for r in (el.content or []))
            if text:
                current.text = (current.text + "\n" + text).strip() if current.text else text
            for r in (el.content or []):
                if getattr(r, "image_data_b64", None):
                    current.images_b64.append(r.image_data_b64)
        elif el.type == ElementType.LIST_ITEM:
            text = " ".join((r.text or "") for r in (el.content or []))
            if text:
                current.text = (current.text + "\n• " + text).strip() if current.text else f"• {text}"
        elif el.type == ElementType.TABLE and el.rows:
            grid: list[list[str]] = []
            for row in el.rows:
                grid.append(
                    [
                        " ".join((r.text or "") for r in (cell.content or []))
                        for cell in row.cells
                    ]
                )
            current.tables.append(grid)
        elif el.type == ElementType.IMAGE and el.data_base64:
            current.images_b64.append(el.data_base64)

    _flush()

    return DraftStructure(metadata=content.metadata, sections=sections)


# ---------------------------------------------------------------------------
# CLI entry point (for standalone testing)
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) < 2:
        print("Usage: python word_ext.py <path-to-docx>")
        sys.exit(1)

    path = sys.argv[1]
    content, styling = extract_word_document(file_path=path)

    # Write outputs
    base = os.path.splitext(path)[0]

    content_path = f"{base}_content.json"
    with open(content_path, "w", encoding="utf-8") as f:
        json.dump(content.model_dump(exclude_none=True), f, indent=2, ensure_ascii=False)
    print(f"Content JSON written to: {content_path}")

    styling_path = f"{base}_styling.json"
    with open(styling_path, "w", encoding="utf-8") as f:
        json.dump(styling.model_dump(exclude_none=True), f, indent=2, ensure_ascii=False)
    print(f"Styling JSON written to: {styling_path}")
