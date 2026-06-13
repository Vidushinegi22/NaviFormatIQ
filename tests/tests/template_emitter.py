"""
Style-preserving Word emission via docxtpl.

Two render modes:

  1. **Jinja-marker templates** — the template's body contains ``{{ slot_id }}``
     placeholders (and optionally ``{%tr for row in table_xxx %}`` row loops).
     We render with ``docxtpl.DocxTemplate``; the template owns every style.

  2. **Heading-driven templates** — no markers. We open the original .docx
     with python-docx and, for each heading whose ``slot_id`` matches a key
     in ``rewritten``, replace the body paragraphs that follow the heading
     (up to the next heading) with the rewritten text. We never touch the
     heading paragraph itself, so its run-level formatting survives.

The output is always ``bytes`` (a complete .docx file).
"""

from __future__ import annotations

import base64
import io
from copy import deepcopy
from typing import Optional

from docx import Document
from docx.text.paragraph import Paragraph

from models import HeadingSlot, TemplateFingerprint


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def render_template(
    fingerprint: TemplateFingerprint,
    rewritten: dict[str, str],
    tables: Optional[dict[str, list[list[str]]]] = None,
) -> bytes:
    """Render the template using ``rewritten`` (slot_id → body text)."""
    if not fingerprint.template_b64:
        # PDF-sourced template or callers stripped the bytes — fall back to
        # a brand-new doc that simulates the structure as best we can.
        return _render_synthesized(fingerprint, rewritten, tables or {})

    raw = base64.b64decode(fingerprint.template_b64)
    if _has_jinja_markers(fingerprint):
        return _render_via_docxtpl(raw, fingerprint, rewritten, tables or {})
    return _render_via_heading_walk(raw, fingerprint, rewritten, tables or {})


# ---------------------------------------------------------------------------
# Mode detection
# ---------------------------------------------------------------------------

def _has_jinja_markers(fingerprint: TemplateFingerprint) -> bool:
    return any(s.placeholder_marker for s in fingerprint.heading_hierarchy)


# ---------------------------------------------------------------------------
# Mode 1 — docxtpl rendering
# ---------------------------------------------------------------------------

def _render_via_docxtpl(
    template_bytes: bytes,
    fingerprint: TemplateFingerprint,
    rewritten: dict[str, str],
    tables: dict[str, list[list[str]]],
) -> bytes:
    from docxtpl import DocxTemplate

    tpl = DocxTemplate(io.BytesIO(template_bytes))
    context = {s.slot_id: rewritten.get(s.slot_id, "") for s in fingerprint.heading_hierarchy}
    # Tables expected as ``table_<slot_id>`` in the Jinja context
    for slot_id, grid in tables.items():
        context[f"table_{slot_id}"] = grid
    tpl.render(context)

    out = io.BytesIO()
    tpl.save(out)
    return out.getvalue()


# ---------------------------------------------------------------------------
# Mode 2 — heading-driven body replacement
# ---------------------------------------------------------------------------

def _render_via_heading_walk(
    template_bytes: bytes,
    fingerprint: TemplateFingerprint,
    rewritten: dict[str, str],
    tables: dict[str, list[list[str]]],
) -> bytes:
    doc = Document(io.BytesIO(template_bytes))
    slots = fingerprint.heading_hierarchy
    slot_by_title = {s.title.strip(): s for s in slots}

    body_paragraphs = list(doc.paragraphs)
    heading_indices: list[tuple[int, HeadingSlot]] = []
    for idx, para in enumerate(body_paragraphs):
        text = (para.text or "").strip()
        if text in slot_by_title:
            heading_indices.append((idx, slot_by_title[text]))

    # Pair each heading with the index of the next heading (or end-of-doc)
    for n, (idx, slot) in enumerate(heading_indices):
        next_idx = (
            heading_indices[n + 1][0]
            if n + 1 < len(heading_indices)
            else len(body_paragraphs)
        )
        new_text = rewritten.get(slot.slot_id, "")
        if not new_text:
            continue
        _replace_body_paragraphs(body_paragraphs[idx + 1 : next_idx], new_text)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _replace_body_paragraphs(paragraphs: list[Paragraph], new_text: str) -> None:
    """Replace the text of ``paragraphs`` (between two headings) with new_text.

    Strategy:
      - Wipe runs from every body paragraph in the range, leaving each
        paragraph's style intact.
      - Write the new content into the first paragraph, splitting on
        newlines. If the new text has more lines than there are body
        paragraphs available, append additional paragraphs cloned from the
        last one to preserve the style.
    """
    if not paragraphs:
        return

    lines = [ln.strip() for ln in (new_text or "").split("\n") if ln.strip()]
    if not lines:
        lines = [""]

    # Clear runs in the target paragraphs
    for p in paragraphs:
        for r in list(p.runs):
            r.text = ""
        # Drop any extra runs by zeroing — python-docx doesn't easily allow
        # full removal, but blanking is sufficient because we re-write text.

    # Anchor paragraph gets the first line (preserve its first run's style).
    anchor = paragraphs[0]
    if anchor.runs:
        anchor.runs[0].text = lines[0]
    else:
        anchor.add_run(lines[0])

    # Remaining lines → re-use available body paragraphs, then clone the
    # anchor for any overflow lines.
    used = 1
    for line in lines[1:]:
        if used < len(paragraphs):
            target = paragraphs[used]
            if target.runs:
                target.runs[0].text = line
            else:
                target.add_run(line)
            used += 1
        else:
            new_p = deepcopy(anchor._p)
            anchor._p.addnext(new_p)
            anchor = Paragraph(new_p, anchor._parent)
            for r in list(anchor.runs):
                r.text = ""
            if anchor.runs:
                anchor.runs[0].text = line
            else:
                anchor.add_run(line)


# ---------------------------------------------------------------------------
# Fallback for PDF-sourced templates (no .docx bytes available)
# ---------------------------------------------------------------------------

def _render_synthesized(
    fingerprint: TemplateFingerprint,
    rewritten: dict[str, str],
    tables: dict[str, list[list[str]]],
) -> bytes:
    """Last-resort rendering: build a new .docx from scratch.

    Style fidelity is necessarily limited — the source template was a PDF or
    its bytes were stripped — so we reuse the existing apply_styling path,
    which produces a reasonable Word document.
    """
    from formater_apply import apply_styling
    from models import (
        ContentElement,
        DocumentContent,
        DocumentMetadata,
        ElementType,
        TextRun,
    )

    elements: list[ContentElement] = []
    for slot in fingerprint.heading_hierarchy:
        elements.append(
            ContentElement(
                type=ElementType.HEADING,
                level=slot.level,
                content=[TextRun(text=slot.title)],
            )
        )
        text = rewritten.get(slot.slot_id, "")
        for para in (text or "").split("\n"):
            if para.strip():
                elements.append(
                    ContentElement(
                        type=ElementType.PARAGRAPH,
                        content=[TextRun(text=para.strip())],
                    )
                )

    content = DocumentContent(
        metadata=DocumentMetadata(
            source_file=fingerprint.metadata.source_file or "synthesized.docx"
        ),
        elements=elements,
    )
    stream = apply_styling(content, fingerprint.style_registry)
    stream.seek(0)
    return stream.read()
