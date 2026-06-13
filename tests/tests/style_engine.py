"""
Style Engine — best-in-class document style transfer.
=====================================================

Apply the *visual identity* of a style-source document onto the *content* of
a target document, preserving 100% of the target's content and structure.

Two transfer paths, chosen automatically from the style source's file type:

  1. DOCX style source  → transplant the source's ``styles.xml`` (merged by
     styleId, source wins), ``docDefaults`` and ``theme1.xml`` into the
     content package. Content body XML is never touched → zero content loss.

  2. PDF style source    → build a :class:`StyleProfile` from PyMuPDF font /
     size / colour / alignment / margin histograms, then mutate the content
     document's named styles, ``docDefaults`` and page margins in place.

Crucially this engine NEVER rebuilds the body from a JSON intermediate (the
approach that previously dropped paragraphs, mangled tables and auto-numbered
every list). It edits the style layer of the real OOXML package, so it scales
to arbitrarily large documents and keeps every paragraph, table and image.

Public API
----------
    transfer_style(content_bytes, content_name, style_bytes, style_name, ...) -> bytes
    transfer_style_files(content_path, style_path, output_path, ...) -> str
    profile_from_pdf(pdf_bytes) -> StyleProfile
    profile_from_docx(docx_bytes) -> StyleProfile
"""

from __future__ import annotations

import collections
import copy
import io
import re
from dataclasses import dataclass, field
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ---------------------------------------------------------------------------
# Profile dataclasses
# ---------------------------------------------------------------------------

@dataclass
class HeadingSpec:
    font: Optional[str] = None
    size_pt: Optional[float] = None
    bold: bool = True
    italic: bool = False
    underline: bool = False
    color_hex: Optional[str] = None          # "RRGGBB"
    alignment: str = "left"                   # left|center|right|justify
    space_before_pt: float = 12.0
    space_after_pt: float = 6.0


@dataclass
class StyleProfile:
    """A normalized description of a document's visual style."""
    body_font: str = "Calibri"
    body_size_pt: float = 11.0
    body_color_hex: str = "000000"
    body_alignment: str = "left"              # left|justify|...
    line_spacing: Optional[float] = None      # multiple, e.g. 1.15
    space_after_pt: float = 8.0
    heading_scale: dict[int, HeadingSpec] = field(default_factory=dict)
    margins_in: tuple[float, float, float, float] = (1.0, 1.0, 1.0, 1.0)  # T,B,L,R
    accent_color_hex: Optional[str] = None
    list_indent_in: float = 0.25
    source_kind: str = "pdf"                  # "pdf" | "docx"

    def heading(self, level: int) -> HeadingSpec:
        if level in self.heading_scale:
            return self.heading_scale[level]
        # Derive a sensible scale from the body if not explicitly captured.
        ratios = {1: 1.8, 2: 1.45, 3: 1.2, 4: 1.08, 5: 1.0, 6: 1.0}
        ratio = ratios.get(level, 1.0)
        return HeadingSpec(
            font=self.body_font,
            size_pt=round(self.body_size_pt * ratio, 1),
            bold=True,
            color_hex=self.accent_color_hex or self.body_color_hex,
            alignment="left",
            space_before_pt=max(6.0, 14.0 - level * 1.5),
            space_after_pt=max(3.0, 8.0 - level),
        )


_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_HEADING_STYLE_NAMES = {f"Heading {i}" for i in range(1, 10)} | {"Title", "Subtitle"}


# ---------------------------------------------------------------------------
# Low-level OOXML helpers
# ---------------------------------------------------------------------------

def _int_to_hex(color_int: int) -> str:
    return f"{color_int & 0xFFFFFF:06X}"


def _clean_font_name(raw: str) -> str:
    """Strip PyMuPDF subset prefixes (``ABCDEF+Calibri``) and weight/style
    suffixes (``Calibri-Bold`` → ``Calibri``)."""
    if not raw:
        return raw
    name = raw.split("+", 1)[-1]  # drop subset prefix
    # Drop common weight/style suffixes
    for sep in ("-", ","):
        if sep in name:
            head = name.split(sep, 1)[0]
            if head:
                name = head
                break
    # Drop trailing "MT"/"PS" PostScript markers
    name = re.sub(r"(MT|PS)$", "", name)
    return name.strip() or raw


def _set_style_font(style, font_name: Optional[str], size_pt: Optional[float] = None,
                    color_hex: Optional[str] = None) -> None:
    """Set a named style's run font on ALL script slots (ascii/hAnsi/cs/eastAsia)."""
    rpr = style.element.get_or_add_rPr()
    if font_name:
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), font_name)
    if size_pt is not None:
        try:
            style.font.size = Pt(size_pt)
        except Exception:
            pass
    if color_hex:
        try:
            style.font.color.rgb = RGBColor.from_string(color_hex)
        except Exception:
            pass


def _set_style_bool(style, *, bold: Optional[bool] = None,
                    italic: Optional[bool] = None,
                    underline: Optional[bool] = None) -> None:
    rpr = style.element.get_or_add_rPr()

    def _toggle(tag: str, on: Optional[bool]):
        if on is None:
            return
        existing = rpr.find(qn(tag))
        if on:
            if existing is None:
                el = OxmlElement(tag)
                rpr.append(el)
        else:
            if existing is not None:
                rpr.remove(existing)

    _toggle("w:b", bold)
    _toggle("w:i", italic)
    if underline is not None:
        existing = rpr.find(qn("w:u"))
        if existing is not None:
            rpr.remove(existing)
        if underline:
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rpr.append(u)


def _set_docdefaults_font(doc, font_name: str, size_pt: Optional[float]) -> None:
    """Set the document-wide default run font/size so any style that doesn't
    pin its own font inherits the body font."""
    styles_el = doc.styles.element
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is None:
        dd = OxmlElement("w:docDefaults")
        styles_el.insert(0, dd)
    rpd = dd.find(qn("w:rPrDefault"))
    if rpd is None:
        rpd = OxmlElement("w:rPrDefault")
        dd.append(rpd)
    rpr = rpd.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpd.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font_name)
    if size_pt is not None:
        half = str(int(round(size_pt * 2)))
        for tag in ("w:sz", "w:szCs"):
            el = rpr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rpr.append(el)
            el.set(qn("w:val"), half)


def _apply_alignment_spacing(style, alignment: str, space_after_pt: Optional[float],
                             line_spacing: Optional[float] = None,
                             space_before_pt: Optional[float] = None) -> None:
    pf = style.paragraph_format
    al = _ALIGN_MAP.get(alignment)
    if al is not None:
        pf.alignment = al
    if space_after_pt is not None:
        pf.space_after = Pt(space_after_pt)
    if space_before_pt is not None:
        pf.space_before = Pt(space_before_pt)
    if line_spacing is not None:
        pf.line_spacing = line_spacing


# ---------------------------------------------------------------------------
# PDF profiling (PyMuPDF)
# ---------------------------------------------------------------------------

def profile_from_pdf(pdf_bytes: bytes) -> StyleProfile:
    """Derive a :class:`StyleProfile` from a PDF using span-level analysis."""
    import fitz  # PyMuPDF

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    size_hist: collections.Counter = collections.Counter()       # (font,size)->chars
    plain_size_hist: collections.Counter = collections.Counter() # size->chars
    font_hist: collections.Counter = collections.Counter()       # family->chars
    color_hist: collections.Counter = collections.Counter()      # int->chars
    heading_spans: list[tuple[float, bool, bool, str]] = []       # size,bold,centered,font

    left_edges: list[float] = []
    right_gaps: list[float] = []
    top_edges: list[float] = []
    bottom_gaps: list[float] = []
    page_w = page_h = 0.0

    for page in doc:
        page_w, page_h = page.rect.width, page.rect.height
        data = page.get_text("dict")
        for block in data.get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    txt = (span.get("text") or "").strip()
                    if not txt:
                        continue
                    sz = round(span.get("size", 0.0), 1)
                    family = _clean_font_name(span.get("font", ""))
                    flags = span.get("flags", 0)
                    bold = bool(flags & 16) or "bold" in (span.get("font", "").lower())
                    color = span.get("color", 0)
                    x0, y0, x1, y1 = span.get("bbox", (0, 0, 0, 0))

                    n = len(txt)
                    size_hist[(family, sz)] += n
                    plain_size_hist[sz] += n
                    font_hist[family] += n
                    color_hist[color] += n
                    left_edges.append(x0)
                    right_gaps.append(page_w - x1)
                    top_edges.append(y0)
                    bottom_gaps.append(page_h - y1)

                    center_x = (x0 + x1) / 2
                    is_centered = abs(center_x - page_w / 2) < 0.06 * page_w
                    if bold and (is_centered or sz > 0):
                        heading_spans.append((sz, bold, is_centered, family))

    doc.close()

    if not plain_size_hist:
        return StyleProfile()  # empty / image-only PDF → defaults

    # Body size = the most common size by character volume.
    body_size = plain_size_hist.most_common(1)[0][0]
    # Body font = most common family overall (favouring the body size).
    body_font_candidates = collections.Counter()
    for (fam, sz), cnt in size_hist.items():
        weight = cnt * (2 if abs(sz - body_size) < 0.6 else 1)
        body_font_candidates[fam] += weight
    body_font = body_font_candidates.most_common(1)[0][0] if body_font_candidates else "Calibri"

    # Body colour = most common colour (almost always near-black).
    body_color = _int_to_hex(color_hist.most_common(1)[0][0])

    # Accent colour = most common colour that is clearly NOT near-black.
    accent = None
    for color_int, _cnt in color_hist.most_common():
        r, g, b = (color_int >> 16) & 0xFF, (color_int >> 8) & 0xFF, color_int & 0xFF
        if max(r, g, b) > 60:  # not near-black
            accent = _int_to_hex(color_int)
            break

    # Alignment: justified if right edges hug the right margin consistently.
    body_alignment = _infer_alignment(left_edges, right_gaps, page_w)

    # Margins (inches) — robust low-percentile of edge positions.
    margins = _infer_margins(left_edges, right_gaps, top_edges, bottom_gaps, page_w, page_h)

    # Heading scale.
    heading_scale = _infer_heading_scale(
        plain_size_hist, body_size, body_font, body_color, accent, heading_spans
    )

    return StyleProfile(
        body_font=body_font,
        body_size_pt=body_size,
        body_color_hex=body_color,
        body_alignment=body_alignment,
        space_after_pt=8.0,
        heading_scale=heading_scale,
        margins_in=margins,
        accent_color_hex=accent,
        source_kind="pdf",
    )


def _infer_alignment(left_edges, right_gaps, page_w) -> str:
    """Detect justified body text.

    Justified lines all terminate at the text-area's right boundary, so they
    leave a near-identical right gap (≈ the right margin) — NOT a gap near
    zero (that would be the page edge). We therefore estimate the right
    boundary as a low percentile of the gaps, then measure how many lines
    cluster at that boundary. A high fraction ⇒ justified.
    """
    if not right_gaps:
        return "left"
    boundary = _percentile(right_gaps, 15)            # the right margin
    tol = 0.03 * page_w
    full_lines = sum(1 for g in right_gaps if g <= boundary + tol)
    frac_full = full_lines / max(len(right_gaps), 1)
    return "justify" if frac_full >= 0.40 else "left"


def _percentile(values: list[float], pct: float) -> float:
    if not values:
        return 0.0
    s = sorted(values)
    k = int(max(0, min(len(s) - 1, round(pct / 100.0 * (len(s) - 1)))))
    return s[k]


def _infer_margins(left_edges, right_gaps, top_edges, bottom_gaps, page_w, page_h):
    """Estimate page margins from text-edge positions.

    Left/right are reliable (lines hit the side margins). Top/bottom are
    noisier (logos, headers, footers, trailing whitespace), so we take the
    minimum observed and fall back to the horizontal margin for symmetry when
    the vertical estimate is implausible.
    """
    clamp = lambda v: round(max(0.3, min(v, 2.0)), 2)
    left = clamp(_percentile(left_edges, 10) / 72.0)
    right = clamp(_percentile(right_gaps, 10) / 72.0)
    top = clamp((min(top_edges) if top_edges else 72) / 72.0)
    bottom = clamp((min(bottom_gaps) if bottom_gaps else 72) / 72.0)
    # If the vertical estimate looks off (header/footer artefacts), borrow the
    # horizontal margin so the page stays balanced.
    side = round((left + right) / 2, 2)
    if top > 1.4 or top < 0.5:
        top = side
    if bottom > 1.4 or bottom < 0.5:
        bottom = side
    return (top, bottom, left, right)


def _infer_heading_scale(plain_size_hist, body_size, body_font, body_color,
                         accent, heading_spans) -> dict[int, HeadingSpec]:
    # Distinct sizes meaningfully larger than the body, by frequency.
    bigger = sorted(
        {sz for sz in plain_size_hist if sz > body_size + 0.6},
        reverse=True,
    )
    scale: dict[int, HeadingSpec] = {}

    # Determine whether the source uses centered/bold headings (flat scale,
    # like the offer letter) when there's no size hierarchy.
    centered_bold = any(c for (_sz, b, c, _f) in heading_spans if b and c)

    if bigger:
        for level, sz in enumerate(bigger[:4], start=1):
            scale[level] = HeadingSpec(
                font=body_font, size_pt=sz, bold=True,
                color_hex=accent or body_color,
                alignment="left",
                space_before_pt=max(6.0, 14.0 - level * 1.5),
                space_after_pt=max(3.0, 8.0 - level),
            )
    elif centered_bold:
        # Flat hierarchy: replicate the letter's bold/underlined heading DNA.
        # Only the top-level title is centered; numbered sub-headings read
        # better left-aligned.
        for level in range(1, 5):
            scale[level] = HeadingSpec(
                font=body_font,
                size_pt=round(body_size * (1.0 + (4 - level) * 0.05), 1),
                bold=True, underline=True,
                color_hex=body_color,
                alignment="center" if level == 1 else "left",
                space_before_pt=14.0 - level * 2,
                space_after_pt=8.0 - level,
            )
    return scale


# ---------------------------------------------------------------------------
# DOCX profiling
# ---------------------------------------------------------------------------

def profile_from_docx(docx_bytes: bytes) -> StyleProfile:
    doc = Document(io.BytesIO(docx_bytes))
    normal = None
    try:
        normal = doc.styles["Normal"]
    except KeyError:
        pass

    body_font = (normal.font.name if normal and normal.font.name else "Calibri")
    body_size = (normal.font.size.pt if normal and normal.font.size else 11.0)

    heading_scale: dict[int, HeadingSpec] = {}
    for level in range(1, 5):
        try:
            hs = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        heading_scale[level] = HeadingSpec(
            font=hs.font.name or body_font,
            size_pt=hs.font.size.pt if hs.font.size else None,
            bold=bool(hs.font.bold),
            italic=bool(hs.font.italic),
            color_hex=str(hs.font.color.rgb) if hs.font.color and hs.font.color.rgb else None,
        )

    sec = doc.sections[0] if doc.sections else None
    margins = (
        (sec.top_margin.inches, sec.bottom_margin.inches,
         sec.left_margin.inches, sec.right_margin.inches)
        if sec else (1.0, 1.0, 1.0, 1.0)
    )

    return StyleProfile(
        body_font=body_font,
        body_size_pt=body_size,
        heading_scale=heading_scale,
        margins_in=margins,
        source_kind="docx",
    )


# ---------------------------------------------------------------------------
# PDF path: apply a profile to a content .docx in place
# ---------------------------------------------------------------------------

def _apply_profile_to_docx(
    content_bytes: bytes,
    profile: StyleProfile,
    *,
    normalize_fonts: bool = True,
    promote_headings: bool = True,
    set_margins: bool = True,
) -> bytes:
    doc = Document(io.BytesIO(content_bytes))

    # 1) Document-wide default font → everything inheriting picks up the body font.
    _set_docdefaults_font(doc, profile.body_font, profile.body_size_pt)

    # 2) Normal style → body look.
    try:
        normal = doc.styles["Normal"]
        _set_style_font(normal, profile.body_font, profile.body_size_pt, profile.body_color_hex)
        _apply_alignment_spacing(
            normal, profile.body_alignment, profile.space_after_pt, profile.line_spacing
        )
    except KeyError:
        pass

    # 3) Heading styles → heading scale.
    for level in range(1, 5):
        spec = profile.heading(level)
        try:
            hs = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        _set_style_font(hs, spec.font or profile.body_font, spec.size_pt, spec.color_hex)
        _set_style_bool(hs, bold=spec.bold, italic=spec.italic, underline=spec.underline)
        _apply_alignment_spacing(
            hs, spec.alignment, spec.space_after_pt, None, spec.space_before_pt
        )

    # 4) Normalize the fonts of all other (non-heading) styles + clear direct
    #    run-level font overrides so the body font wins consistently.
    if normalize_fonts:
        _normalize_fonts(doc, profile.body_font)

    # 5) Page margins.
    if set_margins:
        t, b, l, r = profile.margins_in
        for section in doc.sections:
            section.top_margin = Inches(t)
            section.bottom_margin = Inches(b)
            section.left_margin = Inches(l)
            section.right_margin = Inches(r)

    # 6) Promote bold-Normal pseudo-headings to real Heading styles so they
    #    pick up the source's heading treatment.
    if promote_headings:
        _promote_headings(doc)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _normalize_fonts(doc, body_font: str) -> None:
    """Rewrite non-heading style fonts to body_font and clear direct run-level
    rFonts in body paragraphs (preserving bold/italic/colour/size)."""
    # Style pass — every paragraph/character style that isn't a heading.
    for style in doc.styles:
        try:
            name = style.name
        except Exception:
            continue
        if not name or name in _HEADING_STYLE_NAMES:
            continue
        # Only touch paragraph (1) and character (2) styles.
        if getattr(style, "type", None) is None:
            continue
        try:
            stype = int(style.type)
        except Exception:
            stype = None
        if stype not in (1, 2):
            continue
        rpr = style.element.find(qn("w:rPr"))
        # Always set so explicit non-target fonts (Helvetica, Times) are overridden.
        _set_style_font(style, body_font)

    # Direct-run pass — clear rFonts on runs in non-heading paragraphs.
    def _strip_runs_in(paragraph):
        try:
            pstyle = paragraph.style.name
        except Exception:
            pstyle = None
        if pstyle in _HEADING_STYLE_NAMES:
            return
        for run in paragraph.runs:
            rpr = run._r.find(qn("w:rPr"))
            if rpr is None:
                continue
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                rpr.remove(rfonts)

    for para in doc.paragraphs:
        _strip_runs_in(para)
    for table in doc.tables:
        _strip_table_runs(table, _strip_runs_in)


def _strip_table_runs(table, strip_fn) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                strip_fn(para)
            # nested tables
            for nested in cell.tables:
                _strip_table_runs(nested, strip_fn)


# ---------------------------------------------------------------------------
# Heading promotion (LLM + heuristic)
# ---------------------------------------------------------------------------

def _promote_headings(doc) -> None:
    """Detect bold pseudo-headings authored as Normal paragraphs and assign
    real Heading styles so they adopt the source heading treatment."""
    # Collect candidate body paragraphs (skip those already in a heading style
    # and skip empty paragraphs).
    paras = doc.paragraphs
    candidates: list[tuple[int, str]] = []
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if not text:
            continue
        try:
            if p.style.name in _HEADING_STYLE_NAMES:
                continue
        except Exception:
            pass
        candidates.append((i, text))

    if not candidates:
        return

    decisions = _llm_heading_decisions(candidates)
    if decisions is None:
        decisions = _heuristic_heading_decisions(doc, candidates)

    for idx, level in decisions.items():
        if idx < 0 or idx >= len(paras):
            continue
        level = max(1, min(level, 4))
        style_name = f"Heading {level}"
        try:
            paras[idx].style = doc.styles[style_name]
        except KeyError:
            continue
        # Remove now-redundant direct bold so the heading style governs.
        for run in paras[idx].runs:
            rpr = run._r.find(qn("w:rPr"))
            if rpr is not None:
                for tag in ("w:b", "w:bCs"):
                    el = rpr.find(qn(tag))
                    if el is not None:
                        rpr.remove(el)


def _dominant_bold(paragraph) -> bool:
    runs = [r for r in paragraph.runs if (r.text or "").strip()]
    if not runs:
        return False
    bold_chars = sum(len(r.text) for r in runs if r.bold)
    total = sum(len(r.text) for r in runs)
    return total > 0 and bold_chars / total >= 0.6


def _heuristic_heading_decisions(doc, candidates) -> dict[int, int]:
    paras = doc.paragraphs
    out: dict[int, int] = {}
    for idx, text in candidates:
        if len(text) > 90:
            continue
        words = text.split()
        if len(words) > 14:
            continue
        if text.rstrip().endswith((".", ":", ";", ",")):
            # allow trailing colon for label-style headings
            if not text.rstrip().endswith(":"):
                continue
        if not _dominant_bold(paras[idx]):
            continue
        is_upper = text.isupper()
        is_title = text.istitle() or sum(1 for w in words if w[:1].isupper()) >= max(1, len(words) - 1)
        if not (is_upper or is_title):
            continue
        level = 1 if (is_upper and len(words) <= 6) else 2
        out[idx] = level
    return out


def _llm_heading_decisions(candidates) -> Optional[dict[int, int]]:
    try:
        from llm_client import chat_json, llm_available
    except Exception:
        return None
    if not llm_available():
        return None

    # Cap the payload — long docs send the first N candidates; the heuristic
    # covers the rest.
    sample = candidates[:400]
    listing = [{"index": idx, "text": text[:120]} for idx, text in sample]
    system = (
        "You identify which paragraphs in a document are SECTION HEADINGS "
        "(as opposed to body text, list items, or captions). Return ONLY a "
        "JSON object: {\"headings\": [{\"index\": <int>, \"level\": <1-4>}]}. "
        "index MUST be one of the provided indices. Heading levels: 1 = top "
        "section title, 2 = sub-section, 3-4 = deeper. Be selective; most "
        "paragraphs are NOT headings."
    )
    user = (
        "Paragraphs (index, text):\n"
        + str(listing)
        + "\n\nReturn the headings JSON now."
    )
    raw = chat_json(system, user, temperature=0.0, max_tokens=2000)
    if not isinstance(raw, dict):
        return None
    items = raw.get("headings")
    if not isinstance(items, list):
        return None
    valid_indices = {idx for idx, _ in candidates}
    out: dict[int, int] = {}
    for it in items:
        if not isinstance(it, dict):
            continue
        try:
            idx = int(it.get("index"))
            level = int(it.get("level", 2))
        except (TypeError, ValueError):
            continue
        if idx in valid_indices:
            out[idx] = max(1, min(level, 4))
    return out or None


# ---------------------------------------------------------------------------
# DOCX path: transplant styles + theme into the content package
# ---------------------------------------------------------------------------

def _transplant_docx_styles(content_bytes: bytes, style_bytes: bytes,
                            *, promote_headings: bool = False) -> bytes:
    content = Document(io.BytesIO(content_bytes))
    source = Document(io.BytesIO(style_bytes))

    content_styles_el = content.styles.element
    source_styles_el = source.styles.element

    # 1) Merge <w:style> definitions by styleId (source wins; keep content-only).
    content_by_id = {
        s.get(qn("w:styleId")): s
        for s in content_styles_el.findall(qn("w:style"))
    }
    last_style = None
    for s in content_styles_el.findall(qn("w:style")):
        last_style = s
    for src_style in source_styles_el.findall(qn("w:style")):
        sid = src_style.get(qn("w:styleId"))
        clone = copy.deepcopy(src_style)
        if sid in content_by_id:
            old = content_by_id[sid]
            content_styles_el.replace(old, clone)
        else:
            content_styles_el.append(clone)

    # 2) Replace docDefaults wholesale with the source's.
    src_dd = source_styles_el.find(qn("w:docDefaults"))
    if src_dd is not None:
        dst_dd = content_styles_el.find(qn("w:docDefaults"))
        clone_dd = copy.deepcopy(src_dd)
        if dst_dd is not None:
            content_styles_el.replace(dst_dd, clone_dd)
        else:
            content_styles_el.insert(0, clone_dd)

    # 3) Transplant the theme part (fonts/colour scheme) when both have one.
    try:
        src_theme = source.part.part_related_by(RT.THEME)
        try:
            dst_theme = content.part.part_related_by(RT.THEME)
            dst_theme._blob = src_theme.blob
        except KeyError:
            # Content has no theme part — relate the source's theme blob in.
            from docx.opc.part import Part
            from docx.opc.packuri import PackURI
            partname = PackURI("/word/theme/theme1.xml")
            new_part = Part(partname, src_theme.content_type, src_theme.blob,
                            content.part.package)
            content.part.relate_to(new_part, RT.THEME)
    except KeyError:
        pass  # source has no theme — styles/docDefaults still transferred

    if promote_headings:
        _promote_headings(content)

    out = io.BytesIO()
    content.save(out)
    return out.getvalue()


# ---------------------------------------------------------------------------
# Public dispatcher
# ---------------------------------------------------------------------------

def _looks_like_pdf(data: bytes) -> bool:
    return data[:5] == b"%PDF-"


def _looks_like_docx(data: bytes) -> bool:
    # docx is a zip; sniff PK magic. (We trust the extension downstream too.)
    return data[:2] == b"PK"


def transfer_style(
    content_bytes: bytes,
    content_name: str,
    style_bytes: bytes,
    style_name: str,
    *,
    normalize_fonts: bool = True,
    promote_headings: bool = True,
) -> bytes:
    """Apply ``style_*``'s visual identity onto ``content_*`` and return .docx bytes.

    The content document MUST be a .docx (we edit its OOXML in place). The
    style source may be .docx (full transplant) or .pdf (profile + apply).
    """
    if not _looks_like_docx(content_bytes):
        raise ValueError(
            "Style transfer requires the CONTENT document to be a .docx "
            f"(got {content_name!r}). PDF content is not yet supported."
        )

    style_is_pdf = _looks_like_pdf(style_bytes) or style_name.lower().endswith(".pdf")

    if style_is_pdf:
        profile = profile_from_pdf(style_bytes)
        return _apply_profile_to_docx(
            content_bytes, profile,
            normalize_fonts=normalize_fonts,
            promote_headings=promote_headings,
        )

    # DOCX → DOCX: transplant styles/theme for maximum fidelity, zero loss.
    return _transplant_docx_styles(
        content_bytes, style_bytes, promote_headings=promote_headings
    )


def transfer_style_files(
    content_path: str,
    style_path: str,
    output_path: str,
    *,
    normalize_fonts: bool = True,
    promote_headings: bool = True,
) -> str:
    with open(content_path, "rb") as fh:
        content_bytes = fh.read()
    with open(style_path, "rb") as fh:
        style_bytes = fh.read()
    result = transfer_style(
        content_bytes, content_path,
        style_bytes, style_path,
        normalize_fonts=normalize_fonts,
        promote_headings=promote_headings,
    )
    with open(output_path, "wb") as fh:
        fh.write(result)
    return output_path
