"""
Document Styling Applier
========================
Takes content JSON + styling JSON and generates a fully formatted .docx file.

Capabilities:
  - Apply page setup (size, margins, orientation)
  - Apply paragraph styles (alignment, spacing, indentation, keep-together)
  - Apply run styles (font, size, bold, italic, underline, color, caps, strikethrough)
  - Create tables with cell widths, shading, borders, merges, vertical alignment
  - Insert images from base64 data at correct dimensions
  - Apply list formatting (bullet/numbered)
  - Create headers and footers
  - Cross-document style transfer (content from doc A + styling from doc B)
"""

from __future__ import annotations

import base64
import io
import os
import re
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Pt, RGBColor

from models import (
    Alignment,
    BorderDef,
    BorderStyle,
    CellStyle,
    ContentElement,
    DocumentContent,
    DocumentStyling,
    ElementType,
    ImagePosition,
    IndentStyle,
    ListType,
    Orientation,
    ParagraphStyle,
    RunStyle,
    TableCell,
    TableRow,
    TableStyle,
    TextRun,
    UnderlineType,
    VerticalAlignment,
)


# ---------------------------------------------------------------------------
# Mapping enums back to python-docx constants
# ---------------------------------------------------------------------------

_ALIGNMENT_TO_DOCX = {
    Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
    Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
    Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
    Alignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_UNDERLINE_TO_DOCX = {
    UnderlineType.NONE: False,
    UnderlineType.SINGLE: WD_UNDERLINE.SINGLE,
    UnderlineType.DOUBLE: WD_UNDERLINE.DOUBLE,
    UnderlineType.DOTTED: WD_UNDERLINE.DOTTED,
    UnderlineType.DASHED: WD_UNDERLINE.DASH,
    UnderlineType.WAVY: WD_UNDERLINE.WAVY,
    UnderlineType.THICK: WD_UNDERLINE.THICK,
    UnderlineType.WORDS_ONLY: WD_UNDERLINE.WORDS,
}

_VERT_ALIGN_TO_DOCX = {
    VerticalAlignment.TOP: WD_ALIGN_VERTICAL.TOP,
    VerticalAlignment.CENTER: WD_ALIGN_VERTICAL.CENTER,
    VerticalAlignment.BOTTOM: WD_ALIGN_VERTICAL.BOTTOM,
}

_LINE_SPACING_RULE_MAP = {
    "single": WD_LINE_SPACING.SINGLE,
    "one_point_five": WD_LINE_SPACING.ONE_POINT_FIVE,
    "double": WD_LINE_SPACING.DOUBLE,
    "at_least": WD_LINE_SPACING.AT_LEAST,
    "exactly": WD_LINE_SPACING.EXACTLY,
    "multiple": WD_LINE_SPACING.MULTIPLE,
}


def _hex_to_rgb(hex_color: Optional[str]) -> Optional[RGBColor]:
    """Convert '#RRGGBB' to RGBColor."""
    if not hex_color:
        return None
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        return None
    try:
        r, g, b = int(hex_color[:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
        return RGBColor(r, g, b)
    except ValueError:
        return None


# ---------------------------------------------------------------------------
# Applier class
# ---------------------------------------------------------------------------

class DocumentApplier:
    """Build a formatted .docx from DocumentContent + DocumentStyling."""

    def __init__(self, content: DocumentContent, styling: DocumentStyling):
        self.content = content
        self.styling = styling
        self.doc = Document()
        self._list_counter: dict[int, int] = {}  # level -> current number
        # Set of style names that exist in the doc, so we can prefer native
        # Word list styles ("List Bullet", "List Number") when available.
        self._available_styles: set[str] = set()
        try:
            self._available_styles = {s.name for s in self.doc.styles}
        except Exception:
            self._available_styles = set()
        # Index → "list group id" annotation, populated in apply().
        # Two consecutive LIST_ITEM elements belong to the same group only if
        # they share list_type AND don't go up in level relative to each other.
        self._group_for_element: dict[int, int] = {}
        # Index of the first element in each group, for "should I restart?"
        self._group_first_index: dict[int, int] = {}
        # Memoized restart-numId per (list_type, group_id) so paragraphs in
        # the same group share their dedicated numId.
        self._group_num_id: dict[tuple[str, int], int] = {}
        # The base abstractNumId we clone restarts from. Resolved lazily.
        self._base_abstract_num_id: dict[str, Optional[str]] = {
            "bullet": None,
            "numbered": None,
        }

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def apply(self) -> io.BytesIO:
        """
        Generate a .docx from content + styling.

        Returns:
            BytesIO stream containing the .docx file.
        """
        # 1) Page setup
        self._apply_page_style()

        # 2) Pre-annotate list groups so consecutive list items render as a
        #    single Word list, but separate runs of items restart numbering.
        self._annotate_list_groups()

        # 3) Process each element
        for idx, element in enumerate(self.content.elements):
            self._current_index = idx
            self._process_element(element)

        # 4) Save to stream
        stream = io.BytesIO()
        self.doc.save(stream)
        stream.seek(0)
        return stream

    # ------------------------------------------------------------------
    # List-group annotation
    # ------------------------------------------------------------------

    def _annotate_list_groups(self) -> None:
        """Walk content.elements once and tag each list item with a group id.

        A new group starts whenever:
          - the previous element wasn't a list item, OR
          - the list_type changed (bullet ↔ numbered), OR
          - the new level is *less than* the previous (climbing back out of
            an indent is a new top-level run).
        """
        group_id = 0
        prev_was_list = False
        prev_type = None
        prev_level: Optional[int] = None

        for idx, el in enumerate(self.content.elements):
            if el.type != ElementType.LIST_ITEM:
                prev_was_list = False
                prev_type = None
                prev_level = None
                continue

            cur_level = el.list_level or 0
            start_new = (
                not prev_was_list
                or el.list_type != prev_type
                or (prev_level is not None and cur_level < prev_level)
            )
            if start_new:
                group_id += 1
                self._group_first_index[group_id] = idx
            self._group_for_element[idx] = group_id

            prev_was_list = True
            prev_type = el.list_type
            prev_level = cur_level

    def save_to_file(self, file_path: str) -> str:
        """Save the generated document to a file."""
        stream = self.apply()
        with open(file_path, "wb") as f:
            f.write(stream.read())
        return file_path

    # ------------------------------------------------------------------
    # Page setup
    # ------------------------------------------------------------------

    def _apply_page_style(self):
        """Apply page dimensions, margins, and orientation."""
        ps = self.styling.page_style
        section = self.doc.sections[0]

        # Orientation
        if ps.orientation == Orientation.LANDSCAPE:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(ps.height_inches)
            section.page_height = Inches(ps.width_inches)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Inches(ps.width_inches)
            section.page_height = Inches(ps.height_inches)

        # Margins
        margins = ps.margins
        section.top_margin = Inches(margins.top_inches)
        section.bottom_margin = Inches(margins.bottom_inches)
        section.left_margin = Inches(margins.left_inches)
        section.right_margin = Inches(margins.right_inches)

    # ------------------------------------------------------------------
    # Element dispatch
    # ------------------------------------------------------------------

    def _process_element(self, element: ContentElement):
        """Route an element to the appropriate handler."""
        handlers = {
            ElementType.HEADING: self._add_heading,
            ElementType.PARAGRAPH: self._add_paragraph,
            ElementType.IMAGE: self._add_image,
            ElementType.TABLE: self._add_table,
            ElementType.LIST_ITEM: self._add_list_item,
            ElementType.PAGE_BREAK: self._add_page_break,
            ElementType.HEADER: self._add_header,
            ElementType.FOOTER: self._add_footer,
            ElementType.HYPERLINK: self._add_paragraph,  # fallback
        }
        handler = handlers.get(element.type, self._add_paragraph)
        handler(element)

    # ------------------------------------------------------------------
    # Heading
    # ------------------------------------------------------------------

    def _add_heading(self, element: ContentElement):
        level = element.level or 1
        level = max(1, min(level, 9))

        para = self.doc.add_heading(level=level)
        # Clear default run that add_heading may create
        para.clear()

        # Apply paragraph style
        self._apply_para_style(para, element.style_ref, element.inline_style)

        # Add runs
        if element.content:
            for text_run in element.content:
                self._add_run(para, text_run)

    # ------------------------------------------------------------------
    # Paragraph
    # ------------------------------------------------------------------

    def _add_paragraph(self, element: ContentElement):
        para = self.doc.add_paragraph()

        # Apply paragraph style
        self._apply_para_style(para, element.style_ref, element.inline_style)

        # Add runs
        if element.content:
            for text_run in element.content:
                # Check if this is an inline image marker
                if text_run.style_ref == "__image__":
                    # Skip image markers in paragraph context — images are separate elements
                    continue
                self._add_run(para, text_run)

    # ------------------------------------------------------------------
    # List item
    # ------------------------------------------------------------------

    # Aggressive leading-marker cleanup: strips ANY common bullet glyph or
    # numeric/alpha-list prefix from the start of a run's text. We don't trust
    # the extractor's bullet_char to match what's actually in the run.
    _LIST_LEADER_RE = re.compile(
        r"^\s*(?:"
        r"[•●○◦▪■□◆◇★☆▶►▷▸▹·\-–—*‐-―]"  # bullet-ish glyphs
        r"|"
        r"\(?\s*(?:\d{1,3}|[ivxlcdmIVXLCDM]+|[a-zA-Z])\s*[.)\]]"  # 1. / a) / (iv)
        r")\s+"
    )

    def _strip_leader(self, text: str) -> str:
        return self._LIST_LEADER_RE.sub("", text or "", count=1)

    def _add_list_item(self, element: ContentElement):
        """Add a list item using a real Word list style when available."""
        level = max(0, min(element.list_level or 0, 8))
        idx = getattr(self, "_current_index", 0)
        group_id = self._group_for_element.get(idx, 0)
        is_first_in_group = self._group_first_index.get(group_id) == idx

        # Reset counters at every group boundary so two separate numbered
        # lists in the same doc both start at "1.".
        if is_first_in_group:
            self._list_counter = {}

        list_type = element.list_type or ListType.BULLET

        # Prefer Word's built-in named list styles. They produce real lists
        # (proper indents, hanging numbers, restart-aware) instead of text
        # with prepended characters.
        if list_type == ListType.BULLET:
            style_name = self._pick_native_list_style(
                ["List Bullet", "List Bullet 2", "List Bullet 3"], level
            )
        else:
            style_name = self._pick_native_list_style(
                ["List Number", "List Number 2", "List Number 3"], level
            )

        para = self.doc.add_paragraph()
        if style_name:
            try:
                para.style = self.doc.styles[style_name]
            except KeyError:
                style_name = None  # fall through to manual rendering

        # Override numPr so this group gets its own numId, restarting at 1.
        # We do this for every list item (not just the first) so all items
        # in the group share the same dedicated numId.
        if style_name:
            kind = "bullet" if list_type == ListType.BULLET else "numbered"
            num_id = self._ensure_group_num_id(kind, group_id)
            if num_id:
                self._apply_num_pr(para, num_id=num_id, ilvl=level)

        self._apply_para_style(para, element.style_ref, element.inline_style)

        # If we couldn't apply a native list style, fall back to manual
        # rendering (indent + prepended leader). Counter is still reset
        # per-group thanks to the block above.
        if not style_name:
            pf = para.paragraph_format
            pf.left_indent = Inches(0.25 + level * 0.5)
            pf.first_line_indent = Inches(-0.25)

            if list_type == ListType.BULLET:
                leader = (element.bullet_char or "•") + " "
            else:
                self._list_counter[level] = self._list_counter.get(level, 0) + 1
                leader = f"{self._list_counter[level]}. "

            leader_run = para.add_run(leader)
            self._apply_run_style_to_run(leader_run, RunStyle(bold=False))

        # Emit the body runs, stripping any leading bullet/number that the
        # extractor preserved in the run text.
        emitted_text = False
        if element.content:
            for run_idx, text_run in enumerate(element.content):
                text = text_run.text or ""
                if run_idx == 0:
                    text = self._strip_leader(text)
                if text:
                    run = para.add_run(text)
                    self._apply_run_style(run, text_run.style_ref, text_run.inline_style)
                    emitted_text = True

        # Empty list items would render as just the bullet/number, which
        # looks broken — drop the paragraph if there's nothing in it.
        if not emitted_text and not style_name:
            try:
                para._p.getparent().remove(para._p)
            except Exception:
                pass

    # ------------------------------------------------------------------
    # numbering.xml manipulation — restart numbering per group
    # ------------------------------------------------------------------

    def _numbering_element(self):
        """Return the root <w:numbering> XML element, or None if absent."""
        try:
            numbering_part = self.doc.part.numbering_part
        except Exception:
            return None
        if numbering_part is None:
            return None
        return numbering_part.element

    def _find_base_abstract_num_id(self, kind: str) -> Optional[str]:
        """Find an existing abstractNumId whose first level matches ``kind``.

        ``kind`` is 'bullet' or 'numbered'. We pick the first abstractNum
        whose <w:lvl ilvl="0"><w:numFmt w:val="..."/></w:lvl> matches.
        """
        cached = self._base_abstract_num_id.get(kind)
        if cached is not None:
            return cached

        numbering = self._numbering_element()
        if numbering is None:
            return None

        want = "bullet" if kind == "bullet" else None  # numbered → any decimal-ish
        for an in numbering.findall(qn("w:abstractNum")):
            an_id = an.get(qn("w:abstractNumId"))
            lvl0 = an.find(qn("w:lvl"))
            if lvl0 is None:
                continue
            fmt_el = lvl0.find(qn("w:numFmt"))
            fmt = fmt_el.get(qn("w:val")) if fmt_el is not None else None
            if kind == "bullet" and fmt == "bullet":
                self._base_abstract_num_id[kind] = an_id
                return an_id
            if kind == "numbered" and fmt and fmt != "bullet":
                self._base_abstract_num_id[kind] = an_id
                return an_id

        # Fallback — return the first abstract num if any.
        any_an = numbering.find(qn("w:abstractNum"))
        if any_an is not None:
            self._base_abstract_num_id[kind] = any_an.get(qn("w:abstractNumId"))
            return self._base_abstract_num_id[kind]
        return None

    def _make_restart_num_id(self, abstract_num_id: str) -> int:
        """Append a new <w:num> to numbering.xml that overrides ilvl=0..3
        to restart at 1. Returns the new numId.
        """
        numbering = self._numbering_element()
        if numbering is None:
            return 0

        existing_ids = [
            int(n.get(qn("w:numId")) or 0)
            for n in numbering.findall(qn("w:num"))
        ]
        new_id = (max(existing_ids) if existing_ids else 0) + 1

        new_num = OxmlElement("w:num")
        new_num.set(qn("w:numId"), str(new_id))

        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), abstract_num_id)
        new_num.append(abs_ref)

        # Restart every level we might use back to 1.
        for ilvl in range(0, 4):
            override = OxmlElement("w:lvlOverride")
            override.set(qn("w:ilvl"), str(ilvl))
            start = OxmlElement("w:startOverride")
            start.set(qn("w:val"), "1")
            override.append(start)
            new_num.append(override)

        numbering.append(new_num)
        return new_id

    def _ensure_group_num_id(self, kind: str, group_id: int) -> Optional[int]:
        """Return the numId to use for this (kind, group_id), creating it
        on first use. ``kind`` is 'bullet' or 'numbered'.
        """
        key = (kind, group_id)
        if key in self._group_num_id:
            return self._group_num_id[key]
        abstract = self._find_base_abstract_num_id(kind)
        if abstract is None:
            return None
        # First group reuses the existing num that already points to this
        # abstract num — we only mint restart-overrides for groups 2+.
        if group_id == 1:
            numbering = self._numbering_element()
            if numbering is not None:
                for n in numbering.findall(qn("w:num")):
                    ref = n.find(qn("w:abstractNumId"))
                    if ref is not None and ref.get(qn("w:val")) == abstract:
                        self._group_num_id[key] = int(n.get(qn("w:numId")) or 0)
                        return self._group_num_id[key]
        new_id = self._make_restart_num_id(abstract)
        self._group_num_id[key] = new_id
        return new_id

    def _apply_num_pr(self, para, num_id: int, ilvl: int) -> None:
        """Attach <w:numPr> to the paragraph so it joins ``num_id`` at ``ilvl``."""
        pPr = para._p.get_or_add_pPr()
        # Remove any existing numPr so we don't double-up
        existing = pPr.find(qn("w:numPr"))
        if existing is not None:
            pPr.remove(existing)
        numPr = OxmlElement("w:numPr")
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), str(ilvl))
        num_el = OxmlElement("w:numId")
        num_el.set(qn("w:val"), str(num_id))
        numPr.append(ilvl_el)
        numPr.append(num_el)
        pPr.append(numPr)

    def _pick_native_list_style(
        self, candidates: list[str], level: int
    ) -> Optional[str]:
        """Return the first candidate list style that exists in the doc.

        candidates[0] is level-1 (e.g. "List Bullet"); candidates[1] is
        level-2 ("List Bullet 2"); etc. Falls back down the list when the
        exact level style isn't in the doc.
        """
        if not self._available_styles:
            return None
        # Try the level-matched name first, then shallower variants.
        order = list(range(min(level, len(candidates) - 1), -1, -1))
        for i in order:
            name = candidates[i]
            if name in self._available_styles:
                return name
        return None

    # ------------------------------------------------------------------
    # Image
    # ------------------------------------------------------------------

    def _add_image(self, element: ContentElement):
        """Insert an image from base64 data."""
        if not element.data_base64:
            return

        try:
            image_bytes = base64.b64decode(element.data_base64)
            image_stream = io.BytesIO(image_bytes)

            width = Inches(element.width_inches) if element.width_inches else Inches(4.0)
            height = Inches(element.height_inches) if element.height_inches else None

            # Limit width to page width minus margins
            ps = self.styling.page_style
            max_width_inches = ps.width_inches - ps.margins.left_inches - ps.margins.right_inches
            if element.width_inches and element.width_inches > max_width_inches:
                width = Inches(max_width_inches)
                # Scale height proportionally
                if element.width_inches and element.height_inches:
                    scale = max_width_inches / element.width_inches
                    height = Inches(element.height_inches * scale)

            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(image_stream, width=width, height=height)

        except Exception as e:
            # If image fails, add a placeholder
            para = self.doc.add_paragraph()
            run = para.add_run(f"[Image: {element.alt_text or 'unable to render'}]")
            run.italic = True

    # ------------------------------------------------------------------
    # Table
    # ------------------------------------------------------------------

    def _add_table(self, element: ContentElement):
        """Create a table with full formatting."""
        if not element.rows:
            return

        num_cols = max(len(row.cells) for row in element.rows)
        num_rows = len(element.rows)

        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"  # Ensures visible borders by default

        # Apply table-level style
        if element.table_style_ref and element.table_style_ref in self.styling.table_styles:
            ts = self.styling.table_styles[element.table_style_ref]
            self._apply_table_style(table, ts)
        elif element.inline_table_style:
            self._apply_table_style(table, element.inline_table_style)

        # Populate cells
        for row_idx, row_data in enumerate(element.rows):
            for col_idx, cell_data in enumerate(row_data.cells):
                if col_idx >= num_cols:
                    break

                cell = table.cell(row_idx, col_idx)

                # Clear default paragraph
                if cell.paragraphs:
                    cell.paragraphs[0].clear()
                    para = cell.paragraphs[0]
                else:
                    para = cell.add_paragraph()

                # Add content
                for text_run in cell_data.content:
                    run = para.add_run(text_run.text)
                    self._apply_run_style(run, text_run.style_ref, text_run.inline_style)

                # Apply cell style
                cell_style = cell_data.inline_style
                if cell_data.style_ref and cell_data.style_ref in self.styling.cell_styles:
                    cell_style = self.styling.cell_styles[cell_data.style_ref]

                if cell_style:
                    self._apply_cell_style(cell, cell_style)

                # Bold header row
                if row_data.is_header:
                    for run in para.runs:
                        run.bold = True

    def _apply_table_style(self, table, ts: TableStyle):
        """Apply table-level formatting."""
        if ts.alignment:
            align_map = {
                Alignment.LEFT: WD_TABLE_ALIGNMENT.LEFT,
                Alignment.CENTER: WD_TABLE_ALIGNMENT.CENTER,
                Alignment.RIGHT: WD_TABLE_ALIGNMENT.RIGHT,
            }
            if ts.alignment in align_map:
                table.alignment = align_map[ts.alignment]

    def _apply_cell_style(self, cell, cs: CellStyle):
        """Apply cell-level formatting."""
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()

        # Width
        if cs.width_inches:
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(int(cs.width_inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            tc_pr.append(tc_w)

        # Shading
        if cs.shading_color_hex:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), cs.shading_color_hex.lstrip("#"))
            tc_pr.append(shd)

        # Vertical alignment
        if cs.vertical_alignment:
            v_align = OxmlElement("w:vAlign")
            va_map = {
                VerticalAlignment.TOP: "top",
                VerticalAlignment.CENTER: "center",
                VerticalAlignment.BOTTOM: "bottom",
            }
            v_align.set(qn("w:val"), va_map.get(cs.vertical_alignment, "top"))
            tc_pr.append(v_align)

        # Cell borders
        if cs.borders:
            tc_borders = OxmlElement("w:tcBorders")
            for side_name in ["top", "bottom", "left", "right"]:
                border_def = getattr(cs.borders, side_name, None)
                if border_def:
                    self._add_border_element(tc_borders, side_name, border_def)
            tc_pr.append(tc_borders)

    def _add_border_element(self, parent, side: str, border_def: BorderDef):
        """Add a border definition element."""
        border = OxmlElement(f"w:{side}")
        style_map = {
            BorderStyle.NONE: "none",
            BorderStyle.SINGLE: "single",
            BorderStyle.DOUBLE: "double",
            BorderStyle.DOTTED: "dotted",
            BorderStyle.DASHED: "dashed",
            BorderStyle.THICK: "thick",
        }
        border.set(qn("w:val"), style_map.get(border_def.style, "single"))
        if border_def.width_pt:
            # Word uses eighth-points
            border.set(qn("w:sz"), str(int(border_def.width_pt * 8)))
        if border_def.color_hex:
            border.set(qn("w:color"), border_def.color_hex.lstrip("#"))
        border.set(qn("w:space"), "0")
        parent.append(border)

    # ------------------------------------------------------------------
    # Page break
    # ------------------------------------------------------------------

    def _add_page_break(self, _element: ContentElement):
        self.doc.add_page_break()

    # ------------------------------------------------------------------
    # Header / Footer
    # ------------------------------------------------------------------

    def _add_header(self, element: ContentElement):
        """Add content to the document header."""
        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False

        if header.paragraphs:
            para = header.paragraphs[0]
        else:
            para = header.add_paragraph()

        self._apply_para_style(para, element.style_ref, element.inline_style)

        if element.content:
            for text_run in element.content:
                self._add_run(para, text_run)

    def _add_footer(self, element: ContentElement):
        """Add content to the document footer."""
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()

        self._apply_para_style(para, element.style_ref, element.inline_style)

        if element.content:
            for text_run in element.content:
                self._add_run(para, text_run)

    # ------------------------------------------------------------------
    # Run creation
    # ------------------------------------------------------------------

    def _add_run(self, para, text_run: TextRun):
        """Add a text run to a paragraph with formatting."""
        run = para.add_run(text_run.text)
        self._apply_run_style(run, text_run.style_ref, text_run.inline_style)

        # Hyperlink
        if text_run.hyperlink_url:
            self._add_hyperlink(para, run, text_run.hyperlink_url)

    def _add_hyperlink(self, paragraph, run, url: str):
        """Add a hyperlink to a run. Since python-docx doesn't support
        hyperlinks natively, we modify the XML directly."""
        try:
            part = paragraph.part
            r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)

            # Move the run element into the hyperlink
            run._element.getparent().remove(run._element)
            hyperlink.append(run._element)
            paragraph._element.append(hyperlink)

            # Style the hyperlink text
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.font.underline = True
        except Exception:
            pass  # If hyperlink fails, the text is still there

    # ------------------------------------------------------------------
    # Style application
    # ------------------------------------------------------------------

    def _apply_para_style(
        self,
        para,
        style_ref: Optional[str],
        inline_style: Optional[ParagraphStyle],
    ):
        """Apply paragraph-level formatting."""
        # Resolve style: inline takes precedence, then ref
        ps = inline_style
        if style_ref and style_ref in self.styling.paragraph_styles:
            ref_style = self.styling.paragraph_styles[style_ref]
            if ps is None:
                ps = ref_style
            else:
                # Merge: inline overrides ref
                ps = self._merge_para_styles(ref_style, ps)

        if ps is None:
            return

        pf = para.paragraph_format

        if ps.alignment is not None:
            pf.alignment = _ALIGNMENT_TO_DOCX.get(ps.alignment)

        # Defensive clamp: a "points" value greater than ~10000 is almost
        # certainly raw EMU sneaking through from an upstream extractor bug.
        # Word's own UI caps paragraph spacing at 1584pt — anything above
        # 2000pt would render as feet of blank space.
        if ps.space_before_pt is not None:
            sb = float(ps.space_before_pt)
            if sb > 2000:
                sb = sb / 12700.0  # treat as EMU and convert to pt
            pf.space_before = Pt(max(0.0, min(sb, 1584.0)))

        if ps.space_after_pt is not None:
            sa = float(ps.space_after_pt)
            if sa > 2000:
                sa = sa / 12700.0
            pf.space_after = Pt(max(0.0, min(sa, 1584.0)))

        if ps.line_spacing is not None:
            if ps.line_spacing_rule:
                rule = _LINE_SPACING_RULE_MAP.get(ps.line_spacing_rule)
                if rule:
                    pf.line_spacing_rule = rule
            pf.line_spacing = ps.line_spacing

        if ps.indent:
            ind = ps.indent
            if ind.left_inches is not None:
                pf.left_indent = Inches(ind.left_inches)
            if ind.right_inches is not None:
                pf.right_indent = Inches(ind.right_inches)
            if ind.first_line_inches is not None:
                pf.first_line_indent = Inches(ind.first_line_inches)
            if ind.hanging_inches is not None:
                pf.first_line_indent = Inches(-ind.hanging_inches)

        if ps.keep_with_next is not None:
            pf.keep_with_next = ps.keep_with_next

        if ps.keep_together is not None:
            pf.keep_together = ps.keep_together

        if ps.widow_control is not None:
            pf.widow_control = ps.widow_control

    def _apply_run_style(
        self,
        run,
        style_ref: Optional[str],
        inline_style: Optional[RunStyle],
    ):
        """Apply run-level formatting."""
        rs = inline_style
        if style_ref and style_ref in self.styling.run_styles:
            ref_style = self.styling.run_styles[style_ref]
            if rs is None:
                rs = ref_style
            else:
                rs = self._merge_run_styles(ref_style, rs)

        if rs is None:
            return

        self._apply_run_style_to_run(run, rs)

    def _apply_run_style_to_run(self, run, rs: RunStyle):
        """Apply a RunStyle directly to a python-docx Run object."""
        font = run.font

        if rs.font_name and not rs.font_name.startswith("__"):
            font.name = rs.font_name

        if rs.font_size_pt is not None:
            # Same defensive clamp as paragraph spacing: a "pt" value above
            # ~1638 (Word's max) is raw EMU leaking through.
            fs = float(rs.font_size_pt)
            if fs > 2000:
                fs = fs / 12700.0
            font.size = Pt(max(1.0, min(fs, 1638.0)))

        if rs.bold is not None:
            font.bold = rs.bold

        if rs.italic is not None:
            font.italic = rs.italic

        if rs.underline is not None:
            docx_underline = _UNDERLINE_TO_DOCX.get(rs.underline)
            if docx_underline is not None:
                font.underline = docx_underline

        if rs.strikethrough is not None:
            font.strike = rs.strikethrough

        if rs.double_strikethrough is not None:
            font.double_strike = rs.double_strikethrough

        if rs.superscript is not None:
            font.superscript = rs.superscript

        if rs.subscript is not None:
            font.subscript = rs.subscript

        if rs.color_hex:
            rgb = _hex_to_rgb(rs.color_hex)
            if rgb:
                font.color.rgb = rgb

        if rs.all_caps is not None:
            font.all_caps = rs.all_caps

        if rs.small_caps is not None:
            font.small_caps = rs.small_caps

    # ------------------------------------------------------------------
    # Style merging
    # ------------------------------------------------------------------

    def _merge_para_styles(self, base: ParagraphStyle, override: ParagraphStyle) -> ParagraphStyle:
        """Merge two paragraph styles — override wins for non-None fields."""
        merged_data = base.model_dump(exclude_none=True)
        override_data = override.model_dump(exclude_none=True)
        merged_data.update(override_data)
        return ParagraphStyle(**merged_data)

    def _merge_run_styles(self, base: RunStyle, override: RunStyle) -> RunStyle:
        """Merge two run styles — override wins for non-None fields."""
        merged_data = base.model_dump(exclude_none=True)
        override_data = override.model_dump(exclude_none=True)
        merged_data.update(override_data)
        return RunStyle(**merged_data)


# ---------------------------------------------------------------------------
# Convenience functions
# ---------------------------------------------------------------------------

def apply_styling(
    content: DocumentContent,
    styling: DocumentStyling,
    *,
    use_llm: Optional[bool] = None,
) -> io.BytesIO:
    """
    Apply styling to content and return a .docx as BytesIO.

    Args:
        content: The document content (headings, paragraphs, images, tables).
        styling: The styling definitions (fonts, sizes, colors, spacing).
        use_llm: If True, run the document-understanding pre-pass that asks
            Azure OpenAI to fix mis-classified elements (heading levels,
            missed list items, etc.) before rendering. If None (default),
            it's enabled automatically whenever the Azure OpenAI deployment
            is configured. Pass False to force the pre-pass off.

    Returns:
        BytesIO stream containing the formatted .docx.
    """
    if use_llm is None:
        try:
            from llm_client import llm_available

            use_llm = llm_available()
        except Exception:
            use_llm = False
    if use_llm:
        try:
            from doc_understanding import refine_document

            content = refine_document(content)
        except Exception:
            # Refinement is best-effort; never block on it.
            pass

    applier = DocumentApplier(content, styling)
    return applier.apply()


def apply_styling_to_file(
    content: DocumentContent,
    styling: DocumentStyling,
    output_path: str,
    *,
    use_llm: Optional[bool] = None,
) -> str:
    """
    Apply styling to content and save as a .docx file.

    Returns:
        The output file path.
    """
    stream = apply_styling(content, styling, use_llm=use_llm)
    with open(output_path, "wb") as f:
        f.write(stream.read())
    return output_path


def style_transfer(
    source_content: DocumentContent,
    target_styling: DocumentStyling,
) -> io.BytesIO:
    """
    DEPRECATED low-fidelity style transfer (content JSON + styling JSON).

    This rebuilds the document from a JSON intermediate and is inherently
    lossy (it can drop paragraphs, mangle tables, and auto-number lists).
    For real document-to-document style transfer use
    ``style_engine.transfer_style`` which edits the content document's OOXML
    package in place and preserves 100% of the content.

    Retained only for backward compatibility with callers that already hold
    extracted JSON objects rather than the original file bytes.
    """
    remapped_content = _remap_style_refs(source_content, target_styling)
    return apply_styling(remapped_content, target_styling)


def _remap_style_refs(
    content: DocumentContent,
    target_styling: DocumentStyling,
) -> DocumentContent:
    """
    Remap style references in content to match available styles in target_styling.

    Strategy:
    - For paragraph styles: map by category (heading → heading, body → body)
    - For run styles: map by similarity (font size, bold/italic flags)
    """
    # Build lookup maps from target styling
    target_para_keys = list(target_styling.paragraph_styles.keys())
    target_run_keys = list(target_styling.run_styles.keys())

    # Categorize target paragraph styles
    heading_para_keys = [k for k in target_para_keys if "heading" in k.lower() or "head" in k.lower()]
    body_para_keys = [k for k in target_para_keys if k not in heading_para_keys]

    def remap_para_ref(ref: Optional[str], element_type: ElementType) -> Optional[str]:
        if ref and ref in target_styling.paragraph_styles:
            return ref  # Already exists in target
        if element_type == ElementType.HEADING:
            return heading_para_keys[0] if heading_para_keys else (target_para_keys[0] if target_para_keys else None)
        return body_para_keys[0] if body_para_keys else (target_para_keys[0] if target_para_keys else None)

    def remap_run_ref(ref: Optional[str]) -> Optional[str]:
        if ref and ref in target_styling.run_styles:
            return ref
        return target_run_keys[0] if target_run_keys else None

    # Deep copy and remap
    import copy
    remapped = copy.deepcopy(content)

    for element in remapped.elements:
        element.style_ref = remap_para_ref(element.style_ref, element.type)

        if element.content:
            for run in element.content:
                run.style_ref = remap_run_ref(run.style_ref)

        if element.rows:
            for row in element.rows:
                for cell in row.cells:
                    for run in cell.content:
                        run.style_ref = remap_run_ref(run.style_ref)

    return remapped


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) < 3:
        print("Usage: python formater_apply.py <content.json> <styling.json> [output.docx]")
        sys.exit(1)

    content_path = sys.argv[1]
    styling_path = sys.argv[2]
    output_path = sys.argv[3] if len(sys.argv) > 3 else "output.docx"

    with open(content_path, "r", encoding="utf-8") as f:
        content_data = json.load(f)
    with open(styling_path, "r", encoding="utf-8") as f:
        styling_data = json.load(f)

    content = DocumentContent(**content_data)
    styling = DocumentStyling(**styling_data)

    result_path = apply_styling_to_file(content, styling, output_path)
    print(f"Formatted document written to: {result_path}")
